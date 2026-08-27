from __future__ import annotations

import argparse
import json
import math
import os
import re
import subprocess
import sys
from datetime import date, timedelta
from html import escape
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    CondPageBreak,
    Image as PdfImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

if str(Path(__file__).resolve().parent) not in sys.path:
    sys.path.insert(0, str(Path(__file__).resolve().parent))

from build_integrated_srs import target_job  # noqa: E402
from lldd_skeleton_be import be_skeleton_blocks  # noqa: E402
from lldd_skeleton_fe import fe_skeleton_blocks  # noqa: E402
from lldd_skeleton_job import job_skeleton_blocks  # noqa: E402


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "LLDD"
LEGACY_JAVA_ROOT = ROOT.parent / "fcsJar"
FORMAT_DIRS = {"md": "md", "docx": "word", "pdf": "pdf"}
IMG = ROOT / "output/srs/screenshots/full"
SLICE = ROOT / "output/srs/screenshots/slices"
FONT = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
# 2026-08-07: re-baseline วันเริ่มเป็นจันทร์ 10/08/2026 (วันเริ่มเดิม 29/07/2026 ผ่านมาแล้ว
# ทำให้แผนมี End Date เป็นอดีตตั้งแต่วันส่งมอบ) · เส้นวิกฤตคือ Aphiwit 130 ชม. = 22 วันทำงาน -> 08/09/2026
LLDD_START_DATE = date(2026, 8, 10)
# 2026-08-10: เลิกใช้เป็น "กรอบบังคับ" — วันจบมาจาก planned_finish_date() ที่คำนวณจาก dependency จริง
LLDD_END_DATE = date(2026, 9, 8)  # คงไว้เป็นวันเป้าหมายเดิมเพื่อเทียบเท่านั้น
WORKDAYS_PER_WEEK = 5
# 2026-08-25: กรอบส่งมอบ "จบใน 4 สัปดาห์ · 5 วัน/สัปดาห์ · 8.5 ชม./วัน"
#             = 170 ชม./คน · ทีม 6 คน = 1,020 ชม. เทียบงาน 824 ชม. (81% utilisation)
#             ชั่วโมงงานรวมไม่เปลี่ยน — เปลี่ยนแค่ตัวหารที่แปลงชั่วโมงเป็นวัน
HOURS_PER_DAY = 8.5
HOURS_PER_WEEK = WORKDAYS_PER_WEEK * HOURS_PER_DAY
# 2026-08-11: ถอดกรอบ work-week ออก — เดิมบังคับให้ทุกคนอยู่ระหว่าง >3 ถึง <=4.5 สัปดาห์
#             ซึ่งทำให้ต้องปรับชั่วโมงให้ "พอดีกรอบ" แทนที่จะประเมินตามเนื้องานจริง
MIN_WORK_WEEKS_EXCLUSIVE = 0.0   # ไม่บังคับแล้ว
MAX_WORK_WEEKS = 0.0             # ไม่บังคับแล้ว
FE_OWNER_KITTISAK = "Kittisak <New> Kaeowika"
# 2026-08-07: Peerakorn ย้ายจากสาย FE ไปสาย BE (รับ Attachment / Report-and-Master-Data /
# Job-Batch-SRM / Job 5,7,9,10) เพื่อเปิดที่ให้ Aphiwit รับ Database-Structure + Data-Migration
BE_OWNER_PEERAKORN = "Peerakorn <Pete> Sakunkaewphithak"
FE_OWNER = "Chidchanok <lin> Saengamnat"
BANK_BE_OWNER = "Aphiwit <Bank> Khammoon"
BE_OWNER = "Tunyatorn <Vava> Kiatkongphongsa"
BE_OWNER_BUTSABA = "Butsaba <But> Podamrong"
ATTACHMENT_ALLOWED_EXTENSIONS = "vsd, dwg, afp, pdf, mda, zip, wav, mp3, gif, jpg, tif, tiff, htm, html, txt, xml, mpg, mov, ivs, doc, docx, xls, xlsx, pps, ppt, pot, csv"
# 2026-08-07: ปรับชั่วโมง/เจ้าของใหม่ · 2026-08-11: รวม 703 ชม. (ถอด buffer ออก)
# Job 8b -> Tunyatorn (job เดียวที่เรียก workflow engine · ถือ Workflow-Engine-Definition อยู่แล้ว)
# Job 5/7/9/10 -> Peerakorn (งาน interface ที่พึ่งพา job อื่นน้อยที่สุด)
JOB_ESTIMATES: dict[str, int] = {
    "2": 14,
    "3": 10,
    "4": 14,
    "5": 16,
    "6": 26,   # +6 (2026-08-21): เขียน fgi_impact_compensations + 5 คอลัมน์รอบชดเชยใน fgi_impact_processes (F8+F1)
    "7": 10,
    "8": 18,
    "8b": 22,   # +6 (2026-08-21): ตัดสินประเภทเคสก่อนเปิด workflow (3 จุดเข้า) + addPreApprover เจ้าของงานคนเดิม
    "9": 11,
    "10": 8,
}

# 2026-08-25 (รอบ 3): Bank รับผิดชอบ migration DB + **batch job ทั้งหมด** + สร้าง workflow (นิยาม)
#   จึงไม่มี override ใด ๆ — ทุก job ตกเป็นของ BANK_BE_OWNER ตาม default
#   ⚠️ ผลรวมของ Bank = 296 ชม. เกินเพดาน 4 สัปดาห์ (170 ชม.) อยู่ 126 ชม. — ดูหมายเหตุใน README
JOB_OWNER_OVERRIDES: dict[str, str] = {}

HIGH_LEVEL_ESTIMATES: dict[str, int] = {
    "FE/LLDD-FE-Integration-Contracts": 16,
    "FE/LLDD-FE-Foundation": 28,
    "FE/LLDD-FE-Document-Lists": 28,
    "FE/LLDD-FE-Create-Document": 6,
    "FE/LLDD-FE-Document-Detail": 60,
    "FE/LLDD-FE-Report": 20,
    "FE/LLDD-FE-Master-Data": 16,
    "FE/LLDD-FE-Testing-Delivery": 12,
    "BE/LLDD-BE-API-Common-Contracts": 18,
    "BE/LLDD-BE-Integration-SBP-Platform": 20,
    "BE/LLDD-BE-API-Document-List-Search": 20,
    "BE/LLDD-BE-API-Document-Create-Update": 24,
    "BE/LLDD-BE-API-Document-Detail-Aggregate": 24,
    "BE/LLDD-BE-API-Document-Workflow-Actions": 28,
    "BE/LLDD-BE-API-Workflow-Instances": 24,
    "BE/LLDD-BE-Workflow-Engine-Definition": 24,   # +8 (2026-08-11): ปิดชื่อ function ที่ขัดกัน 3 ชุด · DP-2 ไม่มี PK/index · ออกแบบลำดับขั้นใหม่เพราะ workflow_state ไม่มี seq · ตัดสิน workflow_part_display ที่ทับซ้อน data-editrole
    "BE/LLDD-BE-API-Attachment-Sales-Timeline": 26,
    "BE/LLDD-BE-API-Lookup": 10,
    "BE/LLDD-BE-API-Report-and-Master-Data": 30,   # +3 (2026-08-11): ต้อง join store -> fr_store -> juristic ของระบบเดิม และดึงประเภทร้าน/ภาคจาก common_code + store/all-regions
    "BE/LLDD-BE-Job-Batch-Email-SRM": 14,
    "BE/LLDD-BE-Database-Structure": 31,   # +3 (2026-08-21): DDL ตาราง fgi_impact_compensations + 5 คอลัมน์รอบชดเชย + index/UK (F8+F1)
    "BE/LLDD-BE-Data-Migration-Cutover": 43,   # +3 (2026-08-21): map/migrate FGI_IMPACT_STORE_COMPENSATE และคอลัมน์รอบชดเชยจาก FGI_IMPACT_STORE_ON_PROCESS · +6 (2026-08-11): งานเพิ่มจาก master จริง — SectionLimitCost 100,000 ต้อง seed ใหม่ · DecisionCode ที่ Excel แปลงเป็นวันที่ · สถานะเดิม 10 ค่า map เหลือ 6 · SectionCode เติมศูนย์ · doc_no เป็น UNIQUE ตาม DP-1
}


@dataclass
class ApiSpec:
    method: str
    path: str
    purpose: str
    request: dict[str, Any] | None = None
    response: dict[str, Any] | None = None
    buttons: list[str] = field(default_factory=list)


@dataclass
class Topic:
    file: str
    title: str
    track: str
    days: float
    hours: int
    owner: str
    objective: str
    screenshots: list[str]
    scope: list[str]
    fields: list[tuple[str, str, str, str]]
    actions: list[tuple[str, str, str, str]]
    apis: list[ApiSpec]
    flow: list[str]
    acceptance: list[str]
    tests: list[str]
    db_tables: list[tuple[str, str, str]] = field(default_factory=list)
    flow_diagram: str | None = None
    base_hours: int = 0     # = hours (คงไว้เพื่อความเข้ากันได้ — ระบบ buffer ถอดออก 2026-08-11)
    buffer: float = 0.0     # ไม่ใช้แล้ว


def p(text: str) -> dict[str, Any]:
    return {"type": "p", "text": text}


def h(level: int, text: str) -> dict[str, Any]:
    return {"type": f"h{level}", "text": text}


def bullets(items: list[str]) -> dict[str, Any]:
    return {"type": "bullets", "items": items}


def table(headers: list[str], rows: list[list[Any]]) -> dict[str, Any]:
    return {"type": "table", "headers": headers, "rows": rows}


def image(path: str, caption: str) -> dict[str, Any]:
    return {"type": "image", "path": path, "caption": caption}


def code(text: str, lang: str = "") -> dict[str, Any]:
    return {"type": "code", "text": text, "lang": lang}


def payload(title: str, text: str) -> dict[str, Any]:
    return {"type": "payload", "title": title, "text": text}


def pagebreak() -> dict[str, Any]:
    return {"type": "pagebreak"}


LEGACY_JOB_SOURCES: dict[str, dict[str, Any]] = {
    "2": {
        "input": "Period year/month, optional zone filter, and ALLMAP SEVEN_IMPACT_VIEW rows.",
        "progress": "query candidate impacted stores, deduplicate by store/month, batch insert impact-store master data, derive related new-store/impact-store records, update verification flags.",
        "output": "FGI_IMPACT_STORE and related impact/new-store tables contain imported candidates for the requested period with duplicate-safe status.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/ImportImpactStore.java", "24-186", "Legacy main entrypoint for impacted-store import."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ImportStoreJdbc.java", "30-84, 170-484", "Query SEVEN_IMPACT_VIEW and insert/update FGI impact/new-store records."],
        ],
    },
    "3": {
        "input": "Period year/month and competitor impact data from ALLMAP COMPETITOR_IMPACT_VIEW.",
        "progress": "validate period, skip when period already exists, query competitor view, insert in chunks inside a transaction, send status mail.",
        "output": "FGI_IMPACT_COMPETITOR rows for the target period; run status is success/no-data/failed with inserted-count reconciliation.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/ImportImpactCompetitor.java", "16-48", "Legacy main entrypoint and notification wrapper."],
            ["fcsJar/src/th/co/gosoft/fgi/controller/ImportController.java", "483-598", "Validate params, skip duplicates, query source, chunk insert competitors."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ImportJdbc.java", "200-241", "Count existing period, query COMPETITOR_IMPACT_VIEW, insert FGI_IMPACT_COMPETITOR."],
        ],
    },
    "4": {
        "input": "FGI_IMPACT_STORE_SALES rows waiting for IAS sales data and EAI S3 bucket/prefix parameters.",
        "progress": "query eligible stores, write outbound IAS request file, upload to EAI S3 outbound prefix, keep local backup, record success/failure and notification.",
        "output": "IAS request file containing store/open-date pairs; run history includes generated file name and exported row count.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/PrepareImpactStoreToIAS.java", "28-243", "Legacy main entrypoint, file generation, upload, backup, notification."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ImportStoreJdbc.java", "99-115", "Query FGI_IMPACT_STORE_SALES rows eligible for IAS request."],
        ],
    },
    "5": {
        "input": "IAS sales response files from configured source path; file name pattern and pipe-delimited daily sales records.",
        "progress": "scan files, validate pattern, parse daily sales windows, derive before/after impact metrics, write transaction rows, update working-day counts and growth status, backup processed files.",
        "output": "FGI_IMPACT_STORE_SALES_TRN and FGI_IMPACT_STORE_SALES updated; confirm-receive rows written; source file moved to backup or error recorded.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/ImportImpactSaleFromIAS.java", "9-19", "Legacy main entrypoint that delegates to import controller."],
            ["fcsJar/src/th/co/gosoft/fgi/controller/ImportController.java", "101-411", "Parse IAS file, compute sales windows, prepare inserts/updates, backup and notify."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ImportJdbc.java", "136-182, 517-804", "Update verification flags, working days, growth-rate calculations, cleanup old files."],
        ],
    },
    "6": {
        "input": "Approved/initial compensation data from FGI impact/new-store tables plus QSSI score lookup and FS export configuration.",
        "progress": "query rows for FS, generate compensation interface payload, insert/update compensate records, upload/export, backup, notify.",
        "output": "FS outbound data and FGI compensation tables synchronized; run summary includes exported counts and file/status.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/ExportImpactStoreToFS.java", "19-68", "Legacy main entrypoint for exporting impact-store compensation to FS."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ExportJdbc.java", "119-180, 386-970", "Query FS export data and insert/update impact/new-store compensation records."],
        ],
    },
    "7": {
        "input": "FGI_IMPACT_COMPETITOR rows linked to active impact-process records and BPM/export confirmation state.",
        "progress": "query latest competitor rows, skip already-confirmed transactions, create outbound payload per competitor, upload/export, insert confirm-receive rows.",
        "output": "Competitor sync payload/output for downstream workflow; confirm-receive rows prevent duplicate export.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/ExportCompetitor.java", "9-20", "Legacy main entrypoint for competitor export."],
            ["fcsJar/src/th/co/gosoft/fgi/controller/ExportController.java", "659-760", "Query competitor data, generate file content, upload, backup, notification."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ExportJdbc.java", "1596-1628", "Query latest competitor rows eligible for export."],
        ],
    },
    "8": {
        "input": "Impact-store compensation rows in initial status with workflow sequence values and no prior confirm-receive output.",
        "progress": "update BPM sequence, query eligible impact-store rows, refresh not-OPT data, generate workflow payload, insert confirm-receive rows, upload/export, notify.",
        "output": "Impact-store workflow create payload/output with generated sequence numbers and duplicate guard.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/ExportImpactStoreFlowToBPM.java", "9-17", "Legacy main entrypoint for exporting impact-store flow data."],
            ["fcsJar/src/th/co/gosoft/fgi/controller/ExportController.java", "518-657", "Build impact-store BPM payload, write file, upload, backup, notification."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ExportJdbc.java", "1654-1692", "Query impact-store rows eligible for workflow export."],
        ],
    },
    "8b": {
        "input": "Impact-store rows waiting to start workflow plus generated workflow/document identifiers.",
        "progress": "select waiting rows, start workflow instance, update generated-flow flag per transaction, log success/failure.",
        "output": "Workflow instances started and source rows marked generated; failed rows remain rerunnable with error detail.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/StartK2WorkFlow.java", "16-51", "Legacy main entrypoint for starting K2 workflow."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/StartFlowJdbc.java", "17-173", "Select rows for workflow start and update generated-flow flags."],
        ],
    },
    "9": {
        "input": "New-store compensation rows linked to active impact-process records (writes to document_new_stores directly; no export file).",
        "progress": "query eligible new-store rows, filter process errors, write outbound new-store payload, insert confirm-receive rows, upload/export, backup, notify.",
        "output": "New-store sync payload/output and confirm-receive rows keyed by NEW_STORE_INFO_ID/month/year.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/ExportOpenStore.java", "1-22", "Legacy main entrypoint; constant job name is ExportNewStoreToBPM."],
            ["fcsJar/src/th/co/gosoft/fgi/controller/ExportController.java", "404-516, 893-961", "Query new stores, create payload content, upload, backup, notification."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ExportJdbc.java", "1558-1594", "Query new-store rows eligible for export."],
        ],
    },
    "10": {
        "input": "FGI_CONFIRM_RECEIVE_DATA rows without return_code after the waiting threshold.",
        "progress": "query missing receive data, group by data_name/direction (To-Be — เดิม Oracle ใช้ interface_type), build notification message, send admin mail, close run.",
        "output": "Notification sent for overdue receive confirmations; run status records grouped counts or no-data success.",
        "sources": [
            ["fcsJar/src/th/co/gosoft/fgi/main/NotifyNoReceiveData.java", "16-37", "Legacy main entrypoint for missing-receive notification."],
            ["fcsJar/src/th/co/gosoft/fgi/controller/ManageCompensateController.java", "748-775", "Build and send notification content for missing receive data."],
            ["fcsJar/src/th/co/gosoft/fgi/dao/jdbc/ExportJdbc.java", "1894-1917", "Query confirm-receive rows without return_code."],
        ],
    },
}


JOB_IMPLEMENTATION_SPECS: dict[str, dict[str, str]] = {
    "2": {
        "repository": "impactStoreRepository",
        "read": """SELECT impacted_store_code, new_store_code, impact_month, distance_km, region_code, zone_code, branch_type
FROM allmap_seven_impact_view
WHERE impact_month = :impact_month
  AND (:zone_code IS NULL OR zone_code = :zone_code)
  AND distance_km <= CASE
        WHEN region_code = ANY(:bangkok_metro_region_codes) THEN 1.000
        ELSE 2.000
      END;""",
        "write": """INSERT INTO fgi_impact_stores
    (impact_process_id, impacted_store_code, new_store_code, impact_month, distance_km, updated_at)
VALUES (:impact_process_id, :impacted_store_code, :new_store_code, :impact_month, :distance_km, CURRENT_TIMESTAMP)
ON CONFLICT (impacted_store_code, new_store_code, impact_month)
DO UPDATE SET distance_km = EXCLUDED.distance_km,
              impact_process_id = EXCLUDED.impact_process_id,
              updated_at = CURRENT_TIMESTAMP;""",
        "idempotency": "UNIQUE(impacted_store_code, new_store_code, impact_month); rerun อัปเดตค่าที่เปลี่ยนแต่ไม่สร้างคู่ร้านซ้ำ",
        "transaction": "สร้าง/หา fgi_impact_processes และ upsert candidate ทีละ chunk ใน transaction; chunk fail rollback เฉพาะ chunk",
        "security": "ALLMAP connection ใช้ datasource secretRef และ TLS verify-full; job parameter เก็บได้เฉพาะ datasource alias ไม่เก็บ username/password",
        "steps": "loadAllmapCandidates|resolveImpactProcesses|upsertImpactPairs|reconcileImportedPairs",
    },
    "3": {
        "repository": "impactCompetitorRepository",
        "read": """SELECT impact_process_id, competitor_code, name_th, branch_th, opened_date, closed_date, period_key
FROM allmap_competitor_impact_view
WHERE period_key = :period_key;""",
        "write": """INSERT INTO fgi_impact_competitors
    (impact_process_id, competitor_code, name_th, branch_th, opened_date, closed_date, period_key, updated_at)
VALUES (:impact_process_id, :competitor_code, :name_th, :branch_th, :opened_date, :closed_date, :period_key, CURRENT_TIMESTAMP)
ON CONFLICT (impact_process_id, competitor_code, period_key)
DO UPDATE SET name_th = EXCLUDED.name_th,
              branch_th = EXCLUDED.branch_th,
              opened_date = EXCLUDED.opened_date,
              closed_date = EXCLUDED.closed_date,
              updated_at = CURRENT_TIMESTAMP;""",
        "idempotency": "UNIQUE(impact_process_id, competitor_code, period_key); source row ซ้ำในไฟล์/วิวต้อง deduplicate ก่อน upsert",
        "transaction": "validate งวดก่อนอ่าน; upsert ทีละ chunk และ commit หลัง reconcile จำนวน input/success/reject ของ chunk ตรงกัน",
        "security": "ALLMAP datasource ใช้ secretRef และ TLS verify-full; จำกัด DB user เป็น SELECT เฉพาะ source view",
        "steps": "loadCompetitorPeriod|deduplicateCompetitors|upsertCompetitors|reconcileCompetitorCount",
    },
    "4": {
        "repository": "iasRequestRepository",
        "read": """SELECT s.id, s.impact_process_id, s.impacted_store_code, s.new_store_code, s.impact_month
FROM fgi_impact_stores s
WHERE s.sales_request_status = 'W'
ORDER BY s.id
FOR UPDATE SKIP LOCKED;""",
        "write": """UPDATE fgi_impact_stores
SET sales_request_status = 'P', updated_at = CURRENT_TIMESTAMP
WHERE id = ANY(:impact_store_ids) AND sales_request_status = 'W';

INSERT INTO interface_transactions
    (run_id, data_name, direction, status, impact_process_id, business_key, period_key,
     file_name, file_checksum, outbox_status, purge_after)
SELECT :run_id, 'IAS_SALES_REQUEST', 'OUT', 'READY', impact_process_id,
       impacted_store_code || ':' || new_store_code, impact_month,
       :file_name, :file_checksum, 'READY', CURRENT_TIMESTAMP + INTERVAL '180 days'
FROM fgi_impact_stores
WHERE id = ANY(:impact_store_ids)
ON CONFLICT (data_name, direction, business_key, period_key) DO NOTHING;""",
        "idempotency": "ชื่อไฟล์ deterministic จาก period+runId และ UNIQUE(data_name,direction,business_key,period_key); outbox retry ใช้ transaction เดิม ไม่สร้าง request ซ้ำ",
        "transaction": "สร้างไฟล์ temp, fsync, atomic rename และคำนวณ checksum ให้สำเร็จก่อน; จากนั้น transaction เดียว lock W, update W→P และ insert outbox READY; ห้าม commit W→P ก่อนมี durable file",
        "security": "สิทธิ์เขียน EAI S3 ใช้ IAM role ของ pod หรือ secretRef=secret/sbpgi/interfaces/eai-s3; จำกัดสิทธิ์เฉพาะ prefix ขาออกของ IAS (PutObject เท่านั้น) และห้าม editable access key ในหน้าจอ/ไฟล์ config",
        "steps": "lockWaitingSalesRequests|writeDurableIasFile|markPendingAndCreateOutbox|dispatchIasOutbox",
    },
    "5": {
        "repository": "iasSalesRepository",
        "read": """SELECT t.sales_summary_id, t.txn_date, t.sales_amount, t.window_no, t.source_checksum
FROM sales_transactions t
JOIN fgi_impact_sales_summaries s ON s.id = t.sales_summary_id
WHERE s.impact_process_id = :impact_process_id
ORDER BY t.sales_summary_id, t.txn_date, t.window_no;""",
        "write": """INSERT INTO sales_transactions
    (sales_summary_id, txn_date, window_no, sales_amount, sales_diff, is_outlier, source_checksum)
VALUES (:sales_summary_id, :txn_date, :window_no, :sales_amount, :sales_diff, :is_outlier, :source_checksum)
ON CONFLICT (sales_summary_id, txn_date, window_no)
DO UPDATE SET sales_amount = EXCLUDED.sales_amount,
              sales_diff = EXCLUDED.sales_diff,
              is_outlier = EXCLUDED.is_outlier,
              source_checksum = EXCLUDED.source_checksum;

UPDATE fgi_impact_sales_summaries
SET total_working_days = :total_working_days,
    growth_rate_before = :growth_rate_before,
    growth_rate_after = :growth_rate_after,
    growth_rate_diff = :growth_rate_diff,
    sales_status = :sales_status,
    updated_at = CURRENT_TIMESTAMP
WHERE id = :sales_summary_id;""",
        "idempotency": "checksum กันไฟล์ซ้ำ + UNIQUE(sales_summary_id,txn_date,window_no); คำนวณ summary ใหม่จาก transaction rows ทุก rerun",
        "transaction": "upsert รายวันและ update summary ของ sales_summary_id เดียวกันใน transaction; checksum/file tracking commit พร้อมกัน",
        "security": "สิทธิ์อ่าน EAI S3 ใช้ IAM role ของ pod หรือ secretRef=secret/sbpgi/interfaces/eai-s3 จำกัดเฉพาะ prefix ขาเข้า/backup ของ IAS (GetObject + PutObject เฉพาะ backup); quarantine อ็อบเจกต์ที่ checksum/รูปแบบไม่ผ่าน แทนการลบทิ้ง",
        "steps": "downloadAndStageIasSales|validateSalesWindows|upsertDailySales|recalculateSalesSummaries",
    },
    "6": {
        "repository": "statementExportRepository",
        "read": """SELECT d.doc_no, d.impact_process_id, s.id AS sales_summary_id,
       d.total_compensation_amount, q.score
FROM compensation_documents d
JOIN fgi_impact_sales_summaries s ON s.impact_process_id = d.impact_process_id
LEFT JOIN fcs_qssi_score q ON q.store_id = d.impacted_store_code AND q.month = d.impact_month
JOIN LATERAL (
    SELECT c.result_category
    FROM consideration_logs c
    WHERE c.doc_no = d.doc_no
    ORDER BY c.action_datetime DESC
    LIMIT 1
) latest_decision ON latest_decision.result_category = 'APPROVE'
WHERE d.status_code = '99'
  AND NOT EXISTS (
      SELECT 1 FROM interface_transactions i
      WHERE i.data_name = 'COMPENSATE_APPROVE_I' AND i.direction = 'OUT'
        AND i.doc_no = d.doc_no AND i.status IN ('READY','SENT','ACKED'));""",
        "write": """INSERT INTO interface_transactions
    (run_id, data_name, direction, status, doc_no, impact_process_id, sales_summary_id,
     business_key, period_key, file_name, file_checksum, outbox_status, purge_after)
VALUES (:run_id, 'COMPENSATE_APPROVE_I', 'OUT', 'READY', :doc_no, :impact_process_id,
        :sales_summary_id, :doc_no, :period_key, :file_name, :file_checksum, 'READY',
        CURRENT_TIMESTAMP + INTERVAL '365 days')
ON CONFLICT (data_name, direction, business_key, period_key) DO NOTHING;

WITH purge_candidates AS (
    SELECT id
    FROM interface_transactions
    WHERE data_name = ANY(:purge_data_names)
      AND status IN ('ACKED','COMPLETED')
      AND purge_after < CURRENT_TIMESTAMP
      AND legal_hold = FALSE
    ORDER BY id
    LIMIT :purge_batch_size
    FOR UPDATE SKIP LOCKED
)
DELETE FROM interface_transactions i
USING purge_candidates p
WHERE i.id = p.id
RETURNING i.id, i.data_name, i.business_key;""",
        "idempotency": "UNIQUE(data_name,direction,business_key,period_key); STA ACK เปลี่ยน transaction เดิมเป็น ACKED ไม่ insert แถวใหม่",
        "transaction": "สร้าง payload/checksum ก่อน แล้ว insert outbox READY; dispatcher ส่งและเปลี่ยน SENT แยก transaction; callback ACK เปลี่ยน ACKED แบบ compare-and-set",
        "security": "RabbitMQ broker ใช้ secretRef=secret/sbpgi/mq/sta, เชื่อมด้วย AMQPS (TLS 1.2+ verify-full); exchange/routing key มาจาก config ไม่ใช่ค่าที่ผู้ใช้แก้ได้; credential rotation ไม่ต้องแก้เอกสารหรือ job param",
        "steps": "loadApprovedCompensations|buildStatementPayload|enqueueStatementOutbox|purgeAcknowledgedTracking",
    },
    "7": {
        "repository": "documentCompetitorRepository",
        "read": """SELECT d.doc_no, c.competitor_code, c.name_th, c.branch_th, c.opened_date, c.closed_date
FROM fgi_impact_competitors c
JOIN compensation_documents d ON d.impact_process_id = c.impact_process_id
WHERE c.period_key = :period_key;""",
        "write": """INSERT INTO document_competitors
    (doc_no, competitor_code, name_th, branch_th, opened_date, closed_date, source_system, updated_at)
VALUES (:doc_no, :competitor_code, :name_th, :branch_th, :opened_date, :closed_date, 'ALLMAP', CURRENT_TIMESTAMP)
ON CONFLICT (doc_no, competitor_code)
DO UPDATE SET name_th = EXCLUDED.name_th, branch_th = EXCLUDED.branch_th,
              opened_date = EXCLUDED.opened_date, closed_date = EXCLUDED.closed_date,
              updated_at = CURRENT_TIMESTAMP;

DELETE FROM document_competitors dc
WHERE dc.doc_no = :doc_no
  AND dc.source_system = 'ALLMAP'
  AND NOT EXISTS (
      SELECT 1
      FROM fgi_impact_competitors src
      JOIN compensation_documents d ON d.impact_process_id = src.impact_process_id
      WHERE d.doc_no = dc.doc_no
        AND src.period_key = :period_key
        AND src.competitor_code = dc.competitor_code
  );""",
        "idempotency": "UNIQUE(doc_no,competitor_code); upsert และ prune เฉพาะ source_system=ALLMAP ให้ target ตรง source ปัจจุบันโดยไม่ลบแถว USER",
        "transaction": "upsert + prune document_competitors และ tracking (direction=INTERNAL) ใน transaction เดียวต่อ doc_no",
        "security": "service account ภายในมีสิทธิ์ SELECT source และ INSERT/UPDATE target เท่านั้น; ไม่มี external credential",
        "steps": "loadLatestDocumentCompetitors|upsertDocumentCompetitors|recordInternalCompetitorSync|reconcileDocumentCompetitors",
    },
    "8": {
        "repository": "compensationDocumentRepository",
        "read": """SELECT p.id AS impact_process_id, p.impacted_store_code, p.impact_month,
       SUM(COALESCE(s.adjust_compensation_amount, s.forecast_compensation_amount, 0)) AS total_compensation_amount
FROM fgi_impact_processes p
JOIN fgi_impact_stores s ON s.impact_process_id = p.id
WHERE p.process_status = 'READY_DOCUMENT'
GROUP BY p.id, p.impacted_store_code, p.impact_month;""",
        "write": """INSERT INTO compensation_documents
    (doc_no, year, running_no, impact_process_id, impacted_store_code, impact_month,
     source, status_code, current_section_code, total_compensation_amount, created_by)
VALUES (:doc_no, :year, :running_no, :impact_process_id, :impacted_store_code, :impact_month,
        'FS', '06', '06', :total_compensation_amount, 'JOB-8')
ON CONFLICT (impact_process_id) DO NOTHING;

INSERT INTO interface_transactions
    (run_id, data_name, direction, status, impact_process_id, doc_no,
     business_key, period_key, outbox_status, purge_after, completed_at)
SELECT :run_id, 'DOCUMENT_CREATE', 'INTERNAL', 'COMPLETED', d.impact_process_id, d.doc_no,
       CAST(d.impact_process_id AS VARCHAR), d.impact_month, 'COMPLETED',
       CURRENT_TIMESTAMP + INTERVAL '365 days', CURRENT_TIMESTAMP
FROM compensation_documents d
WHERE d.impact_process_id = :impact_process_id
ON CONFLICT (data_name, direction, business_key, period_key) DO NOTHING;""",
        "idempotency": "UNIQUE(impact_process_id) และ UNIQUE(year,running_no); lock running number ต่อปีใน transaction; conflict ต้องคืน/อ้าง doc_no เดิม และยอมให้เลขที่จองกระโดดโดยห้าม reuse",
        "transaction": "lock เลขรัน + insert document + update process + tracking (direction=INTERNAL) ใน transaction เดียว",
        "security": "internal service account เท่านั้น; ห้ามสร้างไฟล์ BPM06001O, ห้าม SFTP และห้ามเก็บ K2 credential",
        "steps": "loadDocumentCandidates|allocateDocumentNumbers|createCompensationDocuments|recordDocumentCreation",
    },
    "8b": {
        "repository": "workflowRepository",
        "read": """WITH locked_process AS (
    SELECT p.id
    FROM fgi_impact_processes p
    JOIN compensation_documents d ON d.impact_process_id = p.id
    WHERE p.workflow_generation_status = 'W'
      -- ⚠️ sps_store.workflow_transaction ไม่มี PK/index (19,283 แถว) → เงื่อนไขนี้เป็น seq-scan · DP-2 ยังไม่ตัดสิน
      -- ✅ DP-1 ปิดแล้ว: reference_id = compensation_documents.id (surrogate) แปลงเป็น text
      AND NOT EXISTS (SELECT 1 FROM sps_store.workflow_transaction w WHERE w.reference_id = d.id::text   -- DP-1 = surrogate id (reference_id เป็น varchar(255)) AND w.version_id = :sbpgi_version_id)   -- @srm/glb-workflow
    ORDER BY p.id
    FOR UPDATE OF p SKIP LOCKED
), gate AS (
    SELECT p.id AS impact_process_id, d.doc_no, d.current_section_code,
           CASE
             WHEN BOOL_OR(ns.store_type IS NULL OR ns.store_type NOT IN ('FAM','FB1','FC1','FB2','FVB','FVC')) THEN 'N'
             WHEN BOOL_OR(pair.distance_km > CASE
                    WHEN impacted.zone_cd = ANY(:bangkok_metro_region_codes) THEN 1.000
                    ELSE 2.000
                  END) THEN 'N'
             WHEN BOOL_OR(pair.distance_km IS NULL) THEN 'W'
             WHEN ist.opt_dv_user_id IS NULL OR BTRIM(ist.opt_dv_user_id) = '' THEN 'N'
             WHEN ij.juristic_name IS NULL OR BOOL_OR(nj.juristic_name IS NULL) THEN 'W'
             WHEN BOOL_OR(ij.juristic_name = nj.juristic_name) THEN 'N'
             WHEN ss.growth_rate_diff IS NULL THEN 'W'
             WHEN ss.growth_rate_diff > -10 THEN 'N'
             WHEN ss.sales_status IS NULL OR ss.sales_status NOT IN ('Y','N') THEN 'W'
             ELSE 'Y'
           END AS gate_decision
    FROM locked_process lp
    JOIN fgi_impact_processes p ON p.id = lp.id
    JOIN compensation_documents d ON d.impact_process_id = p.id
    JOIN impacted_stores ist ON ist.store_code = p.impacted_store_code
    JOIN store impacted ON impacted.store_id = p.impacted_store_code
    JOIN fgi_impact_stores pair ON pair.impact_process_id = p.id
    JOIN store ns ON ns.store_id = pair.new_store_code
    -- นิติบุคคลไม่ได้อยู่บน store — ต้องผ่าน fr_store.juristic_id -> juristic.juristic_name
    LEFT JOIN fr_store ifs ON ifs.store_id = impacted.store_id
    LEFT JOIN juristic ij  ON ij.juristic_id = ifs.juristic_id
    LEFT JOIN fr_store nfs ON nfs.store_id = ns.store_id
    LEFT JOIN juristic nj  ON nj.juristic_id = nfs.juristic_id
    LEFT JOIN fgi_impact_sales_summaries ss ON ss.impact_process_id = p.id
    GROUP BY p.id, d.doc_no, d.current_section_code, ist.opt_dv_user_id,
             ij.juristic_name, ss.growth_rate_diff, ss.sales_status
)
SELECT * FROM gate;""",
        "write": """UPDATE fgi_impact_processes
SET workflow_generation_status = 'N', updated_at = CURRENT_TIMESTAMP
WHERE id = :impact_process_id
  AND workflow_generation_status = 'W'
  AND :gate_decision = 'N';

-- gate_decision='Y': เปิด workflow ผ่าน @srm/glb-workflow ของระบบ SBP เดิม (ไม่ INSERT ตารางเอง)
-- ✅ ชื่อ function ยึดชีต Detail ของ LLDD lib (ปิด 2026-08-14) — API 8 ตัว:
--    initializeWorkflow / eventWorkflow / getPermissionEvents / getHistory /
--    getTransaction / getPendingFlowByUser / getWorkflowsByUser / addPreApprover
--   initializeWorkflow({ versionId: :sbpgi_version_id, referenceId: :reference_id, userId: 'JOB-8B' })
--   addPreApprover({ versionId, referenceId: :reference_id, stateId: '06', approver, seq: 1 })
-- ✅ DP-1 ปิดแล้ว 2026-08-17: referenceId = compensation_documents.id (surrogate · ส่งเป็น string)
-- library จะเขียน sps_store.workflow_transaction / workflow_approver / workflow_history ให้เอง
UPDATE fgi_impact_processes
SET workflow_generation_status = 'Y', updated_at = CURRENT_TIMESTAMP
WHERE id = :impact_process_id
  AND workflow_generation_status = 'W'
  AND :gate_decision = 'Y';

-- gate_decision='W' ไม่เปลี่ยนสถานะ; บันทึก reason ลง application log (structured) เพื่อ rerun — ไม่มีตาราง job_run_histories แล้ว (2026-08-06).""",
        "idempotency": ("กันซ้ำระดับ application — ตรวจว่ามี transaction เดิมของ reference นี้อยู่แล้วหรือไม่ ก่อนเรียก initialize แล้ว skip "
                        "· ⚠️ **ไม่มี UNIQUE(version_id, reference_id) จริงใน `sps_store.workflow_transaction`** (ตารางนี้ไม่มีทั้ง PK และ index "
                        "ทั้งที่มี 19,283 แถว — ตรวจแล้วที่ `SBP/db-schema-sps_store.md`) จึงพึ่ง constraint ฝั่ง DB ไม่ได้ และ query ตาม reference_id เป็น seq-scan "
                        "· จะขอ sign-off เพิ่ม PK/index กับทีมเจ้าของ library หรือยอมรับสภาพ **ยังไม่ตัดสิน (DP-2)**"),
        "transaction": "lock process + evaluate gate + branch N/W/Y; เฉพาะ Y จึงเรียก initializeWorkflow + addPreApprover ของ @srm/glb-workflow (ชื่อ function ตามชีต Detail ของ LLDD lib) และ W→Y ใน transaction เดียว, N ต้อง persist ถาวร, W คงเดิมเพื่อ rerun",
        "security": "internal service token จาก workload identity/secretRef; ห้าม Basic Auth หรือ K2 REST credential เดิม",
        "steps": "lockWorkflowCandidates|evaluateGenerationGate|startInternalWorkflows|notifyWorkflowOwners",
    },
    "9": {
        "repository": "documentNewStoreRepository",
        "read": """SELECT d.doc_no, s.new_store_code,
       COALESCE(s.adjust_compensate_percent, s.forecast_compensate_percent) AS compensate_percent,
       COALESCE(s.adjust_compensation_amount, s.forecast_compensation_amount) AS compensation_amount
FROM fgi_impact_stores s
JOIN compensation_documents d ON d.impact_process_id = s.impact_process_id
WHERE s.impact_month = :impact_month
  AND COALESCE(s.adjust_compensate_percent, s.forecast_compensate_percent) IS NOT NULL
  AND COALESCE(s.adjust_compensate_percent, s.forecast_compensate_percent) BETWEEN 0 AND 100;""",
        "write": """-- validateAllocationValues ต้องยืนยัน source_row_count = valid_row_count ก่อนคำสั่งนี้;
-- ถ้าค่า percent เป็น NULL/นอกช่วง ให้ throw COMPENSATE_PERCENT_INVALID และ rollback ก่อน upsert/prune.
INSERT INTO document_new_stores
    (doc_no, new_store_code, compensate_percent, compensation_amount, source_system, updated_at)
SELECT :doc_no, :new_store_code, :compensate_percent, :compensation_amount, 'FGI', CURRENT_TIMESTAMP
WHERE :compensate_percent IS NOT NULL
  AND :compensate_percent BETWEEN 0 AND 100
ON CONFLICT (doc_no, new_store_code)
DO UPDATE SET compensate_percent = EXCLUDED.compensate_percent,
              compensation_amount = EXCLUDED.compensation_amount,
              updated_at = CURRENT_TIMESTAMP
RETURNING doc_no, new_store_code;

-- Service ต้องได้ RETURNING 1 แถวต่อ source row; ไม่ครบให้ rollback และห้าม prune.

DELETE FROM document_new_stores dns
WHERE dns.doc_no = :doc_no
  AND dns.source_system = 'FGI'
  AND NOT EXISTS (
      SELECT 1
      FROM fgi_impact_stores src
      JOIN compensation_documents d ON d.impact_process_id = src.impact_process_id
      WHERE d.doc_no = dns.doc_no
        AND src.impact_month = :impact_month
        AND src.new_store_code = dns.new_store_code
  );

SELECT CASE WHEN ABS(SUM(compensate_percent) - 100) <= 0.0001 THEN TRUE ELSE FALSE END AS allocation_valid
FROM document_new_stores
WHERE doc_no = :doc_no;""",
        "idempotency": "UNIQUE(doc_no,new_store_code); upsert + prune เฉพาะ source_system=FGI ให้ target ตรง impact set ปัจจุบัน โดยไม่ลบแถว USER",
        "transaction": "validate source percent ต้องไม่เป็น NULL และอยู่ 0..100 ก่อน upsert; จากนั้น upsert + prune ร้านของ doc_no, validate ผลรวม 100% และ tracking (direction=INTERNAL) ใน transaction เดียว; invalid/ไม่ครบให้ rollback ก่อน prune",
        "security": "internal service account least privilege; ไม่มี SFTP/BPM credential หรือ editable external endpoint",
        "steps": "loadNewStoreAllocations|validateAllocationValues|upsertDocumentNewStores|reconcileAllocationTotals",
    },
    "10": {
        "repository": "pendingAckRepository",
        "read": """SELECT id, data_name, business_key, file_name, sent_at
FROM interface_transactions
WHERE direction = 'OUT'
  AND status = 'SENT'
  AND acked_at IS NULL
  AND sent_at < CURRENT_TIMESTAMP - (:threshold_hours * INTERVAL '1 hour')
  AND (last_ack_notified_on IS NULL OR last_ack_notified_on < CURRENT_DATE)
ORDER BY sent_at;""",
        "write": """-- ยกเลิกตาราง audit_logs แล้ว (2026-08-07) — marker กันส่งซ้ำย้ายมาไว้บน interface_transactions เอง
-- คอลัมน์ last_ack_notified_on DATE มีอยู่ใน DDL ของ interface_transactions แล้ว (ดู LLDD-Database 5.x)
UPDATE interface_transactions
   SET last_ack_notified_on = CURRENT_DATE
 WHERE id = ANY(:transaction_ids)
   AND (last_ack_notified_on IS NULL OR last_ack_notified_on < CURRENT_DATE)
RETURNING id;""",
        "idempotency": "คอลัมน์ last_ack_notified_on บน interface_transactions เป็น marker ต่อรายการต่อวัน; rerun วันเดียวกันไม่ส่งอีเมลซ้ำ (ย้ายมาจาก audit_logs ที่ถูกยกเลิก 2026-08-07)",
        "transaction": "อ่าน pending แบบ read-only; reserve notification marker ก่อนส่ง; ส่งล้มเหลว mark FAILED และ retry ด้วย marker เดิม",
        "security": "SBPGI เรียก sendEmail() ของ email-lib เอง (ปิด DP-5 · 2026-08-14) — เลข template มาจาก workflow_route.email_id · credential SMTP/SES และตาราง email_template/email_sent เป็นของระบบ SBP เดิม",
        "steps": "loadOverdueAcknowledgements|reserveNotificationMarkers|sendPendingAckDigest|closeNotificationMarkers",
    },
}


def api_json(obj: Any) -> str:
    return json.dumps(obj, ensure_ascii=False, indent=2)


API_REQUIRED_QUERY_FIELDS: dict[str, set[str]] = {
    "/api/v1/sbpgi/document": {"year"},
    "/api/v1/sbpgi/report/status-summary": {"year", "status"},
    "/api/v1/sbpgi/report/status-summary/export": {"year", "status"},
}


def api_value_type(value: Any) -> str:
    if value is None:
        return "string | null"
    if isinstance(value, bool):
        return "boolean"
    if isinstance(value, int):
        return "integer"
    if isinstance(value, float):
        return "number"
    if isinstance(value, list):
        if not value:
            return "array<object>"
        return f"array<{api_value_type(value[0])}>"
    if isinstance(value, dict):
        return "object"
    return "string"


def api_field_constraint(field_path: str, value: Any) -> str:
    name = field_path.split(".")[-1].replace("[]", "").lower()
    if name in {"page"}:
        return ">= 1; default 1"
    if name in {"size"}:
        return "1..100; default 20"
    if name.endswith("storecode"):
        return "exactly 5 digits; preserve leading zero"
    if name in {"docno"}:
        return "ค.ศ. YYYY/xxxxx"
    if name.endswith("datetime") or name.endswith("at") or name.endswith("date") or name.endswith("month"):
        return "ISO-8601 ค.ศ.; nullable only when type includes null"
    if name.endswith("percent"):
        return "number 0..100 with 2 decimals"
    if "amount" in name:
        return "number >= 0 with 2 decimals"
    if name in {"statuscode", "nextsection", "sectioncode", "rolecode"}:
        return "canonical code; do not replace with display label"
    if name in {"reason", "comment"}:
        return "trimmed UTF-8 Thai text; required by operation/business rule"
    if isinstance(value, list):
        return "JSON array; element type shown in Type column"
    if isinstance(value, dict):
        return "JSON object; nested fields listed below"
    return "UTF-8; use value domain described by endpoint purpose"


def api_schema_rows(spec: ApiSpec, body: Any, direction: str) -> list[list[str]]:
    if body is None:
        return [["-", "none", "No", "Endpoint has no JSON body/query object"]]

    rows: list[list[str]] = []
    required_query = API_REQUIRED_QUERY_FIELDS.get(spec.path, set())

    def walk(value: Any, path: str = "") -> None:
        if isinstance(value, dict):
            if path:
                required = "Yes" if direction == "response" or spec.method.upper() != "GET" else ("Yes" if path in required_query else "No")
                rows.append([path, "object", required, api_field_constraint(path, value)])
            for key, child in value.items():
                child_path = f"{path}.{key}" if path else key
                walk(child, child_path)
            return
        if isinstance(value, list):
            required = "Yes" if direction == "response" or spec.method.upper() != "GET" else ("Yes" if path in required_query else "No")
            rows.append([path, api_value_type(value), required, api_field_constraint(path, value)])
            if value and isinstance(value[0], dict):
                for key, child in value[0].items():
                    walk(child, f"{path}[].{key}")
            return
        if direction == "response":
            required = "No" if value is None else "Yes"
        elif spec.method.upper() == "GET":
            required = "Yes" if path in required_query else "No"
        else:
            required = "No" if value is None else "Yes"
        rows.append([path or "value", api_value_type(value), required, api_field_constraint(path or "value", value)])

    walk(body)
    return rows or [["-", "none", "No", "No fields"]]


def sanitize_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "-", name).strip("-")


def fmt_days(days: float) -> str:
    return f"{days:.1f}".rstrip("0").rstrip(".")


def fmt_date(d: date) -> str:
    return d.strftime("%d/%m/%Y")


def is_workday(d: date) -> bool:
    return d.weekday() < 5


def next_workday(d: date) -> date:
    current = d
    while not is_workday(current):
        current += timedelta(days=1)
    return current


def add_workdays(start: date, workdays: int) -> date:
    current = next_workday(start)
    remaining = max(1, workdays) - 1
    while remaining:
        current += timedelta(days=1)
        if is_workday(current):
            remaining -= 1
    return current


# --------------------------------------------------------------------------------------
# ลำดับงานจริงระหว่างเอกสาร (finish-to-start) — ผู้ผลิตต้องจบก่อนผู้บริโภคเริ่ม
#
# 2026-08-10: เดิม build_topic_schedule คิดแค่ "ชั่วโมงสะสมรายคน" จึงไม่รู้จัก dependency เลย
# ผลคือตารางเดิมละเมิดลำดับงาน 14 เส้น ที่หนักที่สุดคือ Job 8 (ผู้สร้าง compensation_documents)
# ถูกจัดไว้ท้ายสุด แต่ Job 7 / 9 / 8b ที่อ่านผลของมันเสร็จไปก่อนแล้ว
#
# แหล่งอ้างอิงของแต่ละเส้น: field `rels` ของแต่ละ job ใน job-batch.html (ดู load_batch_jobs)
# --------------------------------------------------------------------------------------
# Job 1 (ImportQSSI) ถูกตัดออก 2026-08-24 — ระบบ SBP เดิมนำเข้า fcs_qssi_score ให้แล้ว SBPGI แค่อ่าน
JOB_PIPELINE_ORDER: list[str] = ["2", "3", "4", "5", "6", "7", "8", "8b", "9", "10"]

# 2026-08-11: ถอดระบบ buffer ออกทั้งหมด — ชั่วโมงในเอกสารคือค่าประเมินตรง ๆ ไม่มีส่วนเผื่อ

JOB_DEPENDENCIES: dict[str, list[str]] = {
    "3": ["2"],       # fgi_impact_stores เป็นแม่ของ fgi_impact_competitors
    "4": ["2"],       # fgi_impact_sales_summaries ต่อจาก fgi_impact_stores
    "5": ["4"],       # Job 5 เติม sales_transactions ใต้หัวตารางที่ Job 4 สร้าง
    "6": ["2"],       # อ่าน fcs_qssi_score (ระบบ SBP เดิมเติมให้) ครบ 6 หมวด + mutate fgi_impact_stores
    "7": ["3", "8"],  # อ่าน fgi_impact_competitors เขียนลง document_competitors.doc_no
    "8": ["6"],       # query impact profile สถานะ COMPENSATE I ที่ Job 6 ปล่อย
    "8b": ["5", "8"], # growth_rate_diff -> Gen Flow Gate · ต้องมี compensation_documents ก่อน
    "9": ["8"],       # document_new_stores.doc_no -> compensation_documents
    "10": ["6"],      # เฝ้าแถว interface_transactions ที่ Job 6 สร้าง
}

DOC_DEPENDENCIES: dict[str, list[str]] = {
    "BE/LLDD-BE-Data-Migration-Cutover": ["BE/LLDD-BE-Database-Structure"],
    "BE/LLDD-BE-API-Common-Contracts": [
        "BE/LLDD-BE-Database-Structure",
        "BE/LLDD-BE-Integration-SBP-Platform",
    ],
    # กันงานซ้ำระหว่าง Vava (BE) กับ lin (FE): ตัวตน/สิทธิ์/เมนู ถูกนิยามฝั่ง BE ก่อนหนึ่งครั้ง
    # แล้ว FE จึงเขียนสัญญาฝั่งบริโภค — ไม่ให้สองคนนิยาม header/permission ซ้ำกันคนละแบบ
    "FE/LLDD-FE-Integration-Contracts": ["BE/LLDD-BE-Integration-SBP-Platform"],
    "FE/LLDD-FE-Foundation": ["FE/LLDD-FE-Integration-Contracts"],
    "FE/LLDD-FE-Create-Document": ["FE/LLDD-FE-Foundation"],
    "FE/LLDD-FE-Document-Lists": ["FE/LLDD-FE-Foundation"],
    "FE/LLDD-FE-Document-Detail": ["FE/LLDD-FE-Foundation", "BE/LLDD-BE-Workflow-Engine-Definition"],
    "FE/LLDD-FE-Report": ["FE/LLDD-FE-Foundation"],
    "FE/LLDD-FE-Master-Data": ["FE/LLDD-FE-Foundation"],
    # role pack 5 ฉบับเป็นมุมมองต่อบทบาทของหน้าเดียวกัน — ต้องรอโครงหน้า Document Detail นิ่งก่อน
    # (เดิมไม่ได้ประกาศไว้ ทำให้กระดาน Kanban แสดงว่าเริ่มได้ทันทีทั้งที่เริ่มไม่ได้จริง)
    "FE/LLDD-FE-Document-Detail-Role-06-SBP-DSA": ["FE/LLDD-FE-Document-Detail"],
    "FE/LLDD-FE-Document-Detail-Role-08-SBP-DSA-Officer": ["FE/LLDD-FE-Document-Detail"],
    "FE/LLDD-FE-Document-Detail-Role-01-Business-Promotion": ["FE/LLDD-FE-Document-Detail"],
    "FE/LLDD-FE-Document-Detail-Role-02-GM-Business-Promotion": ["FE/LLDD-FE-Document-Detail"],
    "FE/LLDD-FE-Document-Detail-Role-03-AVP-SBP": ["FE/LLDD-FE-Document-Detail"],
}


def document_dependencies(topics_list: list[Topic]) -> dict[str, set[str]]:
    """แผนที่ เอกสาร -> เอกสารที่ต้องจบก่อน (เฉพาะที่มีอยู่จริงใน topics_list)"""
    files = {t.file for t in topics_list}
    job_file = {
        t.file.split("Job-")[1].split("-")[0]: t.file
        for t in topics_list
        if "/Jobs/" in t.file
    }
    preds: dict[str, set[str]] = {t.file: set() for t in topics_list}
    for consumer, producers in DOC_DEPENDENCIES.items():
        if consumer in files:
            preds[consumer] |= {p for p in producers if p in files}
    for consumer, producers in JOB_DEPENDENCIES.items():
        if consumer in job_file:
            preds[job_file[consumer]] |= {job_file[p] for p in producers if p in job_file}
    # สองฉบับนี้อธิบายการต่อกับ "ระบบเดิม" (BFF/auth/mas_param และ @srm/glb-workflow)
    # ไม่ได้อ้าง DDL ของ SBPGI จึงเขียนขนานกับ Database-Structure ได้ตั้งแต่สัปดาห์แรก
    platform_docs = {"BE/LLDD-BE-Integration-SBP-Platform", "BE/LLDD-BE-Workflow-Engine-Definition"}
    fe_leaf = "FE/LLDD-FE-Testing-Delivery"
    if fe_leaf in files:
        preds[fe_leaf] |= {f for f in files if f.startswith("FE/") and f != fe_leaf}
    for topic in topics_list:
        # เอกสารฐานราก: DDL ต้องนิ่งก่อนเขียนเอกสาร BE อื่น · envelope กลางต้องนิ่งก่อนเขียน BE-API
        if (
            topic.file.startswith("BE/")
            and topic.file != "BE/LLDD-BE-Database-Structure"
            and topic.file not in platform_docs
        ):
            preds[topic.file] |= {"BE/LLDD-BE-Database-Structure"} & files
        if topic.file.startswith("BE/LLDD-BE-API") and topic.file != "BE/LLDD-BE-API-Common-Contracts":
            preds[topic.file] |= {"BE/LLDD-BE-API-Common-Contracts"} & files
        if "/Jobs/" in topic.file:
            preds[topic.file] |= {"BE/LLDD-BE-API-Common-Contracts"} & files
    for topic in topics_list:
        preds[topic.file].discard(topic.file)
    return preds


def build_topic_schedule(topics_list: list[Topic], start_date: date = LLDD_START_DATE) -> dict[str, tuple[date, date]]:
    """จัดตารางแบบรู้จัก dependency — เจ้าของหนึ่งคนทำได้ทีละฉบับ และผู้บริโภคเริ่มหลังผู้ผลิตจบ

    ยังคงการอัดชั่วโมงต่อคนแบบเดิม (เอกสาร 9 ชม. กินวันครึ่ง ฉบับถัดไปเริ่มกลางวันได้)
    แต่เมื่อ dependency บังคับให้เริ่มช้ากว่านั้น จะดันชั่วโมงสะสมของเจ้าของไปที่ต้นวันนั้นเลย
    """
    preds = document_dependencies(topics_list)
    priority = {topic.file: index for index, topic in enumerate(topics_list)}
    by_file = {topic.file: topic for topic in topics_list}
    used_hours_by_owner: dict[str, int] = {}
    start_day: dict[str, int] = {}
    end_day: dict[str, int] = {}
    pending = set(by_file)
    while pending:
        ready = [f for f in pending if not (preds[f] & pending)]
        if not ready:
            raise ValueError(f"Cyclic document dependency among: {sorted(pending)}")

        def earliest_start(file_key: str) -> int:
            owner_free = int(used_hours_by_owner.get(by_file[file_key].owner, 0) // HOURS_PER_DAY)
            blocked = max((end_day[p] + 1 for p in preds[file_key]), default=0)
            return max(owner_free, blocked)

        # เลือกงานที่ "เริ่มได้เร็วที่สุดจริง" ก่อน แล้วจึงใช้ลำดับใน MAIN_INDEX_ORDER ตัดสินเสมอ
        # (ถ้าเรียงตามลำดับประกาศอย่างเดียว งานที่ติด dependency จะถูกดันไปท้ายแถวโดยไม่จำเป็น)
        file_key = min(ready, key=lambda f: (earliest_start(f), priority[f]))
        topic = by_file[file_key]
        used_hours = used_hours_by_owner.get(topic.owner, 0)
        start_offset = earliest_start(file_key)
        used_hours = max(used_hours, start_offset * HOURS_PER_DAY)
        start_day[file_key] = start_offset
        # ใช้ชั่วโมงรวม (implementation + unit test) — unit test เป็นงานที่ต้องทำจริง
        # ถ้าคิดแค่ implementation แผนจะสั้นกว่าความจริง 154 ชั่วโมงทั้งโครงการ
        end_day[file_key] = int((used_hours + total_hours(topic) - 1) // HOURS_PER_DAY)
        used_hours_by_owner[topic.owner] = used_hours + total_hours(topic)
        pending.discard(file_key)
    schedule: dict[str, tuple[date, date]] = {}
    for file_key in by_file:
        start = add_workdays(start_date, start_day[file_key] + 1)
        end = add_workdays(start_date, end_day[file_key] + 1)
        if start < start_date or end < start:
            raise ValueError(f"Invalid schedule window for {file_key}: {start} - {end}")
        schedule[file_key] = (start, end)
    return schedule


def schedule_finish_date(topics_list: list[Topic], start_date: date = LLDD_START_DATE) -> date:
    """วันจบจริงของแผน — คำนวณจากตาราง ไม่ใช่ค่าคงที่ที่ตั้งไว้ล่วงหน้า"""
    return max(end for _, end in build_topic_schedule(topics_list, start_date).values())


def dependency_steps(topics_list: list[Topic]) -> dict[str, int]:
    """ลำดับขั้นของงาน = ความลึกที่ยาวที่สุดใน dependency graph + 1

    เอกสารที่อยู่ "ขั้น" เดียวกันไม่มีใครรอใคร จึงลงมือพร้อมกันได้
    ใช้แทนวันที่ในเอกสารส่งมอบ — สื่อลำดับงานโดยไม่ผูกกับปฏิทิน
    """
    preds = document_dependencies(topics_list)
    steps: dict[str, int] = {}

    def depth(file_key: str, seen: frozenset[str] = frozenset()) -> int:
        if file_key in steps:
            return steps[file_key]
        if file_key in seen:
            raise ValueError(f"Cyclic document dependency at {file_key}")
        parents = preds.get(file_key, set())
        steps[file_key] = 1 + max((depth(p, seen | {file_key}) for p in parents), default=0)
        return steps[file_key]

    for topic in topics_list:
        depth(topic.file)
    return steps


_PLANNED_FINISH: date | None = None


def planned_finish_date() -> date:
    """วันจบของแผนทั้งชุด (cache ไว้ เพราะถูกเรียกจากหลายที่ตอนเรนเดอร์)"""
    global _PLANNED_FINISH
    if _PLANNED_FINISH is None:
        counted = [t for t in main_index_ordered(topics()) if not is_document_detail_role_doc(t.file)]
        _PLANNED_FINISH = schedule_finish_date(counted)
    return _PLANNED_FINISH


def image_path(name: str) -> Path:
    path = IMG / name
    if path.exists():
        return path
    path = SLICE / name
    if path.exists():
        return path
    return ROOT / "output/srs/screenshots" / name


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: Any, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    if bold:
        runs = [para.add_run(str(text) if text is not None else "")]
        runs[0].bold = True
    else:
        # แปลง markdown inline ในเซลล์ตารางด้วย (หมายเหตุของ skeleton ใช้ **bold** / `code` เยอะ)
        add_md_runs(para, text)
        runs = list(para.runs) or [para.add_run("")]
    for run in runs:
        set_run_fonts(run, "Courier New" if run.font.name == "Courier New" else "Arial")
        run.font.size = Pt(9)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.keep_together = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_docx_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, head in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], head, True)
        set_cell_shading(t.rows[0].cells[i], "E8EEF5")
    set_row_cant_split(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        normalized = list(row[: len(headers)]) + [""] * max(0, len(headers) - len(row))
        for i, val in enumerate(normalized):
            set_cell_text(cells[i], val)
        set_row_cant_split(t.rows[-1])
    doc.add_paragraph()


def add_docx_payload(doc: Document, title: str, text: str) -> None:
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    header = t.rows[0].cells[0]
    set_cell_text(header, title, True)
    set_cell_shading(header, "DDEBFF")
    body = t.rows[1].cells[0]
    body.text = ""
    para = body.paragraphs[0]
    para.paragraph_format.keep_together = True
    run = para.add_run(text)
    set_run_fonts(run, "Courier New")
    run.font.size = Pt(8)
    para.paragraph_format.space_after = Pt(0)
    set_cell_shading(body, "F7FAFE")
    for row in t.rows:
        set_row_cant_split(row)
    doc.add_paragraph()


def add_docx_image(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        doc.add_paragraph(f"[ไม่พบรูปภาพ: {path.name}]")
        return
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.keep_with_next = True
    pic.add_run().add_picture(str(path), width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_together = True
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.italic = True


_MD_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`", re.S)


def add_md_runs(paragraph, text: Any) -> None:
    """เขียน text ลง paragraph โดยแปลง markdown inline (**bold** / `code`) เป็น run จริง

    ใช้เฉพาะ block ชนิด paragraph/bullet — block ชนิด code ต้องคงข้อความดิบไว้
    """
    raw = str(text if text is not None else "")
    pos = 0
    for m in _MD_INLINE_RE.finditer(raw):
        if m.start() > pos:
            paragraph.add_run(raw[pos:m.start()])
        if m.group(1) is not None:
            paragraph.add_run(m.group(1)).bold = True
        else:
            body = m.group(2)
            run = paragraph.add_run(body)
            # Courier New ไม่การันตี glyph ไทย — ใช้เฉพาะ ASCII ล้วน (เหตุผลเดียวกับฝั่ง PDF)
            try:
                body.encode("ascii")
                set_run_fonts(run, "Courier New")
            except UnicodeEncodeError:
                run.font.color.rgb = RGBColor.from_string("8A3324")
        pos = m.end()
    if pos < len(raw):
        paragraph.add_run(raw[pos:])


def build_docx(title: str, blocks: list[dict[str, Any]], out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    set_style_fonts(styles["Normal"], "Arial")
    styles["Normal"].font.size = Pt(10)
    for style_name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11, "1F4D78"),
        ("Heading 4", 10, "2E74B5"),
    ]:
        st = styles[style_name]
        st.font.name = "Arial"
        set_style_fonts(st, "Arial")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    set_run_fonts(title_run, "Arial")
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    doc.add_paragraph("SBP Mall - ระบบประกันรายได้ | Low Level Design Document")

    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_fonts(run, "Arial")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("66717F")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("หน้า ")
    set_run_fonts(footer_run, "Arial")
    footer_run.font.size = Pt(8)
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    footer._p.append(page_field)

    heading_pending = False
    page_break_pending = False
    figure_no = 0
    for block in blocks:
        btype = block["type"]
        if btype == "h1":
            para = doc.add_heading(block["text"], level=1)
            para.paragraph_format.page_break_before = page_break_pending
            page_break_pending = False
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.keep_together = True
            heading_pending = True
        elif btype == "h2":
            para = doc.add_heading(block["text"], level=2)
            para.paragraph_format.page_break_before = page_break_pending
            page_break_pending = False
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.keep_together = True
            heading_pending = True
        elif btype in ("h3", "h4"):
            para = doc.add_heading(block["text"], level=3 if btype == "h3" else 4)
            para.paragraph_format.page_break_before = page_break_pending
            page_break_pending = False
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.keep_together = True
            heading_pending = True
        elif btype == "p":
            para = doc.add_paragraph()
            add_md_runs(para, block["text"])
            if heading_pending:
                para.paragraph_format.keep_with_next = False
            heading_pending = False
        elif btype == "bullets":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Bullet")
                add_md_runs(para, item)
                para.paragraph_format.keep_together = True
            heading_pending = False
        elif btype == "table":
            add_docx_table(doc, block["headers"], block["rows"])
            heading_pending = False
        elif btype == "image":
            figure_no += 1
            add_docx_image(doc, ROOT / block["path"], f"รูปที่ {figure_no}: {block['caption']}")
            heading_pending = False
        elif btype == "code":
            para = doc.add_paragraph()
            run = para.add_run(block["text"])
            set_run_fonts(run, "Courier New")
            run.font.size = Pt(7 if block.get("lang") == "java" else 8)
            para.paragraph_format.left_indent = Inches(0.2)
            para.paragraph_format.keep_together = False
            heading_pending = False
        elif btype == "payload":
            add_docx_payload(doc, block["title"], block["text"])
            heading_pending = False
        elif btype == "pagebreak":
            page_break_pending = True
            heading_pending = False

    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        trailing = doc.paragraphs[-1]._element
        trailing.getparent().remove(trailing)
    doc.save(out_path)


def init_pdf_styles():
    if os.path.exists(FONT):
        pdfmetrics.registerFont(TTFont("TH", FONT))
        pdfmetrics.registerFont(TTFont("TH-Bold", FONT))
        base = "TH"
    else:
        base = "Helvetica"
    ss = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("TitleTH", parent=ss["Title"], fontName=base, fontSize=18, leading=24, textColor=colors.HexColor("#0B2545"), alignment=TA_LEFT),
        "h1": ParagraphStyle("H1TH", parent=ss["Heading1"], fontName=base, fontSize=14, leading=18, textColor=colors.HexColor("#1F4D78"), spaceBefore=12, spaceAfter=6),
        "h2": ParagraphStyle("H2TH", parent=ss["Heading2"], fontName=base, fontSize=12, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=8, spaceAfter=4),
        "h3": ParagraphStyle("H3TH", parent=ss["Heading3"], fontName=base, fontSize=10.5, leading=14, textColor=colors.HexColor("#1F4D78"), spaceBefore=6, spaceAfter=3),
        "h4": ParagraphStyle("H4TH", parent=ss["Heading4"], fontName=base, fontSize=9.5, leading=13, textColor=colors.HexColor("#2E74B5"), spaceBefore=5, spaceAfter=2),
        "body": ParagraphStyle("BodyTH", parent=ss["BodyText"], fontName=base, fontSize=9.2, leading=13, spaceAfter=5),
        "small": ParagraphStyle("SmallTH", parent=ss["BodyText"], fontName=base, fontSize=8, leading=10, spaceAfter=3),
        "code": ParagraphStyle("CodeTH", parent=ss["Code"], fontName=base, fontSize=7.4, leading=9, backColor=colors.HexColor("#F4F6F9"), borderPadding=4),
        "java": ParagraphStyle("JavaTH", parent=ss["Code"], fontName=base, fontSize=6.4, leading=7.8, backColor=colors.HexColor("#F4F6F9"), borderPadding=4),
    }


_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*", re.S)
_MD_CODE_RE = re.compile(r"`([^`]+)`")


def para(text: Any, style: ParagraphStyle) -> Paragraph:
    esc = str(text if text is not None else "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    esc = esc.replace("\n", "<br/>")
    # แปลง markdown inline ที่ generator ใช้อยู่ (ไม่งั้น PDF จะโชว์ดอกจันดิบ)
    esc = _MD_BOLD_RE.sub(r"<b>\1</b>", esc)
    # inline code: Courier ไม่มี glyph ภาษาไทย ถ้าสลับ face จะได้ ■■■ ทั้งก้อน
    # จึงใช้ Courier เฉพาะเนื้อหาที่เป็น ASCII ล้วน · ที่เหลือคงฟอนต์ไทยไว้แล้วใช้สีแทนเพื่อยังเห็นว่าเป็น code
    def _inline_code(m: "re.Match[str]") -> str:
        body = m.group(1)
        try:
            body.encode("ascii")
        except UnicodeEncodeError:
            return f"<font color='#8A3324'>{body}</font>"
        return f"<font face='Courier'>{body}</font>"

    esc = _MD_CODE_RE.sub(_inline_code, esc)
    return Paragraph(esc, style)


def code_para(text: Any, style: ParagraphStyle) -> Paragraph:
    escaped_lines = []
    for line in str(text if text is not None else "").splitlines():
        leading = len(line) - len(line.lstrip(" "))
        body = line[leading:].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        escaped_lines.append("&nbsp;" * leading + body)
    esc = "<br/>".join(escaped_lines)
    return Paragraph(esc, style)


def add_pdf_table(story: list[Any], styles: dict[str, ParagraphStyle], headers: list[str], rows: list[list[Any]]) -> None:
    data = [[para(h, styles["small"]) for h in headers]]
    for row in rows:
        normalized = list(row[: len(headers)]) + [""] * max(0, len(headers) - len(row))
        data.append([para(v, styles["small"]) for v in normalized])
    col_count = max(1, len(headers))
    widths = [7.0 * inch / col_count] * col_count
    t = Table(data, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#C9D2DC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    story.append(Spacer(1, 8))


def add_pdf_payload(story: list[Any], styles: dict[str, ParagraphStyle], title: str, text: str) -> None:
    data = [
        [para(title, styles["small"])],
        [code_para(text, styles["code"])],
    ]
    t = Table(data, colWidths=[7.0 * inch])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BFD0E6")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DDEBFF")),
        ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#F7FAFE")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 8))


def add_pdf_image(story: list[Any], styles: dict[str, ParagraphStyle], path: Path, caption: str) -> None:
    if not path.exists():
        story.append(para(f"[ไม่พบรูปภาพ: {path.name}]", styles["body"]))
        return
    with Image.open(path) as im:
        w, h = im.size
    max_w = 7.0 * inch
    max_h = 4.4 * inch
    scale = min(max_w / w, max_h / h, 1.0)
    story.append(KeepTogether([PdfImage(str(path), width=w * scale, height=h * scale), para(caption, styles["small"]), Spacer(1, 8)]))


def build_pdf(title: str, blocks: list[dict[str, Any]], out_path: Path) -> None:
    styles = init_pdf_styles()
    doc = SimpleDocTemplate(str(out_path), pagesize=A4, leftMargin=0.55 * inch, rightMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch)
    story: list[Any] = [Paragraph(title, styles["title"]), para("SBP Mall - ระบบประกันรายได้ | Low Level Design Document", styles["body"]), Spacer(1, 8)]
    figure_no = 0
    for block in blocks:
        btype = block["type"]
        if btype in ("h1", "h2", "h3", "h4"):
            min_space = {"h1": 1.6 * inch, "h2": 1.15 * inch, "h3": 0.9 * inch, "h4": 0.8 * inch}[btype]
            story.append(CondPageBreak(min_space))
            story.append(Paragraph(block["text"], styles[btype]))
        elif btype == "p":
            story.append(para(block["text"], styles["body"]))
        elif btype == "bullets":
            story.append(ListFlowable([ListItem(para(i, styles["body"])) for i in block["items"]], bulletType="bullet", leftIndent=16))
        elif btype == "table":
            add_pdf_table(story, styles, block["headers"], block["rows"])
        elif btype == "image":
            figure_no += 1
            add_pdf_image(story, styles, ROOT / block["path"], f"รูปที่ {figure_no}: {block['caption']}")
        elif btype == "code":
            story.append(code_para(block["text"], styles["java"] if block.get("lang") == "java" else styles["code"]))
        elif btype == "payload":
            before = len(story)
            add_pdf_payload(story, styles, block["title"], block["text"])
            story[before:] = [KeepTogether(story[before:])]
        elif btype == "pagebreak":
            story.append(PageBreak())

    def add_page_number(canvas, pdf_doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#66717F"))
        canvas.drawCentredString(A4[0] / 2, 0.28 * inch, f"{pdf_doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


def md_table(headers: list[str], rows: list[list[Any]]) -> str:
    def cell(value: Any) -> str:
        return str(value).replace("|", "\\|").replace("\n", "<br>")

    out = ["| " + " | ".join(cell(header) for header in headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        normalized = list(row[: len(headers)]) + [""] * max(0, len(headers) - len(row))
        out.append("| " + " | ".join(cell(value) for value in normalized) + " |")
    return "\n".join(out)


ARTIFACT_LABELS = {
    "workflow.md": "Workflow design",
    "api.md": "API design",
    "database.md": "Database design",
    "plan-api.html": "API specification screen",
    "plan-database.html": "Database design screen",
    "plan-flow.html": "Integrated flow screen",
    "index.html": "Portal screen",
    "k2-create.html": "Create Document screen",
    "k2-document.html": "Document Detail screen",
    "k2-report.html": "Status Report screen",
    "k2-list-waiting.html": "Task Inbox screen",
    "k2-list-related.html": "Related Documents screen",
    "k2-list-abnormal.html": "Abnormal Data screen",
    "k2-operators.html": "Operator Master screen",
    "k2-factors.html": "External Factor Master screen",
    "k2-permissions.html": "RBAC Matrix screen",
}


def scrub_pdf_reference(match: re.Match[str]) -> str:
    stem = Path(match.group(1)).name
    if stem.startswith("LLDD-"):
        return f"{stem}.pdf"
    return "ไฟล์" + Path(stem).stem.replace("-", " ")


def scrub_lldd_text(value: Any) -> str:
    text = str(value)
    text = re.sub(
        r"(?i)(?:[A-Za-z0-9_.-]+/)*(LLDD-[A-Za-z0-9-]+)(?:\.(?:md|docx|pdf))?",
        lambda m: f"{m.group(1)}.pdf",
        text,
    )
    for source, target in ARTIFACT_LABELS.items():
        text = text.replace(source, target)
    text = re.sub(r"(?i)\bSRS\b", "ข้อกำหนดทางธุรกิจ", text)
    text = re.sub(r"([\w./\-\u0E00-\u0E7F]+)\.md\b", lambda m: Path(m.group(1)).name.replace("-", " "), text, flags=re.IGNORECASE)
    text = re.sub(r"([\w./\-\u0E00-\u0E7F]+)\.html\b", lambda m: Path(m.group(1)).name.replace("-", " ") + " screen", text, flags=re.IGNORECASE)
    text = re.sub(r"([\w./\-\u0E00-\u0E7F]+)\.docx\b", lambda m: Path(m.group(1)).stem.replace("-", " "), text, flags=re.IGNORECASE)
    text = re.sub(r"([\w./\-\u0E00-\u0E7F]+)\.pdf\b", scrub_pdf_reference, text, flags=re.IGNORECASE)
    return text


def scrub_lldd_block(block: dict[str, Any], preserve_java: bool = False) -> dict[str, Any]:
    cleaned = dict(block)
    if preserve_java and cleaned.get("type") == "code" and cleaned.get("lang") == "java":
        return cleaned
    for key in ("text", "title", "caption"):
        if key in cleaned:
            cleaned[key] = scrub_lldd_text(cleaned[key])
    for key in ("headers", "items"):
        if key in cleaned:
            cleaned[key] = [scrub_lldd_text(item) for item in cleaned[key]]
    if "rows" in cleaned:
        cleaned["rows"] = [[scrub_lldd_text(cell) for cell in row] for row in cleaned["rows"]]
    return cleaned


def scrub_lldd_blocks(blocks: list[dict[str, Any]], preserve_java: bool = False) -> list[dict[str, Any]]:
    return [scrub_lldd_block(block, preserve_java=preserve_java) for block in blocks]


def delivery_intro_blocks(is_job: bool) -> list[dict[str, Any]]:
    rows = [
        ["วัตถุประสงค์", "ใช้เป็นรายละเอียดระดับพัฒนาสำหรับออกแบบ ลงมือพัฒนา ตรวจทาน และทดสอบขอบเขตที่ระบุในฉบับนี้"],
        ["ลำดับการอ่าน", "เริ่มจาก Overview และ Scope จากนั้นตรวจ Field/Validation, Implementation, Contract, Processing Flow, Acceptance Criteria และ Test Checklist ตามลำดับ"],
        ["ผลลัพธ์ที่คาดหวัง", "ผู้อ่านสามารถระบุ input, ขั้นตอนประมวลผล, output, เงื่อนไขผิดพลาด และหลักฐานการทดสอบได้จากเนื้อหาในฉบับเดียว"],
    ]
    if is_job:
        rows.append(["Java เดิม", "ภาคผนวกท้ายฉบับระบุ source file, ช่วงบรรทัด และ code Java เดิมสำหรับตรวจเทียบ behavior ก่อนย้ายระบบ"])
    return [
        h(1, "แนวทางการใช้เอกสาร"),
        table(["หัวข้อ", "คำอธิบาย"], rows),
        p("คำว่า Input, Progress และ Output ในเอกสารนี้หมายถึงข้อมูลตั้งต้น ลำดับการทำงาน และผลลัพธ์ที่ตรวจสอบได้ของขอบเขตที่กำลังพัฒนา"),
        pagebreak(),
    ]


def parse_line_ranges(value: str) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    for part in value.split(","):
        match = re.fullmatch(r"\s*(\d+)\s*-\s*(\d+)\s*", part)
        if match:
            ranges.append((int(match.group(1)), int(match.group(2))))
    return ranges


def compact_excerpt_ranges(start: int, end: int, max_lines: int = 24) -> list[tuple[int, int]]:
    if end - start + 1 <= max_lines:
        return [(start, end)]
    half = max_lines // 2
    return [(start, start + half - 1), (end - half + 1, end)]


def legacy_java_appendix_blocks(document_key: str) -> list[dict[str, Any]]:
    if "/Jobs/" not in document_key:
        return []
    job_no = document_key.split("LLDD-BE-Job-", 1)[1].split("-", 1)[0]
    legacy = LEGACY_JOB_SOURCES.get(job_no)
    if not legacy:
        return []

    blocks: list[dict[str, Any]] = [
        pagebreak(),
        h(1, "ภาคผนวก: Java Source เดิม"),
        p("Code ในส่วนนี้คัดจาก Java source เดิมโดยตรงและคงข้อความตามต้นฉบับ ใช้เลขบรรทัดที่ระบุหน้าทุก snippet เพื่อตรวจเทียบ business behavior, SQL, transaction และ error handling"),
    ]
    for source_index, (relative_path, line_ranges, responsibility) in enumerate(legacy["sources"], start=1):
        source_path = ROOT.parent / relative_path
        if not source_path.exists():
            blocks.extend([
                h(2, f"J{source_index}. {relative_path}"),
                p(f"ไม่พบ Java source ที่ {relative_path}; ต้องตรวจ path ก่อนเริ่มพัฒนา"),
            ])
            continue
        source_lines = source_path.read_text(encoding="utf-8", errors="replace").splitlines()
        first_excerpt = True
        for declared_start, declared_end in parse_line_ranges(line_ranges):
            actual_start = max(1, declared_start)
            actual_end = min(len(source_lines), declared_end)
            for excerpt_start, excerpt_end in compact_excerpt_ranges(actual_start, actual_end):
                if not first_excerpt:
                    blocks.append(pagebreak())
                heading = f"J{source_index}. {Path(relative_path).name} — lines {excerpt_start}-{excerpt_end}"
                blocks.extend([
                    h(2, heading),
                    table(["รายการ", "รายละเอียด"], [
                        ["Source file", relative_path],
                        ["Original lines", f"{excerpt_start}-{excerpt_end}"],
                        ["Responsibility", responsibility],
                    ]),
                    code("\n".join(source_lines[excerpt_start - 1:excerpt_end]), "java"),
                ])
                first_excerpt = False
    return blocks


def prepare_delivery_blocks(title: str, blocks: list[dict[str, Any]], document_key: str) -> tuple[str, list[dict[str, Any]]]:
    is_job = "/Jobs/" in document_key
    combined = [*delivery_intro_blocks(is_job), *blocks, *legacy_java_appendix_blocks(document_key)]
    return scrub_lldd_text(title), scrub_lldd_blocks(combined, preserve_java=is_job)


def build_md(title: str, blocks: list[dict[str, Any]], out_path: Path) -> None:
    lines = [f"# {title}", "", "SBP Mall - ระบบประกันรายได้ | Low Level Design Document", ""]
    figure_no = 0
    for block in blocks:
        btype = block["type"]
        if btype == "h1":
            lines.extend([f"## {block['text']}", ""])
        elif btype == "h2":
            lines.extend([f"### {block['text']}", ""])
        elif btype == "h3":
            lines.extend([f"#### {block['text']}", ""])
        elif btype == "h4":
            lines.extend([f"##### {block['text']}", ""])
        elif btype == "p":
            lines.extend([block["text"], ""])
        elif btype == "bullets":
            lines.extend([f"- {i}" for i in block["items"]])
            lines.append("")
        elif btype == "table":
            lines.extend([md_table(block["headers"], block["rows"]), ""])
        elif btype == "image":
            figure_no += 1
            caption = f"รูปที่ {figure_no}: {block['caption']}"
            rel = os.path.relpath(ROOT / block["path"], out_path.parent)
            lines.extend([f"![{caption}]({rel})", "", f"_{caption}_", ""])
        elif btype == "code":
            lines.extend([f"```{block.get('lang', '')}", block["text"], "```", ""])
        elif btype == "payload":
            lines.extend([f"#### {block['title']}", "", "```json", block["text"], "```", ""])
    out_path.write_text("\n".join(lines), encoding="utf-8")


def render_all(title: str, blocks: list[dict[str, Any]], base: Path, formats: set[str]) -> None:
    relative_base = base.relative_to(OUT)
    if "md" in formats:
        md_path = (OUT / FORMAT_DIRS["md"] / relative_base).with_suffix(".md")
        md_path.parent.mkdir(parents=True, exist_ok=True)
        build_md(title, blocks, md_path)
    delivery_title, delivery_blocks = prepare_delivery_blocks(title, blocks, str(relative_base))
    if "docx" in formats:
        docx_path = (OUT / FORMAT_DIRS["docx"] / relative_base).with_suffix(".docx")
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        build_docx(delivery_title, delivery_blocks, docx_path)
    if "pdf" in formats:
        pdf_path = (OUT / FORMAT_DIRS["pdf"] / relative_base).with_suffix(".pdf")
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        build_pdf(delivery_title, delivery_blocks, pdf_path)


# เอกสารที่ไม่ผ่าน skeleton generator: ไม่มี REST endpoint ของตัวเอง (skeleton จะได้แต่โครงว่าง)
SKELETON_SKIP_FILES = {
    "FE/LLDD-FE-Testing-Delivery",
    "BE/LLDD-BE-Database-Structure",
    "BE/LLDD-BE-Data-Migration-Cutover",
    "BE/LLDD-BE-Integration-SBP-Platform",
    "BE/LLDD-BE-Workflow-Engine-Definition",
}
_SKELETON_SQL_CACHE: dict[str, str] | None = None
_BATCH_JOBS_BY_NO: dict[str, dict[str, Any]] | None = None


def api_sql_map() -> dict[str, str]:
    """SQL_BY_PATH จาก plan-api.html (อ่านครั้งเดียวแล้ว cache — การอ่านต้องเรียก node)"""
    global _SKELETON_SQL_CACHE
    if _SKELETON_SQL_CACHE is None:
        try:
            _SKELETON_SQL_CACHE = plan_api_sql_by_path()
        except Exception:
            _SKELETON_SQL_CACHE = {}
    return _SKELETON_SQL_CACHE


def batch_job_by_no(job_no: str) -> dict[str, Any] | None:
    """dict ของ job ตัวนั้นจาก JOBS ใน job-batch.html (cache ทั้ง array ไว้ครั้งเดียว)"""
    global _BATCH_JOBS_BY_NO
    if _BATCH_JOBS_BY_NO is None:
        try:
            _BATCH_JOBS_BY_NO = {str(job.get("no")): job for job in load_batch_jobs()}
        except Exception:
            _BATCH_JOBS_BY_NO = {}
    return _BATCH_JOBS_BY_NO.get(str(job_no))


def job_no_from_file(file_key: str) -> str:
    match = re.search(r"LLDD-BE-Job-([0-9]+[A-Za-z]?)-", file_key)
    return match.group(1) if match else ""


def is_job_doc(file_key: str) -> bool:
    return file_key.startswith("BE/Jobs/")


def map_block_text(block: dict[str, Any], fn) -> dict[str, Any]:
    """apply fn กับทุก string ที่แสดงผลใน block เดียว (ใช้ renumber หัวข้อของ skeleton module)"""
    for key in ("text", "title", "caption"):
        if isinstance(block.get(key), str):
            block[key] = fn(block[key])
    for key in ("headers", "items"):
        if isinstance(block.get(key), list):
            block[key] = [fn(item) if isinstance(item, str) else item for item in block[key]]
    if isinstance(block.get("rows"), list):
        block["rows"] = [
            [fn(cell) if isinstance(cell, str) else cell for cell in row] if isinstance(row, list) else row
            for row in block["rows"]
        ]
    return block


def promote_skeleton_heading(block: dict[str, Any], numbers: set[str]) -> dict[str, Any]:
    """หัวข้อหลักของ module (h2 ขึ้นต้นด้วยเลข section) -> h1 พร้อมจุดให้ตรงกับหัวข้ออื่นในเอกสาร"""
    if block.get("type") != "h2":
        return block
    match = re.match(r"^(\d+)\s+(.*)$", str(block.get("text", "")))
    if match and match.group(1) in numbers:
        block["type"] = "h1"
        block["text"] = f"{match.group(1)}. {match.group(2)}"
    return block


def skeleton_code_blocks(topic: Topic, section_no: int) -> list[dict[str, Any]]:
    """Skeleton Code section ของเอกสารหนึ่งฉบับ โดยหัวข้อหลักเริ่มที่ `section_no`

    - เอกสาร Job (`BE/Jobs/...`) -> lldd_skeleton_job (ctx: job dict จาก JOBS ใน job-batch.html)
    - เอกสาร BE อื่น            -> lldd_skeleton_be  (ctx: sql_by_path จาก SQL_BY_PATH ใน plan-api.html)
    - เอกสาร FE (ยกเว้น testing/delivery) -> lldd_skeleton_fe
    ใช้ h1 เป็นหัวข้อหลักเสมอ เพื่อให้นับต่อกับหัวข้ออื่นของ topic_blocks ได้
    """
    file_key = topic.file
    if file_key in SKELETON_SKIP_FILES:
        return []
    try:
        if is_job_doc(file_key):
            job_no = job_no_from_file(file_key)
            raw = job_skeleton_blocks(topic, {"job": batch_job_by_no(job_no)})
            if not raw:
                return []

            def renumber(text: str) -> str:
                return re.sub(
                    r"\b5\.(9[4-9])(\.\d+)?\b",
                    lambda m: f"{section_no}.{int(m.group(1)) - 93}{m.group(2) or ''}",
                    text,
                )

            raw = [map_block_text(dict(block), renumber) for block in raw]

            def normalize(block: dict[str, Any]) -> dict[str, Any]:
                """ระดับหัวข้อย่อยของ skeleton ให้เท่ากันทั้ง 3 track (FE/BE ใช้ h3, h4)"""
                if block.get("type") == "h2":
                    block["type"] = "h3"
                elif block.get("type") == "h3":
                    block["type"] = "h4"
                if isinstance(block.get("text"), str):
                    block["text"] = block["text"].replace("Skeleton Code — ", "")
                return block

            raw = [normalize(block) for block in raw]
            title = f"{section_no}. Skeleton Code (Batch Job {job_no})" if job_no else f"{section_no}. Skeleton Code"
            return [h(1, title)] + raw
        if topic.track == "BE":
            raw = be_skeleton_blocks(topic, {
                "sql_by_path": api_sql_map(),
                "skeleton_section": str(section_no),
                "sql_section": str(section_no + 1),
            })
            numbers = {str(section_no), str(section_no + 1)}
            return [promote_skeleton_heading(dict(block), numbers) for block in raw]
        if topic.track == "FE":
            raw = fe_skeleton_blocks(topic, {"section_prefix": str(section_no), "heading_level": 1})
            blocks = [dict(block) for block in raw]
            for block in blocks:
                if block.get("type") == "h1":
                    block["text"] = re.sub(r"^(\d+)\s+", r"\1. ", str(block.get("text", "")))
                    break
            return blocks
    except Exception as error:  # generator ต้องไม่ล้มทั้งชุดเพราะ skeleton ฉบับเดียว
        return [
            h(1, f"{section_no}. Skeleton Code"),
            p(f"ไม่สามารถสร้าง skeleton code อัตโนมัติสำหรับเอกสารนี้ได้ ({type(error).__name__}: {error})"),
        ]
    return []


def count_top_sections(blocks: list[dict[str, Any]]) -> int:
    return sum(1 for block in blocks if block.get("type") == "h1")


# ---------------------------------------------------------------------------
# ชั่วโมง Unit Test (เพิ่ม 2026-08-18 ตามข้อกำหนด: LLDD ต้องประเมิน unit test ของงานนั้นด้วย)
# ---------------------------------------------------------------------------
# อัตราคิดจากลักษณะงาน ไม่ใช่ตัวเลขลอย:
#   BE / Job  30%  — business rule + SQL + transaction/rollback ต้องมี test คลุม
#                    (jest + repository mock; job ต้องมี fixture ของไฟล์ interface)
#   FE        25%  — component/hook test ด้วย React Testing Library + mock API layer
#   เอกสารสัญญา/ออกแบบ  0%  — ไม่มีโค้ดของตัวเอง หรือเป็นเอกสารที่กำกับการทดสอบอยู่แล้ว
NO_UNIT_TEST_DOCS = {
    "BE/LLDD-BE-API-Common-Contracts",      # สัญญากลาง — ทดสอบผ่าน endpoint ที่ใช้สัญญานี้
    "BE/LLDD-BE-Database-Structure",        # DDL/index — ตรวจด้วย migration verification ไม่ใช่ unit test
    "BE/LLDD-BE-Data-Migration-Cutover",    # reconcile script — มีเกณฑ์ตรวจรับของตัวเอง
    "BE/LLDD-BE-Workflow-Engine-Definition",# config ของ engine ระบบเดิม — ไม่ได้เขียนโค้ดเอง
    "BE/LLDD-BE-Integration-SBP-Platform",  # สัญญาต่อระบบเดิม — ทดสอบผ่าน integration ของเส้นที่ใช้
    "FE/LLDD-FE-Integration-Contracts",     # สัญญากลางฝั่ง FE
    "FE/LLDD-FE-Testing-Delivery",          # เอกสารกำกับการทดสอบเอง
}


def unit_test_hours(topic: "Topic") -> int:
    """ชั่วโมงเขียน unit test ของเอกสารนั้น (ปัดขึ้นเป็นจำนวนเต็มชั่วโมง)"""
    if topic.file in NO_UNIT_TEST_DOCS:
        return 0
    ratio = 0.25 if topic.track == "FE" else 0.30
    return math.ceil(topic.hours * ratio)


def total_hours(topic: "Topic") -> int:
    return topic.hours + unit_test_hours(topic)


def estimate_cell(topic: "Topic") -> str:
    ut = unit_test_hours(topic)
    if not ut:
        return f"{topic.hours} ชั่วโมง (ไม่มี unit test แยก — ดูเหตุผลใน NO_UNIT_TEST_DOCS)"
    return (f"**{total_hours(topic)} ชั่วโมง** = implementation {topic.hours} + "
            f"unit test {ut} ({int((0.25 if topic.track == 'FE' else 0.30) * 100)}%)")


def unit_test_scope_blocks(topic: "Topic", section_no: int) -> list[dict[str, Any]]:
    """ขอบเขต unit test ของเอกสารนั้น — ผูกกับชั่วโมงที่ประเมินไว้ใน Overview

    รายการทดสอบ derive จากข้อมูลของเอกสารเอง (field/validation · acceptance ·
    endpoint · ตารางที่เขียน) ไม่ใช่ checklist สำเร็จรูป จึงต่างกันจริงทุกฉบับ
    ต่างจาก "Developer Test Checklist" ที่เป็น scenario ระดับ end-to-end/manual
    """
    ut = unit_test_hours(topic)
    if not ut:
        return []
    is_fe = topic.track == "FE"
    is_job = "/Jobs/" in topic.file

    rows: list[list[str]] = []

    # 1) validation ของแต่ละ field ที่มี "กฎ" จริง
    #    ตาราง field ของเอกสาร Job ใช้คอลัมน์นี้เป็นธง แก้ไขได้/แก้ไม่ได้ ไม่ใช่กฎ validation
    #    จึงต้องคัดออก ไม่งั้นจะได้เคสไร้สาระแบบ "ผ่านเมื่อถูกกฎ — กฎ: แก้ไขได้"
    EDITABILITY = {"แก้ไขได้", "ค่าคงที่/แก้ผ่านหน้าจอไม่ได้", "-", "response only",
                   "optional", "read-only", "response only"}
    RULE_HINT = ("กฎ", "เงื่อนไข", "เกณฑ์", "rule", "format", "pattern")
    import re as _re
    # ค่าที่ไม่ใช่กฎ: ตำแหน่งคอลัมน์ (column 3) · ธงไม่บังคับล้วน (optional …) ที่ไม่มีเงื่อนไขต่อท้าย
    NOT_A_RULE = _re.compile(r"^(column\s*\d+|row\s*key|display\s*only|optional(\s+\w+){0,3})$", _re.I)
    for name, fmt, validation, behavior in topic.fields:
        v = (validation or "").strip()
        if NOT_A_RULE.match(v):
            continue
        if v in EDITABILITY or "แก้ผ่านหน้าจอไม่ได้" in v or not v:
            # ไม่มีกฎ validation — แต่ถ้าชื่อ field บอกว่าเป็นกฎ/เงื่อนไข ให้ทดสอบตัวกฎจาก format แทน
            if any(k in name.lower() or k in name for k in RULE_HINT) and fmt:
                rows.append([f"`{name}`", "rule", f"ใช้กฎกับข้อมูลตัวอย่างแล้วได้ผลตามที่ระบุ — {fmt}"])
            continue
        rows.append([
            f"`{name}`",
            "validation",
            f"ผ่านเมื่อถูกกฎ / โยน error เมื่อผิด — กฎ: {v}" + (f" · รูปแบบ: {fmt}" if fmt else ""),
        ])

    # 2) กฎธุรกิจจาก Acceptance Criteria
    for rule in topic.acceptance:
        rows.append(["business rule", "logic", rule])

    # 3) endpoint: success + error envelope
    for spec in topic.apis:
        if is_fe:
            rows.append([
                f"`{spec.method} {spec.path}`",
                "api client",
                "hook/service เรียกเส้นนี้ด้วยพารามิเตอร์ถูกต้อง · map {success:true,data} เป็น state ที่หน้าจอใช้ "
                "· เจอ {success:false,error} แล้วแสดงข้อความไทย verbatim (mock ด้วย msw)",
            ])
        else:
            rows.append([
                f"`{spec.method} {spec.path}`",
                "handler",
                "คืน {success:true,data} ตามรูปแบบที่ระบุ และคืน {success:false,error:{code,message}} "
                "เมื่อ input ผิด — mock repository/lib ไม่แตะ DB จริง",
            ])

    # 4) เฉพาะฝั่งที่เขียนข้อมูล: transaction / rollback
    writes = [r[0] for r in topic.db_tables if str(r[1]).upper().startswith(("W", "R/W"))]
    if writes:
        rows.append([
            ", ".join(f"`{w}`" for w in writes[:3]),
            "transaction",
            "จำลอง error กลางทาง แล้วยืนยันว่า rollback ครบ ไม่เหลือแถวค้าง (mock DataSource/QueryRunner)",
        ])

    # 5) รายการเฉพาะ track
    if is_job:
        rows.append(["runner", "idempotency", "รันซ้ำด้วย fixture เดิมต้องไม่เกิดแถวซ้ำ (ON CONFLICT / business unique key ทำงาน)"])
        rows.append(["runner", "lock", "เรียกซ้อนขณะกำลังรัน ต้องถูกปฏิเสธด้วย advisory lock"])
    elif is_fe:
        rows.append(["component", "render", "render ด้วย React Testing Library แล้วเห็น element ตาม field/action contract ของเอกสารนี้"])
        rows.append(["hook/state", "interaction", "ยิง action แล้ว state เปลี่ยนตามที่ระบุ และเรียก API layer ที่ mock ไว้ด้วยพารามิเตอร์ถูกต้อง"])
        rows.append(["error path", "ui", "API ตอบ error envelope แล้วหน้าจอต้องแสดงข้อความไทย verbatim ไม่ crash"])
    else:
        rows.append(["service", "error mapping", "แปลง error ของ repository/lib เป็น error code ตามสัญญากลาง (LLDD-BE-API-Common-Contracts)"])

    ratio = 25 if is_fe else 30
    tool = ("Jest + React Testing Library + msw (mock API layer)" if is_fe
            else "Jest + mock repository/DataSource (ไม่ต่อ DB จริง)")

    return [
        h(1, f"{section_no}. Unit Test Scope"),
        p(f"**{ut} ชั่วโมง** ({ratio}% ของ implementation {topic.hours} ชั่วโมง) · เครื่องมือ: {tool}"),
        p("หัวข้อนี้คือ **unit test** ที่ต้องเขียนคู่กับโค้ด — ต่างจาก *Developer Test Checklist* "
          "ซึ่งเป็น scenario ระดับ end-to-end/manual ที่ใช้ตอนตรวจรับ · "
          "รายการด้านล่าง derive จาก field/validation, acceptance criteria, endpoint และตารางที่เอกสารนี้เขียน"),
        table(["สิ่งที่ทดสอบ", "ประเภท", "เกณฑ์ผ่าน"], rows),
        bullets([
            "ทุกเคสต้องรันได้โดยไม่ต่อ DB/บริการภายนอกจริง — mock ที่ขอบ repository/client เสมอ",
            "ข้อความไทยที่ยืนยันในเทสต้องเป็น verbatim ตาม SRS ห้ามพิมพ์ใหม่",
            "เกณฑ์ผ่านของ CI: ทุกเคสในตารางนี้มี test จริงและผ่านทั้งหมด",
        ]),
    ]


# DOCX: ภาษาไทยเป็น "complex script" ใน Word — ถ้าไม่ตั้ง w:cs Word จะใช้ฟอนต์ default
# ของ theme ซึ่งมักไม่มี glyph ไทย แล้วแสดงเป็น ■■■ · Arial เองก็ไม่มีอักษรไทย
# Tahoma มีทั้งบน Windows และ macOS และครอบคลุมไทย จึงใช้เป็นฟอนต์ complex script
DOCX_THAI_FONT = "Tahoma"


def set_run_fonts(run: Any, latin: str = "Arial") -> None:
    """ตั้งฟอนต์ให้ครบทั้ง latin / eastAsia / complex-script (ไทย)"""
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), latin)
    rfonts.set(qn("w:cs"), DOCX_THAI_FONT)


def set_style_fonts(style: Any, latin: str = "Arial") -> None:
    style.font.name = latin
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), latin)
    rfonts.set(qn("w:cs"), DOCX_THAI_FONT)


def target_repo_row(topic: "Topic") -> list[str]:
    """repo ปลายทางของโค้ดตามเอกสารฉบับนั้น — ทุกฉบับต้องบอกให้ชัดว่าโค้ดไปวางที่ไหน

    FE  -> SBP/srm-sps-spsap-web-frontend (sbp-portal · Next.js · portal target sbpm)
    BE  -> SBP/srm-sps-spsap-store-backend (NestJS + TypeORM · schema sps_store)
           + SBP/srm-sps-spsap-sbp-bff (proxy/forward · ไม่มี DB) สำหรับเส้นที่ FE เรียก
    Job -> SBP/srm-sps-spsap-store-backend (runner/cron ฝั่ง backend · ไม่ผ่าน BFF)
    """
    if topic.track == "FE":
        return ["Target repository", "`SBP/srm-sps-spsap-web-frontend` (sbp-portal · Next.js · `NEXT_PUBLIC_APP_TARGET=sbpm`) — "
                "เรียก API ผ่าน `SBP/srm-sps-spsap-sbp-bff` เท่านั้น ห้ามยิง store-backend ตรง"]
    if "/Jobs/" in topic.file:
        return ["Target repository", "`SBP/srm-sps-spsap-store-backend` (NestJS + TypeORM · schema `sps_store`) — "
                "batch runner ฝั่ง backend **ไม่ผ่าน BFF** · cron/พารามิเตอร์อยู่ใน backend config (env/config file)"]
    return ["Target repository", "`SBP/srm-sps-spsap-store-backend` (NestJS + TypeORM · schema `sps_store`) "
            "+ `SBP/srm-sps-spsap-sbp-bff` (forward ผ่าน client service · ไม่มี DB) สำหรับเส้นที่ FE เรียก"]


def topic_blocks(topic: Topic) -> list[dict[str, Any]]:
    if topic.file == "FE/LLDD-FE-Testing-Delivery":
        return testing_delivery_blocks(topic)
    is_batch_monitor = is_batch_monitor_doc(topic.file)
    blocks: list[dict[str, Any]] = [
        h(1, "1. Overview"),
        table(["รายการ", "รายละเอียด"], [
            ["Track", topic.track],
            ["Estimate", estimate_cell(topic)],
            ["Owner", topic.owner],
            target_repo_row(topic),
            ["Objective", topic.objective],
        ]),
        h(1, "2. Screen / Functional Scope"),
        bullets(topic.scope),
    ]
    if not is_batch_monitor:
        blocks.insert(2, p("Common contract reference: ทุกหัวข้อ API/FE ต้องยึด LLDD-BE-API-Common-Contracts และ LLDD-FE-Integration-Contracts สำหรับ error/auth/format/pagination/action/RBAC ก่อนลงรายละเอียดเฉพาะหน้าหรือเฉพาะ endpoint"))
    # หัวข้อ 3 ต้องปรากฏทุกฉบับ ไม่งั้นเลขหัวข้อกระโดด 2 → 4 (ผู้ตรวจรับจะถามหา)
    blocks.append(h(1, "3. Screenshot Reference"))
    if topic.screenshots:
        for shot in topic.screenshots:
            blocks.append(image(str(image_path(shot).relative_to(ROOT)), f"Screenshot: {shot}"))
    else:
        blocks.append(p("ไม่มีภาพหน้าจอสำหรับหัวข้อนี้ — เป็นเอกสารฝั่ง Backend/Batch ที่ไม่มี UI "
                        "(ภาพหน้าจอทั้งหมดอยู่ในเอกสารชุด FE)"))
    if topic.flow_diagram and not is_batch_monitor:
        draw_flow_diagram(topic.title, topic.flow, topic.flow_diagram)
        blocks.extend([
            h(1, "4. Implementation Flow Diagram (Reference)"),
            image(topic.flow_diagram, f"Implementation flow reference: {topic.title}"),
        ])
    blocks.extend([
        h(1, "4. Field, Format, and Validation" if is_batch_monitor else "5. Field, Format, and Validation"),
        table(["Field / UI", "Format", "Validation", "Behavior"], topic.fields),
    ])
    blocks.extend(topic_extra_blocks(topic.file))
    blocks.extend(topic_io_contract_blocks(topic))
    blocks.extend(implementation_detail_blocks(topic))
    blocks.extend([
        h(1, "5. Button / User Action Mapping" if is_batch_monitor else "6. Button / User Action Mapping"),
        table(["Action", "Trigger", "UI Area", "Expected Result"] if is_batch_monitor else ["Action", "Trigger", "API / Service", "Expected Result"], topic.actions),
    ])
    blocks.append(h(1, "6. API Contract" if is_batch_monitor else "7. API Contract"))
    if not topic.apis:
        blocks.append(p(
            "**เอกสารฉบับนี้ไม่มี endpoint ของตัวเอง** — เป็นสัญญา/งานภายในที่เอกสารอื่นเรียกใช้ "
            "(ดูขอบเขตใน 5.90 Endpoint Implementation Contract) · "
            "รายการ endpoint ทั้ง 29 เส้นของ SBPGI อยู่ที่ **LLDD-API** และ `api.md`"
        ))
    for spec in topic.apis:
        blocks.extend([
            h(2, f"{spec.method} {spec.path}"),
            p(spec.purpose),
        ])
        if spec.buttons:
            blocks.append(table(["Triggered by"], [[b] for b in spec.buttons]))
        if spec.request is not None:
            request_title = "Query Params" if spec.method.upper() == "GET" else "Request"
            blocks.append(payload(request_title, api_json(spec.request)))
        blocks.append(h(3, "Request Field Schema"))
        blocks.append(table(["Field", "Type", "Required", "Constraint / Meaning"], api_schema_rows(spec, spec.request, "request")))
        if spec.response is not None:
            blocks.append(payload("Response", api_json(spec.response)))
        blocks.append(h(3, "Response Field Schema"))
        blocks.append(table(["Field", "Type", "Required", "Constraint / Meaning"], api_schema_rows(spec, spec.response, "response")))
    if is_batch_monitor:
        next_no = 7
        skeleton = skeleton_code_blocks(topic, next_no)
        blocks.extend(skeleton)
        next_no += count_top_sections(skeleton)
        blocks.extend([
            h(1, f"{next_no}. Tab Interaction Flow"),
            table(["Step", "Description"], [[i + 1, s] for i, s in enumerate(topic.flow)]),
            h(1, f"{next_no + 1}. Acceptance Criteria"),
            bullets(topic.acceptance),
            h(1, f"{next_no + 2}. Developer Test Checklist"),
            table(["No", "Test"], [[i + 1, t] for i, t in enumerate(topic.tests)]),
        ])
        blocks.extend(unit_test_scope_blocks(topic, next_no + 3))
        return blocks
    next_no = 8
    if topic.db_tables:
        blocks.extend([
            h(1, f"{next_no}. Reference DB Mapping (No Database Page Work)"),
            p("ส่วนนี้เป็นข้อมูลอ้างอิงสำหรับการ implement API/Job เท่านั้น ไม่ใช่งานสร้างหน้า Database, ไม่ใช่งานออกแบบ DB page และไม่ถูกนับเป็น deliverable แยกของ FE/BE"),
            table(["Table / Object", "R/W", "Usage"], topic.db_tables),
        ])
        next_no += 1
    skeleton = skeleton_code_blocks(topic, next_no)
    blocks.extend(skeleton)
    next_no += count_top_sections(skeleton)
    blocks.extend([
        h(1, f"{next_no}. Processing Flow"),
        table(["Step", "Description"], [[i + 1, s] for i, s in enumerate(topic.flow)]),
        h(1, f"{next_no + 1}. Acceptance Criteria"),
        bullets(topic.acceptance),
        h(1, f"{next_no + 2}. Developer Test Checklist"),
        table(["No", "Test"], [[i + 1, t] for i, t in enumerate(topic.tests)]),
    ])
    blocks.extend(unit_test_scope_blocks(topic, next_no + 3))
    return blocks


def common_doc_fields() -> list[tuple[str, str, str, str]]:
    return [
        ("docNo", "YYYY/xxxxx", "required when opening existing document", "ใช้ปี **ค.ศ.** และ running 5 หลัก (มติ 2026-08-06)"),
        ("storeCode", "string 5 digits", "numeric length = 5", "แสดง leading zero"),
        ("amount", "number, 2 decimals", ">= 0", "format `#,##0.00` บาท"),
        ("percent", "number, 2 decimals", "0-100", "ใช้ `%` และรวม allocation ต้องเท่ากับ 100 — **B5: เพิ่ม/ลบร้านที่กระทบเพิ่มเมื่อไร ต้องเกลี่ยใหม่ทั้งชุดแล้วคำนวณ `compensateAmount` ของทุกแถวใหม่ ไม่ใช่เฉพาะแถวที่เพิ่ม**"),
        ("sourceSystem", "enum", "ALLMAP / USER", "**B5** ที่มาของแถวร้านเปิดใหม่ — `ALLMAP` ระบบ default ให้อัตโนมัติ (Job 9) · `USER` เจ้าหน้าที่ SBP DSA คีย์เองจากเอกสารแจ้งของหน่วยงานส่งเสริม (ผัง To-Be · SDD สไลด์ 7) · ซ้ำ `(doc_no, new_store_code)` ให้คืน `409`"),
        ("date", "DD/MM/YYYY", "valid date", "payload เป็น ISO ค.ศ. · FE แสดง ค.ศ. เป็นค่าเริ่มต้น (DatePicker buddhistEra=false) แสดง พ.ศ. เฉพาะจุดที่เปิด flag"),
        ("attachment", "file", "<= 5 MB", f"รองรับ {ATTACHMENT_ALLOWED_EXTENSIONS}"),
    ]


def sbpgi_namespace_blocks() -> list[dict[str, Any]]:
    """5.80 — namespace + การจัดกลุ่ม path ของงานประกันรายได้ (มติ 2026-08-25)"""
    return [
        h(2, "5.80 Namespace + กลุ่ม path ของงานประกันรายได้ (มติ 2026-08-25)"),
        p(
            "SBPGI ไม่ได้แยก backend/พอร์ทัลใหม่ (มติ **DP-10** ให้อยู่ใน `srm-sps-spsap-store-backend` "
            "และโมดูลใน `srm-sps-spsap-web-frontend` เดิม) ทุกอย่างของงานประกันรายได้จึงอยู่ใต้ "
            "**ชื่อเดียวกันทั้ง 3 ชั้น** แล้วแตกเป็น **6 กลุ่มย่อยตามกลุ่มงาน** — ตรงกับ 6 กลุ่มใน `api.md` แบบ 1:1"
        ),
        table(["ชั้น", "รูปแบบ", "ตัวอย่าง"], [
            ["URL ของ API", "`/api/v1/sbpgi/<กลุ่ม>/<resource>`", "`/api/v1/sbpgi/document/{docNo}/actions`"],
            ["route ของหน้าจอ", "`/sbpgi/<กลุ่ม>/<หน้า>`", "`/sbpgi/document/waiting` · `/sbpgi/report/status-summary`"],
            ["โฟลเดอร์ไฟล์", "`**/sbpgi/*`", "`src/app/(main)/sbpgi/*` · `src/services/sbpgi/*` · `src/types/sbpgi/*`"],
        ]),
        h(3, "6 กลุ่มย่อยใต้ `sbpgi`"),
        table(["กลุ่ม", "prefix", "เส้น", "ครอบคลุมอะไร"], [
            ["งาน & เอกสารประกันรายได้", "`/sbpgi/document/*`", "11",
             "`/tasks` (กล่องงาน) · ค้นหา/สร้าง/แก้เอกสาร · `/{docNo}/actions` · `/timeline` · `/attachments` · `/sales`"],
            ["ข้อมูลอ้างอิง (Lookup)", "`/sbpgi/lookup/*`", "2",
             "`/document-statuses` · `/workflow-sections` — อ่านอย่างเดียว ไม่มีหน้าจอดูแล"],
            ["Master Data", "`/sbpgi/master/*`", "8",
             "`/factors` (CRUD 4) · `/competitors` (CRUD 4) — master ที่มีหน้าจอดูแลของตัวเอง"],
            ["รายงาน", "`/sbpgi/report/*`", "2", "`/status-summary` · `/status-summary/export`"],
            ["Workflow ภายใน", "`/sbpgi/workflow/*`", "3", "`/instances` · `/instances/{id}` · `/summary`"],
            ["Interface (tracking / ACK)", "`/sbpgi/interface/*`", "3", "`/tracking` · `/pending-ack` · `/sta/ack`"],
        ]),
        p(
            "**Batch job ไม่มีกลุ่ม path ของตัวเอง** — Jobs 2-10 + 8b รันด้วย cron/CLI ไม่ได้เปิด endpoint "
            "(กลุ่ม Batch Job Admin 6 เส้นถูกตัดทิ้ง 2026-08-06) · หน้าต่างที่มองเห็นผลของ job คือ "
            "**`/sbpgi/interface/*`** (tracking + ACK ของ `interface_transactions`) กับ application log เท่านั้น"
        ),
        h(3, "ทำไมต้องมี prefix (ไม่ใช่แค่ความสวยงาม)"),
        table(["ระบบเดิมมีอยู่แล้ว", "ของ SBPGI ถ้าไม่ใส่ prefix", "ผล"], [
            ["`/document` · `/statement/...`", "`/documents`", "ชนเชิงความหมาย อ่าน routing แล้วสับสน"],
            ["`/report` · `/performance-report` · `/statement/report/ej`", "`/reports/status-summary`", "ชนเชิงความหมาย"],
            ["**`/interface/sta/upload-cmadd`** · `/interface/add`", "**`/interfaces/sta/ack`**", "🔴 เกือบเหมือนกัน — เสี่ยงยิงผิดเส้นจริง"],
            ["`/common` · `/master` · `/store`", "`/factors` `/competitors` `/document-statuses`", "ปนกับ master ของโมดูลอื่น"],
        ]),
        bullets([
            "ฝั่ง NestJS: **`SbpgiModule` เดียว** ผูก prefix ที่ระดับโมดูล "
            "(`RouterModule.register([{ path: 'sbpgi', module: SbpgiModule }])`) แล้วแตกเป็น 6 controller ตามกลุ่ม "
            "(`DocumentController` `LookupController` `MasterController` `ReportController` `WorkflowController` `InterfaceController`) — "
            "**ห้ามเติม `sbpgi/` ในแต่ละ `@Controller()`**",
            "ในกลุ่ม `document` ต้องประกาศ route คงที่ (`/tasks`) **ก่อน** route ที่มีพารามิเตอร์ (`/{docNo}`) "
            "และ `docNo` เป็น `YYYY/xxxxx` จึงต้อง `encodeURIComponent` ทุกครั้งที่ประกอบ URL",
            "เส้นที่ **ไม่ใช่ของ SBPGI ห้ามใส่ prefix และห้ามแตะ** — `GET /store/search` · `GET /store/all-regions` · "
            "`GET /common/common-code` · `GET /menus` · `GET /groups/current-user/permissions` · "
            "`POST /statement/upload-file-aws` · `GET /api/workflow/pending` เป็นของระบบ SBP เดิม",
            "BFF ส่งต่อทั้ง prefix (`/api/v1/sbpgi/*`) โดยไม่ตัดคำ · สิทธิ์เมนูผูกกับ URL ของ **หน้าจอ** (`/sbpgi/<กลุ่ม>/...`) ไม่ใช่ URL ของ API",
        ]),
    ]


def topic_extra_blocks(file_key: str) -> list[dict[str, Any]]:
    if file_key == "FE/LLDD-FE-Create-Document":
        return create_document_fs_iframe_blocks()
    if file_key == "FE/LLDD-FE-Master-Data":
        return master_config_screen_blocks()
    if file_key == "FE/LLDD-FE-Document-Detail":
        return document_detail_role_blocks()
    role_profile = document_detail_role_profile(file_key)
    if role_profile:
        return document_detail_single_role_blocks(role_profile)
    if file_key == "FE/LLDD-FE-Integration-Contracts":
        return sbpgi_namespace_blocks()
    if file_key == "BE/LLDD-BE-API-Common-Contracts":
        return sbpgi_namespace_blocks() + common_contract_extra_blocks()
    if file_key == "BE/LLDD-BE-API-Document-Create-Update":
        return document_create_update_extra_blocks()
    if file_key == "BE/LLDD-BE-API-Document-Workflow-Actions":
        return workflow_action_transition_blocks()
    if file_key == "BE/LLDD-BE-API-Document-Detail-Aggregate":
        return document_detail_aggregate_extra_blocks()
    if file_key == "BE/LLDD-BE-API-Attachment-Sales-Timeline":
        return attachment_storage_extra_blocks()
    if file_key == "BE/LLDD-BE-Database-Structure":
        return database_structure_extra_blocks()
    if file_key == "BE/LLDD-BE-Data-Migration-Cutover":
        return data_migration_extra_blocks()
    if file_key == "BE/LLDD-BE-Integration-SBP-Platform":
        return integration_sbp_platform_extra_blocks()
    if file_key == "BE/LLDD-BE-Workflow-Engine-Definition":
        return workflow_engine_definition_extra_blocks()
    if file_key == "BE/Jobs/LLDD-BE-Job-8b-StartInternalWorkflow":
        return workflow_engine_unconfirmed_warning_blocks()
    return []


def workflow_engine_unconfirmed_warning_blocks() -> list[dict[str, Any]]:
    """คำเตือน F9 + F1/F2/F3 สำหรับ Job 8b — job เดียวที่เรียก workflow engine โดยตรง

    เอกสารฉบับนี้เคยระบุชื่อ function ตายตัว (ปนสองชุด) และอ้าง UNIQUE constraint
    ที่ไม่มีอยู่จริง โดยไม่มีคำเตือนเหมือนที่เอกสาร BE-API มี
    """
    return [
        h(1, "4a. จุดเข้า flow ตามประเภทเคส — Job 8b เป็นคนตัดสินว่าเปิด workflow ที่ state ไหน"),
        p("ผัง To-Be 12/02/2026 กำหนดว่า **เอกสารไม่ได้เริ่มที่ state 06 เสมอไป** · Job 8b ต้องอ่านข้อมูลรอบชดเชย (คอลัมน์ที่รับเข้าโครง 2026-08-21 · gap F8) แล้วเลือก state เริ่มต้นก่อนเรียก initializeWorkflow/addPreApprover"),
        table(["เคส", "เงื่อนไขที่ Job 8b ต้องอ่าน", "เปิด workflow ที่ state", "ผู้รับผิดชอบขั้นแรก"], [
            ["① เปิดเรื่องใหม่", "fgi_impact_processes.last_compensate_seq_no = 1", "**06**", "group ฝ่าย SBP DSA (ปกติ)"],
            ["② ชดเชยต่อเนื่อง", "last_compensate_seq_no > 1 และ flag_action = 'Y'", "**08** (Auto Approve — ข้ามขั้น 06)", "**เจ้าหน้าที่ SBP DSA คนเดิม** ผ่าน addPreApprover"],
            ["③ ยอดชดเชย 0 ติดกัน <= 3 เดือน", "COALESCE(adjust_amount, forecast_amount) = 0 ใน fgi_impact_compensations งวดที่ 1-3", "**01** (ข้ามทั้ง 06 และ 08)", "group หน่วยงานส่งเสริมธุรกิจฯ"],
            ["③ ยอดชดเชย 0 ติดกัน > 3 เดือน", "งวดที่ 4 ขึ้นไป", "**ไม่เปิด workflow** — ปิดเอกสารเป็นเสร็จสิ้น (หยุดชดเชยประกันรายได้)", "-"],
        ]),
        p("**ที่มาของค่าที่ใช้ตัดสิน** — ทุกค่าอยู่ในโซน A (FGI/FCS) ที่ batch เขียนไว้ก่อนเปิดเอกสาร ไม่ใช่ค่าที่ Job 8b คำนวณเอง"),
        table(["ค่าที่ใช้ในเงื่อนไข", "ระบบเดิม (Oracle FCS_FRN)", "ตาราง SBPGI", "คอลัมน์ · ชนิด", "เขียนโดย"], [
            ["`LAST_COMPENSATE_SEQ_NO`", "`FGI_IMPACT_STORE_ON_PROCESS.LAST_COMPENSATE_SEQ_NO`", "`fgi_impact_processes`", "`last_compensate_seq_no` · INTEGER", "Job 2 — `ImportJdbc` (`SEQ_NO + 1` เมื่อเป็นรอบต่อเนื่อง)"],
            ["`FLAG_ACTION`", "`FGI_IMPACT_STORE_ON_PROCESS.FLAG_ACTION` (โดเมน Y/W/N)", "`fgi_impact_processes`", "`flag_action` · CHAR(1)", "Job 2 เขียน `'Y'` · Job 6 ปิดรอบ `Y->N` / พัก `Y->W`"],
            ["`DATASOURCE`", "`FGI_IMPACT_STORE_ON_PROCESS.DATASOURCE` (เดิมมี ALM/STA/HRS)", "`fgi_impact_processes`", "`datasource` · VARCHAR(5)", "Job 2/3 = `ALM` · Job 5 = `STA` · **`PRO` เชิงรุก / `REA` เชิงรับ = คนคีย์** (รหัสใหม่ 2026-08-24)"],
            ["`forecast`", "`FGI_IMPACT_STORE_COMPENSATE.COMPENSATE_FORECAST`", "`fgi_impact_compensations`", "`forecast_amount` · NUMERIC(14,2)", "Job 5 — นำเข้ายอดจาก IAS/MIS"],
            ["`adjust`", "`FGI_IMPACT_STORE_COMPENSATE.COMPENSATE_ADJUST`", "`fgi_impact_compensations`", "`adjust_amount` · NUMERIC(14,2)", "เจ้าหน้าที่ SBP DSA ปรับยอดในเอกสาร"],
        ]),
        p("> ยอดที่ใช้จริงทุกที่คือ `COALESCE(adjust_amount, forecast_amount)` — ค่าที่คนปรับชนะค่าที่ระบบคำนวณเสมอ  \n"
          "> `datasource` ไม่ได้เปลี่ยน state เริ่มต้นของ workflow — มันบอกแค่ว่า **ใครคีย์ข้อมูล** (`ALM`/`STA` = ระบบส่งงานมาให้เลือก · `PRO`/`REA` = เจ้าของงานคีย์เอง · SDD GI สไลด์ 17 · 47 · 49)  \n"
          "> ⚠️ ทั้งสองตารางเป็น gap **F8/F1** ที่เพิ่งรับเข้าโครงเมื่อ 2026-08-21 — ต้อง migrate ครบก่อน Job 8b จึงทำงานตามผัง To-Be ได้"),
        code("""-- ตัดสินประเภทเคสก่อนเปิด workflow (Job 8b)
SELECT p.last_compensate_seq_no,
       p.flag_action,
       (SELECT COUNT(*) FROM fgi_impact_compensations c
         WHERE c.impact_process_id = p.id
           AND COALESCE(c.adjust_amount, c.forecast_amount) = 0
           AND c.compensate_seq = p.last_compensate_seq) AS zero_months
FROM fgi_impact_processes p
WHERE p.id = :impactProcessId;
-- zero_months >= 4            -> ไม่เปิด workflow · ปิดเอกสารเป็น 99 พร้อม result = หยุดชดเชยประกันรายได้
-- zero_months BETWEEN 1 AND 3 -> initializeWorkflow แล้ว addPreApprover ที่ state 01
-- seq_no > 1 AND flag_action='Y' -> state 08 + approver = เจ้าหน้าที่คนเดิม (จาก consideration_logs รอบก่อน)
-- นอกนั้น                      -> state 06 ตามปกติ""", "sql"),
        p("**ทุกเส้นทางอัตโนมัติต้องบันทึกลง `consideration_logs` ด้วยผู้ดำเนินการ `SYSTEM`** เพื่อไม่ให้ timeline ของเอกสารขาดช่วง · รายละเอียดกติกาเต็มดู `workflow.md` หัวข้อจุดเข้า flow ตามประเภทเคส"),
        h(1, "4b. ข้อค้างที่ต้องยืนยันก่อนเขียนโค้ด (workflow engine)"),
        p(
            "✅ **ชื่อ function ของ engine — ยึด LLDD ของ lib (ปิดข้อค้าง 2026-08-14)** · API จริงคือ 8 ตัวตามชีต `Detail` ของ `SBP/TSM-SRM-LLDD SBP workflow 1.2.xlsx` (เอกสารของ lib เอง): `initializeWorkflow` · `eventWorkflow` · `getPermissionEvents` · `getHistory` · `getTransaction` · `getPendingFlowByUser` · `getWorkflowsByUser` · `addPreApprover` · ชื่อที่เคยขัดกันไม่ใช่ชื่อ API — *Trigger Event* เป็นชื่อหัวข้อขั้นตอนภายใน `eventWorkflow` และ `*UseCase` เป็น class ที่ store-backend ห่อไว้ใช้เอง (ดู `LLDD-BE-Workflow-Engine-Definition` หัวข้อ 5.3)"
        ),
        table(
            ["ข้อค้าง", "ข้อเท็จจริงที่ตรวจแล้ว", "ผลต่อ Job 8b", "สถานะ"],
            [
                [
                    "DP-1 · `referenceId` ของ workflow ✅ ปิดแล้ว 2026-08-17",
                    "ระบบเดิม (cooperation-request · inform-evaluate) ใช้ surrogate id ทุกจุด",
                    "ค่าที่ส่งเข้า initialize และคีย์ที่ใช้เช็คซ้ำเปลี่ยนตามข้อนี้",
                    "✅ ปิดแล้ว 2026-08-17 — เลือก surrogate id (`compensation_documents.id` ส่งเป็น string) ตามที่ cooperation-request / inform-evaluate ทำจริง",
                ],
                [
                    "DP-2 · `sps_store.workflow_transaction` ไม่มี PK/index",
                    "19,283 แถว · ไม่มีทั้ง PK และ index (`SBP/db-schema-sps_store.md`) ต่างจาก `sps_auth` ที่มี PK ปกติ",
                    "กันซ้ำด้วย DB constraint ไม่ได้ ต้องกันที่ application · query ตาม reference_id เป็น seq-scan",
                    "ยังไม่ตัดสิน — ขอ sign-off เพิ่ม index กับทีมเจ้าของ library หรือยอมรับสภาพ",
                ],
                [
                    "schema ของ engine",
                    "engine ตัวจริงมี **13 ตาราง** อยู่ใน schema **`sps_store`** — `sps_auth` มีชื่อตารางชุดเดียวกันแต่เป็นสำเนาของ auth-backend คนละเวอร์ชัน",
                    "ทุก SQL ในเอกสารนี้ต้อง prefix `sps_store.`",
                    "ข้อเท็จจริง ไม่ใช่ข้อค้าง",
                ],
            ],
        ),
    ]


# --------------------------------------------------------------------------------------
# เนื้อหาเฉพาะของเอกสาร BE ที่เพิ่มใหม่ 4 ฉบับ (2026-08-07)
#
# ข้อเท็จจริงเชิงตัวเลขทุกตัวในหัวข้อนี้มาจากการตรวจฐานข้อมูลจริง 2026-08-07 และบันทึกไว้ที่
# `SBP/db-schema-sps_store.md` · `SBP/db-schema-sps_auth.md` · `SBP/SBPGI-vs-existing-system.md`
# ส่วนนิยาม engine มาจาก `SBP/TSM-SRM-LLDD-SBP-workflow-1.2.md`
# ⚠️ ข้อค้างตัดสินใจ 12 ข้อ (DP-1..DP-12) ในเอกสารนี้ **ยังไม่ตัดสิน** — ห้ามเลือกแทนเจ้าของโครงการ
# --------------------------------------------------------------------------------------
DECISION_DOC = "SBP/SBPGI-vs-existing-system.md หัวข้อ 4"


def pending_decision_blocks(heading: str, rows: list[list[str]]) -> list[dict[str, Any]]:
    """ตารางข้อค้างตัดสินใจ — บันทึกทางเลือกไว้เฉย ๆ ไม่เลือกให้"""
    return [
        h(2, heading),
        p(
            f"รายการต่อไปนี้ **ยังไม่ตัดสิน** ทั้งหมด · ทางเลือกและผลกระทบเต็มอยู่ที่ `{DECISION_DOC}` "
            "การตัดสินใจขั้นสุดท้ายเป็นของเจ้าของโครงการ — เอกสาร LLDD ฉบับนี้บันทึกไว้เป็นข้อค้าง "
            "และห้าม dev เลือกทางใดทางหนึ่งเองระหว่าง implement"
        ),
        table(["ข้อค้าง", "ทางเลือก A", "ทางเลือก B", "สถานะ"], rows),
    ]


def database_structure_extra_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 ขอบเขตตารางในโครง SBPGI (20 ตาราง — CREATE จริง 19 + reuse 1)"),
        p(
            "DDL เต็มอยู่ที่เอกสาร `LLDD-Database` หัวข้อ Executable DDL · เอกสารฉบับนี้เป็นเจ้าของ "
            "**สคริปต์ deploy จริง** และกติกาว่าอะไรสร้างได้/สร้างไม่ได้"
        ),
        p(
            "⚠️ **20 = จำนวนตารางในโครง ไม่ใช่จำนวนที่ต้อง CREATE** — `fcs_qssi_score` นับอยู่ในโครงโซน A "
            "แต่ใช้ตารางเดิมของ `sps_store` (23,958,780 แถว) จึง **ห้าม CREATE TABLE** ดูหัวข้อ 5.1.1 · "
            "จำนวนที่ต้อง CREATE จริงคือ **19 ตาราง** (20 ตารางในโครง ลบ fcs_qssi_score ที่ reuse) · สถานะ reuse ของ `fcs_qssi_score` ยังผูกกับข้อค้าง "
            f"**DP-4 ✅ ปิดแล้ว 2026-08-24** — reuse ตารางเดิมแบบ **อ่านอย่างเดียว** ไม่มีการเขียนจากฝั่ง SBPGI จึงไม่ต้องแก้ constraint/index ของตารางเดิมและไม่ต้องขอ sign-off (`{DECISION_DOC}`)"
        ),
        table(
            ["โซน", "จำนวน", "ตาราง"],
            [
                ["A — FGI/FCS pipeline", "8 (CREATE 7 + reuse 1)", "fgi_impact_processes, **fgi_impact_compensations**, fgi_impact_stores, fgi_impact_sales_summaries, sales_transactions, fgi_impact_competitors, interface_transactions · **+ fcs_qssi_score = reuse ห้าม CREATE (ดู 5.1.1 · DP-4)**"],
                ["B — เอกสาร/ประวัติ", "9", "compensation_documents, document_new_stores, document_competitors, document_external_factors, consideration_logs, document_attachments, compensation_histories, document_cost_details, document_running_numbers"],
                ["C — master ที่ SBPGI เป็นเจ้าของ", "3", "impacted_stores, external_factors, competitors (decisions ย้ายไป common_code · DP-9 · status_email_rules ตัดตาม DP-5 — SBPGI เรียก email-lib เองโดยใช้เลข template จาก workflow_route.email_id)"],
                ["รวม", "**20 (CREATE 19 + reuse 1)**", "ตรงกับ `database.md` และผลรวมของโซน A 8 + B 9 + C 3 · ประวัติ: 34 → 24 (2026-08-06 reuse ของระบบเดิม) → 22 (ตัดกลุ่ม batch) → 21 (ยกเลิก `audit_logs` 2026-08-07) → **20 (มติ DP-9 2026-08-10 ย้าย `decisions` ไป `common_code`)** → คงที่หลังรับ F8+F1 เข้าโครง 2026-08-21 (เพิ่ม `fgi_impact_compensations` แทน `status_email_rules` ที่ตัดตาม DP-5)"],
            ],
        ),
        h(3, "5.1.1 ตารางที่ระบบ SBP เดิมมีอยู่แล้ว — ห้าม CREATE TABLE"),
        p(
            "ตรวจฐานข้อมูลจริง 2026-08-07 (`SBP/db-schema-sps_store.md`) ทุกตารางในตารางนี้อยู่ใน schema "
            "**`sps_store`** และมีข้อมูลจริงใช้งานอยู่ การสร้างซ้ำใน SBPGI = ข้อมูลสองชุดที่ไม่มีวันตรงกัน"
        ),
        table(
            ["ตาราง/กลุ่มตาราง", "schema", "จำนวนแถวจริง", "หมายเหตุ"],
            [
                ["workflow engine 13 ตาราง", "sps_store", "workflow_transaction 19,283 · workflow_history 38,010 · workflow_approver 96,542", "ดู LLDD-BE-Workflow-Engine-Definition"],
                ["fcs_qssi_score (เอกพจน์)", "sps_store", "23,958,780", "มี import pipeline ใช้งานอยู่ (POST /performance/import-qssi · staging fcs_tmp_qssi_score) — ห้ามสร้างใหม่"],
                ["mas_param", "sps_store", "93,752", "ค่ากำหนดกลาง"],
                ["common_code / common_code_type", "sps_store", "2,609 / 376", "วงเงินอนุมัติ code_type = SBPGI_APPROVE_LIMIT"],
                ["email_template / email_sent", "sps_store", "85 / 5,214", "เทมเพลตอีเมลและ log การส่ง"],
                ["business_user", "sps_store", "12,752", "ตัวตนผู้ใช้/ผู้อนุมัติ"],
                ["store / mas_store", "sps_store", "19,402 / 19,647", "master ร้าน"],
                ["fcs_monthly_sales", "sps_store", "711,384", "ยอดขาย**รายเดือน** (key store_id+year+month) — ใช้แทน sales_transactions รายวันไม่ได้ ย้อนกลับเป็นรายวันไม่ได้ · ใช้ cross-check ได้"],
            ],
        ),
        h(3, "5.1.2 แกนธุรกิจที่ยืนยันแล้วว่าต้องสร้างเอง"),
        p(
            "ค้นทั้งฐาน 276 ตาราง / 4,396 คอลัมน์ ด้วยคำ `impact` · `compensat` · `guarantee` · `income` · "
            "`competitor` · `growth` · `outlier` · `distance` · `radius` · `latitude` · `longitude` · `window_no` "
            "ได้ **0 hit ทุกคำ** → ตารางโซน A และแกนเอกสารโซน B ไม่มีของเดิมให้ reuse ต้องสร้างเองทั้งหมด"
        ),
        h(2, "5.2 ลำดับไฟล์ deploy"),
        table(
            ["ไฟล์", "เนื้อหา", "รันเมื่อไร"],
            [
                ["01_schema.sql", "CREATE TABLE 19 ตาราง เรียงตาม dependency (C master -> A pipeline -> B document) — ไม่รวม fcs_qssi_score ที่ reuse ของเดิม", "ครั้งเดียวต่อ environment"],
                ["02_index.sql", "index, unique/partial index, check constraint", "หลัง 01 · rerun ได้เมื่อเพิ่ม index"],
                ["03_seed.sql", "external_factors, competitors (01-11) — ไม่มี decisions แล้ว (DP-9 ย้ายไป common_code · seed ที่ระบบเดิม)", "หลัง 02"],
                ["04_grant.sql", "GRANT ให้ role ของ application (แยก read/write)", "หลัง 03"],
                ["99_rollback.sql", "DROP TABLE ย้อนลำดับ เฉพาะตารางของ SBPGI", "เฉพาะกรณี rollback"],
            ],
        ),
        code(
            """-- 01_schema.sql (ตัวอย่างส่วนหัว — DDL เต็มอยู่ที่ LLDD-Database)
-- ห้ามมี CREATE TABLE ของตาราง reuse: ตรวจด้วยคำสั่งนี้ก่อน commit
--   grep -nE 'CREATE TABLE (workflow_|fcs_qssi_score|mas_param|common_code|business_user|store|mas_store|email_template|decisions)' 01_schema.sql
BEGIN;
SET search_path TO sps_store;

-- โซน C: master ที่ SBPGI เป็นเจ้าของ (ต้องมาก่อนเพราะโซน A/B อ้างถึง)
-- ❌ ไม่มี CREATE TABLE decisions — มติ DP-9 (2026-08-10) ย้ายไป common_code ของระบบเดิม
--    (code_type = 'SBPGI_DECISION') · FE อ่านผ่าน GET /common/common-code?codeType=SBPGI_DECISION
CREATE TABLE external_factors (...);
CREATE TABLE competitors (...);
CREATE TABLE impacted_stores (...);
-- ❌ ไม่สร้างตาราง status_email_rules ใน SBPGI (ปิด DP-5 · แก้มติ 2026-08-14)
--    workflow ให้ 'เลข template' ผ่าน workflow_route.email_id → SBPGI เรียก sendEmail() ของ email-lib เอง
--    ผู้รับ resolve จาก workflow_approver → business_user.email · ⚠️ คอลัมน์จริงคือ email_sent.send_by ไม่ใช่ sent_by
--    SBPGI อ่าน workflow_route.email_id แล้วเรียก sendEmail() ของ email-lib · template อยู่ที่ email_template ของระบบ SBP เดิม (85 แถว)
--    lib เขียน log ให้เองที่ email_sent (5,214 แถว · mail_to/mail_cc/is_sent/error · ⚠️ คอลัมน์ผู้ส่งคือ send_by)


-- โซน A: pipeline
CREATE TABLE fgi_impact_processes (...);
-- ...

-- โซน B: เอกสาร
CREATE TABLE compensation_documents (...);
-- ...
COMMIT;""",
            "sql",
        ),
        h(2, "5.3 Seed ที่ต้องมีตั้งแต่วันแรก"),
        table(
            ["ตาราง", "ข้อมูล seed", "ที่มา"],
            [
                ["(ไม่สร้าง) decisions", "ย้ายไป common_code ของระบบเดิม (code_type = SBPGI_DECISION) — มติ DP-9 2026-08-10", "MSSQL DecisionProfile"],
                ["competitors", "แบรนด์คู่แข่ง 11 รายการ รหัส 01-11 (ไทย+อังกฤษ)", "หน้าจอ K2 เดิม (k2-competitors.html)"],
                ["external_factors", "ปัจจัยภายนอกที่ใช้อยู่", "MSSQL FactorProfile"],
                
                ["common_code (ระบบเดิม)", "SBPGI_APPROVE_LIMIT: THRESHOLD=100000 (เกณฑ์เดียว)", "มติประชุม 2026-08-18 — เขียนที่ common_code ของระบบเดิม ไม่ใช่ตารางของ SBPGI"],
                ["common_code (ระบบเดิม)", "SBPGI_DATASOURCE: ALM=ระบบ (ALLMAP) · STA=ระบบ (Statement) · PRO=เชิงรุก · REA=เชิงรับ", "SDD GI สไลด์ 17 — 3 แหล่งข้อมูลร้านที่ต้องชดเชย · รหัส PRO/REA ตั้งใหม่ 2026-08-24 ตามแพตเทิร์น 3 ตัวอักษรของ DATASOURCE เดิม (ALM/STA/HRS) เพราะ SDD และระบบเดิมไม่ได้กำหนดไว้"],
            ],
        ),
        *pending_decision_blocks(
            "5.4 ข้อค้างตัดสินใจที่กระทบ DDL (ยังไม่ตัดสิน)",
            [
                ["DP-3 ✅ ตัดสินแล้ว 2026-08-10 = ทางเลือกที่ 3 (snapshot เฉพาะร้านที่เคยเข้ารอบชดเชย · เติมตอนสร้าง fgi_impact_processes)", "view จากระบบเดิม (`v_sbpgi_sp_store`) — ไม่ต้อง sync แต่ร้านที่ยกเลิกเกิน 1 เดือนหายจาก view ทำให้เอกสารย้อนหลังหาร้านไม่เจอ", "ตาราง snapshot ของ SBPGI — เอกสารย้อนหลังหาร้านเจอเสมอ แต่ต้อง sync (มีทางเลือกที่ 3: snapshot เฉพาะร้านที่เคยเข้ารอบชดเชย)", "✅ ตัดสินแล้ว 2026-08-10 = ทางเลือกที่ 3"],
                ["DP-4 ✅ ปิดแล้ว 2026-08-24 · `fcs_qssi_score` reuse แบบอ่านอย่างเดียว", "reuse ตารางเดิม 23,958,780 แถว — ระบบ SBP เดิมนำเข้าให้แล้วผ่าน `POST /performance/import-qssi`", "สร้างตารางของ SBPGI เอง — ตกไป (จะมีข้อมูล QSSI สองชุด)", "✅ **reuse อ่านอย่างเดียว** — ตัด Job 1 (ImportQSSI) ทั้ง job · SBPGI ไม่เขียนตารางนี้ จึงไม่ต้อง backfill / SET NOT NULL / sign-off เจ้าของ `performance.service.ts` · ห้ามสร้างตารางชื่อ `fcs_qssi_scores` (พหูพจน์)"],
                ["DP-9 ✅ ตัดสินแล้ว 2026-08-10 = แยกตัดสิน (decisions → common_code · external_factors/competitors ยังเป็นตารางของ SBPGI)", "ยัดลง `common_code` ของระบบเดิม", "ตารางเล็กของ SBPGI ตามที่ DDL ปัจจุบันเขียนไว้", "✅ ตัดสินแล้ว 2026-08-10 = แยกตัดสิน (`decisions` → `common_code` · `external_factors`/`competitors` ยังเป็นตารางของ SBPGI)"],
                ["DP-1 · `reference_id` ของ workflow", "`doc_no` — ตกไป", "**เลือก surrogate id** (`compensation_documents.id` · ส่งเป็น string เพราะ `reference_id` เป็น varchar(255)) แบบที่ cooperation-request/inform-evaluate ทำจริง", "✅ ปิดแล้ว 2026-08-17 — ยืนยันตามระบบเดิม"],
                ["DP-7 ✅ ปิดแล้ว 2026-08-24 · `consideration_logs`", "**เลือกข้อนี้ — ตารางของ SBPGI เอง** (ผูก `transaction_id` ของ engine) ตามที่ DDL ปัจจุบันเขียนไว้", "ตารางส่วนขยายบน `sps_store.workflow_history` ของ engine — ตกไป (engine ไม่มี decision code / ไฟล์แนบ / ความเห็น)", "✅ ปิดแล้ว 2026-08-24 · กระทบ DDL ของตารางนี้และ response ของ `GET /sbpgi/document/{docNo}/timeline`"],
                ["DP-12 · audit ของ master", "เอากลับมาโดยใช้กลไกของระบบเดิม", "ไม่มีเลยตามมติ 2026-08-07 (สถานะปัจจุบันของ DDL)", "ยังไม่ตัดสิน"],
            ],
        ),
    ]


def data_migration_extra_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 Source-to-Target Mapping ระดับตาราง"),
        table(
            ["ต้นทาง", "ระบบ", "ปลายทาง (SBPGI)", "กฎแปลงที่ต้องระวัง"],
            [
                ["FGI_IMPACT_STORE_ON_PROCESS", "ORA FCS_FRN", "fgi_impact_processes", "PK IMPACT_PROCESS_ID (seq SEQ_FGI_IMPACT_PROCESS) เป็น hub ของทั้งโซน A · **ต้อง migrate คอลัมน์รอบชดเชยด้วย (gap F8 · รับเข้าโครง 2026-08-21)**: `LAST_COMPENSATE_SEQ/_SEQ_NO -> last_compensate_seq/_seq_no` · `START/END_COMPENSATE_MONTH-YEAR -> start/end_compensate_month/year` · `FLAG_ACTION -> flag_action` · `DATASOURCE -> datasource` · ⚠️ `FLAG_ACTION` โดเมนจริงคือ **Y/W/N** (active = `IN ('Y','W')`) ไม่ใช่ Y/N — Job 6 เขียน `Y->W` ตอนพัก/รอจ่าย ถ้า CHECK ปลายทางรับแค่ Y/N แถวกลุ่มนี้จะ migrate ไม่ผ่าน · ทั้ง 4 กลุ่มนี้คือค่าที่ Job 8b ใช้ตัดสินจุดเข้า flow"],
                ["FGI_IMPACT_STORE", "ORA FCS_FRN", "fgi_impact_stores + impacted_stores", "แถวฝั่ง `_I` ทำ distinct เข้า impacted_stores · ที่เหลือเป็นคู่ร้าน"],
                ["FGI_IMPACT_STORE_COMPENSATE", "ORA FCS_FRN", "**fgi_impact_compensations** (รับเข้าโครง 2026-08-21 · gap F1)", "`COMPENSATE_FORECAST -> forecast_amount` · `COMPENSATE_ADJUST -> adjust_amount` · `COMPENSATE_SEQ/_SEQ_NO -> compensate_seq/_seq_no` · UK (impact_process_id, compensate_month) · ใช้นับยอด 0 ติดกันกี่งวดด้วย `COALESCE(adjust_amount, forecast_amount) = 0` — เป็น input ของ Job 8b เคส ③"],
                ["FGI_IMPACT_STORE_SALES", "ORA FCS_FRN", "fgi_impact_sales_summaries", "key STORECODE_I + MONTH + YEAR"],
                ["FGI_IMPACT_STORE_SALES_TRN", "ORA FCS_FRN", "sales_transactions", "4 หน้าต่าง × 15 วัน — ห้ามใช้ fcs_monthly_sales แทน (รายเดือน ย้อนกลับเป็นรายวันไม่ได้)"],
                ["FGI_IMPACT_COMPETITOR", "ORA FCS_FRN", "fgi_impact_competitors", "data_source = ALM"],
                ["FGI_CONFIRM_RECEIVE_DATA", "ORA FCS_FRN", "interface_transactions", "TRANSACTION_PK เป็น polymorphic — ต้องแตกตาม DATA_NAME เป็น typed FK"],
                ["FCS_QSSI_SCORE", "ORA FCS_FRN", "fcs_qssi_score (sps_store)", "ปลายทางมีข้อมูลอยู่แล้ว 23,958,780 แถว — ต้องเทียบก่อนว่าจะโหลดทับหรือไม่ (ผูกกับ DP-4)"],
                ["CompensateFlow", "MSSQL CPA_FRN_FGI", "compensation_documents", "CompDocumentID -> doc_no · เก็บ round_no/loop_no/allmap_url/statement_id/approver_snapshot"],
                ["CompensateHistory", "MSSQL CPA_FRN_FGI", "consideration_logs", "PK ActionID · เติม result_category (APPROVE/REJECT/CANCELLED/PENDING)"],
                ["ImpactProfile", "MSSQL CPA_FRN_FGI", "document_new_stores", "ฝั่ง `_N` + %ชดเชย/ยอดต่อร้าน"],
                ["CompetInCompenProfile", "MSSQL CPA_FRN_FGI", "document_competitors", "คู่แข่งที่ผูกกับเอกสาร · competitor_code อ้าง master competitors (11 รหัส 01-11) · แถวที่มาจาก ALLMAP ตั้ง data_source = ALM"],
                ["FactorInCompenProfile", "MSSQL CPA_FRN_FGI", "document_external_factors", "ปัจจัยภายนอกที่ผูกกับเอกสาร · factor_code อ้าง master external_factors + ช่วงวันที่มีผล"],
                ["FGI_IMPACT_STORE_COMPENSATE + CompensateFlow", "ORA + MSSQL", "compensation_histories", "ประวัติชดเชยต่อร้าน/รอบ · submit_account_month (งวดที่ Job 6 ส่งไป STA ผ่าน RabbitMQ) · ⚠️ **ต้องปิด DP-11 ก่อน** — ยังไม่ตัดสินว่า SBPGI เป็นต้นทางตัวเลขเงิน หรือ fr_store_insure ยังคีย์มือ"],
                ["ImpactCostDetail", "MSSQL CPA_FRN_FGI", "document_cost_details", "ยอดชดเชยแยกรายเดือน/รายร้านใหม่"],
                ["RunningNumber", "MSSQL CPA_FRN_FGI", "document_running_numbers", "ตั้ง last_running_no ต่อปีให้ตรงกับเลขสูงสุดที่ย้ายมา"],
                ["CompDocAttachment / CompTempAttachment / AttachFileProfile", "MSSQL CPA_FRN_FGI", "document_attachments", "metadata เท่านั้น · ไฟล์จริงต้องย้ายขึ้น S3 ของระบบเดิม"],
                ["FactorProfile / CompetitionProfile", "MSSQL CPA_FRN_FGI", "external_factors / competitors", "เป็น master ที่ SBPGI เป็นเจ้าของ · **DecisionProfile ไม่ย้ายมาแล้ว** — มติ DP-9 (2026-08-10) ให้ seed ลง common_code ของระบบเดิม (code_type = SBPGI_DECISION) ไม่สร้างตาราง decisions"],
            ],
        ),
        h(2, "5.2 กฎแปลงข้อมูลที่ผิดบ่อย"),
        table(
            ["เรื่อง", "อาการถ้าไม่ทำ", "กฎที่ต้องใช้"],
            [
                ["leading zero ของรหัสร้าน", "ร้าน 00788 กลายเป็น 788 แล้ว join ไม่ติด", "lpad(store_code, 5, '0') ทุกจุด · ปลายทางเป็น VARCHAR(5)"],
                ["ปี พ.ศ./ค.ศ.", "วันที่เพี้ยน 543 ปี", "เก็บ ค.ศ. ใน DB และ `doc_no` เป็นปี **ค.ศ.** ด้วย (มติ 2026-08-06) · ถ้าของเดิมเป็น พ.ศ. ต้องแปลงตอน migrate ด้วย toAD()"],
                ["polymorphic key", "FK ชี้ผิดตาราง", "แตก TRANSACTION_PK ตาม DATA_NAME เป็น impact_process_id / sales_summary_id / doc_no"],
                ["เลขเอกสารซ้ำ", "ออกเลขใหม่ทับของเก่า", "หลังโหลด ตั้ง document_running_numbers.last_running_no = MAX(running) ต่อปี"],
                ["ยอดขายรายวัน", "ข้อมูล 60 วันไม่ครบ ทำให้ธงผิดปกติเพี้ยน", "ต้องมาจาก FGI_IMPACT_STORE_SALES_TRN เท่านั้น · fcs_monthly_sales (711,384 แถว) ใช้ cross-check ได้อย่างเดียว"],
            ],
        ),
        h(2, "5.3 แผน Cutover"),
        table(
            ["รอบ", "กิจกรรม", "เกณฑ์ผ่าน"],
            [
                ["T-14 วัน", "Profiling ต้นทาง + dry-run รอบที่ 1", "อธิบายแถวที่ reject ได้ทุก reason code"],
                ["T-7 วัน", "Full load บน staging + reconcile", "จำนวนแถว/ยอดเงินตรง หรืออธิบายส่วนต่างได้"],
                ["T-2 วัน", "ซ้อม cutover เต็มรูปแบบรวม rollback", "rollback สำเร็จอย่างน้อย 1 ครั้ง"],
                ["T-0 (freeze)", "หยุดใช้ระบบเดิม -> delta load -> reconcile รอบสุดท้าย -> ย้าย workflow ที่ยังวิ่ง", "ทุกเอกสารที่ยังไม่จบ flow เปิดในระบบใหม่ได้ที่ state เดิม"],
                ["T+1..T+7", "เฝ้าระวัง · เก็บ snapshot ก่อน cutover ไว้", "ไม่มีเอกสารที่หาไม่เจอ/สถานะเพี้ยน"],
            ],
        ),
        h(2, "5.4 การย้าย workflow ที่ยังวิ่งอยู่"),
        p(
            "เอกสารที่ยังไม่จบ flow ต้องถูกเปิด transaction ใหม่ใน `@srm/glb-workflow` ให้อยู่ state ปัจจุบัน — "
            "ไม่ใช่เริ่มต้นที่ state แรก ขั้นตอนที่ต้องทำต่อเอกสาร: `initializeWorkflow` -> เดิน event จนถึง state ปัจจุบัน "
            "หรือ set `current_state_id`/`current_status_id`/`current_approver` โดยตรง แล้วเติม `workflow_history` "
            "ย้อนหลังจาก `CompensateHistory` เพื่อให้ timeline ไม่ขาด · **วิธีที่จะใช้จริงต้องยืนยันกับทีมเจ้าของ library ก่อน** "
            "เพราะ engine ไม่มี API สำหรับ set state ตรง ๆ"
        ),
        *pending_decision_blocks(
            "5.5 ข้อค้างตัดสินใจที่กระทบ migration (ยังไม่ตัดสิน)",
            [
                ["DP-4 ✅ ปิดแล้ว 2026-08-24 · `fcs_qssi_score`", "reuse ตารางเดิมแบบอ่านอย่างเดียว — ระบบ SBP เดิมนำเข้าให้แล้ว", "สร้างตารางของ SBPGI แล้วโหลดใหม่ — ตกไป", "✅ **ไม่มีอะไรต้อง migrate** — SBPGI อ่านอย่างเดียว ไม่ต้อง dedup/backfill (ตัด Job 1 พร้อมกัน)"],
                ["DP-3 ✅ ตัดสินแล้ว 2026-08-10 = ทางเลือกที่ 3", "view (ไม่มีอะไรให้ migrate)", "ตาราง snapshot (ต้อง migrate + sync job)", "✅ ตัดสินแล้ว — migrate เฉพาะร้านที่เคยเข้ารอบชดเชยเป็น snapshot"],
                ["DP-1 · `reference_id`", "`doc_no` — ตกไป", "**เลือก surrogate id** (`compensation_documents.id` · ส่งเป็น string เพราะ `reference_id` เป็น varchar(255))", "✅ ปิดแล้ว 2026-08-17 — ยืนยันตามระบบเดิม"],
                ["DP-11 · ตัวเลขเงินประกันรายได้", "SBPGI เป็นต้นทาง", "`fr_store_insure` ยังคีย์มือ", "ยังไม่ตัดสิน (เป็นคำถามเชิงธุรกิจ)"],
                ["retention/purge ของเอกสารเก่า", "ย้ายทั้งหมด", "ย้ายเฉพาะช่วงปีที่ตกลง แล้ว archive ที่เหลือ", "ยังไม่ตัดสิน · ระบบเดิมมี ListDocumentsPendingRemoval แต่โครงใหม่ยังไม่มี data retention plan"],
            ],
        ),
    ]


def integration_sbp_platform_extra_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 User Context จาก BFF"),
        p(
            "SBPGI **ไม่มีระบบ login ของตัวเอง** — login จริงอยู่ที่ **AWS Cognito ฝั่ง BFF** · FE ไม่แตะ token · "
            "BFF ยืนยันตัวเองกับ backend ด้วย `x-api-key` แล้วส่งบริบทผู้ใช้ต่อเป็น header · "
            "guard ของ store-backend แปลง header เป็น user context แล้วส่งต่อให้ service ทุกชั้น "
            "(รูปแบบเดียวกับที่ `export-data.service.ts` / `relation.service.ts` / `backlog.service.ts` ของ BFF ใช้อยู่แล้ว)"
        ),
        h(3, "5.1.1 ตัวอย่าง request จริงที่ SBPGI ได้รับ"),
        p("**HTTP request จาก BFF → store-backend (SBPGI)**"),
        code(
            """POST /api/v1/sbpgi/document/2026%2F00123/actions HTTP/1.1
Host: store-backend:3004
Content-Type: application/json
accept-language: th

x-api-key: 8f2b1c94-6d5e-4a70-b1c3-9ee27a4f0d51
x-user-id: 0000123456
x-user-group-id: 08
x-user-full-name: %E0%B8%AA%E0%B8%A1%E0%B8%8A%E0%B8%B2%E0%B8%A2%20%E0%B9%83%E0%B8%88%E0%B8%94%E0%B8%B5
x-user-permissions: [{"url":"/sbpgi/document/waiting","canView":true,"canManage":true,"canExport":false,"canOther":false},{"url":"/sbpgi/report/status-summary","canView":true,"canManage":false,"canExport":true,"canOther":false}]

{"result":"ส่งหน่วยงานส่งเสริมธุรกิจ SBP","comment":"ตรวจยอดชดเชยแล้ว"}""",
            "http",
        ),
        h(3, "5.1.2 แต่ละ header คืออะไร ใช้ทำอะไร"),
        table(
            ["Header", "ตัวอย่างค่า", "มาจากไหน", "SBPGI ใช้ทำอะไร", "ถ้าไม่มี/ผิด"],
            [
                ["`x-api-key`", "`8f2b1c94-6d5e-4a70-b1c3-9ee27a4f0d51`",
                 "env ของ BFF ต่อ backend (`API_STORE_BACKEND_KEY_VALUE`) เทียบกับ `X_API_KEY` ของ store-backend",
                 "พิสูจน์ว่า request มาจาก BFF จริง ไม่ใช่ใครยิงตรง — **ไม่ใช่ตัวตนผู้ใช้**",
                 "**401** `ไม่พบสิทธิ์การเข้าใช้งาน` · `HttpHeaderGuard` ของระบบเดิมเทียบแบบ `===` ตรง ๆ"],
                ["`x-user-id`", "`0000123456`",
                 "`sub`/employee id จาก JWT ของ Cognito (BFF ถอดจาก cookie)",
                 "🔴 **ตัวตนผู้ใช้** — ใส่ใน `created_by`/`updated_by`, `consideration_logs.actor_user_id`, "
                 "และส่งเป็น `userId` เข้า `eventWorkflow` / `initializeWorkflow` ของ engine",
                 "**401** — ห้ามให้ผ่านโดยไม่มี userId เพราะ audit trail จะขาด"],
                ["`x-user-group-id`", "`08`",
                 "auth-backend (ABS) — กลุ่มสิทธิ์ของผู้ใช้",
                 "map เป็น **section_code** ของ workflow (06/08/01/02/03) เพื่อกรองกล่องงานและตัดสินว่ากดปุ่มไหนได้",
                 "**403** เมื่อ endpoint ต้องรู้ section · endpoint อ่านอย่างเดียวยอมให้ผ่านได้"],
                ["`x-user-full-name`", "`%E0%B8%AA%E0%B8%A1%E0%B8%8A%E0%B8%B2%E0%B8%A2%20%E0%B9%83%E0%B8%88%E0%B8%94%E0%B8%B5`  → `สมชาย ใจดี`",
                 "employee backend (`/employees/{empId}/profile`)",
                 "แสดงชื่อผู้ทำรายการใน timeline/อีเมล · **ต้อง `decodeURIComponent` ก่อนใช้เสมอ** (BFF encode มา)",
                 "ไม่บล็อก — fallback เป็น `x-user-id` แล้วเติมชื่อทีหลังจาก `business_user`"],
                ["`x-user-permissions`", "`[{\"url\":\"/sbpgi/document/waiting\",\"canView\":true,\"canManage\":true,\"canExport\":false,\"canOther\":false}]`",
                 "auth-backend `GET /groups/current-user/permissions` (ชุดเดียวกับที่ FE ใช้)",
                 "กันเรียก API ตรงโดยข้ามหน้าจอ — เทียบ `url` ของหน้าที่เป็นเจ้าของ endpoint นั้น + `canManage` ก่อนยอมให้เขียน",
                 "**403** สำหรับ endpoint ที่เขียนข้อมูล · ⚠️ ดูข้อควรระวังด้านล่าง"],
                ["`accept-language`", "`th`", "BFF ส่งต่อจาก browser",
                 "เลือกภาษาข้อความ error — SBPGI ใช้ **ไทย verbatim ตาม SRS** เป็นค่าตั้งต้นเสมอ", "ไม่บล็อก — default `th`"],
            ],
        ),
        p(
            "⚠️ **ข้อควรระวัง `x-user-permissions` (ต้องยืนยันกับทีม BFF ก่อนลงมือ):** "
            "เอกสารวิเคราะห์ระบบเดิมยืนยันแค่ว่า *มี* header ตัวนี้และเนื้อหาคือชุดสิทธิ์ต่อ URL "
            "(`canView` / `canManage` / `canExport` / `canOther`) แต่**ยังไม่ยืนยัน 2 เรื่อง** — "
            "(1) รูปแบบที่ serialize มา (JSON ตรง ๆ · base64 · หรือย่อเป็น CSV) และ "
            "(2) พฤติกรรมเมื่อสิทธิ์เยอะจน header ยาวเกินลิมิตของ proxy (ปกติ ~8 KB) · "
            "**จนกว่าจะยืนยัน ห้ามใช้ header นี้เป็นด่านเดียว** — ให้ตัดสินสิทธิ์เขียนจาก "
            "`x-user-group-id` + สถานะเอกสาร + ผู้ถืองานจาก `getTransaction()` ของ engine เป็นหลัก "
            "แล้วใช้ `x-user-permissions` เป็นด่านเสริม"
        ),
        h(3, "5.1.3 Guard ที่ต้องเขียน"),
        code(
            """// src/common/guards/bff-user.guard.ts (ยึด convention ของ store-backend)
export interface SbpgiUser {
  userId: string;          // x-user-id            เช่น '0000123456'
  groupId: string;         // x-user-group-id      เช่น '08'
  fullName: string;        // x-user-full-name     decode แล้ว เช่น 'สมชาย ใจดี'
  permissions: UrlPermission[];   // x-user-permissions  (ดูข้อควรระวัง 5.1.2)
}

@Injectable()
export class BffUserGuard implements CanActivate {
  constructor(private readonly config: ConfigService) {}

  canActivate(ctx: ExecutionContext): boolean {
    const req = ctx.switchToHttp().getRequest();

    // 1) ยืนยันว่ามาจาก BFF จริง — เทียบกับ X_API_KEY (มาจาก Secret Manager ห้าม hardcode/commit)
    const apiKey = req.headers['x-api-key'];
    if (!apiKey || apiKey !== this.config.get('X_API_KEY')) {
      throw new UnauthorizedException('ไม่พบสิทธิ์การเข้าใช้งาน');
    }

    // 2) ตัวตนผู้ใช้ — ไม่มี userId = ไม่ให้ผ่าน เพราะ audit trail จะขาด
    const userId = req.headers['x-user-id'];
    if (!userId) throw new UnauthorizedException('ไม่พบสิทธิ์การเข้าใช้งาน');

    req.user = {
      userId,
      groupId: req.headers['x-user-group-id'] ?? '',
      // BFF encodeURIComponent มา -> ต้อง decode ก่อนใช้ ไม่งั้นชื่อไทยกลายเป็น %E0%B8...
      fullName: safeDecode(req.headers['x-user-full-name']),
      permissions: parsePermissions(req.headers['x-user-permissions']),
    } satisfies SbpgiUser;
    return true;
  }
}

/** header อาจว่าง/พัง — ห้ามให้ทั้ง request ล้มเพราะ decode ไม่ผ่าน */
function safeDecode(v?: string): string {
  if (!v) return '';
  try { return decodeURIComponent(v); } catch { return v; }
}

/** ⚠️ รูปแบบ serialize ยังไม่ยืนยัน (ดู 5.1.2) — parse ไม่ผ่านให้คืน [] แล้วตกไปใช้ด่าน group/สถานะแทน */
function parsePermissions(v?: string): UrlPermission[] {
  if (!v) return [];
  try { const j = JSON.parse(v); return Array.isArray(j) ? j : []; } catch { return []; }
}""",
            "ts",
        ),
        h(3, "5.1.4 ใช้ใน controller / ทดสอบเอง"),
        code(
            """// controller — อ่าน user ที่ guard แปะไว้ ห้ามอ่าน header ตรงในทุก service
// ⚠️ ไม่มี 'sbpgi/' ใน @Controller() เพราะ prefix ผูกที่ระดับโมดูล:
//    RouterModule.register([{ path: 'sbpgi', module: SbpgiModule }])
//    -> URL จริงคือ /api/v1/sbpgi/document/... (ดู LLDD-BE-API-Common-Contracts 5.80)
@UseGuards(BffUserGuard)
@Controller('document')
export class SbpgiDocumentController {
  @Post(':docNo/actions')
  submit(@Param('docNo') docNo: string, @Body() dto: ActionDto, @Req() req: { user: SbpgiUser }) {
    // docNo มาเป็น '2026%2F00123' -> Nest decode ให้แล้วเป็น '2026/00123'
    return this.service.submit(docNo, dto, req.user);
  }
}""",
            "ts",
        ),
        code(
            """# ยิงทดสอบเองตอน dev (ไม่ผ่าน BFF) — ใส่ header ให้ครบเหมือนที่ BFF ส่งจริง
curl -X POST 'http://localhost:3004/api/v1/sbpgi/document/2026%2F00123/actions' \\
  -H 'x-api-key: '"$X_API_KEY" \\
  -H 'x-user-id: 0000123456' \\
  -H 'x-user-group-id: 08' \\
  -H 'x-user-full-name: %E0%B8%AA%E0%B8%A1%E0%B8%8A%E0%B8%B2%E0%B8%A2%20%E0%B9%83%E0%B8%88%E0%B8%94%E0%B8%B5' \\
  -H 'accept-language: th' \\
  -H 'Content-Type: application/json' \\
  -d '{"result":"ส่งหน่วยงานส่งเสริมธุรกิจ SBP","comment":"ตรวจยอดชดเชยแล้ว"}'""",
            "bash",
        ),
        bullets([
            "🔴 **ห้ามให้ FE ส่ง header เหล่านี้เอง** — ต้องมาจาก BFF เท่านั้น · store-backend ต้องอยู่หลัง network layer ที่เปิดให้เฉพาะ BFF เข้าถึง",
            "🔴 **ห้าม log ค่า `x-api-key`** ลง application log / error message ทุกกรณี",
            "`x-user-id` ที่ส่งเข้า engine ต้องเป็นตัวเดียวกับที่บันทึกใน `consideration_logs` — ไม่งั้น timeline ของ SBPGI กับ `workflow_history` ของ engine จะชี้คนละคน",
            "unit test ต้องครอบ: ไม่มี `x-api-key` → 401 · `x-api-key` ผิด → 401 · ไม่มี `x-user-id` → 401 · "
            "`x-user-full-name` เป็น %-encoded → decode ถูก · `x-user-permissions` พัง/ว่าง → ไม่ throw แต่ตกไปใช้ด่าน group",
        ]),
        h(2, "5.2 Response Envelope"),
        code(
            """// ResponseInterceptor ของ store-backend ห่อให้แล้ว — service ห้ามห่อซ้ำ
// success : { success: true,  data: <payload> }
// error   : { success: false, data: null, error: { code, message } }
// message ต้องเป็นภาษาไทย verbatim ตาม SRS และโยนผ่าน HttpException เท่านั้น""",
            "ts",
        ),
        h(2, "5.3 ไฟล์แนบผ่าน S3 ของระบบเดิม"),
        table(
            ["ขั้นตอน", "ปลายทาง", "สิ่งที่ SBPGI เก็บเอง"],
            [
                ["อัปโหลด", "POST /statement/upload-file-aws (ระบบ SBP เดิม)", "objectKey + ชื่อไฟล์ + ขนาด + content type + section_code"],
                ["ดาวน์โหลด", "POST /statement/download-file-aws (ระบบ SBP เดิม)", "SBPGI แปลงเป็น **binary stream** ก่อนคืนให้ FE · ห้ามคืน objectKey ให้ FE"],
                ["ลบ/purge", "lifecycle ของ S3 + flag ใน document_attachments", "purge_flag / storage_delete_status"],
            ],
        ),
        p(
            "🔴 **ข้อจำกัดที่ต้องรู้ก่อนเขียนโค้ด (ตรวจ `store-backend` 2026-08-26):** `AwsService` ของระบบเดิมเป็น "
            "**wrapper แบบ base64** ไม่ใช่ stream — `upload-file-aws` รับไฟล์เป็น base64 และ `download-file-aws` "
            "**คืนไฟล์เป็น base64 ใน JSON** สายส่งจริงจึงเป็น "
            "`FE ← binary stream ← SBPGI BE ← base64 JSON ← /statement/download-file-aws ← S3`"
        ),
        table(["ผลกระทบ", "ตัวเลข", "ต้องทำอย่างไร"], [
            ["base64 ทำให้ payload โตขึ้น ~33%", "ไฟล์ 5 MB → **~6.7 MB** ใน JSON",
             "ยังไม่ชน body limit ของ store-backend (**100 MB** ที่ `main.ts:33`) แต่กิน memory ต่อ request จริง"],
            ["ปุ่ม **ดาวน์โหลดทั้งหมด (.zip)** ต้องดึงหลายไฟล์", "n ไฟล์ × 1.33 พร้อมกัน",
             "🔴 ห้ามโหลดทุกไฟล์เข้า memory พร้อมกัน — ดึงทีละไฟล์แล้ว **stream เข้า zip ทันที** (archiver แบบ streaming)"],
            ["FE ไม่ควรรู้ว่าใต้ท้องเป็น base64", "—",
             "สัญญาฝั่ง FE ยังเป็น **binary stream + Content-Type / Content-Disposition** ตาม `LLDD-BE-API-Attachment-Sales-Timeline` — SBPGI เป็นคนแปลง"],
        ]),
        p(
            "⚠️ **ต้องยืนยันกับทีม store-backend:** wrapper รองรับ **range request / partial download** หรือไม่ · "
            "ถ้าไม่รองรับ ไฟล์ใหญ่จะ resume ไม่ได้ และปุ่มดาวน์โหลดทั้งหมดต้องกำหนดเพดานจำนวน/ขนาดรวม"
        ),
        p(
            "**แก้ความเข้าใจผิดเดิม (ตรวจ schema จริง 2026-08-07):** เอกสารรุ่นก่อนเคยเขียนว่าถ้าใช้ตาราง "
            "`upload_general` ของระบบเดิมจะติด FK `job_id` — **ไม่จริง** `job_id` และ `audit_log_id` เป็น "
            "**nullable ทั้งคู่** · เหตุผลจริงที่ SBPGI ต้องมี `document_attachments` ของตัวเองคือ `upload_general` "
            "**ไม่มีคอลัมน์** `file_size` · `content_type` · `section_code` · `upload_status` · `purge_flag`"
        ),
        h(2, "5.4 อีเมล"),
        p(
            "ส่งผ่าน `@gosoft-sbp/email-lib` โดยอ่านเนื้อหาจาก `email_template` (85 แถว) และบันทึกผลที่ "
            "`email_sent` (5,214 แถว)"
        ),
        p(
            "**แก้ความเข้าใจผิดเดิม (ตรวจ schema จริง 2026-08-07):** เอกสารรุ่นก่อนเคยเขียนว่าระบบเดิม "
            "**ไม่มีที่เก็บ CC ของอีเมล** — **ไม่จริง** มีอยู่ 3 ที่: `email_sent.mail_cc` · "
            "`fcs_reminder_log.reminder_cc` · `fml_email_account`"
        ),
        h(2, "5.5 ค่ากำหนดกลาง — `mas_param` กับ `common_code` คืออะไร"),
        p(
            "🔴 **สองตารางนี้เป็นของระบบ SBP เดิม อยู่ใน schema `sps_store` เท่านั้น** และ "
            "**ค่าของ SBPGI ยังไม่มีอยู่จริง — เป็นแถวที่เราต้อง seed เองตอน setup** "
            "(เอกสารรอบก่อนเขียนกำกวมจนอ่านเหมือนมีข้อมูลอยู่แล้ว · แก้ 2026-08-25)"
        ),
        table(
            ["ตาราง", "คืออะไร", "โครงคีย์", "ของจริงตอนนี้ (ตรวจ 07/08/2026)"],
            [
                ["`sps_store.mas_param`",
                 "**ตาราง config กลางของ store-backend** — คู่ชื่อ/ค่าแบบอิสระ ที่ทั้งระบบเดิมใช้ร่วมกัน "
                 "(เช่น `GROUP_ID_VIEW_ALL_STMT` คุมว่ากลุ่มไหนเห็นใบแจ้งยอดทั้งหมด · ช่วงวันที่ของไฟล์อากรแสตมป์)",
                 "`param_name` · `param_value`(4000) · `ref_name` · `description` · `is_config` · `active_flag`",
                 "**93,752 แถว** · ⚠️ **ไม่มี PK ไม่มี unique** มีแค่ btree `(param_name, param_value)` → "
                 "ชื่อพารามิเตอร์ซ้ำได้ ต้องกันเองที่ระดับแอปและ `WHERE active_flag = 'Y'` เสมอ"],
                ["`sps_store.common_code`",
                 "**lookup กลาง** ของทั้งระบบเดิม — ชุดรหัส/ชื่อที่ใช้ทำ dropdown",
                 "`code_type`(**20**) · `seq_no` · `code_value`(100) · `code_name`(1000) · `other_value`(50) · `code_mapping`(100) · `active_flag`",
                 "**2,609 แถว** · ⚠️ **ไม่มี PK ไม่มี unique** บน (`code_type`,`code_value`) · "
                 "`code_type` ต้องลงทะเบียนที่ **`common_code_type`** (376 แถว) ก่อน"],
            ],
        ),
        h(3, "5.5.1 ทำไมค้นแล้วไม่เจอข้อมูล (2 กับดักที่เจอจริง)"),
        table(["กับดัก", "ข้อเท็จจริง", "ต้องทำอย่างไร"], [
            ["ค้นผิด schema",
             "`mas_param` มี **เฉพาะ `sps_store`** — ใน `sps_auth` **ไม่มีตารางนี้เลย** · "
             "ส่วน `common_code` มี **ทั้งสอง schema แต่เป็นคนละตาราง**: `sps_store` 14 คอลัมน์ 2,609 แถว vs "
             "`sps_auth` **13 คอลัมน์ 2,594 แถว** (ชุดเก่าของ auth-backend)",
             "🔴 SBPGI ใช้ **`sps_store` เท่านั้น** · เขียน schema นำหน้าทุกครั้งใน SQL "
             "(กับดักเดียวกับตาราง `workflow_*` ที่มีสองชุด — ดู 5.4)"],
            ["คิดว่าค่าของ SBPGI มีอยู่แล้ว",
             "`SBPGI_APPROVE_LIMIT` · `SBPGI_DECISION` · `SBPGI_DATASOURCE` **ยังไม่มีสักแถวในระบบจริง** — "
             "เป็นค่าที่การออกแบบ *วางแผนจะเพิ่ม* ไม่ใช่ของเดิมที่ reuse ได้ทันที",
             "ต้อง **seed เองตอน setup** (ดู 5.5.2) และนับเป็นงานของ `LLDD-BE-Data-Migration-Cutover`"],
        ]),
        h(3, "5.5.2 ค่าที่ SBPGI ต้อง seed เอง"),
        table(["ค่า", "ลงที่ไหน", "คีย์ที่ใช้", "สถานะ"], [
            ["วงเงินอนุมัติ เกณฑ์เดียว **100,000**", "`sps_store.common_code`",
             "`code_type = 'SBPGI_APPROVE_LIMIT'` · `code_value = 'THRESHOLD'` · `code_name = '100000'`",
             "🔴 **ยังไม่มี — ต้อง seed**"],
            ["ผลการพิจารณา 6 ค่า (มติ DP-9)", "`sps_store.common_code`",
             "`code_type = 'SBPGI_DECISION'`", "🔴 **ยังไม่มี — ต้อง seed**"],
            ["ต้นทาง `PRO` (เชิงรุก) · `REA` (เชิงรับ)", "`sps_store.common_code`",
             "`code_type = 'SBPGI_DATASOURCE'`", "🔴 **ยังไม่มี — ต้อง seed** (เพิ่มจากของเดิมที่มี `ALM`/`STA`)"],
            ["รัศมีผลกระทบ 1 กม. (กทม./ปริมณฑล) · 2 กม. (ต่างจังหวัด)", "`sps_store.mas_param`",
             "`param_name = 'SBPGI_IMPACT_RADIUS_BKK' / '..._UPC'`", "🔴 **ยังไม่มี — ต้อง seed** · อ่านตอนคำนวณ ห้าม hardcode"],
            ["เกณฑ์ยอดขายไม่ครบ **60 วัน** · growth rate **-10%**", "`sps_store.mas_param`",
             "`param_name = 'SBPGI_SALES_DAYS_MIN' / 'SBPGI_GROWTH_RATE_MAX'`",
             "🔴 **ยังไม่มี — ต้อง seed** · ใช้กับธงข้อมูลผิดปกติและ Gen Flow Gate"],
        ]),
        code(
            """-- seed ตอน setup (idempotent) — ⚠️ ทั้งสองตารางไม่มี unique จึงต้อง guard ด้วย NOT EXISTS เอง
-- 1) ลงทะเบียน code_type ก่อนเสมอ ไม่งั้น dropdown ของระบบเดิมจะไม่รู้จัก
INSERT INTO sps_store.common_code_type (code_type, code_type_name, active_flag, create_date, create_user)
SELECT 'SBPGI_APPROVE_LIMIT', 'วงเงินอนุมัติ ประกันรายได้', 'Y', CURRENT_TIMESTAMP, 'SBPGI-SETUP'
WHERE NOT EXISTS (SELECT 1 FROM sps_store.common_code_type WHERE code_type = 'SBPGI_APPROVE_LIMIT');

-- 2) ค่าจริง · code_type เป็น varchar(20) -> 'SBPGI_APPROVE_LIMIT' = 19 ตัว เหลือที่ว่าง 1 ตัวเท่านั้น
INSERT INTO sps_store.common_code (code_type, seq_no, code_value, code_name, active_flag, create_date, create_user)
SELECT 'SBPGI_APPROVE_LIMIT', 1, 'THRESHOLD', '100000', 'Y', CURRENT_TIMESTAMP, 'SBPGI-SETUP'
WHERE NOT EXISTS (SELECT 1 FROM sps_store.common_code
                  WHERE code_type = 'SBPGI_APPROVE_LIMIT' AND code_value = 'THRESHOLD');

-- 3) ค่ากำหนดกลางที่ไม่ใช่ lookup -> mas_param
INSERT INTO sps_store.mas_param (param_name, param_value, description, is_config, active_flag, create_by, create_date)
SELECT 'SBPGI_SALES_DAYS_MIN', '60', 'จำนวนวันยอดขายขั้นต่ำก่อนถือว่าข้อมูลครบ', 'Y', 'Y', 'SBPGI-SETUP', CURRENT_TIMESTAMP
WHERE NOT EXISTS (SELECT 1 FROM sps_store.mas_param
                  WHERE param_name = 'SBPGI_SALES_DAYS_MIN' AND active_flag = 'Y');

-- อ่านค่ากลับมาใช้ — ต้องกรอง active_flag เสมอ และ LIMIT 1 เพราะไม่มี unique กันซ้ำ
SELECT param_value FROM sps_store.mas_param
WHERE param_name = :name AND active_flag = 'Y'
ORDER BY update_date DESC NULLS LAST, create_date DESC LIMIT 1;""",
            "sql",
        ),
        h(3, "5.5.3 ของเดิมอยู่ที่ไหน — ทำไมค้นใน `mas_param`/`common_code` แล้วไม่เจอ"),
        p(
            "🔴 **ข้อมูลของงานประกันรายได้เดิม ไม่เคยอยู่ใน `mas_param` หรือ `common_code` เลย** — สองตารางนั้นเป็นของ "
            "**store-backend (SBP Mall)** ส่วนของเดิมอยู่คนละฐานข้อมูล คือ **SQL Server `CPA_FRN_FGI`** (ฝั่ง K2 · 47 ตาราง) "
            "กับ **Oracle** (ฝั่ง FGI/FCS) · ที่เอกสารเขียนว่า *\"ใช้ `common_code` แทน\"* หมายถึง **ปลายทางที่จะย้ายไป** "
            "ไม่ใช่ที่ที่ข้อมูลอยู่ตอนนี้"
        ),
        table(["ค่า", "ของเดิมอยู่ที่ (ฐานข้อมูลเดิม)", "ปลายทางใหม่", "ต้องทำอะไร"], [
            ["วงเงินอนุมัติ", "**MSSQL** `SectionProfile.SectionLimitCost` — มีค่าเดียวคือ section 02 (GM) = 100,000 · AVP เป็น NULL",
             "`sps_store.common_code` `SBPGI_APPROVE_LIMIT`",
             "🔴 **ห้าม migrate ค่าเดิมมาตรง ๆ** — เกณฑ์เก่าไม่ตรง SDD GI · **seed ใหม่** เป็นเกณฑ์เดียว 100,000"],
            ["ผลการพิจารณา", "**MSSQL** `DecisionProfile`", "`sps_store.common_code` `SBPGI_DECISION`",
             "แปลงชื่อ 3 ชุด (ปุ่ม/flow/ผลลัพธ์) ลง `code_name` / `code_mapping` / `other_value` แล้ว seed"],
            ["ปัจจัยภายนอก", "**MSSQL** `FactorProfile`", "**ตาราง `external_factors` ของ SBPGI** (ไม่ได้ไป `common_code` · มติ DP-9)",
             "migrate เข้าตารางของเราเอง เพราะมีหน้าจอ CRUD และช่องข้อความของ `common_code` ไม่พอ"],
            ["ร้านคู่แข่ง 11 แบรนด์", "**MSSQL** `CompetitionProfile` + **ORA** `MAS_STORE_COMPETITOR`",
             "**ตาราง `competitors` ของ SBPGI** (มติ DP-9)", "migrate เข้าตารางของเราเอง"],
            ["รัศมี/เกณฑ์คำนวณ (1-2 กม. · 60 วัน · -10%)", "**hardcode อยู่ในโค้ด Java เดิม** ไม่ได้อยู่ในตารางไหน",
             "`sps_store.mas_param`", "🔴 **ไม่มีของเดิมให้ migrate** — ยกค่าจากโค้ดมา seed เป็น data"],
        ]),
        p(
            "**สรุปสั้น ๆ:** ค้นใน `mas_param`/`common_code` แล้วไม่เจอเป็นเรื่อง**ปกติและถูกต้อง** — "
            "(1) ของเดิมอยู่คนละฐานข้อมูล (MSSQL/Oracle) · (2) ค่าของ SBPGI ยังไม่ถูก seed · "
            "(3) ถ้าเปิดผิด schema (`sps_auth`) จะยิ่งไม่เจอเพราะ `mas_param` ไม่มีในนั้นเลย"
        ),
        bullets([
            "🔴 **`code_type` เป็น `varchar(20)`** (ขณะที่ `common_code_type.code_type` เป็น `varchar(50)`) — "
            "`SBPGI_APPROVE_LIMIT` ยาว 19 ตัว **เหลือที่ว่างแค่ 1 ตัวอักษร** · ตั้งชื่อ `code_type` ใหม่ห้ามเกิน 20",
            "ระบบเดิม**ไม่มี POST/PUT/DELETE ของ `common_code`** (module `common` มีแต่ GET) — "
            "SBPGI จะเขียนลง lookup กลางที่ทุกโมดูลใช้ร่วม ต้องทำผ่าน migration script ที่ review ได้ ไม่ใช่หน้าจอ",
            "SBPGI **อ่านอย่างเดียวในเวลาปกติ** — แก้ค่าทำที่ระบบ SBP เดิม (หน้าจอ Global Config ของ SBPGI ถูกลบไปแล้ว 2026-08-06) · "
            "การเขียนเกิดเฉพาะตอน **seed/cutover** เท่านั้น",
        ]),
        *pending_decision_blocks(
            "5.6 ข้อค้างตัดสินใจที่กระทบ integration (ยังไม่ตัดสิน)",
            [
                ["DP-5 · อีเมล ✅ ปิดแล้ว 2026-08-14", "ให้ engine ส่งเอง — **ตกไป** เพราะ `triggerEvent` ไม่มี `mailTo`/`mailCc`/`param` ที่ `sendEmail` บังคับ", "**เลือกทางนี้:** workflow ให้เลข template ผ่าน `workflow_route.email_id` แล้ว **SBPGI เรียก `sendEmail()` ของ email-lib เอง** · reminder/escalation ที่ไม่ใช่ transition เก็บเลข template ที่ `mas_param`", "ปิดแล้ว · เหลือยืนยันกับทีม engine ว่าไม่ส่งซ้ำ"],
                ["DP-8 ✅ ปิดแล้ว 2026-08-24 · `document_attachments`", "**เลือกข้อนี้ — ตารางของ SBPGI เก็บ metadata เอง** แล้วใช้ service S3 ของระบบเดิม ไม่เขียน storage layer เอง", "ต่อยอด `upload_general` ของระบบเดิม — ตกไป เพราะไม่มีคอลัมน์ `file_size` · `content_type` · `section_code` · `upload_status` · `purge_flag`", "✅ ปิดแล้ว 2026-08-24 — เหตุผลเต็มอยู่ที่ 5.3"],
                ["DP-10 ✅ ปิดแล้ว 2026-08-21 · ที่อยู่ของ SBPGI", "**เลือกข้อนี้** — โมดูลใน `srm-sps-spsap-store-backend` เดิม", "backend ใหม่แยกต่างหาก — ตกไป", "✅ ปิดแล้ว — ใช้ guard/interceptor/response envelope ของ store-backend เดิมได้ทันที ไม่ต้องเขียนใหม่"],
                ["DP-6 · `interface_transactions`", "ออกแบบใหม่ตาม DDL ปัจจุบัน", "ลอกแพตเทิร์น `statement_summary` ของระบบเดิม", "ยังไม่ตัดสิน"],
            ],
        ),
    ]


def workflow_engine_definition_extra_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.0 ทำไมเอกสารฉบับนี้ต้องปิดเป็นฉบับแรก"),
        p(
            "เอกสารฉบับนี้ **ไม่มี endpoint ของตัวเอง** — ผลลัพธ์คือชุดนิยาม state/status/route/part ที่เอกสารอื่น "
            "เอาไปใช้ต่อ จึงต้องจบก่อนผู้บริโภคทั้งหมดเริ่ม (ปรับลำดับ 2026-08-10: เดิมถูกจัดไว้ท้ายกลุ่ม API "
            "ทำให้ `BE-API-Document-Workflow-Actions` และ `BE-API-Workflow-Instances` เริ่มก่อนเอกสารที่นิยาม "
            "สิ่งที่มันต้องใช้)"
        ),
        table(
            ["เอกสารที่รอ", "รออะไรจากฉบับนี้"],
            [
                ["BE-API-Document-Workflow-Actions", "รหัส event ต่อปุ่ม · route ของแต่ละ state · เงื่อนไขแตกสายตามวงเงิน"],
                ["BE-API-Workflow-Instances", "โครง version/state/status ที่จะ query และรูปแบบ payload ของ engine"],
                ["BE-Job-8b-StartInternalWorkflow", "ลำดับเรียก initialize -> addPreApprover และค่า `referenceId`"],
                ["FE-Document-Detail (5 ฉบับ role)", "`workflow_part_display` READ/WRITE ต่อ state ที่คุมการแสดงผลรายส่วน"],
            ],
        ),
        h(2, "5.1 Engine คือของกลาง 13 ตาราง ใน schema `sps_store`"),
        p(
            "`@srm/glb-workflow` เป็น library กลางที่ทุกระบบใน SBP platform import ไปใช้ "
            "(ต้นฉบับ: `SBP/TSM-SRM-LLDD-SBP-workflow-1.2.md` v1.2 ลงวันที่ 29/04/2026) · "
            "SBPGI ใช้ engine ตัวนี้ตามมติ 2026-08-06 ที่ตัดตาราง workflow ของตัวเองทิ้งทั้งหมด"
        ),
        p(
            "**ตัวเลขที่เอกสารรุ่นก่อนเขียนผิด 2 จุด (แก้แล้ว 2026-08-07):** (1) engine มี **13 ตาราง ไม่ใช่ 10** · "
            "(2) engine ตัวที่ใช้งานจริงอยู่ schema **`sps_store` ไม่ใช่ `sps_auth`** — ทั้งสอง schema มีครบ 13 ตาราง"
            "เหมือนกันแต่เป็นคนละชุดและคนละเวอร์ชัน (`workflow_state` ของ `sps_auth` มี 3 คอลัมน์ · ของ `sps_store` "
            "มี 4 คอลัมน์) · `sps_auth.workflow_transaction` มีแค่ 55 แถว (route 41 · state 10) ซึ่งเป็นชุดของ "
            "auth-backend คนละเรื่องกัน"
        ),
        table(
            ["กลุ่ม", "ตาราง", "หน้าที่"],
            [
                ["นิยาม flow (config)", "workflow · workflow_version · workflow_state · workflow_status · workflow_event · workflow_route", "ตั้งครั้งเดียวต่อระบบ · `workflow_version.url_main` / `url_param_mapping` ทำให้ inbox กลางลิงก์กลับหน้าเอกสารของ SBPGI ได้"],
                ["กลุ่มผู้อนุมัติ", "workflow_group · workflow_group_map", "`map_table` ว่าง = เทียบกับ field ของ user ตรง ๆ · ระบุ map_table = ต้องเป็น view ที่ where ด้วย user_id/group_id ได้"],
                ["ข้อมูลรันไทม์", "workflow_transaction · workflow_history · workflow_approver", "19,283 / 38,010 / 96,542 แถวใน sps_store (ตรวจ 2026-08-07)"],
                ["คุมการแสดงผล", "workflow_part · workflow_part_display", "`part_display_type` = READ / WRITE ต่อ state — คืนมากับ getPermissionEvents"],
            ],
        ),
        h(2, "5.2 ความเสี่ยงที่ต้องคุยกับทีมเจ้าของ library"),
        table(
            ["ความเสี่ยง", "ข้อเท็จจริงที่ตรวจแล้ว", "ผลกระทบต่อ SBPGI", "สิ่งที่ต้องทำ"],
            [
                [
                    "`sps_store.workflow_transaction` ไม่มี PK และไม่มี index เลย",
                    "มี 19,283 แถวแต่ schema dump ไม่พบ PK/index ใด ๆ (ตารางชื่อเดียวกันใน `sps_auth` มี PK `transaction_id` ปกติ) · `workflow_state` / `workflow_event` / `workflow_part_display` ของ `sps_store` ก็ไม่มี PK เช่นกัน",
                    "ทุกครั้งที่เปิดเอกสารหรือกด action ต้อง seq-scan 19,283 แถวเพื่อหา `reference_id` · ไม่มีอะไรกัน initialize ซ้ำแม้ระดับ application · จะแย่ลงเมื่อ SBPGI เพิ่มอีกราวหมื่นแถวต่อปี",
                    "ยื่นเรื่องขอ sign-off เพิ่ม PK + UNIQUE(version_id, reference_id) + index กับทีมเจ้าของ `@srm/glb-workflow` · ระหว่างรอ ให้กันซ้ำ + เก็บ mapping ที่ฝั่ง SBPGI (**ทางเลือกที่จะใช้จริงยังไม่ตัดสิน — DP-2**)",
                ],
                [
                    "`part_display_type` สะกดว่า `WRTIE` ในไฟล์ต้นฉบับ",
                    "สะกดผิดทุกแถวของชีต `sample data`",
                    "ถ้า SBPGI เขียนค่า `WRITE` แล้ว engine เทียบกับ `WRTIE` การแสดงผลจะเพี้ยนทั้งหน้า",
                    "ยืนยันค่าจริงในระบบกับทีม library ก่อนลงทะเบียน part",
                ],
                [
                    "`workflow_route` มี 2 นิยามในไฟล์เดียวกัน",
                    "ชีต `sample data` มีคอลัมน์ `group_id` แต่ entity ที่แนบมาใช้ `approver` และตั้งชื่อ property ว่า `approverRoleId`",
                    "เขียนโค้ดผูกผู้อนุมัติผิดคอลัมน์",
                    "ยืนยัน schema จริงของ route กับทีม library",
                ],
                [
                    "ไม่มี API ถอน/แก้ผู้อนุมัติล่วงหน้า",
                    "มีแต่ `addPreApprover`",
                    "เคสเปลี่ยนตัวผู้อนุมัติ (ลาออก/รักษาการ) ทำไม่ได้ผ่าน library",
                    "ถามทีม library ว่าจะเพิ่มให้หรือให้ SBPGI แก้ตารางตรง",
                ],
            ],
        ),
        h(2, "5.3 API ของ engine — 8 function (ยึด LLDD ของ lib · ปิดข้อค้าง 2026-08-14)"),
        p(
            "แหล่งความจริงคือชีต `Detail` ของ `SBP/TSM-SRM-LLDD SBP workflow 1.2.xlsx` "
            "ซึ่งเป็น **เอกสารของ lib เอง** · ชื่อที่เคยนับว่าขัดกันไม่ใช่ชื่อ API: "
            "*Trigger Event* เป็น**ชื่อหัวข้อของขั้นตอนภายใน** `eventWorkflow` ในชีต 2 "
            "ส่วน `TriggerEventUseCase` / `AddPreparedApproverUseCase` / `GetPendingFlowUseCase` "
            "เป็น **UseCase class ที่ store-backend ห่อไว้ใช้เอง** ไม่ใช่ API ของ lib"
        ),
        table(
            ["#", "function", "พารามิเตอร์ (ชีต Detail)", "SBPGI ใช้ที่ไหน"],
            [
                ["1", "`initializeWorkflow`", "version, userId, referenceId",
                 "เปิด flow ให้เอกสารใหม่ (Job 8b · `POST /sbpgi/workflow/instances`)"],
                ["2", "`eventWorkflow`",
                 "version, referenceId, event, eventParam, remark, userId **+ userData · userFullname · nextApproverId** "
                 "(ส่วนขยาย 29/04 · 20/05 · 16/06/2026 — ยึดชุดนี้เวลาเขียนโค้ด)",
                 "`POST /sbpgi/document/{docNo}/actions`"],
                ["3", "`getPermissionEvents`", "version, referenceId, userData",
                 "ปุ่ม/ผลพิจารณาที่ user กดได้ในหน้าเอกสาร"],
                ["4", "`getHistory`", "version, referenceId", "`GET /sbpgi/document/{docNo}/timeline`"],
                ["5", "`getTransaction`", "version, referenceId", "สถานะ + ผู้ถืองานปัจจุบันของเอกสาร"],
                ["6", "`getPendingFlowByUser`", "userData",
                 "**หน้า เอกสาร → รอดำเนินการ** + reminder รายสัปดาห์"],
                ["7", "`getWorkflowsByUser`", "userData",
                 "**หน้า เอกสาร → ที่เกี่ยวข้อง** (รวมที่ยังไม่ถึงคิวและที่อนุมัติไปแล้ว)"],
                ["8", "`addPreApprover`", "version, userId, referenceId, state_id, approver, seq",
                 "ตั้งผู้อนุมัติล่วงหน้าของขั้นถัดไป"],
            ],
        ),
        h(2, "5.4 นิยาม flow ของ SBPGI ที่ต้อง register"),
        table(
            ["state", "ชื่อสถานะเอกสาร", "event ที่ทำได้", "ปลายทาง"],
            [
                ["06", "รอฝ่าย SBP DSA ดำเนินการ", "submit (ส่งเจ้าหน้าที่ SBP DSA) · reject (เห็นควรไม่ชดเชย) · cancel (หยุดชดเชย) · submit (ส่งหน่วยงานส่งเสริมธุรกิจ SBP)", "08 หรือ 01 หรือจบ flow"],
                ["08", "รอเจ้าหน้าที่ SBP DSA ดำเนินการ", "submit (คำนวณเงินชดเชยเรียบร้อย)", "01"],
                ["01", "รอหน่วยงานส่งเสริมธุรกิจ SBP ดำเนินการ", "approve (เห็นควรชดเชย) · reject (เห็นควรไม่ชดเชย → จบ flow ทันที) · sendback (ฝ่าย SBP DSA ดำเนินการ)", "02 · จบ flow · 06"],
                ["02", "รอ GM ส่งเสริมธุรกิจ SBP ดำเนินการ", "approve (เห็นควรชดเชย) · reject (เห็นควรไม่ชดเชย → จบ flow ทันที) · sendback", "จบ flow เมื่อยอด < 100,000 · ไป 03 เมื่อ ≥ 100,000 · 01"],
                ["03", "รอผู้บริหารสำนักบริหาร SBP ดำเนินการ", "approve (เห็นควรชดเชย) · sendback", "จบ flow · 02"],
            ],
        ),
        code(
            """-- ⚠️ ตัวอย่างนี้คือ **ทางเลือก B ของข้อค้าง 5.6 (ยังไม่ตัดสิน) — ห้าม seed ลงจริงก่อนได้ข้อสรุป**
-- มติเดิม (ทางเลือก A) คือเก็บวงเงินที่ `common_code` (code_type = SBPGI_APPROVE_LIMIT) แล้ว "อ่านทุกครั้ง ห้าม hardcode"
-- ตามที่ LLDD-BE-Integration-SBP-Platform / LLDD-Database ระบุไว้ · กติกาที่แน่นอนคือ **ห้ามเก็บสองที่**
-- ถ้าเลือกทางเลือก A: route ยังแตกสองเส้นเหมือนเดิม แต่ SBPGI เป็นผู้เทียบยอดกับ common_code
--   แล้วส่งผลลัพธ์ (เช่น eventParam = {"limitTier":"GM"|"AVP"}) ให้ engine เลือก route โดยไม่ฝังตัวเลขใน condition_json
--
-- ตัวอย่างทางเลือก B (ฝังวงเงินใน condition_json ตามความสามารถของ engine):
-- SBPGI ส่ง eventParam = {"amount": <ยอดชดเชยรวมของเอกสาร>} แล้วให้ engine เลือก route เอง
-- seq = ลำดับที่ engine ใช้ไล่ตรวจ condition_json (ตัวแรกที่ตรงชนะ)
-- ตัวเลข 100000 ด้านล่างเป็นค่า **ตัวอย่าง** ของเกณฑ์เดียว (มติ 2026-08-18) ไม่ใช่ค่าที่ตกลงให้ hardcode
INSERT INTO sps_store.workflow_route
  (version_id, from_state_id, event, to_state_id, to_status_id, seq, condition_json, approver_type, group_id)
VALUES
  (:v, :state_02, 'approve', :state_end, :status_done, 1,
   '{"field":"amount","operator":"<","value":100000}', 'group', :group_none),
  (:v, :state_02, 'approve', :state_03,  :status_wait_avp, 2,
   '{"field":"amount","operator":">=","value":100000}', 'group', :group_avp);
-- ✅ เกณฑ์เดียวจึงไม่มีช่องโหว่ปลายบน: ทุกยอดตั้งแต่ 100,000 ขึ้นไปวิ่งเข้า AVP เส้นเดียว""",
            "sql",
        ),
        h(2, "5.5 `workflow_part_display` ทับซ้อนกับกลไกของ prototype"),
        p(
            "หน้า `k2-document.html` ของ prototype คุมสิทธิ์แก้ไขรายส่วนด้วย `data-editrole` / `data-roleonly` / "
            "`.edit-only` ฝั่ง client เอง · engine มีกลไกเดียวกันให้อยู่แล้วผ่าน `workflow_part` + "
            "`workflow_part_display` ที่คืนมาใน `display[]` ของ `getPermissionEvents` "
            "(รูปแบบ `{partId, partName, stateId, partDisplayType, partSeq}`)"
        ),
        p(
            "**บันทึกเป็นข้อสังเกต ยังไม่ตัดสิน** ว่า SBPGI จะลงทะเบียน part ของทุกส่วนในหน้าเอกสารแล้วให้ FE อ่าน "
            "`display[]` แทนการ hardcode สิทธิ์ต่อ role หรือไม่ · ถ้าเลือกทางนี้จะกระทบ "
            "`LLDD-FE-Document-Detail` + role pack 5 ฉบับ ที่ปัจจุบันอ่าน `visibleSections`/`editableSections` "
            f"จาก API ของ SBPGI เอง · ทางเลือกเต็มอยู่ที่ `{DECISION_DOC}`"
        ),
        *pending_decision_blocks(
            "5.6 ข้อค้างตัดสินใจที่กระทบ engine (ยังไม่ตัดสิน)",
            [
                ["DP-1 · `reference_id`", "`doc_no` — ตกไป (บังคับออกเลขตั้งแต่ initialize และแก้ภายหลังไม่ได้)", "**เลือก surrogate id** (`compensation_documents.id` · ส่งเป็น string เพราะ `reference_id` เป็น varchar(255)) — ตรงกับที่ cooperation-request/inform-evaluate ทำจริงทุกจุด", "✅ ปิดแล้ว 2026-08-17 — ยืนยันตามระบบเดิม"],
                ["DP-2 · `workflow_transaction` ไม่มี PK/index", "ขอ sign-off จากทีม library ให้เพิ่ม PK + UNIQUE + index", "ไม่แตะตารางของ library · กันซ้ำและทำ index ที่ฝั่ง SBPGI", "ยังไม่ตัดสิน 🔴"],
                ["วงเงินอนุมัติเก็บที่ไหน", "`common_code` (SBPGI_APPROVE_LIMIT) ตามมติเดิม", "`workflow_route.condition_json` ตามความสามารถของ engine", "ยังไม่ตัดสิน · กติกาที่แน่นอนคือ **ห้ามเก็บสองที่**"],
                ["DP-5 · ใครเรียก email-lib ✅ ปิดแล้ว 2026-08-14", "engine ส่งเอง — ตกไป (ไม่มี `mailTo`/`param` ใน `triggerEvent`)", "**SBPGI ส่งเอง** โดยใช้เลข template จาก `workflow_route.email_id`", "ปิดแล้ว"],
                ["ผู้อนุมัติของ SBPGI", "`workflow_group_map` ผ่าน view (ต้องยืนยันว่า view where ด้วย user_id/group_id ได้)", "`addPreApprover` ระบุรายคน", "ยังไม่ตัดสิน"],
                ["`workflow_part_display` แทน `data-editrole`", "ลงทะเบียน part แล้วให้ FE อ่าน `display[]`", "คงกลไกของ SBPGI เอง (`visibleSections`/`editableSections`)", "ยังไม่ตัดสิน · กระทบ role pack 5 ฉบับ"],
            ],
        ),
    ]


def workflow_action_transition_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 Canonical Workflow Transition Matrix"),
        p("BE ต้องคำนวณ transition จาก currentSection, result และ totalCompensationAmount ภายใน transaction; FE ส่งเพียง result/comment และห้ามส่ง nextSection เอง"),
        table(["Current", "Result / condition", "statusCode", "nextSection", "Task effect"], [
            ["06", "ส่งเจ้าหน้าที่ SBP DSA ดำเนินการ (เส้นทางปกติ — ให้คำนวณยอดก่อน)", "08", "08", "close 06; open 08"],
            ["06", "ส่งหน่วยงานส่งเสริมธุรกิจ SBP (SDD GI · **เส้นทางข้ามขั้น 08**)", "01", "01", "close 06; open 01"],
            ["06", "เห็นควรไม่ชดเชย หรือ หยุดชดเชยประกันรายได้", "99", "null", "close 06; complete instance"],
            ["08", "คำนวณเงินชดเชยเรียบร้อย", "01", "01", "close 08; open 01"],
            ["01", "เห็นควรชดเชย", "02", "02", "close 01; open 02"],
            ["01", "เห็นควรไม่ชดเชย (SDD GI — **จบ flow ทันที** ไม่ตีกลับให้ 06)", "99", "null", "close 01; complete instance"],
            ["02", "เห็นควรชดเชย และ totalCompensationAmount >= 100,000 (มติ 2026-08-18)", "03", "03", "close 02; open 03"],
            ["02", "เห็นควรชดเชย และ totalCompensationAmount < 100,000 (มติ 2026-08-18)", "99", "null", "close 02; complete instance"],
            ["02", "เห็นควรไม่ชดเชย (SDD GI — **จบ flow ทันที** ไม่ตีกลับเป็นทอด ๆ)", "99", "null", "close 02; complete instance"],
            ["03", "เห็นควรชดเชย", "99", "null", "close 03; complete instance"],
            ["03", "เห็นควรไม่ชดเชย ⏳ *SDD GI ไม่ได้ระบุขั้น AVP — คงพฤติกรรมเดิม (ตีกลับ 06) รอ confirm*", "06", "06", "close 03; reopen 06"],
            ["ทุก section ที่รองรับ", "ส่งกลับ", "รหัส section ปลายทางตาม action option (08→06 · 01→06 · 02→01 · 03→02)", "section ปลายทาง", "close current; reopen target with new task id"],
        ]),
        h(2, "5.1b Auto-assign เจ้าของงานคนเดิม (SDD สไลด์ 46 · 48 · 64)"),
        p("สองปุ่มที่จบเอกสารเหมือนกันแต่พฤติกรรมหน้ารายการตรงข้ามกัน — BE ต้อง implement แยกกันให้ชัด ห้ามรวมเป็นเส้นเดียว"),
        table(["ปุ่มที่กดที่ขั้น 06", "เดือนที่กด", "เดือนถัดไป", "ผู้ดำเนินการ (เจ้าของงาน)"], [
            ["เห็นควรไม่ชดเชยรายได้", "ปิดเอกสาร (99) และ GET /sbpgi/document/tasks ของ 06 ต้อง **ไม่คืน** เอกสารนี้ในเดือนนั้น", "ระบบตั้งงานรอบเดือนถัดไปของร้านเดิมอัตโนมัติ", "**คนเดิม** ที่พิจารณาเอกสารรอบก่อนในขั้นเดียวกัน"],
            ["หยุดชดเชยประกันรายได้", "ปิดเอกสาร (99) แต่ GET /sbpgi/document/tasks ของ 06 **ต้องคืนทันที** พร้อม stoppedReopenable=true", "ไม่มีการตั้งงานอัตโนมัติ", "ฝ่าย SBP DSA (06)"],
            ["เคสต่อเนื่อง (ไม่ใช่ปุ่ม — เงื่อนไขของงานรอบถัดไป)", "ระบบสร้างงานให้เอง ไม่ต้องแจกงานด้วยมือ", "เหมือนกันทุกเดือนที่ยังต่อเนื่อง", "**คนเดิม** — เจ้าหน้าที่ SBP DSA รอบก่อนหน้า"],
        ]),
        p("**วิธี resolve เจ้าของงานคนเดิม** — ไม่มีคอลัมน์ assignee ในตารางของ SBPGI (ตาราง workflow_tasks ถูกตัดออกจากโครง 20 ตารางแล้ว) ผู้รับผิดชอบเป็นข้อมูลของ engine"),
        table(["ขั้น", "การทำงาน"], [
            ["1", "หาเอกสารรอบก่อนหน้าของร้านเดียวกัน (impacted_store_code เดิม · round_no/loop_no ก่อนหน้า)"],
            ["2", "อ่าน consideration_logs แถวล่าสุดของเอกสารนั้นที่ section_code = ขั้นที่จะมอบหมาย -> consider_by (คอลัมน์ผู้ดำเนินการ · อ้าง business_user ของระบบเดิม)"],
            ["3", "ผูกเป็นผู้รับผิดชอบผ่าน addPreApprover(versionId, referenceId, stateId, approver, seq) ของ @srm/glb-workflow"],
            ["4", "Fallback: รอบก่อนไม่เคยผ่านขั้นนั้น หรือพนักงานไม่อยู่ในกลุ่มแล้ว -> มอบหมายตาม group ของ auth-backend ตามปกติ"],
            ["5", "พนักงานลาออกยังต้องเปิด SR เพื่อแก้ชื่อผู้ดำเนินการ (ข้อจำกัดที่ SDD สไลด์ 48 ระบุ ไม่แก้ในเฟสนี้)"],
        ]),
        code("""-- resolve เจ้าของงานคนเดิมของขั้น :sectionCode จากเอกสารรอบก่อนของร้านเดียวกัน
SELECT cl.consider_by
FROM compensation_documents d
JOIN consideration_logs cl ON cl.doc_no = d.doc_no
WHERE d.impacted_store_code = :impactedStoreCode
  AND d.doc_no <> :currentDocNo
  AND cl.section_code = :sectionCode
ORDER BY d.round_no DESC, d.loop_no DESC, cl.action_datetime DESC
LIMIT 1;
-- ได้ค่าแล้วส่งเข้า addPreApprover(...) ตอนเปิดงานรอบใหม่ ห้าม INSERT sps_store.workflow_approver เอง
-- NULL -> fallback group ของ auth-backend""", "sql"),
        h(2, "5.2 Action Response Type"),
        table(["Field", "Type", "Required", "Rule"], [
            ["statusCode", "enum 06|08|01|02|03|99", "Yes", "ค่าหลัง commit; 99 = เสร็จสิ้น"],
            ["nextSection", "enum 06|08|01|02|03 | null", "Yes", "null เมื่อ workflow จบ"],
            ["message", "string", "Yes", "ข้อความผล mutation สำหรับแสดงผู้ใช้"],
        ]),
        *pending_decision_blocks(
            "5.3 ข้อค้างตัดสินใจที่กระทบ endpoint ของเอกสารนี้ (ยังไม่ตัดสิน)",
            [
                ["DP-7 ✅ ปิดแล้ว 2026-08-24 · แหล่งข้อมูลของ `GET /sbpgi/document/{docNo}/timeline`", "**เลือกข้อนี้ — อ่าน `consideration_logs` ของ SBPGI เป็น timeline เต็ม** (ผูก `transaction_id` ของ engine)", "อ่าน `getHistory()` / `sps_store.workflow_history` ของ engine แล้ว join — ตกไป · เป็นตารางส่วนขยาย (decision code · ไฟล์แนบ · ความเห็น ซึ่ง engine ไม่มี)", "ยังไม่ตัดสิน · กระทบทั้ง DDL ของ `consideration_logs` และรูปแบบ response"],
                ["DP-1 · `referenceId` ที่ส่งเข้า engine", "`doc_no` — ตกไป", "**เลือก surrogate id** (`compensation_documents.id` · ส่งเป็น string เพราะ `reference_id` เป็น varchar(255)) แบบที่ cooperation-request / inform-evaluate ทำจริงทุกจุด", "✅ ปิดแล้ว 2026-08-17 — ยืนยันตามระบบเดิม"],
                ["DP-2 · `sps_store.workflow_transaction` ไม่มี PK/index", "ขอ sign-off ให้ทีมเจ้าของ library เพิ่ม PK + UNIQUE + index", "กันซ้ำและทำ index ที่ฝั่ง SBPGI", "ยังไม่ตัดสิน 🔴 · ทุก action ต้อง seq-scan 19,283 แถว"],
                ["DP-5 ✅ ปิดแล้ว (แก้มติ 2026-08-14) — **workflow ให้เลข template · SBPGI เรียก lib ส่งเอง**", "`SBP/TSM-SRM-LLDD SBP EMAIL1.0.xlsx` — lib เสร็จแล้ว รับ `{emailId, mailTo, mailCc, param, fileAttach, userId}` · input ของ `triggerEvent` ไม่มี `mailTo`/`param` engine จึงเรียกแทนไม่ได้ · บรรทัด 'เรียก function ส่งเมล์จาก lib .....' ยังเป็น placeholder", "SBPGI อ่าน `workflow_route.email_id` → เรียก `sendEmail()` **นอก transaction** · ไม่มีตาราง `status_email_rules`", "ปิดแล้ว"],
            ],
        ),
    ]


def master_config_screen_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 Screen Boundary and Route Matrix"),
        p("หัวข้อนี้ประกอบด้วย **2 หน้าจออิสระ** แต่ละหน้ามี route, state, validation และ endpoint ของตนเอง "
          "ห้าม implement เป็น form/table เดียวที่สลับชนิดข้อมูลด้วยเงื่อนไขใน component · "
          "SCR-08 ผู้ปฏิบัติงาน และ SCR-10 สิทธิ์เมนู **ถูกตัดออก 2026-08-05** (ใช้ของระบบ SBP เดิม)"),
        table(["Screen", "Route / Component", "Primary model", "Main operations"], [
            ["SCR-09 ปัจจัยภายนอก", "/admin/external-factors / ExternalFactorPage", "ExternalFactor",
             "list, add, edit, delete, duplicate-code guard"],
            ["รายชื่อร้านคู่แข่ง (master แบรนด์)", "/admin/competitors / CompetitorPage", "Competitor",
             "list, add, edit, delete, duplicate-code guard, active toggle"],
            ["~~SCR-08 ผู้ปฏิบัติงาน~~ **ตัด 2026-08-05**", "ไม่มีหน้าจอใน SBPGI", "-",
             "ใช้ group + scope ของ auth-backend และ prepared approver ของ @srm/glb-workflow · "
             "จัดการที่หน้า /setting/manage-user-rights ของระบบเดิม — **ไม่มี field/endpoint ให้ implement**"],
            ["~~SCR-10 สิทธิ์เมนู~~ **ตัด 2026-08-05**", "ไม่มีหน้าจอใน SBPGI", "-",
             "สิทธิ์เมนูจัดการที่หน้า /setting/manage-user-rights ของระบบเดิม · "
             "SBPGI อ่านผ่าน header x-user-permissions จาก BFF — **ไม่มี field/endpoint ให้ implement**"],
        ]),
        h(2, "5.2 SCR-09 External Factor"),
        table(["Field", "Type", "Required / Rule", "UI behavior"], [
            ["factorCode", "string", "required; unique; immutable after create", "uppercase and trim before submit"],
            ["factorName", "string", "required; 1..200 chars", "Thai UTF-8 supported"],
            ["description", "string | null", "optional; max 1000 chars", "multiline editor"],
            ["active", "boolean", "required", "inactive rows remain visible under filter แต่ไม่ขึ้นใน dropdown ของหน้าเอกสาร"],
        ]),
        p("⚠️ **ไม่มีฟิลด์ `reason`** — ยกเลิกระบบ audit ของ master ทั้งหมด 2026-08-07 · "
          "mutation ไม่ต้องส่ง `reason` และไม่มี audit dialog ก่อน submit"),
        h(2, "5.3 Competitor Brand Master"),
        table(["Field", "Type", "Required / Rule", "UI behavior"], [
            ["competitorCode", "string(30)", "required; unique; immutable after create", "รหัส 01-11 ของ master เดิม · trim ก่อน submit"],
            ["nameTh", "string(200)", "required", "ชื่อแบรนด์ภาษาไทย — แสดงใน dropdown ร้านคู่แข่งของหน้าเอกสาร"],
            ["nameEn", "string(200)", "required", "ชื่อแบรนด์ภาษาอังกฤษ (ระบบเดิมเก็บทั้งสองภาษา)"],
            ["remark", "string(500) | null", "optional", "คอลัมน์ รายละเอียดเพิ่มเติม ของหน้า k2-competitors.html"],
            ["active", "boolean", "required", "ปิด active แทนการลบเมื่อถูกอ้างใน document_competitors แล้ว"],
        ]),
        p("**คนละระดับกับ `document_competitors`** ซึ่งเก็บคู่แข่ง *รายสาขา* ที่ import จาก ALLMAP "
          "พร้อมรหัสของตัวเอง (เช่น `4832`, `TD58_08`) — หน้านี้ดูแลเฉพาะ **แบรนด์**"),
        h(2, "5.6 Screen-level Acceptance"),
        bullets([
            "แต่ละหน้ามี route/component/state แยกและสามารถ test/release แยกกันได้",
            "mutation refresh เฉพาะ resource ที่เปลี่ยน — **ไม่ส่ง reason และไม่มี audit dialog** (ยกเลิก audit ของ master 2026-08-07)",
            "SCR-09 กัน factorCode ซ้ำทั้ง client response handling และ BE error (409 CODE_DUPLICATE)",
            "หน้าคู่แข่งกัน competitorCode ซ้ำ และบังคับ nameTh + nameEn ครบทั้งคู่",
            "ลบรายการที่ถูกอ้างในเอกสารแล้วต้องได้ 409 และ UI ต้องเสนอให้ปิด active แทน",
            "(SCR-08 และ SCR-10 ตัดออก 2026-08-05 — ไม่มีเกณฑ์ตรวจรับสำหรับสองหน้านี้)",
        ]),
    ]


def testing_delivery_blocks(topic: Topic) -> list[dict[str, Any]]:
    return [
        h(1, "1. Overview"),
        table(["รายการ", "รายละเอียด"], [
            ["Track", topic.track], ["Estimate", estimate_cell(topic)], ["Owner", topic.owner],
            target_repo_row(topic),
            ["Document type", "FE verification and release handover specification; not an application screen"],
            ["Objective", topic.objective],
        ]),
        h(1, "2. Delivery Scope"),
        bullets([
            "Regression suites for Dashboard, document lists/create/detail/actions, report, master/config, batch monitor and email template",
            "Contract verification against the endpoint schemas embedded in each feature LLDD",
            "Responsive and browser checks for supported viewports",
            "UAT defect triage, retest evidence and release handover",
            "No screen route, UI field table or synthetic API endpoint is created by this work item",
        ]),
        h(1, "3. Test Suite Matrix"),
        table(["Suite", "Coverage", "Entry condition", "Required evidence"], [
            ["FE-SMOKE", "app bootstrap, menus, dashboard, open list/detail", "deploy reachable and test user available", "timestamped run result and failed-step detail"],
            ["FE-DOC", "create, edit section, attachment, action, timeline and role views", "fixture documents for sections 06/08/01/02/03", "case ID, docNo, requestId and screenshots for failures"],
            ["FE-REPORT", "required status filter, 14 columns (SDD slide 60), totals, Excel parity", "known report fixture and expected aggregate", "query snapshot, row count, totals and exported checksum"],
            ["FE-MASTER", "SCR-09 ปัจจัยภายนอก + รายชื่อคู่แข่ง (SCR-08/10/11 และ email template ตัดออกแล้ว)", "admin role and reversible test data", "before/after values"],
            ["FE-BATCH", "job selection, editable params, locked params, run history", "job metadata/run fixtures", "request/response capture and UI state"],
            ["FE-RESP", "desktop 1440, tablet 768, mobile 390", "latest supported browsers", "page checklist with overflow/modal/navigation result"],
        ]),
        h(1, "4. Environment and Fixture Contract"),
        table(["Item", "Required content", "Control"], [
            ["Build identity", "commit SHA, build number, deploy timestamp", "freeze before regression"],
            ["API identity", "base URL and contract version", "no production credentials in evidence"],
            ["Role users", "one account per tested RBAC role/profile", "masked identifiers in shared evidence"],
            ["Document fixtures", "docNo per current section plus <100,000 (GM) and >=100,000 (AVP) cases per มติ 2026-08-18", "resettable or uniquely generated"],
            ["File fixtures", "valid type, >5MB, unsupported type, AV-blocked stub", "checksum recorded"],
            ["Job fixtures", "SUCCESS/FAILED/RUNNING/QUEUED histories", "read-only unless manual-run case"],
        ]),
        h(1, "5. Execution and Defect Flow"),
        table(["Step", "Action", "Exit rule"], [
            [1, "Record build/environment and run FE-SMOKE", "all smoke cases pass before broad regression"],
            [2, "Execute feature suites using fixed fixtures", "each case has pass/fail and evidence reference"],
            [3, "Log defects with severity, route, role, data key, steps and expected/actual", "defect is reproducible or explicitly closed as non-reproducible"],
            [4, "Retest fixes and run impacted regression", "no Critical/High open; Medium has accepted disposition"],
            [5, "Run responsive/browser matrix and release checklist", "all mandatory cells pass"],
            [6, "Produce handover summary", "build identity, known limitations, evidence index and rollback note complete"],
        ]),
        h(1, "6. Release Gate"),
        table(["Gate", "Pass condition"], [
            ["Functional", "All Critical/High feature and workflow cases pass"],
            ["Contract", "No request/response field mismatch against feature LLDD schema tables"],
            ["Visual", "No blocked action, clipped modal/table or unusable navigation at required viewports"],
            ["Security", "Unauthorized routes/actions fail closed; evidence contains no token/secret"],
            ["Data", "Report totals/export parity and action transitions reconcile with persisted result"],
            ["Handover", "Known limitations, rollback steps and test evidence index are complete"],
        ]),
        h(1, "7. Developer / QA Checklist"),
        table(["No", "Check"], [[i + 1, test] for i, test in enumerate(topic.tests)]),
    ]


def create_document_fs_iframe_blocks() -> list[dict[str, Any]]:
    return [
        p(
            "🔴 **หัวข้อ 5.1-5.6 ต่อจากนี้เป็นดีไซน์ *ก่อน* มติ 2026-08-06 — ยังไม่อยู่ในขอบเขตที่ประเมินไว้ 8 ชั่วโมง** · "
            "ของจริงที่ต้องทำคือ **iframe ของหน้าสร้างเอกสารระบบ FS ตรง ๆ + หมายเหตุ 4 ขั้นตอนใต้ iframe** (ดูโครงไฟล์ในหัวข้อ 8) · "
            "เก็บ SBP mirror form + FS bridge ไว้เป็น **ทางเลือกสำรอง** เผื่อ FS ไม่ยอมให้ฝัง iframe หรือ origin ไม่ผ่าน — "
            "ถ้าจะทำจริงต้องตั้งงบใหม่ ไม่ใช่ 6 ชม. · error code `FS_BRIDGE_*` ใน `LLDD-BE-API-Common-Contracts` ผูกกับทางเลือกสำรองนี้เท่านั้น"
        ),
        h(2, "5.1 Tab Structure *(ทางเลือกสำรอง — ไม่อยู่ในขอบเขตปัจจุบัน)*"),
        p("หน้า Create Document ต้องมี tab แยกสำหรับสร้างเอกสารจาก FS โดย UI หลักยังเป็น form ของ SBP Mall แต่มี hidden iframe ของ FS เป็น source/submit target จริง"),
        table(
            ["Tab", "Purpose", "Render behavior"],
            [
                ["สร้างเอกสารทั่วไป", "สร้างเอกสาร MANUAL/out-of-condition ผ่าน API ของ SBPGI", "ใช้ form ปกติและ submit POST /api/v1/sbpgi/document"],
                ["เอกสารจาก FS", "สร้างเอกสารโดยอ้าง field/form ของ FS เดิม", "โหลด FS iframe แบบ hidden แล้วสร้าง SBP form mirror ตาม field ที่พบใน iframe"],
            ],
        ),
        h(2, "5.2 FS iframe Integration Contract"),
        table(
            ["Item", "Required behavior", "Dev note"],
            [
                ["iframe element", "`<iframe id=\"fsCreateFrame\" hidden>` อยู่ในหน้า Create Document", "iframe ต้อง load ก่อน render field mirror; แสดง loading state ระหว่างรอ"],
                ["iframe source", "URL มาจาก config เช่น `fs.createDocumentUrl`", "ห้าม hardcode URL ใน component"],
                ["Access model", "ถ้า same-origin ให้ใช้ DOM adapter; ถ้า cross-origin ให้ใช้ SBP-FS Bridge Protocol v1 ด้านล่างเท่านั้น", "ตรวจ event.origin และ event.source ทุกข้อความ; protocol ไม่พร้อมให้ fail closed พร้อม code FS_BRIDGE_UNAVAILABLE"],
                ["Field discovery", "อ่าน input/select/textarea ใน FS form แล้ว map เป็น SBP field model", "ใช้ name/id/data-label/required/type/options จาก FS เป็น metadata"],
                ["Hidden source of truth", "FS iframe เป็น submit target จริง; SBP form เป็น mirror สำหรับ UX/validation", "ห้าม submit API โดยตรงแทน FS ใน tab นี้ เว้นแต่ FS callback ระบุให้ทำ"],
                ["Submit target", "เมื่อ user กดส่ง ให้ sync values ทั้งหมดเข้า iframe แล้ว trigger submit ของ FS form", "ป้องกัน double submit และรอ iframe load/callback หลัง submit"],
            ],
        ),
        h(2, "5.3 FS Field Mapping"),
        table(
            ["SBP mirror field", "FS iframe field source", "Mapping rule"],
            [
                ["impactedStoreCode", "input[name=impactedStoreCode] หรือ field ที่ FS ระบุเป็นร้านถูกกระทบ", "คง string 5 digits; leading zero ต้องไม่หาย"],
                ["newStoreCode", "input[name=newStoreCode] หรือ field ร้านเปิดใหม่ของ FS", "คง string 5 digits; validate ก่อน sync"],
                ["impactMonth", "month/date field ของ FS", "SBP ใช้ ค.ศ. · sync เป็น format ที่ FS field ต้องการ"],
                ["statementPeriod", "period field ของ FS", "required สำหรับ FS tab"],
                ["roundNo", "round/sequence field ของ FS", "default 1 ถ้า FS field ว่างและ metadata อนุญาต"],
                ["reason/remark", "textarea/input remark ของ FS", "trim ก่อน sync; preserve Thai text"],
                ["dynamicFields[]", "field เพิ่มเติมที่พบใน FS form", "render ตาม type/options/required จาก FS และเก็บ mapping ไว้ใน form state"],
            ],
        ),
        h(2, "5.4 Change and Submit Flow"),
        table(
            ["Step", "FE behavior", "Failure handling"],
            [
                ["1. Open FS tab", "โหลด hidden iframe จาก config และรอ iframe load", "timeout แสดง error พร้อม retry; ไม่ render empty form"],
                ["2. Discover fields", "อ่าน field metadata จาก iframe form แล้วสร้าง SBP mirror form", "field required แต่ไม่รู้ label ให้ใช้ name/id เป็น fallback"],
                ["3. User changes value", "update SBP state แล้ว sync ค่าเข้า iframe field ทันที", "ถ้า sync field ไม่พบ ให้ mark fieldMappingError และห้าม submit"],
                ["4. Client validate", "validate required/type/range ตาม metadata จาก FS และ validation กลางของ SBP", "แสดง inline error ใน SBP form"],
                ["5. Submit", "sync all values อีกครั้ง, dispatch input/change event ใน iframe, แล้ว submit FS form", "disable submit จนกว่า iframe submit result/callback กลับมา"],
                ["6. Handle result", "รับ FS_SUBMIT_RESULT ที่ requestId ตรงกับคำขอ; success navigate ไป detail เมื่อมี docNo", "timeout หรือ schema ไม่ถูกต้องให้ปลด submitting state และแสดง error ที่ retry ได้; ห้ามเดาสถานะสำเร็จ"],
            ],
        ),
        h(2, "5.5 SBP-FS Bridge Protocol v1"),
        table(["Envelope field", "Type", "Required", "Rule"], [
            ["protocolVersion", "literal `1.0`", "Yes", "reject version อื่นด้วย FS_PROTOCOL_VERSION_UNSUPPORTED"],
            ["type", "message enum", "Yes", "FS_FORM_READY | SBP_FIELD_DISCOVERY_REQUEST | FS_FIELD_SCHEMA | SBP_SET_VALUES | SBP_SUBMIT | FS_SUBMIT_RESULT | FS_ERROR"],
            ["requestId", "UUID string", "Yes", "สร้างใหม่ต่อ request และใช้ correlate response"],
            ["correlationId", "UUID string | null", "Response only", "ต้องเท่ากับ requestId ของ message ที่ตอบ"],
            ["timestamp", "ISO-8601 string", "Yes", "ใช้ตรวจ stale message; ไม่ใช้เป็น authorization"],
            ["source", "literal `SBP` | `FS`", "Yes", "ต้องสอดคล้องกับ window ฝั่งผู้ส่ง"],
            ["payload", "object", "Yes", "validate ตาม type ก่อนใช้"],
        ]),
        h(3, "Message payload schemas"),
        table(["Message type", "Payload fields", "Response / rule"], [
            ["FS_FORM_READY", "formId:string, capabilities:string[], schemaVersion:string", "SBP ส่ง SBP_FIELD_DISCOVERY_REQUEST เมื่อ capabilities มี FIELD_SCHEMA"],
            ["SBP_FIELD_DISCOVERY_REQUEST", "formId:string", "FS ตอบ FS_FIELD_SCHEMA ด้วย correlationId"],
            ["FS_FIELD_SCHEMA", "formId:string, fields:FsFieldDescriptor[]", "descriptor ทุกตัวต้องผ่าน schema ด้านล่าง"],
            ["SBP_SET_VALUES", "formId:string, values:Record<string,string|number|boolean|null>", "FS validate key ที่รู้จักและตอบ FS_ERROR เมื่อ map ไม่ได้"],
            ["SBP_SUBMIT", "formId:string, values:Record<...>, clientReference:string", "idempotent ต่อ requestId; ห้าม submit ซ้ำ"],
            ["FS_SUBMIT_RESULT", "success:boolean, fsReference:string|null, docNo:string|null, fieldErrors:FieldError[]", "success=true ต้องมี fsReference; docNo เป็น optional"],
            ["FS_ERROR", "code:string, message:string, retryable:boolean, field:string|null", "FE แสดง message และเปิด retry เฉพาะ retryable=true"],
        ]),
        table(["FsFieldDescriptor field", "Type", "Required", "Constraint"], [
            ["name", "string", "Yes", "unique within form; key used by values map"],
            ["label", "string", "Yes", "UTF-8 display label"],
            ["type", "enum text|number|date|month|select|radio|checkbox|textarea|hidden", "Yes", "unknown type is rejected"],
            ["required", "boolean", "Yes", "drives client validation"],
            ["readOnly", "boolean", "Yes", "read-only field is never overwritten by SBP"],
            ["value", "string|number|boolean|null", "Yes", "initial value"],
            ["options", "array<{value:string,label:string}>|null", "For select/radio", "selected value must exist in options"],
            ["constraints", "{min,max,minLength,maxLength,pattern}|null", "No", "FE and FS both validate"],
        ]),
        h(3, "Handshake, security and timeout"),
        table(["Phase", "Required behavior", "Timeout / failure"], [
            ["Origin setup", "allowlist มาจาก config และ targetOrigin ต้องเป็น origin เฉพาะ ห้ามใช้ `*`", "origin ไม่ตรงให้ ignore และ security log โดยไม่ log payload"],
            ["Ready", "รอ FS_FORM_READY จาก iframe window เดียวกัน", "10s -> FS_BRIDGE_TIMEOUT; retry reload iframe ได้ 1 ครั้ง"],
            ["Schema", "ส่ง discovery และ validate FS_FIELD_SCHEMA", "5s หรือ schema invalid -> FS_FIELD_SCHEMA_INVALID"],
            ["Value sync", "ส่ง SBP_SET_VALUES พร้อม requestId ใหม่และ debounce 150ms", "FS_ERROR ผูก correlationId กลับ field"],
            ["Submit", "ส่ง SBP_SUBMIT หนึ่งครั้งและ disable submit", "30s -> FS_SUBMIT_TIMEOUT; user retry สร้าง requestId ใหม่"],
            ["Result", "ยอมรับเฉพาะ correlationId ที่ pending และ source/origin ถูกต้อง", "late/duplicate result ถูก ignore แบบ idempotent"],
        ]),
        h(3, "Protocol example"),
        payload("FS_FIELD_SCHEMA", api_json({
            "protocolVersion": "1.0", "type": "FS_FIELD_SCHEMA", "requestId": "6f6c8cf0-7df1-4a1a-9e7f-4d953667a824",
            "correlationId": "a8e88f2a-e83b-47ce-99f5-fdcad1876095", "timestamp": "2026-07-22T10:15:00+07:00", "source": "FS",
            "payload": {"formId": "income-guarantee-create", "fields": [{"name": "impactedStoreCode", "label": "รหัสร้านถูกกระทบ", "type": "text", "required": True, "readOnly": False, "value": "00788", "options": None, "constraints": {"pattern": "^[0-9]{5}$"}}]}
        })),
        h(2, "5.6 Acceptance Criteria for FS Tab"),
        bullets([
            "tab เอกสารจาก FS ต้องโหลด hidden iframe และสร้าง mirror form จาก field metadata ได้",
            "เมื่อ user เปลี่ยนค่าใน SBP form ค่าเดียวกันต้องถูก sync เข้า iframe field ที่ map ไว้",
            "กด submit ต้อง sync ทุก field อีกครั้งก่อน submit FS iframe form",
            "field ที่ required ใน FS ต้องแสดง required ใน SBP mirror form",
            "store code 5 หลักต้องไม่สูญเสีย leading zero ระหว่าง SBP form -> iframe",
            "cross-origin ต้อง handshake, discover schema, sync, submit และรับผลผ่าน SBP-FS Bridge Protocol v1 ครบ",
            "message ที่ origin/source/version/correlationId ไม่ถูกต้องต้องถูก ignore หรือ reject แบบ fail closed",
            "timeout และ FS_ERROR ต้องออกจาก loading/submitting state และ retry ได้ตาม retryable flag",
        ]),
    ]


def is_document_detail_role_doc(file_key: str) -> bool:
    return file_key.startswith("FE/LLDD-FE-Document-Detail-Role")


def is_batch_monitor_doc(file_key: str) -> bool:
    return False  # ตัดเอกสาร Batch Monitor ออกจากชุดส่งมอบ 2026-08-06


def role_profile_code(profile: dict[str, Any]) -> str:
    return f"P-{profile['code']}"


def implementation_detail_blocks(topic: Topic) -> list[dict[str, Any]]:
    if is_document_detail_role_doc(topic.file):
        return role_doc_implementation_blocks(topic)
    if is_batch_monitor_doc(topic.file):
        return batch_monitor_implementation_blocks()
    if "/Jobs/" in topic.file:
        return job_implementation_blocks(topic)
    if topic.track == "FE":
        return fe_implementation_blocks(topic)
    return be_implementation_blocks(topic)


def topic_io_contract_blocks(topic: Topic) -> list[dict[str, Any]]:
    if is_batch_monitor_doc(topic.file):
        return [
            h(2, "4.93 Input / Progress / Output Contract"),
            table(
                ["Stage", "Contract for implementation"],
                [
                    ["Input", "Selected jobNo, editable parameter form values, run-history filters, and current operator permission."],
                    ["Progress", "Load job list, select job, render params/history tabs, validate changed params, save with audit, refresh history after manual run."],
                    ["Output", "Updated job parameter snapshot, visible run-history status, validation messages, and audit reference for saved changes."],
                ],
            ),
        ]
    if "/Jobs/" in topic.file:
        job_no = topic.file.split("LLDD-BE-Job-", 1)[1].split("-", 1)[0]
        legacy = LEGACY_JOB_SOURCES.get(job_no, {})
        rows = [
            ["Input", legacy.get("input", "Job parameters, source files/tables, schedule period, and service-token context defined by this job.")],
            ["Progress", legacy.get("progress", "Create run record, load input, process in chunks, update checkpoint/summary, and expose current state through run history.")],
            ["Output", legacy.get("output", "Target rows/files/interfaces updated according to DB mapping; run history stores status, counts, output reference, and error detail.")],
        ]
        return [
            h(2, "5.9 Input / Progress / Output Contract"),
            table(["Stage", "Contract for implementation"], rows),
        ]

    request_sources = [f"{api.method} {api.path}" for api in topic.apis[:3]]
    db_outputs = [row[0] for row in topic.db_tables if str(row[1]).upper() in {"W", "R/W"}][:3]
    flow_summary = "; ".join(topic.flow[:4]) if topic.flow else "Validate request, apply business rule, persist or render result, and return normalized status."
    rows = [
        # เอกสารที่ไม่มี endpoint ของตัวเอง (สัญญา/งานภายใน) ต้องไม่ได้ประโยคกลาง ๆ ภาษาอังกฤษของหน้าจอ
        ["Input", "; ".join(request_sources) if request_sources
         else "ไม่มี endpoint ของตัวเอง — input คือ request ที่เอกสารอื่นส่งเข้ามา พร้อม user context จาก BFF header (ดู 5.1) และค่ากำหนดกลางที่อ่านจากระบบเดิม"],
        ["Progress", flow_summary],
        ["Output", "; ".join(db_outputs) if db_outputs
         else "ไม่มีตารางที่เอกสารนี้เขียนเอง — output คือ response ตาม envelope กลาง `{success, data}` และร่องรอยที่ตรวจย้อนได้ (log / consideration_logs / workflow_history ของ engine)"],
    ]
    return [
        h(2, "5.9 Input / Progress / Output Contract"),
        table(["Stage", "Contract for implementation"], rows),
    ]


def legacy_job_source_blocks(job_no: str) -> list[dict[str, Any]]:
    legacy = LEGACY_JOB_SOURCES.get(job_no)
    if not legacy:
        return []
    return [
        h(2, "5.92 Legacy Java Source Reference"),
        table(
            ["Legacy file", "Line range", "Responsibility to carry forward"],
            legacy["sources"],
        ),
        p("Line ranges refer to the legacy Java implementation under /Users/bank_mac/gosoft/java/SBP/fcsJar. Use these ranges to preserve business behavior while implementing the target Node job."),
    ]


def node_job_skeleton(job_no: str, topic: Topic) -> str:
    function_name = re.sub(r"[^A-Za-z0-9]+", " ", topic.title).title().replace(" ", "")
    function_name = re.sub(r"^\d+", "", function_name)
    spec = JOB_IMPLEMENTATION_SPECS[job_no]
    steps = spec["steps"].split("|")
    step_lines = []
    for index, step in enumerate(steps, start=1):
        previous = "undefined" if index == 1 else f"step{index - 1}"
        step_lines.append(f"    const step{index} = await services.{step}(ctx, {previous});")
    return f"""export async function run{function_name}(ctx, services) {{
  const run = await services.jobRuns.acquire({{
    jobNo: "{job_no}", period: ctx.period, triggeredBy: ctx.triggeredBy
  }});

  try {{
    ctx = {{ ...ctx, runId: run.id, repository: services.{spec['repository']} }};
{chr(10).join(step_lines)}
    const result = step{len(steps)};
    await services.jobRuns.finish(run.id, "SUCCESS", result.metrics);
    return {{ runId: run.id, status: "SUCCESS", ...result }};
  }} catch (error) {{
    await services.jobRuns.finish(run.id, "FAILED", {{
      errorCode: error.code ?? "JOB_FAILED",
      errorMessage: error.message
    }});
    throw error;
  }}
}}"""


def batch_monitor_implementation_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "4.90 Developer Implementation Scope"),
        table(
            ["Area", "Implementation detail", "Definition of done"],
            [
                ["Page shell", "แสดงรายการ job เพื่อเลือก job ที่ต้องดูรายละเอียด และเปิด detail panel ของ job ที่เลือก", "เลือก job แล้ว panel แสดงชื่อ job, คำอธิบาย, tags และ tab default เป็นแบบฟอร์มพารามิเตอร์"],
                ["Tab set", "สร้างเฉพาะ 2 tab ที่ต้องใช้งานจริง: แบบฟอร์มพารามิเตอร์ และประวัติการรัน", "ไม่มี requirement ให้ dev ทำ tab Flowchart การทำงาน หรือ Database ที่ใช้ใน scope นี้"],
                ["Form parameter", "render field ตาม metadata ของแต่ละ job โดยแยก editable/read-only ให้ชัด", "field ที่ read-only แก้ไม่ได้, editable field validate ก่อนบันทึก และส่งเฉพาะค่าที่แก้ได้"],
                ["History run", "แสดงประวัติ run ของ job ที่เลือก พร้อม status, เวลาเริ่ม, เวลาจบ, duration, trigger, ผู้สั่ง และ summary/error", "sort ล่าสุดก่อน, filter status ได้ถ้ามี control บนหน้า, เปิด row เพื่อดู log/detail ได้"],
                ["Reference material", "flowchart การทำงานและฐานข้อมูลที่ใช้เป็นเอกสารอ้างอิงสำหรับ dev เท่านั้น", "ไม่ถูกนับเป็น UI deliverable ของ Batch Monitor และไม่ต้องทำรายละเอียดเทคนิค backend/storage ในเอกสารนี้"],
            ],
        ),
        h(2, "4.91 Two-tab Behavior"),
        table(
            ["Tab", "Visible content", "Editable / Action rule"],
            [
                ["แบบฟอร์มพารามิเตอร์", "ข้อมูลรอบการรัน, cron/schedule, source/target path, file prefix, encoding, batch size, manual run period และ note/runbook ที่เกี่ยวข้อง", "แก้ได้เฉพาะ field ที่ metadata ระบุ editable; save button disabled จนกว่าจะมีการแก้ไขและ validate ผ่าน"],
                ["ประวัติการรัน", "ตาราง run history, status badge, start/end time, duration, trigger type, operator, result summary และปุ่มดูรายละเอียด log", "เป็น read-only; action หลักคือเปิดรายละเอียด run/log จาก row ที่เลือก"],
            ],
        ),
        h(2, "4.92 UI States and Error Handling"),
        table(
            ["State", "Trigger", "UI behavior"],
            [
                ["No selected job", "เปิดหน้าครั้งแรกก่อนเลือก job", "แสดง placeholder ให้เลือก job จากรายการ และไม่แสดง form/history ของ job ใด"],
                ["Loading selected job", "เลือก job หรือ refresh detail", "แสดง loading placeholder เฉพาะ detail panel โดยไม่ clear รายการ job ด้านบน"],
                ["Dirty parameter form", "แก้ editable field แล้ว", "แสดง unsaved indicator, เปิดปุ่มบันทึก และเตือนเมื่อเปลี่ยน job/tab ออกจาก form หากยังไม่บันทึก"],
                ["Validation error", "required/format/range ไม่ผ่าน", "แสดง inline error ใต้ field และไม่บันทึกค่า"],
                ["Empty history", "job ยังไม่มี run history ใน filter ปัจจุบัน", "แสดง empty state ใน tab ประวัติการรัน โดยไม่ซ่อน tab"],
            ],
        ),
    ]


FE_COMPONENT_DETAILS: dict[str, list[tuple[str, str]]] = {
    "FE/LLDD-FE-Integration-Contracts": [
        ("สร้าง shared API client ตัวเดียวสำหรับ base URL, trace header, timeout และ response envelope", "ทุก feature import client กลางและไม่มี axios/fetch instance แยก"),
        ("อ่าน access token จาก platform auth store, แนบ Bearer token และทำ refresh แบบ single-flight", "401 พร้อมกัน refresh ครั้งเดียว, replay request เดิม และไม่สร้างหน้า Login ใหม่"),
        ("แปลง HTTP/Axios failure เป็น ApiError พร้อม code, message, fieldErrors และ traceId โดยไม่แก้ข้อความจาก BE", "validation banner/inline error แสดงข้อความและ traceId จาก response ได้ครบ"),
        ("ให้ formatter กลางสำหรับวันที่ (ค.ศ.), เดือน, เงิน, percent และ docNo โดยไม่เปลี่ยนค่าที่ส่ง API", "payload และ UI ใช้ ค.ศ. เป็นค่าเริ่มต้น (buddhistEra=false); รูปแบบเงิน/docNo ตรงกันทุกหน้า"),
        ("กำหนด PageResponse<T> และ state loading/empty/error/retry สำหรับ list ทุกชนิด", "DataTable/Pager รักษา page/filter เดิมและไม่มี list shape เฉพาะหน้า"),
        ("กำหนด typed action request/response และ consume statusCode/nextSection ที่ BE คำนวณ", "FE ส่งเฉพาะ result/comment และไม่มี client-side workflow routing"),
        ("สร้าง sidebar, route guard, visibleSections, editableSections และ actionOptions จาก platform/menu API", "ไม่ hardcode RBAC role เป็นสิทธิ์เมนูหรือ section ที่แก้ไขได้"),
    ],
    "FE/LLDD-FE-Foundation": [
        ("ประกอบ app bootstrap, environment validation, providers และ error boundary โดยไม่สร้าง business screen", "เปิด application shell ได้เมื่อ config ครบ และ fail-fast พร้อมข้อความเมื่อ config ขาด"),
        ("ลงทะเบียน route/module ของ SBP Mall และเชื่อม route guard กับ menuCode จาก API", "ทุก route เข้าได้เฉพาะเมื่อ menu contract อนุญาตและ unknown route ไป not-found"),
        ("จัดโครงสร้าง DTO, API adapter และ query key กลางให้ response typing ตรงกับ contract", "TypeScript build ผ่านและ feature ไม่ cast unknown response แบบ ad hoc"),
        ("รวม status/menu/action constants และ label resolver โดยให้ API dictionary เป็น source of truth", "unknown code แสดง fallback ที่ trace ได้และไม่เพิ่มสถานะเองใน component"),
        ("สร้าง fixture/mock ให้ใช้ schema เดียวกับ response จริง รวม success, empty และ error", "สลับ mock/real adapter ได้โดยไม่แก้ component props หรือ table mapping"),
        ("กำหนด token และ shared UI สำหรับ table, form, modal, badge และ responsive breakpoints", "shared component ใช้งานได้บน desktop/tablet/mobile โดยข้อความและ control ไม่ล้น"),
    ],
    "FE/LLDD-FE-Document-Lists": [
        ("โหลดงานของผู้ใช้จาก /sbpgi/document/tasks และ map 9 คอลัมน์หลักพร้อม task owner/status", "waiting list แสดง 9 คอลัมน์ตรง type และรักษา leading zero ของรหัสร้าน"),
        ("ค้นหาเอกสารจาก /sbpgi/document โดยบังคับปีและแสดงเอกสารที่เกี่ยวข้องตาม permission", "ไม่ call API เมื่อไม่มีปี และ empty result ไม่แสดงข้อมูลจาก query ก่อนหน้า"),
        ("serialize docNo/year/status/store filters ลง query state และ restore เมื่อย้อนกลับจาก detail", "Search/Clear/refresh ให้ผลซ้ำได้และ pagination ใช้ filter ชุดเดียวกัน"),
        ("ควบคุม page/size/sort และ row navigation โดยใช้ docNo เป็น stable key · เพิ่มคอลัมน์ checkbox แรกสุดสำหรับเลือกหลายเอกสาร (SDD GI สไลด์ 48) — checkbox ต้อง stopPropagation ไม่ให้ทริกเกอร์ row navigation และ \"เลือกทั้งหมด\" ครอบเฉพาะแถวที่แสดงในหน้านั้น ไม่ใช่ทั้งชุดผลลัพธ์", "เปลี่ยนหน้าไม่ reset filter และเปิด detail ของ row ที่เลือกถูกเลขเอกสาร · เลือกหลายรายการแล้วกด \"ดำเนินการที่เลือก\" ต้องเปิด popup ยืนยันพร้อมรายการเลขที่เอกสารก่อนส่ง และเคลียร์การเลือกหลังส่งสำเร็จ"),
        ("คำนวณ presentation flag จาก salesDataDays < 60 โดยไม่ใช้ waitingDays แทน และ render เอกสารที่จบด้วยผลปฏิเสธทั้ง 2 แบบ — หยุดชดเชยประกันรายได้ (stoppedReopenable=true) และ เห็นควรไม่ชดเชยรายได้ (notCompensated=true) — เฉพาะบทบาท section 06", "แถวผิดปกติเป็นสีแดงพร้อม accessible label เฉพาะเมื่อยอดขายไม่ครบ 60 วัน · บทบาท 06 เห็น 3 กลุ่มในหน้าเดียว (มติ 2026-08-24): (1) รอฝ่าย SBP DSA ดำเนินการ (2) เสร็จสิ้นดำเนินการ + ชิป หยุดชดเชยฯ (3) เสร็จสิ้นดำเนินการ + ชิป เห็นควรไม่ชดเชยฯ · บทบาท 08/01/02/03 ต้องไม่เห็นกลุ่ม (2) และ (3) · ชิปทั้งสองเป็นผลการพิจารณาสุดท้าย ไม่ใช่สถานะที่ 7/8 — สถานะจริงยังเป็น เสร็จสิ้นดำเนินการ ตามชุด 6 ค่า และมีตัวกรองแยก 2 ตัวจาก dropdown สถานะ · กลุ่ม (2) คลิกแล้วเปิดเอกสารในโหมดเปิดพิจารณาใหม่ · หมายเหตุที่มา: กลุ่ม (3) กว้างกว่าตัวอักษรของ SDD สไลด์ 46/64 ที่ให้แสดงเฉพาะรอบเดือนถัดไปในหน้างานค้างของ เจ้าหน้าที่ SBP DSA"),
    ],
    "FE/LLDD-FE-Create-Document": [
        ("หน้าเดียว ไม่มี state ของฟอร์ม — ถือแค่ config URL ของ FS และสถานะโหลด iframe", "ไม่มี draft/unsaved-change guard เพราะไม่มีฟอร์มฝั่ง SBP"),
        ("render กรอบ iframe ของหน้าสร้างเอกสารระบบ FS (สไตล์ `.fs-frame` เดียวกับ k2-document)", "iframe load/error/timeout มี state ชัดเจนและมีข้อความบอกผู้ใช้เมื่อโหลดไม่ขึ้น"),
        ("render การ์ดหมายเหตุ 4 ขั้นตอน **verbatim จากหน้าจอ K2 เดิม** ใต้ iframe (นอกกรอบ)", "ข้อความตรงต้นฉบับทุกตัวอักษร ห้าม paraphrase"),
        ("ลิงก์กลับไปหน้ารายการเอกสารและหน้าเอกสารเมื่อ SBP Statement ส่งข้อมูลกลับแล้ว (~1 วัน)", "ผู้ใช้เข้าใจว่าเอกสารจะมาเองไม่ต้องกดสร้างซ้ำ"),
        ("route guard/เมนูของหน้านี้มาจาก `GET /menus` ของระบบเดิม ไม่ hardcode", "ผู้ใช้ที่ไม่มีสิทธิ์เมนูเข้าหน้านี้ไม่ได้"),
        ("⚠️ *(ทางเลือกสำรอง — ไม่อยู่ในขอบเขต 8 ชม.)* SBP mirror form + FS bridge ตามหัวข้อ 5.1-5.6", "ใช้เมื่อ FS ไม่ยอมให้ฝัง iframe หรือ origin ไม่ผ่าน — ต้องตั้งงบใหม่ก่อนทำ"),
    ],
    "FE/LLDD-FE-Document-Detail": [
        ("โหลดและแสดง docNo, status, impacted store, impact month และ current operator จาก aggregate response", "header refresh หลัง mutation และ status badge resolve จาก statusCode"),
        ("render new-store, competitor และ factor collections ด้วย row key และ typed value mapping", "ข้อมูลอ่าน/แก้/ลบตรง editableSections และ percent รวมตรวจได้ 100"),
        ("ใช้ visibleSections/editableSections/canAction เป็น source of truth สำหรับ DOM และ focusable controls", "section ที่ซ่อนไม่อยู่ใน DOM และ read-only section ไม่มี mutation control"),
        ("สร้าง action radio/comment/confirm จาก actionOptions และ requireComment ที่ API ส่งมา", "ไม่ hardcode route/nextSection และ block submit เมื่อ result/comment ไม่ครบ"),
        ("รวม consideration history, workflow timeline และ invalidate หลัง save/upload/action", "ลำดับเวลาใหม่สุดถูกต้องและข้อมูลหลัง submit ไม่ค้างจาก cache เดิม"),
        ("upload ด้วย allowlist/5MB/scan state และ download ผ่าน authorized BE stream", "BLOCKED/PENDING ดาวน์โหลดไม่ได้และ success แสดงชื่อ/ขนาดไฟล์จาก metadata"),
        ("เปิด ALLMAP/map และ sales detail ด้วย doc/store context โดยไม่ expose credential", "link/adapter ส่ง identifier ถูกตัวและ failure กลับสู่หน้า detail ได้"),
    ],
    "FE/LLDD-FE-Testing-Delivery": [
        ("จัด regression matrix ครบ route, role profile, happy path และ typed error path", "ทุก route หลักมีผลทดสอบพร้อม browser/viewport/evidence"),
        ("ตรวจ desktop/tablet/mobile สำหรับ table, modal, form, chart และ navigation", "ไม่มี overflow/overlap และ control สำคัญใช้งานได้ทุก viewport ที่กำหนด"),
        ("เทียบ request/response fixture กับ API field schema และ error catalog", "schema mismatch เป็นศูนย์และไม่มี toy payload ที่ขาด required field"),
        ("ผูก defectId กับ test case, retest build ที่แก้ และเก็บ before/after evidence", "Critical/High defect ปิดพร้อมหลักฐานและ regression รอบเกี่ยวข้องผ่าน"),
        ("ประเมิน build/typecheck, secret scan, contract parity และ unresolved blocker ก่อน release", "release gate fail-closed เมื่อข้อบังคับข้อใดไม่ผ่าน"),
        ("จัด delivery note, test summary, known limitations และ reproducible verification commands", "ผู้รับมอบตรวจซ้ำได้โดยไม่มี token/secret หรือไฟล์ QA ชั่วคราว"),
    ],
    "FE/LLDD-FE-Report": [
        ("จัดการ filter 7 ตัวตาม SDD สไลด์ 60 (status, impacted/new store code, store type, period statement, region, result) พร้อม dependency validation", "status required, คู่รหัสร้านต้องมาด้วยกัน, period statement บังคับเมื่อสถานะ = เสร็จสิ้นดำเนินการ และช่วง from-to ตรวจผ่านก่อนค้นหา/Export"),
        ("map response เป็น summary line และตาราง 14 คอลัมน์ (SDD สไลด์ 60) ด้วย formatter กลาง", "คอลัมน์/ยอดรวม/วันที่ (ค.ศ.)/leading zero ตรง response และข้อมูลยอดขายผิดปกติใช้ salesDataDays"),
        ("ส่ง filter snapshot ล่าสุดไป export endpoint และจัดการ download/error state · SDD GI สไลด์ 62 กำหนดปุ่มออกผล 3 ตัว — Preview Report (ดูตัวอย่างก่อนออกไฟล์) · Export Excel (ทีมบัญชีเทียบ SAP) · Export CSV to Batch (ส่งเข้าคิว batch ประมวลผลต่อ)", "ทั้งสามปุ่มใช้เงื่อนไขค้นหาชุดเดียวกับตารางผลลัพธ์ และชื่อไฟล์/content type ตรง response (.xlsx สำหรับ Excel · .csv สำหรับ CSV to Batch)"),
        ("รองรับ fixture สำหรับ 0 แถว, หลาย region/type, เกิน threshold และยอดขายไม่ครบ 60 วัน", "sample verification ครอบคลุม table/export parity 14 คอลัมน์ โดยไม่ฝังข้อมูลทดสอบใน production"),
    ],
    "FE/LLDD-FE-Master-Data": [
        ("จัดการปัจจัยภายนอก (SCR-09) CRUD ครบ — DELETE เฉพาะรายการที่ยังไม่ถูกอ้างในเอกสาร",
         "factorCode ซ้ำไม่ได้ (409 CODE_DUPLICATE) · ลบรายการที่ถูกอ้างแล้วต้องได้ 409 · deleted row หายหลัง refresh"),
        ("จัดการ master แบรนด์คู่แข่ง (รหัส 01-11) พร้อม nameTh/nameEn/remark",
         "nameTh และ nameEn บังคับทั้งคู่ · competitorCode ซ้ำไม่ได้ · แถวใหม่ขึ้นใน dropdown ร้านคู่แข่งของหน้าเอกสาร"),
        ("ใช้ modal mode ADD/EDIT/DELETE แยก initial values, validation และ confirm copy",
         "เปลี่ยน mode ไม่ทิ้ง stale field และปุ่ม submit กัน double request"),
        ("toggle active แทนการลบเมื่อรายการถูกอ้างในเอกสารแล้ว",
         "แถวที่ปิด active ต้องไม่ขึ้นใน dropdown ของหน้าเอกสาร แต่เอกสารเก่ายังแสดงชื่อเดิมได้"),
    ],
}


def fe_implementation_blocks(topic: Topic) -> list[dict[str, Any]]:
    feature_name = topic.title.replace("LLDD FE - ", "")
    api_rows = []
    for index, spec in enumerate(topic.apis):
        endpoint_path = spec.path.split("?", 1)[0]
        matching_actions = [
            f"{action} ({trigger})"
            for action, trigger, service, _ in topic.actions
            if endpoint_path in service or service in spec.path
        ]
        if spec.buttons:
            invoked_by = ", ".join(spec.buttons)
        elif matching_actions:
            invoked_by = "; ".join(matching_actions)
        elif topic.actions:
            action, trigger, _, _ = topic.actions[index % len(topic.actions)]
            invoked_by = f"{action} ({trigger})"
        else:
            invoked_by = "contract verification step ที่ระบุใน test matrix ของเอกสารนี้"
        api_rows.append([f"{spec.method} {spec.path}", spec.purpose, invoked_by])
    if not api_rows:
        api_rows = [["No direct endpoint", "งานนี้ไม่สร้าง endpoint จำลอง", "ใช้ผลทดสอบ/บริการจาก feature ที่ตรวจ"]]
    component_rows = []
    component_details = FE_COMPONENT_DETAILS.get(topic.file)
    if component_details and len(component_details) != len(topic.scope):
        raise ValueError(f"FE component detail count mismatch for {topic.file}")
    for index, scope_item in enumerate(topic.scope, start=1):
        if component_details:
            responsibility, done_rule = component_details[index - 1]
        else:
            field_hint = topic.fields[(index - 1) % len(topic.fields)][0] if topic.fields else "localState"
            responsibility = f"จัดการ {scope_item} ใน {feature_name} ด้วย typed state `{field_hint}` ตาม field/action contract ของเอกสารนี้"
            done_rule = topic.acceptance[(index - 1) % len(topic.acceptance)] if topic.acceptance else f"{scope_item} ผ่าน component และ interaction test"
        component_rows.append([f"C{index:02d}", scope_item, responsibility, done_rule])
    action_rows = [[action, trigger, service, result] for action, trigger, service, result in topic.actions]
    if not action_rows:
        action_rows = [["Render/read-only", "feature load", "local/shared state", "แสดงผลตาม scope โดยไม่มี mutation"]]
    failure_rows = []
    for index, case in enumerate(topic.tests[:6], start=1):
        expected = topic.acceptance[(index - 1) % len(topic.acceptance)] if topic.acceptance else "UI ต้องอยู่ใน state ที่ retry หรือแก้ข้อมูลได้"
        failure_rows.append([f"FE-{index:02d}", case, expected])
    return [
        h(2, f"5.90 {feature_name} Component Contract"),
        table(["ID", "Component / Scope", "Single responsibility", "Definition of done"], component_rows),
        h(2, f"5.91 {feature_name} API Adapter Map"),
        table(["Endpoint", "Typed adapter purpose", "Invoked by"], api_rows),
        h(2, f"5.92 {feature_name} Interaction State Machine"),
        table(["Action", "Trigger", "API / State transition", "Expected visible result"], action_rows),
        h(2, f"5.93 {feature_name} Feature Failure Checks"),
        table(["Case", "Feature-specific scenario", "Expected evidence"], failure_rows),
    ]


# ---------------------------------------------------------------------------
# Workflow trigger-event contract (มติ 2026-08-25)
# ทุกเอกสาร BE / Job ต้องบอกให้ชัดว่า "ต้องเรียก engine ตัวไหน ที่จุดไหน" หรือ
# "ไม่เรียก" เพื่อให้คน BE ที่ไม่ได้ทำ workflow เองรู้ภาระของตัวเองตั้งแต่อ่านเอกสาร
# ชื่อ function ยึดชีต Detail ของ SBP/TSM-SRM-LLDD-SBP-workflow-1.2.md (8 ตัว)
# ---------------------------------------------------------------------------

WORKFLOW_TRIGGER_CONTRACTS: dict[str, list[list[str]]] = {
    # ---- ฝั่ง BE API: มีเฉพาะบางเส้นที่แตะ engine ----
    "LLDD-BE-API-Document-Workflow-Actions": [
        ["ก่อนแสดงปุ่ม / ก่อนรับ action", "`getPermissionEvents`", "versionId, referenceId, userData", "ถ้า event ที่ส่งมาไม่อยู่ใน event[] ที่คืนมา ต้องตอบ 403 ห้ามเรียก eventWorkflow ต่อ"],
        ["กันกดซ้ำ / กันงานถูกคนอื่นเดินไปแล้ว", "`getTransaction`", "versionId, referenceId", "เทียบ state ปัจจุบันกับ state ที่ FE ส่งมา ไม่ตรงตอบ 409 (optimistic guard)"],
        ["กดผลพิจารณา (trigger event)", "`eventWorkflow`", "versionId, referenceId, userId, event, eventParam (amount สำหรับ route 100,000)", "🔴 หัวใจของเอกสารนี้ — เขียน consideration_logs + แนบไฟล์ + เรียก eventWorkflow ใน transaction เดียว; engine fail ต้อง rollback ฝั่ง SBPGI ทั้งหมด"],
        ["ผูกผู้รับผิดชอบขั้นถัดไป", "`addPreApprover`", "versionId, referenceId, stateId, approver, seq", "ใช้เมื่อ route ระบุตัวบุคคล (เช่น ตีกลับหาเจ้าของงานคนเดิม) — เรียกหลัง eventWorkflow สำเร็จ ใน transaction เดียวกัน"],
    ],
    "LLDD-BE-API-Workflow-Instances": [
        ["เปิด instance ให้เอกสาร", "`initializeWorkflow`", "versionId, userId, referenceId = `compensation_documents.id` (DP-1 ปิดแล้ว)", "idempotent — referenceId เดิมต้องไม่เกิด workflow_transaction ที่สอง"],
        ["ระบุผู้อนุมัติล่วงหน้า", "`addPreApprover`", "versionId, referenceId, stateId, approver, seq, userId", "เรียกต่อทันทีหลัง initialize ภายใน transaction เดียว"],
        ["อ่านสถานะ instance", "`getTransaction`", "versionId, referenceId", "ใช้ยืนยันว่า initialize สำเร็จจริงก่อนคืน 201"],
    ],
    "LLDD-BE-API-Document-Create-Update": [
        ["หลัง insert เอกสารสำเร็จ (Open first task)", "`initializeWorkflow` แล้วต่อด้วย `addPreApprover`", "versionId, userId, referenceId, stateId = `06`", "อยู่ใน transaction boundary เดียวกับการสร้างเอกสาร — engine fail ต้อง rollback เอกสาร"],
    ],
    "LLDD-BE-API-Document-Detail-Aggregate": [
        ["ประกอบหน้าเอกสาร", "`getPermissionEvents`", "versionId, referenceId, userData", "คืน event[] เป็นปุ่ม และ display[] เป็น READ/WRITE ต่อ part — FE ห้ามคำนวณสิทธิ์เอง"],
        ["สถานะ + ผู้ถืองานปัจจุบัน", "`getTransaction`", "versionId, referenceId", "ใช้เป็นค่าอ้างอิงให้ FE ส่งกลับมาตอนกดปุ่ม (optimistic guard)"],
    ],
    "LLDD-BE-API-Document-List-Search": [
        ["กล่องงานรอดำเนินการ", "`getPendingFlowByUser`", "userData, versionId", "เป็นแหล่งความจริงของรายการรอดำเนินการ · section 06 ต้อง union เอกสารที่จบด้วย หยุดชดเชยฯ (stoppedReopenable) เพิ่มเอง"],
    ],
    "LLDD-BE-API-Attachment-Sales-Timeline": [
        ["แท็บประวัติ (timeline)", "`getHistory`", "versionId, referenceId", "⚠️ ขึ้นกับ DP-7 — ถ้าเลือกอ่าน engine ต้อง join `consideration_logs` เพิ่ม decision code / ไฟล์แนบ / ความเห็นที่ engine ไม่มี"],
    ],
    "LLDD-BE-API-Common-Contracts": [
        ["ตัวห่อกลาง (WorkflowGateway)", "ทั้ง 8 ตัวของ `@srm/glb-workflow`", "userData มาจาก BFF header (`x-user-id`, `x-user-group-id`)", "🔴 งานของเอกสารนี้คือ **ทำตัวห่อกลางให้ทุกคนเรียก** — map error ของ engine เข้า envelope `{success:false, error:{code,message}}` และบังคับ timeout/retry ที่เดียว"],
    ],
    "LLDD-BE-Data-Migration-Cutover": [
        ["ย้ายเอกสารที่ค้างกลางทาง", "`initializeWorkflow` แล้ว `eventWorkflow` ซ้ำจนถึง state ปัจจุบัน", "versionId, referenceId, ลำดับ event ตามสถานะเดิมใน K2", "🔴 ห้าม INSERT `workflow_transaction` ตรงเพื่อ 'ตั้ง state ให้ตรง' — ต้องเดิน event จริงเพื่อให้ history ครบ · ต้อง rerun ได้ (referenceId เดิมไม่สร้าง instance ซ้ำ)"],
    ],
    "LLDD-BE-Integration-SBP-Platform": [
        ["ส่ง identity ให้ engine", "ทุก function ที่รับ `userData`", "แปลง BFF header → userData ที่ lib ต้องการ", "ถ้า mapping ผิด `getPermissionEvents` จะคืนปุ่มว่างทั้งหน้า — ต้องมี contract test ครอบ"],
    ],
    # ---- ฝั่ง Job: มีแค่ Job 8b ตัวเดียวที่แตะ workflow engine ----
    # Job อื่น (2/3/4/5/6/7/8/9/10) ไม่เรียก engine จึง **ไม่มีหัวข้อนี้ในเอกสาร**
    # (มติผู้ใช้ 2026-08-25: "job อันไหนต้อง trigger event ก็ใส่ อันไหนไม่มีก็ไม่ต้องใส่")
    "LLDD-BE-Job-8b-StartInternalWorkflow": [
        ["หลังผ่าน gate (เฉพาะเคส Y)", "`initializeWorkflow`", "versionId, userId = `JOB-8B`, referenceId = `compensation_documents.id`", "🔴 หัวใจของ job นี้ · เรียกใน transaction เดียวกับ update `fgi_impact_processes.workflow_generation_status = 'Y'`"],
        ["เลือก state เริ่มต้นตามประเภทเคส", "`addPreApprover`", "stateId = `06` (เปิดเรื่องใหม่) / `08` (ชดเชยต่อเนื่อง) / `01` (ยอด 0 เดือน 1-3), approver, seq = 1", "เคสชดเชยต่อเนื่องต้องผูก **เจ้าหน้าที่ SBP DSA คนเดิม** — ดู 5.2 ของเอกสารนี้"],
        ["ดันเอกสารไปยัง state เริ่มต้นที่ไม่ใช่ state แรก", "`eventWorkflow`", "versionId, referenceId, event ตามผัง To-Be 12/02/2026", "เคส 08 / 01 ต้องเดิน event จาก state แรกจริง ๆ ห้าม INSERT `workflow_transaction` ให้เริ่มที่ state กลาง"],
        ["rerun / กันเปิดซ้ำ", "`initializeWorkflow` (idempotent)", "referenceId เดิม", "referenceId เดิมต้องไม่เกิด workflow_transaction ที่สอง · เคส N persist ถาวร เคส W คงเดิมเพื่อ rerun"],
    ],
}


def workflow_trigger_contract_blocks(topic: "Topic", section_no: str) -> list[dict[str, Any]]:
    """5.92 — บอกให้ชัดต่อเอกสารว่าต้องเรียก @srm/glb-workflow ตัวไหน ที่จุดไหน หรือไม่เรียก"""
    doc_key = topic.file.rsplit("/", 1)[-1]
    rows = WORKFLOW_TRIGGER_CONTRACTS.get(doc_key)
    if not rows:
        # เอกสารที่ไม่ได้แตะ workflow engine — ไม่ต้องมีหัวข้อนี้เลย (มติ 2026-08-25)
        return []
    return [
        h(2, f"{section_no} Workflow Trigger Event Contract"),
        p(
            "งานชิ้นนี้ **ต้องเรียก workflow engine** ตามตารางด้านล่าง · "
            "ชื่อ function ยึด API 8 ตัวของ `@srm/glb-workflow` ตามชีต `Detail` ของ "
            "`SBP/TSM-SRM-LLDD-SBP-workflow-1.2.md` — รายละเอียด signature และตารางที่ engine เขียน "
            "ดู **LLDD-BE-Workflow-Engine-Definition** หัวข้อ 5.3"
        ),
        table(["จุดที่เรียก (call site)", "Engine function", "พารามิเตอร์หลัก", "กติกา / transaction boundary"], rows),
        bullets([
            "🔴 กติกาเหล็ก: ตาราง `sps_store.workflow_*` (13 ตาราง) เป็นของ lib — SBPGI **R เท่านั้น** ห้าม INSERT/UPDATE/DELETE ตรงในทุกกรณี",
            "ทุกการเรียก engine ต้องผ่านตัวห่อกลาง `WorkflowGateway` ที่นิยามใน **LLDD-BE-API-Common-Contracts** (timeout · retry · map error เข้า envelope) ห้าม import lib ตรงจาก service",
            "unit test ต้อง mock engine และครอบอย่างน้อย: เรียกสำเร็จ · engine โยน error แล้ว rollback ฝั่ง SBPGI ครบ · เรียกซ้ำด้วย referenceId เดิมไม่เกิดผลซ้ำ",
        ]),
    ]


def be_implementation_blocks(topic: Topic) -> list[dict[str, Any]]:
    endpoint_rows = []
    for index, spec in enumerate(topic.apis, start=1):
        endpoint_rows.append([
            f"{spec.method} {spec.path}",
            spec.purpose,
            topic.flow[(index - 1) % len(topic.flow)] if topic.flow else "validate -> execute -> map response",
            topic.acceptance[(index - 1) % len(topic.acceptance)] if topic.acceptance else "contract test ผ่าน",
        ])
    if not endpoint_rows:
        endpoint_rows = [["Internal service", topic.objective, "เรียกจาก use case ภายในเท่านั้น", topic.acceptance[0] if topic.acceptance else "service test ผ่าน"]]
    sequence_rows = []
    for index, step in enumerate(topic.flow, start=1):
        # ห้ามวน tests ซ้ำ — step ที่ไม่มี test ตรงตัวให้เขียนว่ายังไม่มี ไม่ใช่ยืมหลักฐานของ step อื่น
        if index - 1 < len(topic.tests):
            failure = topic.tests[index - 1]
        else:
            failure = "— (ยังไม่มี test เฉพาะขั้นนี้ · ครอบด้วย test รวมของเอกสารในหัวข้อ 11)"
        sequence_rows.append([index, step, failure])
    if not sequence_rows:
        sequence_rows = [[1, "อ่านข้อมูลตาม DB Mapping และคืนผลตาม contract", "ไม่พบข้อมูลคืน typed error"]]
    return [
        h(2, "5.90 Endpoint Implementation Contract"),
        table(["Endpoint", "Use-case owner", "Service/repository behavior", "Definition of done"], endpoint_rows),
        h(2, "5.91 Backend Execution Sequence"),
        table(["Step", "Behavior specific to this LLDD", "Failure/test evidence"], sequence_rows),
    ] + workflow_trigger_contract_blocks(topic, "5.92")


def job_implementation_blocks(topic: Topic) -> list[dict[str, Any]]:
    job_no = topic.file.split("LLDD-BE-Job-", 1)[1].split("-", 1)[0]
    spec = JOB_IMPLEMENTATION_SPECS[job_no]
    legacy = LEGACY_JOB_SOURCES[job_no]
    stage_rows = []
    for index, step_name in enumerate(spec["steps"].split("|"), start=1):
        stage_rows.append([index, step_name, spec["repository"], "คืน metrics และ throw typed error; transaction/rerun ใช้ contract ด้านล่าง"])
    blocks = [
        h(2, f"5.90 Job {job_no} Execution Stages"),
        p(legacy["progress"]),
        table(["Order", "Service step", "Repository", "Output / failure contract"], stage_rows),
        h(2, f"5.91 Job {job_no} Run Evidence"),
        table(["Evidence", "Job-specific value", "Acceptance"], [
            ["Input identity", legacy["input"], "snapshot input file/business key/period in run record"],
            ["Output identity", legacy["output"], "reconcile input, success, reject and skipped counts"],
            ["Dedup proof", spec["idempotency"], "rerun fixture produces no duplicate target business key"],
            ["Transaction proof", spec["transaction"], "injected failure leaves no partial committed state outside documented boundary"],
            ["Security proof", spec["security"], "config/log/error contains no plaintext secret"],
        ]),
    ]
    blocks.extend(legacy_job_source_blocks(job_no))
    blocks.extend([
        h(2, "5.93 Target Repository and SQL Contract"),
        table(
            ["Contract", "Target implementation"],
            [
                ["Repository", spec["repository"]],
                ["Idempotency / dedup", spec["idempotency"]],
                ["Transaction boundary", spec["transaction"]],
                ["Security", spec["security"]],
            ],
        ),
        h(3, "Input / candidate query"),
        code(spec["read"], "sql"),
        h(3, "Write / upsert query"),
        code(spec["write"], "sql"),
        h(2, "5.94 Target Node Implementation"),
        p("โครงสร้างนี้ระบุ service/repository เฉพาะงานและต้อง implement ตาม SQL, transaction, idempotency และ security contract ด้านบน โดยทุกขั้นต้องคืน metrics สำหรับ reconcile และ run history"),
        code(node_job_skeleton(job_no, topic), "js"),
    ])
    if job_no == "4":
        blocks.extend([
            h(2, "5.95 Job 4 Atomic File / Outbox Sequence"),
            table(
                ["Order", "Required action", "Failure behavior"],
                [
                    [1, "lock candidate W ด้วย FOR UPDATE SKIP LOCKED และสร้าง payload ใน memory", "validation fail: rollback lock; สถานะยัง W"],
                    [2, "เขียน temporary file, fsync, atomic rename และคำนวณ SHA-256", "write/rename/checksum fail: ลบ temp; สถานะยัง W; ไม่สร้าง outbox"],
                    [3, "transaction เดียว update W→P และ insert interface_transactions/outbox READY", "DB fail: rollback W→P และ outbox; durable file คงไว้ให้ cleanup/reconcile โดย checksum"],
                    [4, "dispatcher อ่าน READY แล้วอัปโหลดขึ้น EAI S3 (prefix ขาออก); compare checksum ก่อนส่ง", "อัปโหลด fail: outbox ยัง READY/FAILED_RETRY; ห้ามเปลี่ยน candidate กลับ W เพื่อไม่ให้สร้างไฟล์ซ้ำ"],
                    [5, "ส่งสำเร็จ mark SENT; callback/import ที่สัมพันธ์กัน mark ACKED", "ใช้ transaction id เดิมตลอด lifecycle"],
                ],
            ),
        ])
    if job_no == "6":
        blocks.extend([
            h(2, "5.96 เขียนข้อมูลรอบชดเชย (รับเข้าโครง 2026-08-21 · gap F8 + F1)"),
            p("Job 6 คือ job เดียวที่เขียนตารางรอบชดเชยในระบบเดิม — `ExportService.manageDBToFs()` เรียก 5 คำสั่งต่อกันเป็นชุด ระบบใหม่ต้องทำครบเหมือนเดิม แต่เขียนลงตารางของ SBPGI"),
            table(["ลำดับใน manageDBToFs()", "ระบบเดิม (Oracle)", "ระบบใหม่ (SBPGI)", "ใช้ทำอะไรต่อ"], [
                ["updateFgiImpactStoreOnProcess(INITDATE)", "FGI_IMPACT_STORE_ON_PROCESS · LAST_COMPENSATE_SEQ_NO + 1 เมื่อ FLAG_ACTION='Y' และเพิ่งชดเชยเดือนที่แล้ว", "fgi_impact_processes.last_compensate_seq_no += 1", "**เคสต่อเนื่อง** (SEQ_NO > 1)"],
                ["insertFgiImpactStoreOnProcess()", "แถวใหม่ · LAST_COMPENSATE_SEQ = MAX+1 · SEQ_NO = 1 · FLAG_ACTION='Y' · DATASOURCE", "fgi_impact_processes แถวใหม่ (last_compensate_seq · last_compensate_seq_no=1 · flag_action · datasource)", "**เปิดเรื่องใหม่** (SEQ_NO = 1)"],
                ["insertFgiImpactStoreCompensate(...)", "FGI_IMPACT_STORE_COMPENSATE · COMPENSATE_FORECAST / COMPENSATE_ADJUST ต่องวด", "**fgi_impact_compensations** (forecast_amount / adjust_amount)", "**นับยอด 0 ติดกันกี่เดือน** (กติกาเดือน 1-3 / เดือนที่ 4)"],
                ["insertFgiNewStoreCompensate(...)", "FGI_NEW_STORE_COMPENSATE", "document_new_stores.compensation_amount / compensate_percent", "ยอดต่อร้านเปิดใหม่"],
                ["updateCompleteImpactStoreOnProcess / FlagYToW", "FLAG_ACTION Y→N / Y→W", "fgi_impact_processes.flag_action", "ปิดรอบ / ส่งกลับรอตรวจ"],
            ]),
            p("⚠️ `ImportJdbc.insertImpactStoreOnProcess()` / `updateImpactStoreOnProcess()` มี SQL ชุดเดียวกันอยู่ในไฟล์ Import แต่ตรวจทั้ง src แล้ว **ไม่มี call site จริง** — เป็นโค้ดตาย ให้ยึด `ExportJdbc` เป็นต้นแบบเท่านั้น"),
            h(2, "5.95 Tracking Retention / Purge SQL"),
            p("Purge ทำได้เฉพาะ ACKED/COMPLETED ที่ครบ purge_after และไม่อยู่ใน legal hold; ต้องรันเป็น batch จำกัดจำนวนเพื่อไม่ lock ตารางยาว"),
            code("""WITH purge_candidates AS (
    SELECT id
    FROM interface_transactions
    WHERE status IN ('ACKED', 'COMPLETED')
      AND purge_after < CURRENT_TIMESTAMP
      AND legal_hold = FALSE
      AND data_name = ANY(:sta_data_names)
    ORDER BY id
    LIMIT :batch_size
    FOR UPDATE SKIP LOCKED
)
DELETE FROM interface_transactions i
USING purge_candidates p
WHERE i.id = p.id
RETURNING i.id, i.data_name, i.business_key;""", "sql"),
        ])
    if job_no == "8":
        blocks.extend([
            h(2, "5.95 Job 8 Document Number Gap and Rerun Policy"),
            p("Job 8 ใช้ running number แบบ monotonic ต่อปี ค.ศ. ช่องว่างของเลขเอกสารจาก concurrent rerun หรือ ON CONFLICT เป็นพฤติกรรมที่ยอมรับได้ เพราะเลขที่มีหน้าที่รับประกัน uniqueness ไม่ได้รับประกันความต่อเนื่อง"),
            table(
                ["Case", "Required behavior", "Evidence / metric"],
                [
                    ["Rerun พบ impact_process_id เดิมก่อนจองเลข", "คืน/ข้ามด้วย doc_no เดิมโดยไม่จอง running_no เพิ่มเมื่อ fast lookup พบข้อมูลแล้ว", "duplicateExistingCount + existingDocNo"],
                    ["Concurrent worker ชน ON CONFLICT หลังจองเลข", "ยอมให้ running_no ที่จองแล้วกลายเป็น gap; ห้ามลด sequence และห้ามนำเลขกลับมาใช้", "numberGapCount + conflictedImpactProcessId"],
                    ["Conflict path", "อ่าน compensation_documents ด้วย impact_process_id แล้วใช้ d.doc_no เดิมสำหรับ tracking/reconcile", "tracking.doc_no ตรงกับเอกสารที่ commit อยู่จริง"],
                    ["New document path", "insert document และ tracking (direction=INTERNAL) ใน transaction เดียว", "createdCount และ trackingCount เพิ่มเท่ากัน"],
                    ["Audit/runbook", "อธิบายว่าเลขอาจไม่ต่อเนื่องแต่ต้องไม่ซ้ำและตรวจสอบย้อนกลับได้", "ไม่มีขั้นตอน manual reuse หรือ renumber"],
                ],
            ),
        ])
    # เลือกเลขหัวข้อถัดไปที่ยังว่าง (job บางตัวใช้ 5.95/5.96 ไปแล้ว) เพื่อไม่ให้เลขกระโดด
    used = {b["text"].split()[0] for b in blocks if str(b.get("type","")).startswith("h") and str(b.get("text","")).startswith("5.9")}
    next_no = next(f"5.{n}" for n in range(95, 100) if f"5.{n}" not in used)
    blocks.extend(workflow_trigger_contract_blocks(topic, next_no))
    return blocks


def role_doc_implementation_blocks(topic: Topic) -> list[dict[str, Any]]:
    profile = document_detail_role_profile(topic.file)
    if not profile:
        raise ValueError(f"Missing role profile for {topic.file}")
    profile_code = f"P-{profile['code']}"
    visible = ", ".join(profile["visible"])
    hidden = ", ".join(profile["hidden"]) if profile["hidden"] else "ไม่มี section ที่ซ่อนเพิ่มจาก profile"
    editable = ", ".join(profile["editable"]) if profile["editable"] else "ไม่มี; business section ทั้งหมด read-only"
    actions = "; ".join(action[0] for action in profile["actions"])
    comment_rules = "; ".join(f"{action}: {rule}" for action, rule in profile["actions"])
    return [
        h(2, f"5.90 {topic.title.replace('LLDD FE - ', '')} Implementation Steps"),
        table(
            ["Step", "Implementation detail", "Check"],
            [
                ["Load exact profile", f"เรียก GET /api/v1/sbpgi/document/{{docNo}} และยืนยัน roleProfileCode={profile_code}, statusCode={profile['code']} ก่อน render action state", f"profile mismatch ต้อง fail closed; ไม่ใช้ role switcher เพื่อจำลอง {profile_code}"],
                ["Render profile sections", f"render เฉพาะ visibleSections ของ {profile_code}: {visible}; ซ่อน: {hidden}", "section ที่ซ่อนต้องไม่อยู่ใน DOM และ section key ที่ไม่รู้จักต้อง log/ignore แบบ fail closed"],
                ["Apply edit boundary", f"เปิด mutation control เฉพาะ editableSections ของ {profile_code}: {editable}", "read-only section ไม่มี focusable input/save/add/delete และ payload ต้องไม่มี field นอก editableSections"],
                ["Attachment control", f"canUploadAttachment={str(profile['upload']).lower()} สำหรับ {profile['short']}; ใช้ allowlist, 5 MB และ scan-status contract", "ปุ่ม upload ตรง flag, FILE_TOO_LARGE/FILE_SCAN_BLOCKED แสดงที่ attachment section"],
                ["Render exact action set", f"แสดง actionOptions ของ {profile_code} เท่านั้น: {actions}; comment rules: {comment_rules}", "radio label/requireComment มาจาก API และ FE ไม่คำนวณ nextSection"],
                ["Submit and reload", f"ส่ง result/comment สำหรับ {profile_code} แล้ว invalidate detail, timeline, task/list cache", f"หลัง submit ต้องโหลด status/actionOptions ใหม่และไม่คง action set ของ {profile_code} เมื่อ workflow เปลี่ยนขั้น"],
            ],
        ),
    ]


def document_detail_role_profiles() -> list[dict[str, Any]]:
    common_read = [
        "doc-header",
        "sec-sales",
        "sec-map",
        "sec-newstore",
        "sec-competitor",
        "sec-factor",
        "sec-attach",
        "sec-comp-history",
        "sec-decision-history",
        "sec-action",
    ]
    return [
        {
            "file": "FE/LLDD-FE-Document-Detail-Role-06-SBP-DSA",
            "code": "06",
            "name": "ฝ่าย SBP DSA",
            "short": "SBP DSA",
            "status": "รอฝ่าย SBP DSA ดำเนินการ",
            "purpose": "ตรวจความครบถ้วนเบื้องต้นและเลือกส่งต่อ/ยุติตามผลพิจารณา",
            "visible": common_read,
            "editable": [],
            "hidden": ["sec-calc"],
            "upload": True,
            "summary": [
                "เห็นข้อมูลเอกสารครบสำหรับตรวจสอบ แต่ทุก section เนื้อหาเป็น read-only",
                "เพิ่มเอกสารแนบประกอบการพิจารณาได้",
                "ไม่เห็น section คำนวณเงินชดเชย",
            ],
            "fields": [
                ["เอกสารแนบ", "file, fileName, attachmentType, remark", "เพิ่มไฟล์ได้; ขนาด <= 5 MB; extension ต้องอยู่ใน allowlist"],
                ["แผงพิจารณา", "result, comment", "result required; comment required เมื่อเลือก เห็นควรไม่ชดเชย"],
            ],
            "actions": [
                ["เห็นควรไม่ชดเชย", "ต้องกรอก comment"],
                ["หยุดชดเชยประกันรายได้", "comment optional"],
                ["ส่งหน่วยงานส่งเสริมธุรกิจ SBP", "comment optional"],
                ["ส่งเจ้าหน้าที่ SBP DSA ดำเนินการ", "comment optional"],
            ],
            "tests": [
                "เปิดด้วย roleProfileCode=P-06 แล้ว sec-calc ต้องไม่ render",
                "section ร้านเปิดใหม่/คู่แข่ง/ปัจจัยต้องไม่มี input/edit/delete/save",
                "ไม่เลือก result แล้วกดส่ง ต้องแสดง popup verbatim",
                "เลือก เห็นควรไม่ชดเชย โดยไม่กรอก comment ต้อง error ACTION_COMMENT_REQUIRED",
                "upload ไฟล์เกิน 5 MB ต้อง error FILE_TOO_LARGE",
            ],
        },
        {
            "file": "FE/LLDD-FE-Document-Detail-Role-08-SBP-DSA-Officer",
            "code": "08",
            "name": "เจ้าหน้าที่ SBP DSA",
            "short": "SBP DSA Officer",
            "status": "รอเจ้าหน้าที่ SBP DSA ดำเนินการ",
            "purpose": "ตรวจ/ยืนยันผลคำนวณเงินชดเชยและส่งผลพิจารณา",
            "visible": common_read + ["sec-calc"],
            "editable": [],
            "hidden": [],
            "upload": True,
            "summary": [
                "เห็น section คำนวณเงินชดเชยเพิ่มเติมจากบทบาทอื่น",
                "section คำนวณเป็น display-only ไม่ใช่ editor",
                "เพิ่มเอกสารแนบและส่ง action ได้",
            ],
            "fields": [
                ["คำนวณเงินชดเชย", "baseCompensationAmount, totalCompensatePercent, totalCompensationAmount, approvalLimitIndicator", "read-only; แสดงเกณฑ์วงเงินอนุมัติจาก API (< 100,000 จบที่ GM · ≥ 100,000 ส่ง AVP · มติ 2026-08-18)"],
                ["เอกสารแนบ", "file, fileName, attachmentType, remark", "เพิ่มไฟล์ได้; ขนาด <= 5 MB; extension ต้องอยู่ใน allowlist"],
                ["แผงพิจารณา", "result, comment", "result required; comment ตาม actionOptions.requireComment"],
            ],
            "actions": [
                ["คำนวณเงินชดเชยเรียบร้อย", "comment optional"],
                ["ส่งกลับฝ่าย SBP DSA", "comment ตาม actionOptions.requireComment"],
            ],
            "tests": [
                "เปิดด้วย roleProfileCode=P-08 แล้ว sec-calc ต้องแสดง",
                "sec-calc ต้องไม่มี input/button บันทึก",
                "section ร้านเปิดใหม่/คู่แข่ง/ปัจจัยต้อง read-only",
                "action radio แสดงเฉพาะ 2 รายการของ role 08",
                "หลัง submit ต้อง reload detail/timeline/status",
            ],
        },
        {
            "file": "FE/LLDD-FE-Document-Detail-Role-01-Business-Promotion",
            "code": "01",
            "name": "หน่วยงานส่งเสริมธุรกิจ SBP",
            "short": "Business Promotion",
            "status": "รอหน่วยงานส่งเสริมธุรกิจ SBP ดำเนินการ",
            "purpose": "ปรับข้อมูลร้านเปิดใหม่ ร้านคู่แข่ง ปัจจัยอื่น และส่งผลพิจารณา",
            "visible": common_read,
            "editable": ["sec-newstore", "sec-competitor", "sec-factor"],
            "hidden": ["sec-calc"],
            "upload": True,
            "summary": [
                "เป็น role profile เดียวที่แก้เนื้อหาเอกสารได้",
                "แก้ % ชดเชย เพิ่ม/แก้/ลบร้านคู่แข่ง และเพิ่ม/แก้/ลบปัจจัยอื่นได้",
                "ไม่เห็น section คำนวณเงินชดเชย",
            ],
            "fields": [
                ["ร้านเปิดใหม่", "newStoreCode, newStoreName, openDate, distanceKm, compensatePercent, calculatedCompensationAmount", "แก้ได้เฉพาะ compensatePercent; ผลรวมต้องเท่ากับ 100"],
                ["ร้านคู่แข่ง", "competitorName, openedImpactDate, detail, remark", "เพิ่ม/แก้/ลบได้; ต้องเลือกร้านคู่แข่งก่อนบันทึก"],
                ["ปัจจัยอื่นๆ", "factorName, startDate, endDate, detail, remark", "เพิ่ม/แก้/ลบได้; endDate ต้องไม่ก่อน startDate"],
                ["เอกสารแนบ", "file, fileName, attachmentType, remark", "เพิ่มไฟล์ได้; ขนาด <= 5 MB; extension ต้องอยู่ใน allowlist"],
                ["แผงพิจารณา", "result, comment", "result required; comment required เมื่อเลือก เห็นควรไม่ชดเชย"],
            ],
            "actions": [
                ["เห็นควรชดเชย", "comment optional"],
                ["เห็นควรไม่ชดเชย", "ต้องกรอก comment"],
                ["ฝ่าย SBP DSA ดำเนินการ (ส่งกลับ)", "comment ตาม actionOptions.requireComment"],
            ],
            "tests": [
                "เปิดด้วย roleProfileCode=P-01 แล้ว sec-newstore/sec-competitor/sec-factor ต้อง editable",
                "แก้ compensatePercent แล้วรวมไม่ครบ 100 ต้อง error COMPENSATE_PERCENT_INVALID",
                "เพิ่มร้านคู่แข่งโดยไม่เลือก competitor ต้อง error COMPETITOR_REQUIRED",
                "เพิ่มปัจจัยอื่นโดยไม่เลือก factor ต้อง error EXTERNAL_FACTOR_REQUIRED",
                "sec-calc ต้องไม่ render สำหรับ role 01",
            ],
        },
        {
            "file": "FE/LLDD-FE-Document-Detail-Role-02-GM-Business-Promotion",
            "code": "02",
            "name": "GM ส่งเสริมธุรกิจฯ",
            "short": "GM Business Promotion",
            "status": "รอ GM ส่งเสริมธุรกิจ SBP ดำเนินการ",
            "purpose": "อ่านข้อมูลประกอบการอนุมัติวงเงินและส่งผลพิจารณา",
            "visible": common_read,
            "editable": [],
            "hidden": ["sec-calc"],
            "upload": True,
            "summary": [
                "เห็นข้อมูลเอกสารทั้งหมดแบบ read-only",
                "เพิ่มเอกสารแนบประกอบการอนุมัติได้",
                "FE แสดงข้อความช่วยตัดสินจากยอดชดเชยรวม แต่ไม่คำนวณปลายทาง action เอง",
            ],
            "fields": [
                ["ข้อมูลประกอบอนุมัติ", "totalCompensationAmount, approvalLimitIndicator", "read-only จาก API; ใช้แสดงเกณฑ์วงเงินอนุมัติ (< 100,000 จบที่ GM · ≥ 100,000 ส่ง AVP ตาม SDD GI)"],
                ["เอกสารแนบ", "file, fileName, attachmentType, remark", "เพิ่มไฟล์ได้; ขนาด <= 5 MB; extension ต้องอยู่ใน allowlist"],
                ["แผงพิจารณา", "result, comment", "result required; comment ตาม actionOptions.requireComment"],
            ],
            "actions": [
                ["เห็นควรชดเชย", "comment optional"],
                ["เห็นควรไม่ชดเชย", "comment ตาม actionOptions.requireComment"],
                ["ส่งกลับหน่วยงานส่งเสริมธุรกิจ SBP", "comment ตาม actionOptions.requireComment"],
            ],
            "tests": [
                "เปิดด้วย roleProfileCode=P-02 แล้วทุก business section ต้อง read-only",
                "sec-calc ต้องไม่ render",
                "action radio แสดงเฉพาะ 3 รายการของ role 02",
                "แสดง totalCompensationAmount/approvalLimitIndicator โดยไม่คำนวณ route ใน FE",
                "หลัง submit ต้อง reload detail/timeline/status",
            ],
        },
        {
            "file": "FE/LLDD-FE-Document-Detail-Role-03-AVP-SBP",
            "code": "03",
            "name": "AVP สำนักบริหาร SBP",
            "short": "AVP SBP",
            "status": "รอผู้บริหารสำนักบริหาร SBP ดำเนินการ",
            "purpose": "อ่านข้อมูลประกอบการอนุมัติระดับสูงและส่งผลพิจารณา",
            "visible": common_read,
            "editable": [],
            "hidden": ["sec-calc"],
            "upload": True,
            "summary": [
                "เห็นข้อมูลเอกสารทั้งหมดแบบ read-only",
                "ต้องเห็นประวัติพิจารณา/timeline เพื่อประกอบการตัดสินใจ",
                "เพิ่มเอกสารแนบและส่ง action ได้",
            ],
            "fields": [
                ["ข้อมูลประกอบอนุมัติ", "doc-header, totalCompensationAmount, considerationHistory, timeline", "read-only ทั้งหมด"],
                ["เอกสารแนบ", "file, fileName, attachmentType, remark", "เพิ่มไฟล์ได้; ขนาด <= 5 MB; extension ต้องอยู่ใน allowlist"],
                ["แผงพิจารณา", "result, comment", "result required; comment ตาม actionOptions.requireComment"],
            ],
            "actions": [
                ["เห็นควรชดเชย", "comment optional"],
                ["เห็นควรไม่ชดเชย", "comment ตาม actionOptions.requireComment"],
                ["ส่งกลับ GM ส่งเสริมธุรกิจฯ", "comment ตาม actionOptions.requireComment"],
            ],
            "tests": [
                "เปิดด้วย roleProfileCode=P-03 แล้วทุก business section ต้อง read-only",
                "sec-calc ต้องไม่ render",
                "history/timeline ต้องแสดงก่อนส่ง action ได้",
                "action radio แสดงเฉพาะ 3 รายการของ role 03",
                "หลัง submit ต้อง reload detail/timeline/status",
            ],
        },
    ]


def document_detail_role_profile(file_key: str) -> dict[str, Any] | None:
    for profile in document_detail_role_profiles():
        if profile["file"] == file_key:
            return profile
    return None


def document_detail_role_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 Role-based Render Contract (ไม่ใช่ Routing Spec)"),
        p("หน้า Document Detail ต้องแสดงผลตาม role profile ที่ API ส่งมาเท่านั้น โดย role profile ระบุ visibleSections, editableSections และ actionOptions สำหรับผู้ใช้ที่ login จริง FE ไม่ต้องมี role switcher และไม่ต้องฝังตาราง action routing ใน production"),
        h(3, "Section Inventory"),
        table(
            ["Section key", "UI section", "Default display", "Editable by"],
            [
                ["doc-header", "ข้อมูลร้านถูกกระทบ", "read-only", "-"],
                ["sec-sales", "แนวโน้มยอดขายรายวัน", "read-only", "-"],
                ["sec-map", "แผนที่ AllMap", "read-only", "-"],
                ["sec-newstore", "ร้านเปิดใหม่", "read-only", "role profile 01"],
                ["sec-competitor", "ร้านคู่แข่งเปิดกระทบ", "read-only", "role profile 01"],
                ["sec-factor", "ปัจจัยอื่นๆ", "read-only", "role profile 01"],
                ["sec-attach", "เอกสารแนบทั้งหมด", "visible + upload", "all action roles upload"],
                ["sec-calc", "คำนวณเงินชดเชย", "hidden", "visible-only role profile 08"],
                ["sec-comp-history", "ประวัติการชดเชย", "read-only", "-"],
                ["sec-decision-history", "ผลการพิจารณา (ประวัติ)", "read-only", "-"],
                ["sec-action", "พิจารณา / ส่งดำเนินการ", "visible", "current action role"],
            ],
        ),
        h(3, "Role × Section Display Matrix"),
        p("E = แก้ไขได้, R = อ่านอย่างเดียว, H = ซ่อน, Upload = เพิ่มเอกสารแนบได้"),
        table(
            ["Section", "06 ฝ่าย SBP DSA", "08 จนท. SBP DSA", "01 หน่วยงานส่งเสริมธุรกิจ SBP", "02 GM ส่งเสริมฯ", "03 AVP สำนักบริหาร SBP"],
            [
                ["doc-header", "R", "R", "R", "R", "R"],
                ["sec-sales", "R", "R", "R", "R", "R"],
                ["sec-map", "R", "R", "R", "R", "R"],
                ["sec-newstore", "R", "R", "E", "R", "R"],
                ["sec-competitor", "R", "R", "E", "R", "R"],
                ["sec-factor", "R", "R", "E", "R", "R"],
                ["sec-attach", "R+Upload", "R+Upload", "R+Upload", "R+Upload", "R+Upload"],
                ["sec-calc", "H", "R", "H", "H", "H"],
                ["sec-comp-history", "R", "R", "R", "R", "R"],
                ["sec-decision-history", "R", "R", "R", "R", "R"],
                ["sec-action", "Action set 06", "Action set 08", "Action set 01", "Action set 02", "Action set 03"],
            ],
        ),
        h(3, "Action Panel Options"),
        table(
            ["Role profile", "Radio options shown", "Required comment rule"],
            [
                ["06 ฝ่าย SBP DSA", "เห็นควรไม่ชดเชย; หยุดชดเชยประกันรายได้; ส่งหน่วยงานส่งเสริมธุรกิจ SBP; ส่งเจ้าหน้าที่ SBP DSA ดำเนินการ", "บังคับเมื่อเลือก เห็นควรไม่ชดเชย"],
                ["08 เจ้าหน้าที่ SBP DSA", "คำนวณเงินชดเชยเรียบร้อย; ส่งกลับฝ่าย SBP DSA", "บังคับเมื่อ actionOptions.requireComment=true"],
                ["01 หน่วยงานส่งเสริมธุรกิจ SBP", "เห็นควรชดเชย; เห็นควรไม่ชดเชย; ฝ่าย SBP DSA ดำเนินการ (ส่งกลับ)", "บังคับเมื่อเลือก เห็นควรไม่ชดเชย"],
                ["02 GM ส่งเสริมธุรกิจฯ", "เห็นควรชดเชย; เห็นควรไม่ชดเชย; ส่งกลับหน่วยงานส่งเสริมธุรกิจ SBP", "บังคับเมื่อ actionOptions.requireComment=true"],
                ["03 AVP สำนักบริหาร SBP", "เห็นควรชดเชย; เห็นควรไม่ชดเชย; ส่งกลับ GM ส่งเสริมธุรกิจฯ", "บังคับเมื่อ actionOptions.requireComment=true"],
            ],
        ),
        h(3, "Role Detail Documents"),
        p("รายละเอียดแบบอ่านง่ายแยกตามบทบาทอยู่ในเอกสารลูก 5 ฉบับด้านล่าง เอกสารหลักนี้เก็บเฉพาะ contract กลางและ matrix รวม"),
        table(
            ["Role", "เอกสารรายละเอียด", "เนื้อหาหลัก"],
            [
                [profile["code"], f"{Path(profile['file']).name}.pdf", profile["purpose"]]
                for profile in document_detail_role_profiles()
            ],
        ),
        h(3, "Validation Popup Text"),
        table(
            ["Condition", "Popup message"],
            [
                ["กดส่งดำเนินการโดยไม่เลือกผลการพิจารณา", "ท่านยังไม่เลือกผลการพิจารณา กรุณาเลือกข้อมูลก่อนกดส่งดำเนินการ"],
                ["result ที่ requireComment=true แต่ comment ว่าง", "กรุณากรอกความคิดเห็นเพิ่มเติม (บังคับกรอกสำหรับผลการพิจารณานี้) ก่อนส่งดำเนินการ"],
                ["ผลรวม %ชดเชยร้านเปิดใหม่ไม่เท่ากับ 100", "โปรดตรวจสอบ %ชดเชย ของท่าน รวมกันแล้วไม่เท่ากับ 100%"],
            ],
        ),
    ]


def document_detail_single_role_blocks(profile: dict[str, Any]) -> list[dict[str, Any]]:
    section_rows = []
    for key, label in [
        ("doc-header", "ข้อมูลร้านถูกกระทบ"),
        ("sec-sales", "แนวโน้มยอดขายรายวัน"),
        ("sec-map", "แผนที่ AllMap"),
        ("sec-newstore", "ร้านเปิดใหม่"),
        ("sec-competitor", "ร้านคู่แข่งเปิดกระทบ"),
        ("sec-factor", "ปัจจัยอื่นๆ"),
        ("sec-attach", "เอกสารแนบทั้งหมด"),
        ("sec-calc", "คำนวณเงินชดเชย"),
        ("sec-comp-history", "ประวัติการชดเชย"),
        ("sec-decision-history", "ผลการพิจารณา (ประวัติ)"),
        ("sec-action", "พิจารณา / ส่งดำเนินการ"),
    ]:
        if key in profile["hidden"]:
            state = "Hidden"
            control = "ไม่ render section"
        elif key in profile["editable"]:
            state = "Editable"
            control = "เปิด input/action เฉพาะ field ที่ระบุในเอกสารนี้"
        elif key == "sec-attach" and profile["upload"]:
            state = "Read-only + Upload"
            control = "ดูรายการไฟล์และเพิ่มไฟล์แนบได้"
        elif key == "sec-action":
            state = "Action"
            control = "แสดง radio result, textarea comment, ปุ่มส่งดำเนินการ"
        elif key in profile["visible"]:
            state = "Read-only"
            control = "แสดงข้อมูลและปิด input/editor"
        else:
            state = "Hidden"
            control = "ไม่ render section"
        section_rows.append([key, label, state, control])

    response = {
        "docNo": "2026/00123",
        "statusCode": profile["code"],
        "viewerRbacRoleCode": "R-XX",
        "roleProfileCode": role_profile_code(profile),
        "visibleSections": profile["visible"],
        "editableSections": profile["editable"],
        "canUploadAttachment": profile["upload"],
        "canAction": True,
        "actionOptions": [{"value": row[0], "label": row[0], "requireComment": "ต้องกรอก" in row[1]} for row in profile["actions"]],
    }
    return [
        h(2, "5.1 Role View Summary"),
        table(
            ["Item", "Value"],
            [
                ["Role profile", f"{role_profile_code(profile)} - {profile['name']}"],
                ["Workflow section/status code", profile["code"]],
                ["Document status shown", profile["status"]],
                ["Purpose on this page", profile["purpose"]],
                ["Editable sections", ", ".join(profile["editable"]) if profile["editable"] else "-"],
                ["Hidden sections", ", ".join(profile["hidden"]) if profile["hidden"] else "-"],
                ["Attachment upload", "Allowed" if profile["upload"] else "Not allowed"],
            ],
        ),
        h(2, "5.2 What This Role Sees"),
        bullets(profile["summary"]),
        h(2, "5.3 Section-by-section Behavior"),
        table(["Section key", "UI section", "State for this role", "Control behavior"], section_rows),
        h(2, "5.4 Editable Form Fields"),
        table(["Area", "Fields", "Validation / Behavior"], profile["fields"]),
        h(2, "5.5 Action Panel"),
        p("FE ต้อง render ตัวเลือกจาก `actionOptions` ที่ API ส่งมาเท่านั้น และส่ง payload `{result,comment}` โดยไม่คำนวณปลายทาง action เอง"),
        table(["Radio option", "Comment rule"], profile["actions"]),
        h(2, "5.6 API Response Example"),
        payload("GET /api/v1/sbpgi/document/{docNo} response", api_json(response)),
        h(2, "5.7 Validation Popup Text"),
        table(
            ["Condition", "Popup message"],
            [
                ["กดส่งดำเนินการโดยไม่เลือกผลการพิจารณา", "ท่านยังไม่เลือกผลการพิจารณา กรุณาเลือกข้อมูลก่อนกดส่งดำเนินการ"],
                ["result ที่ requireComment=true แต่ comment ว่าง", "กรุณากรอกความคิดเห็นเพิ่มเติม (บังคับกรอกสำหรับผลการพิจารณานี้) ก่อนส่งดำเนินการ"],
                ["ผลรวม %ชดเชยร้านเปิดใหม่ไม่เท่ากับ 100", "โปรดตรวจสอบ %ชดเชย ของท่าน รวมกันแล้วไม่เท่ากับ 100%"],
            ],
        ),
        h(2, "5.8 Role-specific Test Checklist"),
        table(["No", "Test"], [[i + 1, item] for i, item in enumerate(profile["tests"])]),
    ]


def document_detail_role_topic(profile: dict[str, Any]) -> Topic:
    forward_action = {
        "06": ("ส่งเจ้าหน้าที่ SBP DSA ดำเนินการ", "08", "08"),
        "08": ("คำนวณเงินชดเชยเรียบร้อย", "01", "01"),
        "01": ("เห็นควรชดเชย", "02", "02"),
        "02": ("เห็นควรชดเชย", "03", "03"),
        "03": ("เห็นควรชดเชย", None, "99"),
    }[profile["code"]]
    return Topic(
        profile["file"],
        f"LLDD FE - Document Detail Role {profile['code']} {profile['short']}",
        "FE",
        1.2,
        10,
        FE_OWNER_KITTISAK,
        f"อธิบายหน้าจอ Document Detail สำหรับ role {profile['code']} - {profile['name']}",
        [],
        [
            f"Role profile {role_profile_code(profile)} - {profile['name']}",
            "Visible/read-only/hidden section behavior",
            "Editable field and validation behavior",
            "Attachment upload behavior",
            "Action panel options and API response sample",
        ],
        [
            ("roleProfileCode", role_profile_code(profile), "must match API response", "ใช้เลือก view profile เฉพาะบทบาทนี้; แยก namespace จาก workflow section code"),
            ("statusCode", profile["code"], "from API", "workflow status/section code ปัจจุบัน ไม่ใช่ role profile"),
            ("visibleSections", "string[]", "from API", "FE แสดงเฉพาะ section ใน array"),
            ("editableSections", "string[]", "from API", "FE เปิด input/button เฉพาะ section ใน array"),
            ("actionOptions", "array", "from API", "FE render radio จาก array โดยไม่ hardcode"),
        ],
        [
            ("Load detail", "เปิดเอกสาร", "GET /api/v1/sbpgi/document/{docNo}", "render role profile"),
            ("Save editable section", "ปุ่มบันทึก", "PUT /api/v1/sbpgi/document/{docNo}", "ใช้เฉพาะ role ที่มี editableSections"),
            ("Upload attachment", "เลือกไฟล์", "POST /api/v1/sbpgi/document/{docNo}/attachments", "append attachment when allowed"),
            ("Submit action", "ปุ่มส่งดำเนินการ", "POST /api/v1/sbpgi/document/{docNo}/actions", "submit selected result"),
        ],
        [
            ApiSpec("GET", "/api/v1/sbpgi/document/{docNo}", f"โหลด role profile {role_profile_code(profile)} สำหรับหน้า detail", {"docNo": "2026/00123"}, {"docNo": "2026/00123", "statusCode": profile["code"], "viewerRbacRoleCode": "R-XX", "roleProfileCode": role_profile_code(profile), "visibleSections": profile["visible"], "editableSections": profile["editable"], "actionOptions": [{"value": row[0], "label": row[0], "requireComment": "ต้องกรอก" in row[1]} for row in profile["actions"]]}),
            ApiSpec("POST", "/api/v1/sbpgi/document/{docNo}/actions", f"ตัวอย่าง positive-path จาก section {profile['code']}; Section 02 ส่งต่อ AVP (03) เมื่อยอดรวม ≥ 100,000 บาท และจบที่ GM เมื่อ < 100,000 บาท (มติ 2026-08-18)", {"result": forward_action[0], "comment": "ส่งดำเนินการตามลำดับ"}, {"statusCode": forward_action[2], "nextSection": forward_action[1], "message": "submitted"}),
        ],
        [
            "Load document detail",
            "Apply visibleSections and editableSections",
            "Render fields/actions for this role only",
            "Validate popup text",
            "Submit action or save allowed section",
            "Reload detail/timeline/status",
        ],
        [
            "ไม่แสดง role switcher ใน production",
            "section ที่ hidden ต้องไม่ render",
            "section ที่ read-only ต้องไม่มี editable control",
            "action panel ตรงกับ actionOptions จาก API",
        ],
        profile["tests"],
        flow_diagram=f"LLDD/assets/flows/{sanitize_filename(profile['file'])}.png",
    )


def document_detail_role_topics() -> list[Topic]:
    return [document_detail_role_topic(profile) for profile in document_detail_role_profiles()]


def common_contract_extra_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 Error and Popup Catalog"),
        p("ทุก endpoint ต้องใช้ code และ message จาก catalog เดียวกันเมื่อเข้าเงื่อนไขเดียวกัน"),
        table(
            ["code", "HTTP / Scope", "Trigger", "message"],
            [
                ["ACTION_RESULT_REQUIRED", "422", "submit action โดยไม่เลือกผลการพิจารณา", "ท่านยังไม่เลือกผลการพิจารณา กรุณาเลือกข้อมูลก่อนกดส่งดำเนินการ"],
                ["ACTION_COMMENT_REQUIRED", "422", "result ที่ต้องมี comment แต่ comment ว่าง", "กรุณากรอกความคิดเห็นเพิ่มเติม (บังคับกรอกสำหรับผลการพิจารณานี้) ก่อนส่งดำเนินการ"],
                ["COMPENSATE_PERCENT_INVALID", "422", "ผลรวม % ชดเชยร้านเปิดใหม่ไม่เท่ากับ 100", "โปรดตรวจสอบ %ชดเชย ของท่าน รวมกันแล้วไม่เท่ากับ 100%"],
                ["COMPETITOR_REQUIRED", "422", "บันทึกร้านคู่แข่งโดยไม่เลือก competitorCode", "กรุณาเลือกร้านคู่แข่งก่อนบันทึก"],
                ["EXTERNAL_FACTOR_REQUIRED", "422", "บันทึกปัจจัยอื่นโดยไม่เลือก factorCode", "กรุณาเลือกปัจจัยอื่นก่อนบันทึก"],
                ["REPORT_DATE_RANGE_INVALID", "422", "impactMonthFrom มากกว่า impactMonthTo", "เดือนเริ่มต้นต้องไม่มากกว่าเดือนสิ้นสุด"],
                ["FILE_TOO_LARGE", "413", "attachment > 5 MB", "ไฟล์แนบมีขนาดเกิน 5 MB"],
                ["FILE_TYPE_UNSUPPORTED", "415", "extension/content type ไม่อยู่ใน allowlist", "ชนิดไฟล์ไม่อนุญาตให้อัปโหลด"],
                ["FILE_SCAN_BLOCKED", "422", "AV scan พบไวรัสหรือ scan failed", "ไฟล์แนบไม่ผ่านการตรวจสอบความปลอดภัย"],
                ["FORBIDDEN", "403", "ไม่มีสิทธิ์เมนู/เอกสาร/task", "กรุณาติดต่อผู้ดูแลระบบ"],
                ["DUPLICATE_DOCUMENT", "409", "business key ซ้ำตอนสร้างเอกสาร", "ร้านนี้ในเดือนนี้มีเอกสารอยู่แล้ว"],
                ["CONFLICT", "409", "resource/task ถูกเปลี่ยนหรือเงื่อนไขปัจจุบันไม่ตรงกับคำขอ", "ข้อมูลมีการเปลี่ยนแปลง กรุณาโหลดข้อมูลล่าสุดแล้วดำเนินการใหม่"],
                ["STALE_VERSION", "409", "versionNo ที่ส่งมาไม่ตรงกับ compensation_documents.version_no", "ข้อมูลถูกแก้ไขโดยผู้ใช้อื่น กรุณาโหลดข้อมูลล่าสุดแล้วลองอีกครั้ง"],
                ["FS_BRIDGE_UNAVAILABLE", "FE", "hidden iframe ไม่ตอบ FS_FORM_READY ภายในเวลาที่กำหนด", "ไม่สามารถเชื่อมต่อแบบฟอร์ม FS ได้ กรุณาลองอีกครั้ง"],
                ["FS_BRIDGE_ORIGIN_INVALID", "FE", "event.origin ไม่ตรง allowlist", "ไม่สามารถยืนยันแหล่งที่มาของแบบฟอร์ม FS ได้"],
                ["FS_BRIDGE_SCHEMA_INVALID", "FE", "FS_FIELD_SCHEMA ไม่ตรง message schema หรือมี field type ที่ไม่รองรับ", "ข้อมูลแบบฟอร์ม FS ไม่ถูกต้อง กรุณาติดต่อผู้ดูแลระบบ"],
                ["FS_BRIDGE_SUBMIT_FAILED", "FE", "FS_SUBMIT_RESULT ไม่สำเร็จหรือ FS_ERROR ตอน submit", "ส่งแบบฟอร์ม FS ไม่สำเร็จ กรุณาตรวจสอบข้อมูลแล้วลองอีกครั้ง"],
            ],
        ),
        h(2, "5.2 Endpoint Role Matrix"),
        p("Matrix นี้เป็น baseline สำหรับ BE authorization guard; menu-level visibility มาจาก permissions ต่อ URL ของ auth-backend (header `x-user-permissions`)"),
        table(
            ["Endpoint group", "Endpoint pattern", "Allowed roles / identity"],
            [
                ["Current user/menu", "ไม่ใช่ endpoint ของ SBPGI — FE เรียกของระบบเดิมผ่าน BFF: GET /auth/profile, GET /users/current, GET /menus, GET /groups/current-user/permissions", "authenticated user"],
                ["Task inbox", "GET /sbpgi/document/tasks", "authenticated user with assigned task access"],
                ["Document read/list/timeline/sales", "GET /sbpgi/document*, GET /sbpgi/document/{docNo}/timeline, GET /sbpgi/document/{docNo}/sales", "document participant or report/admin role explicitly granted"],
                ["Document create", "POST /sbpgi/document", "🔴 **service token / pipeline เท่านั้น** — มติ 2026-08-06 ตัดฟอร์มสร้างเอกสารใน FE ออกแล้ว (ต้นทางสร้างที่ระบบ FS แล้ว SBP Statement ส่งข้อมูลกลับ) · ห้ามระบุเป็นรหัสกลุ่มสิทธิ์ เพราะเลข 01/02/03 ชนกับ section_code ของ workflow"],
                ["Document update/action/attachment upload", "PUT /sbpgi/document/{docNo}, POST /sbpgi/document/{docNo}/actions, POST /sbpgi/document/{docNo}/attachments", "current action owner; admin override only with policy and audit reason"],
                ["Attachment download", "GET /sbpgi/document/{docNo}/attachments/{attachId}/download", "สิทธิ์เท่ากับอ่านเอกสาร + attachment ต้องเป็นของ docNo นั้น · ⚠️ เงื่อนไข `scan_status` ขึ้นกับนโยบาย AV ที่ยังไม่เคาะ (ดู `LLDD-BE-API-Attachment-Sales-Timeline` 5.1) — บังคับ CLEAN อย่างเดียวตอนนี้จะดาวน์โหลดไม่ได้เลย"],
                ["Lookup", "/sbpgi/lookup/document-statuses, /sbpgi/lookup/workflow-sections (ร้าน/ภาค/ประเภทสาขา ใช้ /store/* + /common/common-code ของระบบ SBP เดิม · 2026-08-06)", "authenticated user with related menu access"],
                ["Master (SBPGI)", "/sbpgi/master/factors*, /sbpgi/master/competitors*", "admin/HQ ตามสิทธิ์เมนูที่มากับ header x-user-permissions"],
                ["RBAC/ผู้ปฏิบัติงาน", "ไม่ใช่ endpoint ของ SBPGI — ตัด /operators* /roles* /menus* /menu-permissions* /employees/search รวม 14 เส้น (2026-08-05) ใช้ auth-backend เดิม จัดการที่หน้า /setting/manage-user-rights", "-"],
                ["Reports", "/sbpgi/report/status-summary*", "admin/HQ/report roles and accounting service user"],
                ["Internal workflow/interface", "/sbpgi/workflow/instances · /sbpgi/interface/* (tracking · pending-ack · sta/ack callback)", "service token หรือ API key เท่านั้น — ไม่ผ่านสิทธิ์เมนูของผู้ใช้"],
            ],
        ),
    ]


def document_detail_aggregate_extra_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 Document Section Keys"),
        p("Aggregate API ต้องคืน key มาตรฐานให้ FE ใช้ render role profile โดยไม่ต้องคำนวณสิทธิ์จากรหัส workflow ใน client"),
        table(
            ["Section key", "UI section", "Render rule"],
            [
                ["doc-header", "ข้อมูลร้านถูกกระทบ / header", "read-only ทุก role"],
                ["sec-sales", "แนวโน้มยอดขายรายวัน", "read-only ทุก role"],
                ["sec-map", "แผนที่ AllMap", "read-only ทุก role"],
                ["sec-newstore", "ร้านเปิดใหม่", "editable เมื่อ BE ส่งใน editableSections"],
                ["sec-competitor", "ร้านคู่แข่งเปิดกระทบ", "editable เมื่อ BE ส่งใน editableSections"],
                ["sec-factor", "ปัจจัยอื่นๆ", "editable เมื่อ BE ส่งใน editableSections"],
                ["sec-attach", "เอกสารแนบทั้งหมด", "upload ได้เมื่อ canUploadAttachment=true"],
                ["sec-calc", "คำนวณเงินชดเชย", "visible เมื่อ BE ส่งใน visibleSections"],
                ["sec-comp-history", "ประวัติการชดเชย", "read-only ทุก role"],
                ["sec-decision-history", "ผลการพิจารณา (ประวัติ)", "read-only ทุก role"],
                ["sec-action", "พิจารณา / ส่งดำเนินการ", "visible เมื่อ canAction=true"],
            ],
        ),
        h(2, "5.2 Role Profile Output"),
        p("BE เป็น source of truth ของ role profile แต่เอกสารนี้ไม่ฝังตาราง route workflow; รายละเอียดการแสดงผลต่อบทบาทอยู่ใน LLDD-FE-Document-Detail"),
        table(
            ["Response field", "Meaning", "FE usage"],
            [
                ["viewerRbacRoleCode", "รหัส role/RBAC ของผู้ใช้ เช่น R-01/R-02/R-10", "แสดง/trace เท่านั้น ไม่ map เป็น section"],
                ["roleProfileCode", "profile สำหรับหน้า Document Detail เช่น P-06/P-08/P-01/P-02/P-03", "เลือกชุด visible/edit/action ที่ BE คำนวณแล้ว; แยก namespace จาก statusCode"],
                ["visibleSections", "section key ที่ต้องแสดง", "ซ่อน section ที่ไม่อยู่ใน array"],
                ["editableSections", "section key ที่แก้ไขได้", "เปิด input/button เฉพาะ section เหล่านี้"],
                ["canUploadAttachment", "boolean", "เปิด/ปิด upload control"],
                ["canAction", "boolean", "เปิด/ปิด action panel"],
                ["actionOptions", "array ของ label + requireComment", "render radio โดยไม่คำนวณปลายทาง"],
            ],
        ),
    ]


def document_create_update_extra_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 docNo Generator and Concurrency Rules"),
        p("เลขเอกสารเป็น business identifier ของระบบ จึงต้อง generate ฝั่ง BE ใน transaction เดียวกับการสร้างเอกสาร และต้องไม่ให้ FE หรือ Job สร้างเลขเอง"),
        table(
            ["Rule", "Required behavior", "Implementation note"],
            [
                ["Format", "YYYY/xxxxx โดย YYYY เป็นปี **ค.ศ.** และ running 5 หลัก (มติ 2026-08-06 · หน้าจอ K2 จริงใช้ ค.ศ. เช่น 2026/01870)", "ตัวอย่าง 2026/00124; เก็บ doc_no เป็น string และเก็บ year/running_no แยกเพื่อ index"],
                ["Sequence scope", "running reset ตามปี ค.ศ.", "unique key `(year, running_no)` และ unique `doc_no`"],
                ["Lock strategy", "lock row sequence ด้วย `SELECT ... FOR UPDATE` หรือ database sequence ต่อปี", "ห้ามอ่าน max(running_no)+1 แบบไม่มี lock"],
                ["Transaction boundary", "generate docNo, insert compensation_documents, insert first workflow task และ audit ใน transaction เดียว", "ถ้าสร้าง task ไม่สำเร็จต้อง rollback ทั้งชุด"],
                ["Gap policy", "เลขที่ถูก commit แล้วห้าม reuse; rollback ก่อน commit ไม่ควรเผยแพร่ docNo ให้ client", "ถ้าใช้ native sequence ที่เกิด gap ได้ต้องบันทึก policy นี้ใน runbook"],
                ["Duplicate guard", "business key ซ้ำต้องคืน 409 ก่อน generate docNo ใหม่เมื่อเป็นไปได้", "business key อย่างน้อย impactedStoreCode+impactMonth+newStoreCode+roundNo+source"],
                ["Idempotency", "requestId ใช้ trace/retry แต่ไม่แทน duplicate business key", "ถ้า retry request เดิมหลัง success ให้คืน docNo เดิมเมื่อจับคู่ requestId ได้"],
            ],
        ),
        h(2, "5.2 Create Document Transaction Flow"),
        table(
            ["Step", "Service behavior", "Rollback / error rule"],
            [
                ["1. Validate input", "ตรวจ required, format, store exists, period, source, roundNo", "invalid คืน 400/422 ก่อน lock sequence"],
                ["2. Check duplicate", "query business key บน compensation_documents", "พบเอกสารเดิมคืน 409 DUPLICATE_DOCUMENT พร้อม docNo เดิมถ้าอนุญาตให้แสดง"],
                ["3. Start transaction", "เปิด transaction และ lock sequence row ของปี ค.ศ.", "lock timeout คืน 409/503 ตามมาตรฐาน platform"],
                ["4. Generate docNo", "เพิ่ม running_no และประกอบ doc_no", "ยังไม่ส่ง response จนกว่า commit สำเร็จ"],
                ["5. Insert document", "insert compensation_documents และ child rows เริ่มต้น", "fail ต้อง rollback sequence/document"],
                ["6. Open first task", "เรียก initializeWorkflow + addPreApprover (state 06) ของ @srm/glb-workflow ภายใน transaction boundary ที่กำหนด — ชื่อ function ตามชีต Detail ของ LLDD lib — ดู LLDD-BE-Workflow-Engine-Definition 5.3", "fail ต้อง rollback document"],
                ["7. Commit", "commit transaction (ไม่มีการเขียน audit ของ master แล้ว · ยกเลิกระบบ audit ของ master 2026-08-07)", "หลัง commit จึง return docNo/statusCode"],
            ],
        ),
        h(2, "5.3 Required Developer Tests for docNo"),
        table(
            ["Test", "Expected result"],
            [
                ["ยิง POST /sbpgi/document พร้อมกัน 20 request ในปีเดียวกัน", "ได้ docNo ไม่ซ้ำ running เรียงตาม commit และไม่มี duplicate key error ที่หลุดเป็น 500"],
                ["สร้าง duplicate business key", "คืน 409 DUPLICATE_DOCUMENT และไม่ consume docNo ใหม่ถ้า duplicate ถูกพบก่อน lock sequence"],
                ["จำลอง error หลัง insert document ก่อนเปิด workflow", "rollback แล้วไม่เหลือ compensation_documents/workflow_transaction/audit partial"],
                ["เปลี่ยนปี ค.ศ.", "running เริ่มที่ 00001 ของปีใหม่"],
            ],
        ),
        h(2, "5.4 docNo Generator SQL Reference"),
        code("""-- ออกเลขเอกสาร YYYY/xxxxx แบบ atomic ต่อ "ปี ค.ศ." (ห้ามใช้ พ.ศ. — ดู api.md มติ 2026-08-06)
-- ตารางจริงคือ document_running_numbers (year · last_running_no · updated_at) ไม่มีคอลัมน์ created_at

-- 1) สร้างแถวของปีนี้ถ้ายังไม่มี (idempotent)
INSERT INTO document_running_numbers (year, last_running_no)
VALUES (:year, 0)
ON CONFLICT (year) DO NOTHING;

-- 2) กินเลขถัดไปในทรานแซกชันเดียวกับการสร้างเอกสาร — UPDATE ... RETURNING ล็อกแถวให้เอง
--    กัน batch (Job 8) กับผู้ใช้สร้างพร้อมกันแล้วได้เลขชนกัน
UPDATE document_running_numbers
SET last_running_no = last_running_no + 1,
    updated_at = CURRENT_TIMESTAMP
WHERE year = :year
RETURNING last_running_no;          -- → :runningNo

-- 3) docNo = :year || \'/\' || lpad(:runningNo::text, 5, \'0\')   เช่น 2026/00123
--    ⚠️ ต้องใส่ impact_process_id ทุกครั้ง — เป็น NOT NULL UNIQUE (หนึ่ง impact process = หนึ่งเอกสาร)
INSERT INTO compensation_documents (
    doc_no, year, running_no,
    impact_process_id, impacted_store_code, impact_month, new_store_code,
    round_no, source, status_code, current_section_code, created_by
) VALUES (
    :docNo, :year, :runningNo,
    :impactProcessId, :impactedStoreCode, :impactMonth, :newStoreCode,
    :roundNo, :source, :statusInit, \'06\', :userId
);
-- created_at / total_compensation_amount / version_no มี DEFAULT อยู่แล้ว ไม่ต้องส่ง""", "sql"),
    ]


def attachment_storage_extra_blocks() -> list[dict[str, Any]]:
    return [
        h(2, "5.1 Attachment Storage and Security Design"),
        p(
            "Attachment API จัดการ binary file จริง ไม่ใช่บันทึกแต่ metadata — แต่ **SBPGI ไม่ได้เป็นเจ้าของ storage layer** · "
            "ตามมติ **DP-8 (ปิด 2026-08-24)** SBPGI เก็บแค่ metadata ใน `document_attachments` ของตัวเอง แล้ว "
            "**ยืม service S3 ของระบบ SBP เดิม** (`POST /statement/upload-file-aws` · `download-file-aws`) — "
            "สิ่งที่ SBPGI เป็นเจ้าของจริงคือ **validation · authorization · metadata · การแปลงเป็น stream ให้ FE**"
        ),
        p(
            "🔴 **wrapper ของระบบเดิมเป็น base64 ไม่ใช่ stream** (ตรวจ `store-backend` 2026-08-26) — สายส่งจริงคือ "
            "`FE ← binary stream ← SBPGI BE ← base64 JSON ← /statement/{upload,download}-file-aws ← S3` · "
            "ไฟล์ 5 MB จะกลายเป็น ~6.7 MB ใน JSON (body limit ของ store-backend คือ 100 MB) · "
            "ปุ่ม **ดาวน์โหลดทั้งหมด (.zip)** ห้ามโหลดทุกไฟล์เข้า memory พร้อมกัน ให้ดึงทีละไฟล์แล้ว stream เข้า zip "
            "— รายละเอียดเต็มอยู่ที่ **LLDD-BE-Integration-SBP-Platform** 5.3"
        ),
        table(
            ["Item", "Required value / convention", "Developer note"],
            [
                ["Storage provider", "**service S3 ของระบบ SBP เดิม** (`AwsService` ผ่าน `/statement/{upload,download}-file-aws`)",
                 "🔴 มติ DP-8 — SBPGI **ห้ามสร้าง storage adapter/ไม่ต่อ S3 SDK เอง** และไม่ต้องเลือก vendor · เก็บ `storage_provider` ไว้เป็น metadata เผื่ออนาคตเท่านั้น"],
                ["Bucket/container", "**bucket ของระบบเดิม** (ทีม SBP เป็นผู้กำหนด)",
                 "⚠️ ต้องยืนยันกับทีม store-backend ว่าไฟล์ของ SBPGI แยก prefix/bucket หรือปนกับของเดิม — lifecycle/backup เป็นของ infra ฝั่งนั้น"],
                ["Object key", "`documents/{year}/{docNoSafe}/{attachId}/{sha256Prefix}-{safeFileName}`", "`docNoSafe` แทน `/` ด้วย `-`; sanitize filename ก่อนใช้ใน key"],
                ["Quarantine / AV", "⚠️ **ยังไม่ยืนยันว่าแพลตฟอร์มมีตัวสแกน**", 
                 "🔴 ไม่พบ AV scanner ในเอกสารวิเคราะห์ระบบเดิมเลย · จนกว่าจะยืนยัน ให้ `scan_status` เริ่มที่ `PENDING` และ **ตัดสินร่วมกับทีม infra** ว่าจะสแกนที่ไหน (ฝั่ง S3 event · ฝั่ง SBPGI · หรือยอมรับความเสี่ยง) — ห้ามสมมติว่ามีของให้ใช้แล้ว"],
                ["Allowed extension", ATTACHMENT_ALLOWED_EXTENSIONS, "ตรวจทั้ง extension และ content type/magic bytes เท่าที่ platform รองรับ"],
                ["AV scan status", "PENDING -> CLEAN หรือ BLOCKED/FAILED", "download อนุญาตเฉพาะ CLEAN; BLOCKED/FAILED คืน FILE_SCAN_BLOCKED"],
                ["Max size", "5 MB ต่อไฟล์", "เกินให้คืน 413 FILE_TOO_LARGE ก่อน upload เข้า storage"],
            ],
        ),
        h(2, "5.2 Attachment Metadata Fields"),
        table(
            ["Field", "Meaning", "Required behavior"],
            [
                ["attachId", "primary key/identifier", "คืนให้ FE หลัง upload"],
                ["docNo", "เลขเอกสาร", "attachment ต้อง belong กับ document นี้เท่านั้น"],
                ["sectionCode", "workflow section ตอน upload", "บันทึกจาก request และ validate กับ current task/permission"],
                ["originalFileName", "ชื่อไฟล์จากผู้ใช้", "เก็บเพื่อแสดงผลและ Content-Disposition"],
                ["contentType", "MIME type", "ใช้ร่วมกับ extension validation"],
                ["fileSizeBytes", "ขนาดไฟล์", "ต้อง <= 5 MB"],
                ["storageProvider/bucketName/objectKey", "ตำแหน่ง binary", "ห้าม expose objectKey ตรงให้ FE"],
                ["sha256", "checksum", "ใช้ตรวจ duplicate/corruption"],
                ["scanStatus/scannedAt/scanMessage", "ผล AV scan", "download ได้เฉพาะ CLEAN"],
                ["uploadedBy/uploadedAt/deletedFlag", "audit metadata", "soft delete เท่านั้นเมื่อมีการลบภายหลัง"],
            ],
        ),
        h(2, "5.3 Upload Flow"),
        table(
            ["Step", "Backend behavior", "Error / response"],
            [
                ["1. Authorize", "ตรวจผู้ใช้มีสิทธิ์อ่านเอกสารและ canUploadAttachment/current task owner", "ไม่มีสิทธิ์คืน 403"],
                ["2. Validate multipart", "ตรวจ file present, size, extension, content type, sectionCode", "คืน 400/413/415 ตาม catalog"],
                ["3. Hash + ส่งขึ้น storage", "คำนวณ sha256 จาก buffer แล้วเรียก `POST /statement/upload-file-aws` (**ส่งเป็น base64**) เก็บ objectKey ที่ได้กลับมา", "service ของระบบเดิม fail คืน 503 และ **ไม่ insert metadata**"],
                ["4. Scan", "⚠️ **ขึ้นกับข้อค้าง AV ด้านบน** — ถ้ายังไม่มีตัวสแกน ให้ค้างที่ `PENDING` ตามนโยบายที่ตกลง", "พบไวรัส (เมื่อมีตัวสแกน) ตั้ง BLOCKED และคืน FILE_SCAN_BLOCKED"],
                ["5. Insert metadata", "insert `document_attachments` พร้อม objectKey · sha256 · scan_status", "🔴 ไม่มีขั้น move/promote — wrapper ของระบบเดิมไม่มี API ย้าย object ให้ SBPGI เรียก"],
                ["6. Respond", "คืน attachId, fileName, fileSizeBytes, scanStatus, uploadedAt", "ไม่คืน bucket/objectKey ให้ FE"],
            ],
        ),
        h(2, "5.4 Download Flow and Authorization"),
        table(
            ["Step", "Backend behavior", "Error / response"],
            [
                ["1. Validate path", "ตรวจ docNo/attachId และ attachment belongs to docNo", "ไม่พบคืน 404"],
                ["2. Authorize read", "สิทธิ์เท่ากับ document read หรือ report/admin ที่ได้รับสิทธิ์", "ไม่มีสิทธิ์คืน 403"],
                ["3. Check scan", "อนุญาตเฉพาะ `scan_status` ที่นโยบายกำหนดว่าดาวน์โหลดได้ และ `deleted_flag = false`",
                 "🔴 **ถ้ายังไม่มี AV scanner** (ดูข้อค้างใน 5.1) การบังคับ `CLEAN` อย่างเดียวจะทำให้ **ดาวน์โหลดไม่ได้เลยทั้งระบบ** — ต้องเคาะนโยบายก่อน go-live: เปิดให้ `PENDING` ดาวน์โหลดได้ หรือรอตัวสแกน · BLOCKED/FAILED คืน 422 FILE_SCAN_BLOCKED เสมอ"],
                ["4. Stream", "เรียก `POST /statement/download-file-aws` ได้ **base64** แล้ว decode เป็น buffer ก่อน stream ออกไป",
                 "🔴 **ไม่มี signed URL ให้ใช้** — wrapper ของระบบเดิมไม่คืน presigned url · ตั้ง Content-Type และ Content-Disposition จาก metadata"],
                ["5. Audit", "บันทึกร่องรอยการดาวน์โหลดที่ **application log** (structured)",
                 "🔴 ตาราง `audit_logs` ถูกตัดไปแล้ว 2026-08-07 — ห้ามอ้างตารางนี้ · ต้อง trace userId/docNo/attachId/requestId ได้จาก log"],
            ],
        ),
        h(2, "5.5 Download Endpoint Contract"),
        table(
            ["Method", "Path", "Response"],
            [
                ["GET", "/api/v1/sbpgi/document/{docNo}/attachments/{attachId}/download", "binary stream; headers Content-Type, Content-Length, Content-Disposition"],
            ],
        ),
        h(2, "5.6 Attachment Repository SQL Reference"),
        code("""-- Insert metadata after storage write and AV scan pass.
INSERT INTO document_attachments (
    doc_no, section_code, file_name, mime_type, file_size,
    storage_provider, bucket, object_key, sha256,
    scan_status, scanned_at, uploaded_by, uploaded_at, deleted_flag
) VALUES (
    :docNo, :sectionCode, :fileName, :mimeType, :fileSize,
    :storageProvider, :bucket, :objectKey, :sha256,
    'CLEAN', CURRENT_TIMESTAMP, :userId, CURRENT_TIMESTAMP, 'N'
)
RETURNING attach_id;

-- Load attachment for download. Authorization is checked in service before streaming.
SELECT
    attach_id, doc_no, file_name, mime_type, file_size,
    storage_provider, bucket, object_key, sha256, scan_status
FROM document_attachments
WHERE doc_no = :docNo
  AND attach_id = :attachId
  AND deleted_flag = 'N';""", "sql"),
    ]


def read_js_value_from_html(html_file: str, var_name: str, open_char: str, close_char: str) -> Any:
    html = (ROOT / html_file).read_text(encoding="utf-8")
    marker = f"var {var_name} = "
    start = html.index(marker) + len(marker)
    bracket_start = html.index(open_char, start)
    depth = 0
    in_str: str | None = None
    escape = False
    end = bracket_start
    for i, ch in enumerate(html[bracket_start:], bracket_start):
        if in_str:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == in_str:
                in_str = None
            continue
        if ch in ("'", '"', "`"):
            in_str = ch
            continue
        if ch == open_char:
            depth += 1
        elif ch == close_char:
            depth -= 1
            if depth == 0:
                end = i + 1
                break
    js_text = html[bracket_start:end]
    node = (
        "let s='';"
        "process.stdin.setEncoding('utf8');"
        "process.stdin.on('data', c => s += c);"
        "process.stdin.on('end', () => {"
        "  const value = (new Function('return ' + s))();"
        "  console.log(JSON.stringify(value));"
        "});"
    )
    proc = subprocess.run(["node", "-e", node], input=js_text, text=True, capture_output=True, check=True)
    return json.loads(proc.stdout)


def read_js_array_from_html(html_file: str, var_name: str) -> list[dict[str, Any]]:
    return read_js_value_from_html(html_file, var_name, "[", "]")


def read_js_object_from_html(html_file: str, var_name: str) -> dict[str, Any]:
    return read_js_value_from_html(html_file, var_name, "{", "}")


def load_batch_jobs() -> list[dict[str, Any]]:
    jobs = read_js_array_from_html("job-batch.html", "JOBS")
    by_no = {str(j["no"]): j for j in jobs}
    selected: list[dict[str, Any]] = []
    for no in ["2", "3", "4", "5", "6", "7", "8", "8b", "9", "10"]:
        if no not in by_no:
            continue
        selected.append(target_job(dict(by_no[no])))
    return selected


def draw_flow_diagram(title: str, steps: list[str], out_rel: str) -> str:
    out_path = ROOT / out_rel
    out_path.parent.mkdir(parents=True, exist_ok=True)
    font_path = FONT if os.path.exists(FONT) else "/System/Library/Fonts/Supplemental/Arial.ttf"
    title_font = ImageFont.truetype(font_path, 24)
    body_font = ImageFont.truetype(font_path, 17)
    small_font = ImageFont.truetype(font_path, 13)
    multi_col = len(steps) > 9
    width = 1600 if multi_col else 1400
    box_w = 680 if multi_col else 1020
    box_h = 74
    gap = 38
    top = 92
    rows = math.ceil(len(steps) / 2) if multi_col else len(steps)
    height = top + rows * (box_h + gap) + 48
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 34), title, fill="#0B2545", font=title_font)
    col_gap = 80
    x_positions = [(width - (box_w * 2 + col_gap)) // 2, (width - (box_w * 2 + col_gap)) // 2 + box_w + col_gap] if multi_col else [(width - box_w) // 2]

    def wrap_text(text: str, max_px: int) -> list[str]:
        words = re.split(r"(\s+)", text)
        lines: list[str] = []
        current = ""
        for token in words:
            candidate = current + token
            if d.textlength(candidate, font=body_font) <= max_px:
                current = candidate
            else:
                if current.strip():
                    lines.append(current.strip())
                current = token.strip()
        if current.strip():
            lines.append(current.strip())
        return lines[:3]

    for idx, step in enumerate(steps):
        col = 1 if multi_col and idx >= rows else 0
        row_idx = idx - rows if col else idx
        x = x_positions[col]
        y = top + row_idx * (box_h + gap)
        fill = "#EEF4FF"
        outline = "#96B4F0"
        if "?" in step or "ผ่าน" in step or "สำเร็จ" in step:
            fill = "#FFFAF0"
            outline = "#E2B45A"
        if any(word in step for word in ["ล้มเหลว", "Rollback", "error", "FAIL", "ค้าง", "ไม่ผ่าน"]):
            fill = "#FDECEC"
            outline = "#D98A8A"
        d.rounded_rectangle([x, y, x + box_w, y + box_h], radius=16, fill=fill, outline=outline, width=3)
        d.ellipse([x + 20, y + 20, x + 54, y + 54], fill="#2F6FED")
        d.text((x + 37, y + 37), str(idx + 1), fill="white", font=small_font, anchor="mm")
        lines = wrap_text(step, box_w - 116)
        line_top = y + 18 if len(lines) <= 2 else y + 10
        for line_idx, line in enumerate(lines):
            d.text((x + 76, line_top + line_idx * 21), line, fill="#2B3440", font=body_font)
        is_last_in_col = multi_col and idx == rows - 1
        if idx < len(steps) - 1 and not is_last_in_col:
            ax = x + box_w // 2
            ay1 = y + box_h + 5
            ay2 = y + box_h + gap - 8
            d.line([ax, ay1, ax, ay2], fill="#8493A5", width=3)
            d.polygon([(ax - 8, ay2 - 2), (ax + 8, ay2 - 2), (ax, ay2 + 12)], fill="#8493A5")
        elif is_last_in_col and idx < len(steps) - 1:
            ax1 = x + box_w + 10
            ay = y + box_h // 2
            ax2 = x_positions[1] - 14
            d.line([ax1, ay, ax2, ay], fill="#8493A5", width=3)
            d.polygon([(ax2 - 2, ay - 8), (ax2 - 2, ay + 8), (ax2 + 12, ay)], fill="#8493A5")
    img.save(out_path)
    return out_rel


def flow_steps_from_job(job: dict[str, Any]) -> list[str]:
    steps = []
    for item in job.get("flow", []):
        text = item.get("t", "")
        if item.get("no"):
            text += f" | No: {item['no']}"
        if item.get("d"):
            text += f" ({item['d']})"
        steps.append(text)
    return steps


def job_topic(job: dict[str, Any]) -> Topic:
    no = str(job["no"])
    owner = JOB_OWNER_OVERRIDES.get(no, BANK_BE_OWNER)
    estimated_hours = JOB_ESTIMATES[no]
    params = job.get("params", [])
    fields = [(p[0], p[1], "แก้ไขได้" if p[3] else "ค่าคงที่/แก้ผ่านหน้าจอไม่ได้", p[4] if len(p) > 4 else "") for p in params]
    db_tables = [tuple(row) for row in job.get("tables", [])]
    flow = flow_steps_from_job(job)
    flow_file = f"LLDD/assets/flows/BE-Job-{no}-{sanitize_filename(job['name'])}.png"
    scope = [
        f"Main class/script: {job.get('cls', '-')} / {job.get('script', '-')}",
        f"Phase: {job.get('phase', '-')}",
        f"Output: {job.get('out', '-')}",
        f"Estimate: {estimated_hours} ชั่วโมง",
        "พารามิเตอร์/cron อ่านจาก backend config (config file/env) — ไม่มีตาราง job_configs และไม่มีหน้าจอควบคุม (หน้า Flow Batch Job ในกลุ่มเมนู Flow เหลือแค่ Flowchart + Database ที่ใช้ · 2026-08-06)",
        "Runbook, rerun rule, risk และ history ตามเอกสาร Batch v4.0 · ผลการรันเขียน application log แบบ structured",
    ]
    if no == "8b":
        scope.append("Depends on LLDD-BE-API-Workflow-Instances; Job 8b เรียก Workflow Engine ภายในและไม่ duplicate Gen Flow Gate logic")
    return Topic(
        f"BE/Jobs/LLDD-BE-Job-{no}-{sanitize_filename(job['name'])}",
        f"LLDD BE - Job {no} {job['name']}",
        "BE",
        round(estimated_hours / HOURS_PER_DAY, 1),
        estimated_hours,
        owner,
        f"{job.get('th', job['name'])}: {job.get('desc', '')}",
        [],
        scope,
        fields or [("jobNo", no, "required", "รหัสงาน Batch")],
        [
            ("รันตามตารางเวลา", "CRON", f"scheduler → runner (job {no})", "อ่าน cron/พารามิเตอร์จาก backend config"),
            ("รันนอกรอบ (manual/rerun)", "CLI", f"CLI/ops runbook → runner (job {no})", "guard ไม่ให้รันซ้อนด้วย distributed lock"),
            ("แก้พารามิเตอร์/เปิด-ปิด job", "CONFIG", "แก้ backend config แล้ว deploy", "ไม่มี endpoint และไม่มีหน้าจอควบคุม — หน้า Flow Batch Job เป็น reference อย่างเดียว (2026-08-06)"),
            ("ตรวจผลการรัน", "LOG", "application log (structured)", "ไม่มีตาราง job_run_histories แล้ว · ไฟล์/ACK ดูที่ interface_transactions"),
        ],
        [],
        flow,
        [
            "พารามิเตอร์และ cron อ่านจาก backend config เท่านั้น — เปลี่ยนค่าโดย deploy config ไม่ใช่ผ่าน API/หน้าจอ",
            "การรันต้องตรวจ enabled flag ใน config และกันรันซ้อนด้วย distributed/advisory lock",
            "ทุกรอบต้องเขียน application log แบบ structured (เวลา/แถว/ไฟล์/ผล) และ error ต้องส่ง EM-07",
            "DB/table mapping ใช้เป็น reference สำหรับ implement Job เท่านั้น ไม่ใช่งานสร้างหน้า Database",
            "รองรับ rerun rule และ risk note ตาม runbook",
        ],
        [
            "รันตามตารางเวลาแล้วผลถูกต้องบน fixture",
            "รันนอกรอบผ่าน CLI ได้ผลเดียวกับ cron",
            "สั่งรันซ้อนขณะกำลังรัน → runner ปฏิเสธ (lock ทำงาน)",
            "แก้ config แล้ว deploy → รอบถัดไปใช้ค่าใหม่",
            "job throw error → EM-07 ออก และ log มีบรรทัด error",
            "ตรวจผลกระทบตารางตาม R/W mapping reference",
        ],
        db_tables=db_tables,
        flow_diagram=flow_file,
    )


def be_job_topics() -> list[Topic]:
    return [job_topic(job) for job in load_batch_jobs()]


# --------------------------------------------------------------------------------------
# 2026-08-07: เอกสาร BE ที่เพิ่มใหม่ 4 ฉบับ (Database-Structure · Data-Migration-Cutover ·
# Integration-SBP-Platform · Workflow-Engine-Definition) — ทั้ง 4 ฉบับไม่มี REST endpoint
# ของตัวเอง จึงไม่ผ่าน skeleton generator (ดู SKELETON_SKIP_FILES) และใช้เนื้อหาจาก
# topic_extra_blocks() แทน
# --------------------------------------------------------------------------------------
def new_be_design_topics() -> list[Topic]:
    return [
        Topic(
            "BE/LLDD-BE-Database-Structure",
            "LLDD BE - Database Structure and Deployment",
            "BE",
            4.0,
            24,
            BANK_BE_OWNER,
            "กำหนด DDL ของ target schema 20 ตาราง พร้อม index/constraint/seed และสคริปต์ deploy ให้ทุกเอกสาร BE อ้างอิงโครงเดียวกัน — เป็น blocker ที่ต้องปิดในสัปดาห์แรก",
            [],
            [
                "DDL ครบ 20 ตารางของ target schema (โซน A 8 · โซน B 9 · โซน C 3)",
                "Index, unique/partial index, check constraint และ FK ที่ต้องมีก่อน SIT",
                "Seed data ที่ต้องมีก่อนเปิดระบบ (external_factors · competitors) — decisions ไป seed ที่ common_code ของระบบเดิม (DP-9)",
                "สคริปต์ deploy/rollback ต่อ environment และลำดับการรันตาม dependency",
                "ตารางที่ห้ามสร้างซ้ำเพราะระบบ SBP เดิมมีอยู่แล้ว (workflow engine 13 ตาราง · store/mas_store · common_code · mas_param · business_user · email_template · fcs_qssi_score)",
                "บันทึกข้อค้างตัดสินใจด้านโครงสร้างข้อมูล — ยังไม่ตัดสิน",
            ],
            [
                ("naming", "lower_snake_case", "บังคับทุกตาราง/คอลัมน์ใหม่", "ห้ามใช้ชื่อไทย/CamelCase หรือชื่อ legacy แบบ FGI_/Comp*"),
                ("store_code / new_store_code", "VARCHAR(5)", "ห้ามเก็บเป็น numeric", "ต้องคง leading zero (00788) ทุกตาราง"),
                ("doc_no", "VARCHAR(12) รูปแบบ YYYY/xxxxx", "unique ต่อปี", "ออกเลขผ่าน document_running_numbers แบบ atomic"),
                ("amount / percent", "NUMERIC(15,2) / NUMERIC(5,2)", "amount >= 0 · percent 0-100", "ผลรวม compensate_percent ต่อเอกสารต้อง = 100"),
                ("period key", "CHAR(7) 'YYYY-MM' (ค.ศ.)", "ค่าคงรูปแบบเดียวทั้ง schema", "แสดงผลเป็น ค.ศ. เช่นกัน"),
                ("fcs_qssi_score", "ตารางเดิมของ sps_store", "ห้าม CREATE TABLE ใหม่", "มีอยู่จริง 23,958,780 แถว + import pipeline ใช้งานอยู่ (POST /performance/import-qssi · staging fcs_tmp_qssi_score)"),
            ],
            [
                ("รัน DDL baseline", "deploy script", "psql -f 01_schema.sql", "สร้าง 19 ตารางตามลำดับ dependency"),
                ("รัน index/constraint", "deploy script", "psql -f 02_index.sql", "index/unique/check ครบก่อนเปิด SIT"),
                ("รัน seed", "deploy script", "psql -f 03_seed.sql", "master ที่ระบบต้องมีตั้งแต่วันแรก"),
                ("Rollback", "deploy script", "psql -f 99_rollback.sql", "DROP ย้อนลำดับ · ห้ามแตะตารางของระบบ SBP เดิม"),
            ],
            [],
            [
                "ยืนยันรายการ 20 ตารางกับ database.md และ LLDD-Database ให้ตรงกันก่อนเขียน DDL",
                "เขียน 01_schema.sql เรียงตาม dependency: โซน C master -> โซน A pipeline -> โซน B document",
                "เขียน 02_index.sql แยกไฟล์ เพื่อให้ rerun/เพิ่ม index ภายหลังได้โดยไม่แตะ schema",
                "เขียน 03_seed.sql เฉพาะ master ที่ระบบต้องมีตั้งแต่วันแรก",
                "ตรวจว่าไม่มี CREATE TABLE ของตารางที่ระบบ SBP เดิมมีอยู่แล้ว",
                "รันบน environment ว่างแล้ว dump schema กลับมาเทียบกับ DDL ต้นฉบับ",
                "ส่งมอบ DDL ให้ Data-Migration-Cutover ใช้เป็นปลายทาง",
            ],
            [
                "DDL รันบนฐานว่างได้ครบในครั้งเดียวโดยไม่มี error ลำดับ FK",
                "จำนวนตารางที่สร้างจริง = 19 ตาราง (20 ในโครง ลบ fcs_qssi_score ที่ reuse) ตรงกับ database.md",
                "ไม่มี CREATE TABLE ของ workflow engine, store master, common_code, mas_param, business_user, email_template หรือ fcs_qssi_score",
                "ทุกตารางมี PK และทุก FK ชี้ไปตารางที่มีอยู่จริงในสคริปต์เดียวกัน",
                "rollback script ลบเฉพาะตารางของ SBPGI",
                "ข้อค้างตัดสินใจถูกบันทึกเป็นข้อค้าง ไม่ถูกตัดสินในเอกสารนี้",
            ],
            [
                "รัน 01+02+03 บนฐานว่าง แล้ว dump schema เทียบกับต้นฉบับ",
                "รัน 01 ซ้ำครั้งที่สอง ต้อง fail แบบชัดเจน ไม่สร้างของซ้ำเงียบ ๆ",
                "ทดสอบ insert เอกสารที่ compensate_percent รวมไม่ครบ 100 ต้องถูก block",
                "ทดสอบ insert store_code '00788' แล้วอ่านกลับได้ leading zero ครบ",
                "ทดสอบออกเลข doc_no พร้อมกัน 20 request ต้องไม่ซ้ำ",
                "grep หา CREATE TABLE ของตาราง reuse ต้องได้ 0 บรรทัด",
            ],
            db_tables=[
                ("20 target tables (โซน A/B/C)", "W", "สร้างจาก DDL baseline ของเอกสารนี้"),
                ("workflow engine 13 ตาราง (sps_store)", "R", "ห้ามสร้างซ้ำ — ใช้ของ @srm/glb-workflow"),
                ("fcs_qssi_score (sps_store)", "R", "ห้ามสร้างซ้ำ — 23,958,780 แถว + import pipeline ใช้งานอยู่"),
                ("mas_param / common_code / business_user / email_template (sps_store)", "R", "ค่ากำหนดกลาง/master/ตัวตน/เทมเพลตอีเมลของระบบ SBP เดิม"),
            ],
        ),
        Topic(
            "BE/LLDD-BE-Data-Migration-Cutover",
            "LLDD BE - Data Migration and Cutover",
            "BE",
            5.0,
            30,
            BANK_BE_OWNER,
            "ออกแบบการย้ายข้อมูลจากระบบเดิม (Oracle FCS_FRN ฝั่ง FGI/FCS + SQL Server CPA_FRN_FGI ฝั่ง K2) เข้าสู่ target schema ของ SBPGI พร้อมแผน cutover, reconcile และ rollback",
            [],
            [
                "Source-to-target mapping ระดับตาราง/คอลัมน์ (ORA FCS_FRN · MSSQL CPA_FRN_FGI -> 20 ตาราง)",
                "การแปลงคีย์: polymorphic TRANSACTION_PK -> typed FK · CompDocumentID -> doc_no · IMPACT_PROCESS_ID -> impact_process_id",
                "แผน cutover เป็นรอบ (dry-run -> delta -> freeze -> final) และ rollback",
                "Reconcile: นับแถว ยอดเงิน และ checksum ต่อโซน",
                "การย้าย workflow ที่ยังวิ่งอยู่เข้าสู่ @srm/glb-workflow",
                "บันทึกข้อค้างตัดสินใจด้าน migration — ยังไม่ตัดสิน",
            ],
            [
                ("source ORA", "Oracle FCS_FRN", "read-only ตอน migrate", "ฝั่ง FGI/FCS pipeline (FGI_IMPACT_* · FCS_QSSI_SCORE · FGI_CONFIRM_RECEIVE_DATA)"),
                ("source MSSQL", "SQL Server CPA_FRN_FGI", "read-only ตอน migrate", "ฝั่ง K2 document (CompensateFlow · CompensateHistory · ImpactProfile · ImpactCostDetail · RunningNumber)"),
                ("business key", "impacted_store_code + month + year", "ต้อง unique หลังแปลง", "ใช้เป็นคีย์ dedup ตอน load โซน A"),
                ("doc_no", "YYYY/xxxxx (**ค.ศ.** · มติ 2026-08-06)", "ต้อง unique", "แปลงจาก CompDocumentID — ถ้าของเดิมเป็น พ.ศ. ต้องแปลงเป็น ค.ศ. ตอน migrate · ตั้งค่า document_running_numbers.last_running_no ต่อปี (ค.ศ.) ให้ตรงกับเลขสูงสุดที่ย้ายมา"),
                ("date", "เก็บเป็น ค.ศ. ใน DB", "แปลงจาก พ.ศ. ของระบบเดิมด้วย toAD()", "FE แสดง ค.ศ. เป็นค่าเริ่มต้น"),
                ("store_code", "VARCHAR(5)", "lpad 5 หลัก", "ระบบเดิมบางตารางเก็บเป็นตัวเลข ทำให้ leading zero หาย"),
            ],
            [
                ("Dry-run migrate", "runbook", "สคริปต์ ETL โหมด --dry-run", "ได้รายงานจำนวนแถว/แถวที่ reject โดยไม่เขียนปลายทาง"),
                ("Full load", "runbook", "สคริปต์ ETL โหมด --full", "โหลดข้อมูลย้อนหลังทั้งหมดเข้า target schema"),
                ("Delta load", "runbook", "สคริปต์ ETL โหมด --delta --since", "โหลดเฉพาะรายการที่เปลี่ยนหลัง full load"),
                ("Reconcile", "runbook", "สคริปต์ reconcile", "เทียบจำนวนแถว/ยอดเงินต้นทาง-ปลายทางต่อโซน"),
                ("Rollback", "runbook", "restore snapshot ก่อน cutover", "กลับไปใช้ระบบเดิมได้ภายในหน้าต่างที่ตกลง"),
            ],
            [],
            [
                "ยืนยันปลายทางกับ LLDD-BE-Database-Structure (DDL ต้องนิ่งก่อน)",
                "ทำ profiling ต้นทาง: นับแถว/ค่า null/ค่าซ้ำของทุกตารางที่จะย้าย",
                "เขียน mapping ต่อคอลัมน์ พร้อมกฎแปลง (พ.ศ.->ค.ศ. · lpad store_code · polymorphic key -> typed FK)",
                "Dry-run บน environment ทดสอบ แล้วแก้ reject rule จนแถวที่ reject อธิบายได้ทุกแถว",
                "Full load + reconcile ครั้งที่ 1",
                "Freeze ระบบเดิม -> delta load -> reconcile ครั้งสุดท้าย",
                "ย้าย workflow ที่ยังวิ่งอยู่: initialize transaction ใน @srm/glb-workflow ให้ตรง state ปัจจุบันของเอกสาร",
                "เปิดระบบใหม่ · เก็บ snapshot ก่อน cutover ไว้สำหรับ rollback ตามหน้าต่างที่ตกลง",
            ],
            [
                "จำนวนแถวปลายทางเท่าต้นทางทุกตาราง หรืออธิบายส่วนต่างได้ทุกแถว",
                "ยอดเงินชดเชยรวมต้นทาง = ปลายทาง (เทียบต่อปีและต่อร้าน)",
                "ไม่มี store_code ที่ leading zero หาย",
                "ไม่มี doc_no ซ้ำ และ document_running_numbers ต่อปีตรงกับเลขสูงสุดที่ย้ายมา",
                "เอกสารที่ยังไม่จบ flow เปิดในระบบใหม่แล้วอยู่ state เดิมและมีผู้อนุมัติปัจจุบันถูกคน",
                "มี rollback plan ที่ทดสอบแล้วอย่างน้อย 1 ครั้ง",
            ],
            [
                "dry-run แล้วรายงาน reject อธิบายได้ครบทุก reason code",
                "full load + reconcile ผ่านบน dataset จริงชุด staging",
                "delta load ซ้ำ 2 รอบต้อง idempotent (ไม่เกิดแถวซ้ำ)",
                "ทดสอบร้านที่ store_code ขึ้นต้นด้วย 0",
                "ทดสอบเอกสารที่มีหลายรอบ (round_no/loop_no) ว่าลำดับไม่สลับ",
                "ทดสอบ rollback: restore snapshot แล้วระบบเดิมกลับมาใช้งานได้",
            ],
            db_tables=[
                ("ORA FCS_FRN (FGI_IMPACT_* · FCS_QSSI_SCORE · FGI_CONFIRM_RECEIVE_DATA)", "R", "ต้นทางฝั่ง FGI/FCS"),
                ("MSSQL CPA_FRN_FGI (CompensateFlow · CompensateHistory · ImpactProfile · ImpactCostDetail · RunningNumber)", "R", "ต้นทางฝั่ง K2 document"),
                ("20 target tables (โซน A/B/C)", "W", "ปลายทางตาม DDL ของ LLDD-BE-Database-Structure"),
                ("workflow_transaction / workflow_approver / workflow_history (sps_store)", "W (ผ่าน lib)", "เปิด transaction ให้เอกสารที่ยังไม่จบ flow ด้วย initializeWorkflow() + addPreApprover() — **ห้าม INSERT ตรง** แม้เป็นสคริปต์ย้ายข้อมูล"),
                ("fcs_monthly_sales (sps_store)", "R", "ใช้ cross-check ยอดขายรายเดือนเท่านั้น — แทนยอดขายรายวันไม่ได้"),
            ],
        ),
        Topic(
            "BE/LLDD-BE-Integration-SBP-Platform",
            "LLDD BE - Integration with SBP Platform",
            "BE",
            3.0,
            18,
            BE_OWNER,
            "กำหนดวิธีที่ SBPGI ต่อกับแพลตฟอร์ม SBP เดิม: BFF header/ตัวตน, response envelope, ไฟล์บน S3, อีเมลผ่าน @gosoft-sbp/email-lib และค่ากำหนดกลางใน mas_param/common_code — เป็น blocker ที่ต้องปิดในสัปดาห์แรก",
            [],
            [
                "ตัวตนผู้ใช้จาก BFF header 6 ตัว (x-api-key · x-user-id · x-user-group-id · x-user-full-name · x-user-permissions · accept-language) — ดูค่าตัวอย่างจริงใน 5.1",
                "Response envelope ของ store-backend: {success, data} / {success:false, data:null, error:{code,message}}",
                "ไฟล์แนบผ่าน service S3 เดิม (POST /statement/upload-file-aws · download-file-aws)",
                "อีเมลผ่าน @gosoft-sbp/email-lib + ตาราง email_template / email_sent",
                "ค่ากำหนดกลางที่ sps_store.mas_param และ sps_store.common_code — 🔴 ค่าของ SBPGI (SBPGI_APPROVE_LIMIT ฯลฯ) ยังไม่มีในระบบจริง ต้อง seed เองตอน setup (ดู 5.5)",
                "การใช้ตาราง master ของระบบเดิม (store/mas_store · business_user · common_code) และปริมาณข้อมูลจริง",
            ],
            [
                ("x-api-key", "string", "required ทุก request จาก BFF", "ตรวจที่ guard ของ store-backend ก่อนเข้า controller"),
                ("x-user-id", "string เช่น `0000123456`", "required ทุก endpoint ของผู้ใช้ — ไม่มี = 401", "created_by/updated_by ของ SBPGI + consideration_logs.actor_user_id + ส่งเป็น userId เข้า engine · 🔴 ห้ามเขียน current_approver เอง (engine เป็นคนเขียน)"),
                ("x-user-group-id", "string เช่น `08`", "required เมื่อ endpoint ต้องรู้ section", "map เป็น section_code ของ workflow (06/08/01/02/03) — เป็นด่านหลักในการตัดสินสิทธิ์เขียน"),
                ("x-user-full-name", "string · **%-encoded**", "ไม่บังคับ", "ชื่อผู้ทำรายการใน timeline/อีเมล · 🔴 ต้อง decodeURIComponent ก่อนใช้เสมอ · ไม่มีให้ fallback เป็น x-user-id"),
                ("x-user-permissions", "string (serialized) · ⚠️ รูปแบบยังไม่ยืนยัน", "**ด่านเสริม ไม่ใช่ด่านเดียว**", "สิทธิ์ต่อ URL จาก auth-backend — SBPGI ไม่คำนวณสิทธิ์เมนูเอง · parse ไม่ผ่านให้ตกไปใช้ x-user-group-id + สถานะเอกสาร (ดู 5.1.2)"),
                ("accept-language", "string เช่น `th`", "ไม่บังคับ", "ภาษาข้อความ error — default th (ไทย verbatim ตาม SRS)"),
                ("envelope", "{success, data}", "บังคับทุก endpoint", "ResponseInterceptor ห่อให้แล้ว — service ห้ามห่อซ้ำ"),
                ("error", "{success:false, data:null, error:{code,message}}", "message ภาษาไทย verbatim ตาม SRS", "โยนผ่าน HttpException เท่านั้น"),
                ("sps_store.mas_param", "key-value ของระบบเดิม", "**runtime = read-only · เขียนเฉพาะตอน seed/cutover**", "93,752 แถว · ไม่มี PK/unique → อ่านต้อง WHERE active_flag='Y' + LIMIT 1 เสมอ · 🔴 ค่า SBPGI_* ยังไม่มี ต้อง seed (5.5.2)"),
                ("sps_store.common_code / common_code_type", "code master ของระบบเดิม", "**runtime = read-only · เขียนเฉพาะตอน seed/cutover**", "2,609 / 376 แถว · code_type เป็น varchar(20) · ต้อง INSERT common_code_type ก่อน · 🔴 SBPGI_APPROVE_LIMIT ยังไม่มี ต้อง seed (5.5.2)"),
            ],
            [
                ("อ่านตัวตนผู้ใช้", "ทุก request", "BffUserGuard อ่าน BFF header (5.1.3)", "req.user = {userId, groupId, fullName, permissions} — fullName decode แล้ว"),
                ("อัปโหลดไฟล์แนบ", "ปุ่มแนบไฟล์", "POST /statement/upload-file-aws (ระบบ SBP เดิม)", "ได้ objectKey กลับมาเก็บใน document_attachments"),
                ("ดาวน์โหลดไฟล์แนบ", "ปุ่มดาวน์โหลด", "POST /statement/download-file-aws (ระบบ SBP เดิม)", "stream ไฟล์ผ่าน BE · ห้ามคืน objectKey ให้ FE"),
                ("ส่งอีเมล", "หลัง action สำเร็จ", "@gosoft-sbp/email-lib + email_template", "บันทึกผลที่ email_sent"),
                ("อ่านค่ากำหนดกลาง", "ตอน bootstrap/ตอนใช้งาน", "mas_param / common_code", "ห้าม hardcode วงเงินอนุมัติในโค้ด"),
            ],
            [],
            [
                "BFF forward request พร้อม header ตัวตนมาที่ store-backend",
                "Guard ตรวจ x-api-key แล้ว map header เป็น user context",
                "Controller/Service ทำงานโดยใช้ user context (ไม่มี JWT ของ SBPGI เอง)",
                "ผลลัพธ์ถูกห่อด้วย ResponseInterceptor เป็น {success, data}",
                "ไฟล์แนบวิ่งผ่าน service S3 เดิม · เก็บเฉพาะ metadata ใน SBPGI",
                "อีเมลส่งผ่าน email-lib โดยอ่าน template จาก email_template และบันทึกผลที่ email_sent",
                "ค่ากำหนด/วงเงินอ่านจาก mas_param และ common_code ทุกครั้ง ไม่ cache ข้ามรอบโดยไม่มี TTL",
            ],
            [
                "ไม่มี endpoint ใดของ SBPGI ออก/ตรวจ JWT เอง",
                "ทุก response ผ่าน envelope เดียวกับ store-backend",
                "ไม่มี credential ของ S3/SMTP อยู่ในโค้ดหรือ config ของ SBPGI",
                "วงเงินอนุมัติ เกณฑ์เดียว 100,000 อ่านจาก common_code (SBPGI_APPROVE_LIMIT) ไม่ hardcode",
                "objectKey ไม่ถูกส่งออกไปที่ FE",
                "x-user-full-name ถูก decodeURIComponent ก่อนใช้ทุกจุด — ไม่มี %E0%B8 หลุดไปที่ timeline/อีเมล",
                "สิทธิ์เขียนตัดสินจาก x-user-group-id + สถานะเอกสาร + getTransaction() ของ engine — ไม่ใช้ x-user-permissions เป็นด่านเดียว (รูปแบบยังไม่ยืนยัน)",
                "ค่า SBPGI_* ใน common_code/mas_param ถูก seed ด้วย script ที่ rerun ได้ (NOT EXISTS guard) ไม่ใช่ INSERT มือ",
                "ข้อค้างที่เหลือจริง (DP-6) ถูกบันทึกเป็นข้อค้าง ไม่ถูกตัดสินในเอกสารนี้ — DP-5/DP-8/DP-10 ปิดแล้ว",
            ],
            [
                "ไม่ส่ง x-api-key ต้องได้ 401 ตาม envelope มาตรฐาน",
                "ส่ง x-api-key ผิดค่า ต้องได้ 401 (เทียบ X_API_KEY ตรง ๆ) และ **ห้ามมีค่า key โผล่ใน log/error**",
                "ไม่ส่ง x-user-id ต้องได้ 401 · ส่ง x-user-group-id ที่ไม่ตรง section ของเอกสาร ต้องได้ 403",
                "error ที่โยนออกมาต้องเป็น {success:false, data:null, error:{code,message}} และ message เป็นไทย verbatim",
                "upload ไฟล์ 6MB ต้องถูก block ก่อนขึ้น S3 · download ไฟล์ของเอกสารที่ผู้ใช้ไม่เกี่ยวข้องต้องถูก block",
                "ส่งอีเมลสำเร็จแล้วมีแถวใน email_sent (คอลัมน์ผู้ส่งคือ send_by)",
                "เปลี่ยนค่า SBPGI_APPROVE_LIMIT ใน common_code แล้ว route อนุมัติเปลี่ยนตามโดยไม่ deploy",
                "x-user-full-name เป็น %-encoded ต้อง decode ถูก · ส่งค่าพัง decode ไม่ผ่านต้องไม่ throw",
                "x-user-permissions ว่าง/parse ไม่ผ่าน ต้องไม่ throw แต่ตกไปใช้ด่าน group + สถานะเอกสาร",
                "objectKey ต้องไม่โผล่ใน response ของทุก endpoint ที่ FE เรียก",
                "รัน seed script ซ้ำ 2 ครั้ง ต้องไม่เกิดแถวซ้ำใน common_code/mas_param (ไม่มี unique กันให้)",
            ],
            db_tables=[
                ("mas_param (sps_store)", "R (+ W ครั้งเดียวตอน seed)", "ค่ากำหนดกลาง 93,752 แถว · runtime อ่านอย่างเดียว · 🔴 ค่า SBPGI_* ยังไม่มี ต้อง seed ตอน setup (5.5.2)"),
                ("common_code / common_code_type (sps_store)", "R (+ W ครั้งเดียวตอน seed)", "2,609 / 376 แถว · 🔴 `SBPGI_APPROVE_LIMIT` / `SBPGI_DECISION` / `SBPGI_DATASOURCE` **ยังไม่มีในระบบจริง** ต้อง seed ตอน setup (5.5.2) · code_type เป็น varchar(20)"),
                ("email_template (sps_store)", "R", "85 แถว · SBPGI/lib อ่านอย่างเดียว — seed 8 แถวของ SBPGI ทำครั้งเดียวตอน migration ไม่ใช่ runtime"),
                ("email_sent (sps_store)", "W (โดย email-lib)", "5,214 แถว · lib เขียน log ให้เอง SBPGI ไม่ INSERT เอง (⚠️ คอลัมน์ผู้ส่งคือ send_by)"),
                ("business_user (sps_store)", "R", "12,752 แถว · ข้อมูลผู้ใช้/ผู้อนุมัติ"),
                ("store / mas_store (sps_store)", "R", "19,402 / 19,647 แถว · master ร้าน"),
                ("document_attachments (SBPGI)", "R/W", "metadata ไฟล์แนบ · ไฟล์จริงอยู่บน S3 ของระบบเดิม"),
            ],
        ),
        Topic(
            "BE/LLDD-BE-Workflow-Engine-Definition",
            "LLDD BE - Workflow Engine Definition",
            "BE",
            2.0,
            12,
            BANK_BE_OWNER,
            "**สร้างข้อมูลนิยาม workflow ลงฐานข้อมูลของ engine** — ระบุว่า flow ของ SBPGI มีกี่ step แต่ละ step ทำอะไร ใครทำได้ กดปุ่มไหนแล้วไป state ใด "
            "โดย register version/state/status/event/route/group/part ของ `@srm/glb-workflow` ตามสัญญาในเอกสารของ lib เอง "
            "(`docs/TSM-SRM-LLDD-SBP-workflow-1.2-full.md` — แปลงจาก `SBP/TSM-SRM-LLDD SBP workflow 1.2.xlsx`) · "
            "**เป็นงานตั้งต้นที่ต้องเสร็จก่อน** ฝั่ง BE คนอื่นจึงจะเรียก `initializeWorkflow` และ `eventWorkflow` (trigger event) ได้ — blocker ของสัปดาห์แรก",
            [],
            [
                "ลงทะเบียน workflow version ของ SBPGI 1 version (url_main + url_param_mapping)",
                "**ผลลัพธ์ที่ส่งมอบคือ seed script/มัยเกรชันของข้อมูลนิยาม** ไม่ใช่โค้ดเรียก engine — ทีมอื่นเรียก engine ต่อจากนิยามชุดนี้",
                "**จำนวน step ที่ต้องสร้าง = 6 state** — 5 ขั้นทำงาน (`06` รอฝ่าย SBP DSA → `08` รอเจ้าหน้าที่ SBP DSA → `01` รอหน่วยงานส่งเสริมธุรกิจ SBP → `02` รอ GM → `03` รอ AVP) + **1 state จบ** (`99` เสร็จสิ้นดำเนินการ) · `state_id` เป็น running ตาม version ตามกติกาของ engine (v1 → 10001+)",
                "**จำนวน route ที่ต้องสร้าง = 12 เส้น** ตาม Canonical Workflow Transition Matrix ใน `LLDD-BE-API-Document-Workflow-Actions` §5.1 (รวมเส้นข้ามขั้น 06→01 · เส้นจบทันทีเมื่อ เห็นควรไม่ชดเชย ที่ 01/02 · เส้นแตกตามวงเงิน 100,000 ที่ 02 และเส้นส่งกลับ)",
                "**ตารางที่ต้อง seed = 10 ตาราง** จาก 13 ตารางของ engine (`sps_store`) ตาม `SBP/TSM-SRM-LLDD-SBP-workflow-1.2.md` §4 — อีก 3 ตารางเป็นรันไทม์ที่ engine เขียนเอง",
                "ขอบเขตหยุดที่ข้อมูลนิยาม: **ไม่รวม** `initializeWorkflow` (เปิด instance · อยู่ใน LLDD-BE-API-Workflow-Instances) และ **ไม่รวม** `eventWorkflow`/trigger event (อยู่ใน LLDD-BE-API-Document-Workflow-Actions)",
                "นิยาม state/status 5 ขั้น 06 -> 08 -> 01 -> 02 -> 03 และปลายทางจบ flow",
                "นิยาม route ของทุกปุ่ม · การแตก route ตามวงเงินอนุมัติ เกณฑ์เดียว 100,000 เขียนเป็น**ตัวอย่างทางเลือก B เท่านั้น** — แหล่งเก็บวงเงินยังไม่ตัดสิน (มติเดิมคือ common_code · ดูข้อค้าง 5.6)",
                "สำรวจทางเลือกผู้อนุมัติ: workflow_group / workflow_group_map เทียบกับ addPreApprover รายคน — **ยังไม่ตัดสิน** (ดูข้อค้าง 5.6)",
                "สำรวจทางเลือก workflow_part / workflow_part_display สำหรับคุมการแสดงผลรายส่วน — **ยังไม่ตัดสิน** ว่าจะใช้แทน data-editrole ของ SBPGI หรือไม่ (ดูข้อค้าง 5.5/5.6)",
                "ความเสี่ยงและข้อค้างของ engine — **ชื่อ function ปิดแล้ว 2026-08-14** (ยึด 8 API ตามชีต Detail ของ LLDD ฝั่ง lib) · ที่ยังค้างคือ DP-2 `workflow_transaction` ไม่มี PK/index",
            ],
            [
                ("versionId", "integer", "1 ระบบ = 1 version", "SBPGI ขอ version ใหม่จากทีมเจ้าของ library"),
                ("referenceId", "string unique", "required ตอน initializeWorkflow", "✅ DP-1 ปิดแล้ว 2026-08-17 = `compensation_documents.id` (surrogate) แปลงเป็น string — `reference_id` ของ engine เป็น varchar(255)"),
                ("state_id", "integer running ตาม version", "1 state มีได้หลาย status", "map 5 ขั้นของ SBPGI: 06/08/01/02/03 + state จบ"),
                ("event", "save|submit|approve|reject|cancel|sendback", "ค่าเริ่มต้นของ engine", "ปุ่มไทยของ SBPGI map ลง event เหล่านี้ผ่าน common_code (code_type=SBPGI_DECISION) — ตาราง decisions ถูกตัดตามมติ DP-9 (2026-08-10)"),
                ("condition_json", '{"field","operator","value"}', "operator: == != > < >= <=", 'ใช้ {"field":"amount","operator":"<","value":100000} แยก route GM/AVP'),
                ("eventParam", "object", "ส่งมาพร้อม event", "SBPGI ส่ง {\"amount\": ยอดชดเชยรวม} ให้ engine เลือก route เอง"),
                ("part_display_type", "READ | WRITE", "ต้องยืนยันค่าจริงกับทีม library", "ไฟล์ต้นฉบับสะกดว่า WRTIE ทุกแถวของชีต sample data"),
                ("url_main / url_param_mapping", "string", "required ตอน register version", "ทำให้ inbox กลาง (GET /api/workflow/pending) ลิงก์กลับหน้าเอกสารของ SBPGI ได้"),
            ],
            [
                ("เปิด workflow", "Job 8b / สร้างเอกสาร", "initializeWorkflow(versionId, userId, referenceId)", "สร้าง workflow_transaction ที่ initial state/status"),
                ("ระบุผู้อนุมัติล่วงหน้า", "หลังเปิด workflow", "addPreApprover(versionId, referenceId, stateId, approver, seq, userId)", "insert workflow_approver (approver_type = user เสมอ)"),
                ("กดผลพิจารณา", "ปุ่มบนหน้าเอกสาร", "eventWorkflow(... event, eventParam ...)", "เดิน state ตาม route ที่ตรง condition_json แล้วบันทึก workflow_history"),
                ("อ่านปุ่ม/สิทธิ์แสดงผล", "เปิดหน้าเอกสาร", "getPermissionEvents(versionId, referenceId, userData)", "คืน event[] + display[] (partId/partDisplayType ต่อ state)"),
                ("อ่านกล่องงาน", "หน้ารายการรอดำเนินการ", "getPendingFlowByUser(userData, versionId)", "คืนงานที่รอ user คนนั้น + url_main"),
                ("อ่านประวัติ", "แท็บประวัติ", "getHistory(versionId, referenceId)", "คืน timeline ต่อแถว"),
            ],
            [],
            [
                "ขอ workflow version ใหม่จากทีมเจ้าของ @srm/glb-workflow (1 ระบบ = 1 version)",
                "ลงทะเบียน state/status 5 ขั้นของ SBPGI + state จบ flow",
                "ลงทะเบียน route ทุกเส้น พร้อม seq · ส่วน condition_json ของวงเงินอนุมัติให้ใส่ **ก็ต่อเมื่อเจ้าของโครงการเลือกทางเลือก B ของข้อค้าง 5.6** (ทางเลือก A = อ่านวงเงินจาก common_code SBPGI_APPROVE_LIMIT ตามมติเดิม แล้วให้ route แยกด้วยค่าที่ SBPGI ส่งมาแทน)",
                "**[conditional — ทำเมื่อเลือกทางเลือก A ของข้อค้าง \"ผู้อนุมัติของ SBPGI\" ใน 5.6]** ลงทะเบียน workflow_group / workflow_group_map สำหรับผู้อนุมัติแบบกลุ่ม · ถ้าเลือกทางเลือก B ให้ใช้ addPreApprover ระบุรายคนแทน",
                "**[conditional — ทำเมื่อเลือกทางเลือก A ของข้อค้าง `workflow_part_display` ใน 5.6]** ลงทะเบียน workflow_part + workflow_part_display ของส่วนต่าง ๆ ในหน้าเอกสาร · ถ้าเลือกทางเลือก B ให้คงกลไก visibleSections/editableSections ของ SBPGI",
                "กรอก url_main / url_param_mapping ให้ inbox กลางลิงก์กลับหน้าเอกสารได้",
                "ทดสอบเดิน flow ครบทุก route บน environment ทดสอบ",
                "ส่งความเสี่ยง/ข้อค้างให้ทีมเจ้าของ library และเจ้าของโครงการตัดสิน",
            ],
            [
                "SBPGI ไม่สร้างตาราง workflow ของตัวเองเลย — ใช้ engine 13 ตารางใน schema sps_store",
                "route ครอบคลุมทุกปุ่มบนหน้าเอกสาร และทุก route มี seq ที่ไม่ชนกัน",
                "วงเงินอนุมัติอยู่ที่เดียว (ยังไม่ตัดสินว่า condition_json หรือ common_code — ห้ามเก็บสองที่)",
                "[conditional] ถ้าเลือกทางเลือก A ของ workflow_part_display — การแสดงผลรายส่วนอ่านจาก display[] ของ getPermissionEvents ได้จริง",
                "ความเสี่ยงเรื่อง workflow_transaction ไม่มี PK/index ถูกยื่นเป็นเรื่องต่อทีมเจ้าของ library แล้ว",
                "ชื่อ function ที่จะใช้จริงถูกยืนยันกับทีมเจ้าของ library ก่อนเขียนโค้ด",
            ],
            [
                "initializeWorkflow ซ้ำด้วย referenceId เดิมต้องไม่เกิด transaction ที่สอง",
                "ยอดชดเชย 99,999 ต้องจบที่ GM · 100,000 ต้องวิ่งต่อ AVP (เกณฑ์เดียว · มติ 2026-08-18)",
                "ผู้ใช้ที่ไม่ใช่ current_approver ต้องไม่ได้ปุ่มใด ๆ จาก getPermissionEvents",
                "[conditional · เฉพาะทางเลือก A ของ workflow_part_display] display[] ของ state 01 ต้องเปิด WRITE เฉพาะส่วนที่ section 01 แก้ได้",
                "getPendingFlowByUser ต้องคืน url_main ที่เปิดกลับหน้าเอกสารได้จริง",
                "เดิน flow จนจบแล้ว workflow_history มีครบทุกขั้น",
            ],
            db_tables=[
                # --- 4.1 ตารางนิยาม flow (config · seed ครั้งเดียวต่อระบบ) — SBP/TSM-SRM-LLDD-SBP-workflow-1.2.md §4.1
                ("workflow (sps_store)", "W ครั้งเดียวตอน setup", "**1 แถว** — `workflow_name` = ระบบประกันรายได้ (SBPGI)"),
                ("workflow_version (sps_store)", "W ครั้งเดียวตอน setup", "**1 แถว** — 1 ระบบ = 1 version · ต้องมี `initial_state_id` (= ขั้น 06), `end_state_id` (= 99), `url_main`, `url_param_mapping` เพื่อให้ inbox กลางลิงก์กลับหน้าเอกสารได้ · **ขอเลข version จากทีมเจ้าของ library**"),
                ("workflow_state (sps_store)", "W ครั้งเดียวตอน setup", "**6 แถว = จำนวน step ของ flow** — 5 ขั้นทำงาน (06 · 08 · 01 · 02 · 03) + 1 ขั้นจบ (99) · `state_id` running ตาม version (v1 → 10001+)"),
                ("workflow_status (sps_store)", "W ครั้งเดียวตอน setup", "**6 แถว** — ชื่อสถานะเอกสารที่ผู้ใช้เห็น 1:1 กับ state (รอฝ่าย SBP DSA ดำเนินการ / รอเจ้าหน้าที่ SBP DSA ดำเนินการ / รอหน่วยงานส่งเสริมธุรกิจ SBP ดำเนินการ / รอ GM ส่งเสริมธุรกิจ SBP ดำเนินการ / รอผู้บริหารสำนักบริหาร SBP ดำเนินการ / เสร็จสิ้นดำเนินการ) · engine รองรับ 1 state หลาย status"),
                ("workflow_event (sps_store)", "R (ใช้ค่า default ของ engine)", "`save` `submit` `approve` `reject` `cancel` `sendback` — ปุ่มไทยของ SBPGI map ลง 6 event นี้ผ่าน `common_code` (`code_type = SBPGI_DECISION`) · **ไม่ต้องเพิ่ม event ใหม่**"),
                ("workflow_route (sps_store)", "W ครั้งเดียวตอน setup", "**12 แถว = ทุกเส้นทางของ flow** ตาม Canonical Workflow Transition Matrix (`LLDD-BE-API-Document-Workflow-Actions` §5.1) — รวมเส้นข้ามขั้น 06→01 · เส้นจบทันทีเมื่อ *เห็นควรไม่ชดเชย* ที่ 01/02 · เส้นแตกตามวงเงิน 100,000 ที่ 02 (`condition_json`) และเส้นส่งกลับทุกขั้น · `seq` ห้ามชนกันภายใน from_state เดียวกัน"),
                # --- 4.2 ตารางกลุ่มผู้อนุมัติ — §4.2
                ("workflow_group (sps_store)", "W ครั้งเดียวตอน setup *(conditional)*", "กลุ่มผู้อนุมัติต่อขั้น — **ทำเมื่อเลือกทางเลือก A ของข้อค้าง “ผู้อนุมัติของ SBPGI” (5.6)** · ถ้าเลือกทางเลือก B (ระบุรายคนด้วย `addPreApprover`) ไม่ต้อง seed"),
                ("workflow_group_map (sps_store)", "W ครั้งเดียวตอน setup *(conditional)*", "map กลุ่ม → ผู้ใช้ · ไม่ระบุ `map_table` = เทียบ `userId`/`groupId` ตรง ๆ · ระบุ `map_table` = ต้องเป็น **view ที่ where ด้วย user_id/group_id ได้**"),
                # --- 4.4 ตารางคุมการแสดงผล — §4.4
                ("workflow_part (sps_store)", "W ครั้งเดียวตอน setup *(conditional)*", "ชื่อ component ของหน้าเอกสาร + `part_seq` — **ทำเมื่อเลือกทางเลือก A ของข้อค้าง `workflow_part_display` (5.6)**"),
                ("workflow_part_display (sps_store)", "W ครั้งเดียวตอน setup *(conditional)*", "`part_display_type` = READ / WRITE ต่อ (state, part) · ⚠️ ไฟล์ต้นฉบับสะกด `WRTIE` ทุกแถว ต้องยืนยันค่าจริงกับทีม library ก่อน seed"),
                # --- 4.3 ตารางรันไทม์ (ไม่ seed) — §4.3
                ("workflow_transaction / workflow_history / workflow_approver (sps_store)", "R เท่านั้น (engine เขียนเอง)", "ข้อมูลรันไทม์ 19,283 / 38,010 / 96,542 แถว (ตรวจ 2026-08-07) — 🔴 **ห้าม INSERT/UPDATE ตรง** ต้องผ่าน `initializeWorkflow` / `eventWorkflow` / `addPreApprover` ของ lib เท่านั้น · DP-2: `workflow_transaction` ไม่มี PK และไม่มี index"),
            ],
        ),
    ]


def topics() -> list[Topic]:
    base = [
        Topic(
            "FE/LLDD-FE-Integration-Contracts",
            "LLDD FE - Integration Contracts",
            "FE",
            2.8,
            24,
            FE_OWNER,
            "กำหนดสัญญากลางฝั่ง Frontend สำหรับการ consume API ทุกหน้า: auth/session, error handling, pagination, format, document action และ RBAC/menu gating",
            [],
            [
                "Shared API client contract",
                "Auth/JWT consumption from platform reference",
                "Error display and validation message mapping",
                "Date/year/money/docNo formatting",
                "Pagination, list empty/loading/error state",
                "Document action result enum and response consumption",
                "RBAC/menu gating and editable section flags",
            ],
            [
                ("Authorization", "Bearer JWT", "required except /auth/login and /auth/refresh", "แนบโดย axios interceptor เท่านั้น; component ห้าม set header เอง"),
                ("ApiError", "{code,message}", "message required", "แสดง message จาก BE ตรง ๆ; fallback ใช้เฉพาะ network/no response"),
                ("PageResponse<T>", "{page,size,total,items}", "page>=1 size<=100", "ใช้กับ DataTable/Pager ทุกหน้า"),
                ("date/month", "ISO ค.ศ. YYYY-MM-DD / YYYY-MM", "payload uses CE", "แสดงผ่าน formatDateThai/formatMonthThai จุดเดียว — ค่าเริ่มต้นเป็น ค.ศ."),
                ("docNo", "YYYY/xxxxx ค.ศ.", "do not split except route params", "route ใช้ /documents/:year/:running แล้วประกอบ docNo"),
                ("result", "verbatim from actionOptions", "required before submit action", "ส่งเป็น payload `{result, comment}` เท่านั้น"),
                ("ActionResponse", "{statusCode,nextSection,message}", "required after action", "invalidate detail/timeline/tasks แล้ว resolve label จาก /sbpgi/lookup/document-statuses"),
                ("MenuItem", "{menuCode,label,route,group}", "จาก GET /menus + GET /groups/current-user/permissions ของระบบเดิม (ผ่าน BFF)", "sidebar filter ด้วย menuCode จาก API; ไม่ hardcode role"),
                ("canEditSections", "string[]", "from document detail", "ใช้เปิด/ปิด section editor; FE ไม่คำนวณสิทธิ์เอง"),
            ],
            [
                ("Attach token", "ทุก API call", "shared/api/client.ts", "Authorization header จาก auth store"),
                ("Refresh token", "401 non-auth endpoint", "POST /api/v1/auth/refresh", "single-flight แล้ว replay request เดิม"),
                ("Show API error", "catch AxiosError", "apiErrorMessage()", "แสดงข้อความไทยจาก BE ตรง ๆ"),
                ("Render list", "GET list endpoint", "PageResponse<T>", "DataTable/Pager ใช้ shape เดียวกัน"),
                ("Submit action", "ปุ่มส่งดำเนินการ", "POST /api/v1/sbpgi/document/{docNo}/actions", "ส่ง `{result, comment}` และ consume `{statusCode,nextSection,message}`"),
                ("Gate route/menu", "login/bootstrap", "GET /menus + GET /groups/current-user/permissions (ระบบเดิม ผ่าน BFF)", "สร้าง sidebar และ route guard จาก menuCode"),
            ],
            [
                ApiSpec("ALL", "/api/v1/sbpgi/*", "Error contract กลางสำหรับ FE ทุกหน้า", None, {"code": "VALIDATION", "message": "ข้อความภาษาไทยตรงตาม SRS"}),
                ApiSpec("GET", "/api/v1/sbpgi/*?page=1&size=20", "List/pagination contract กลาง", {"page": 1, "size": 20}, {"page": 1, "size": 20, "total": 0, "items": []}),
                ApiSpec("POST", "/api/v1/sbpgi/document/{docNo}/actions", "Document action contract ตัวอย่างเมื่อ currentSection=01 จึงเปลี่ยนไป 02; FE ห้ามส่งหรือคำนวณปลายทางเอง", {"result": "เห็นควรชดเชย", "comment": "เห็นควรชดเชยตามหลักเกณฑ์"}, {"statusCode": "02", "nextSection": "02", "message": "submitted"}),
            ],
            [
                "Bootstrap env and API client",
                "Login or restore session with refresh token",
                "Load GET /auth/profile + GET /users/current + GET /menus + GET /groups/current-user/permissions (ทั้งหมดเป็นของระบบเดิมผ่าน BFF)",
                "Render routes/sidebar from menu contract",
                "All feature hooks use shared API client and PageResponse/Error types",
                "Document action sends `{result, comment}` only and consumes `{statusCode,nextSection,message}`",
                "All display formatting goes through shared/lib/format.ts",
            ],
            [
                "ไม่มี feature ใดสร้าง axios instance เอง",
                "ทุก API error แสดง message จาก BE โดยไม่ paraphrase",
                "ทุก list endpoint ใช้ PageResponse shape เดียวกัน",
                "วันที่ใน payload และหน้าจอเป็น ค.ศ. จาก formatter กลาง (แสดง พ.ศ. เฉพาะจุดที่เปิด flag)",
                "Sidebar และ route access มาจาก GET /menus ของระบบเดิม ไม่ hardcode role",
                "FE ไม่คำนวณ action routing เอง; ใช้ role profile และ actionOptions จาก API",
            ],
            ["401 refresh single-flight", "403 route guard", "error message passthrough", "pagination pager mapping", "date BE display", "action response invalidation", "menu filtering by API"],
        ),
        Topic(
            "FE/LLDD-FE-Foundation",
            "LLDD FE - Application Foundation and Shared UI",
            "FE",
            7.1,
            60,
            FE_OWNER,
            "เตรียม foundation ฝั่ง Frontend สำหรับ SBP Mall: routing, API client, constants, shared state, formatters, mock mapping และ shared UI primitives; เอกสารนี้ไม่ใช่หน้าจอ Dashboard",
            [],
            ["Non-screen technical foundation", "Route/module registry เฉพาะ SBP Mall", "API client และ response typing", "Shared constants/menu/status mapping", "Mock data mapping", "CSS/tokens สำหรับ table/form/modal/responsive"],
            [
                ("routePath", "string", "required", "ต้อง map กับเมนู SBP Mall"),
                ("apiBaseUrl", "URL", "required by env", "ใช้กับทุก API call"),
                ("statusCode", "string", "must map to status dictionary", "ใช้ร่วมกับ StatusBadge"),
                ("mockData", "JSON", "schema compatible with API response", "ใช้ก่อน BE พร้อม"),
            ],
            [
                ("Register module route", "bootstrap", "client router", "route guard รู้จักหน้า SBP Mall"),
                ("Call API", "React Query hook", "shared API client", "standard loading/error handling"),
            ],
            [
                ApiSpec("GET", "/api/v1/sbpgi/lookup/document-statuses", "โหลดสถานะเอกสารสำหรับ dropdown/badge", {}, {"items": [{"code": "06", "label": "รอฝ่าย SBP DSA ดำเนินการ"}]}),
            ],
            ["Initialize app config", "Register SBP Mall routes", "Create shared API client", "Prepare constants/formatters", "Wire shared UI primitives"],
            ["ไม่มี screenshot หรือ dashboard behavior ในเอกสารนี้", "ทุก route ถูก register ผ่าน module registry", "API error shape ใช้ร่วมกัน", "ไม่มี dependency กับ Login/Auth ใหม่", "CSS responsive base พร้อม"],
            ["route registration", "API base missing", "status unknown", "mock response compatible", "shared formatter output"],
        ),
        Topic(
            "FE/LLDD-FE-Document-Lists",
            "LLDD FE - Document Lists",
            "FE",
            8.2,
            70,
            FE_OWNER,
            "สร้างหน้ารายการเอกสารรอดำเนินการและเอกสารที่เกี่ยวข้อง",
            ["k2-list-waiting-01.png", "k2-list-waiting-02.png", "k2-list-related-01.png"],
            ["Waiting list", "Related document list", "Search/filter/status filter", "Pagination/row action + เลือกหลายเอกสาร (bulk)", "Red flag (sales < 60 days) + rejected-ending rows ที่บทบาท 06 ต้องเห็น"],
            [
                ("docNo", "YYYY/xxxxx", "optional search", "ถ้าคลิก row ส่งไป detail"),
                ("year", "ค.ศ. YYYY", "required สำหรับ /sbpgi/document", "default current year (ค.ศ.)"),
                ("status", "status code/string", "optional single select", "ใช้ filter chip"),
                ("table.roundNo", "integer", "column 1", "ครั้งที่ (รอบชดเชยของร้าน)"),
                ("table.docNo", "YYYY/xxxxx", "column 2", "เลขที่เอกสารและลิงก์เปิด detail"),
                ("table.impactedStoreCode", "string 5 digits", "column 3", "รหัสร้านถูกกระทบ; คง leading zero"),
                ("table.impactedStoreName", "string", "column 4", "ชื่อร้านถูกกระทบ"),
                ("table.regionCode", "string", "column 5", "ภาค"),
                ("table.salesDeclinePercent", "decimal", "column 6", "ยอดขายที่ลดลง (%)"),
                ("table.totalCompensationAmount", "decimal", "column 7; >=0", "จำนวนเงินที่ชดเชย; format #,##0.00"),
                ("table.statusCode/statusName", "code + label", "column 8", "สถานะ; เก็บ code และ resolve label จาก dictionary"),
                ("table.daysPending", "integer", "column 9; >=0", "รอ (วัน)"),
                ("table.salesDataDays", "integer", "internal (ไม่ใช่คอลัมน์แสดง)", "<60 = แถวผิดปกติสีแดง (red-flag)"),
            ],
            [
                ("Search", "ปุ่มค้นหา", "GET /api/v1/sbpgi/document/tasks หรือ /sbpgi/document", "reload table"),
                ("Clear", "ปุ่มเคลียร์", "client state", "reset filters"),
                ("Open detail", "click row", "navigate /sbpgi/document/:docNo", "เปิดเอกสาร"),
            ],
            [
                ApiSpec("GET", "/api/v1/sbpgi/document/tasks", "รายการเอกสารรอดำเนินการ", {"page": 1, "size": 20, "status": "06"}, {"page": 1, "size": 20, "total": 24, "items": [{"roundNo": 1, "docNo": "2026/00123", "impactedStoreCode": "01234", "impactedStoreName": "สาขาตัวอย่าง", "regionCode": "BE", "salesDeclinePercent": 12.5, "statusCode": "06", "statusName": "รอฝ่าย SBP DSA ดำเนินการ", "totalCompensationAmount": 48200.0, "daysPending": 3, "salesDataDays": 58}]}),
                ApiSpec("GET", "/api/v1/sbpgi/document", "ค้นหาเอกสารที่เกี่ยวข้อง ต้องระบุปี", {"year": 2026, "page": 1, "size": 20}, {"page": 1, "size": 20, "total": 342, "items": [{"roundNo": 2, "docNo": "2026/00124", "impactedStoreCode": "01235", "impactedStoreName": "สาขาตัวอย่าง 2", "regionCode": "BS", "salesDeclinePercent": 18.0, "statusCode": "99", "statusName": "เสร็จสิ้น", "totalCompensationAmount": 72500.0, "daysPending": 0, "salesDataDays": 60}]}),
            ],
            ["Read route mode", "Bind filter values", "Call list API", "Render table", "Apply abnormal row style", "Navigate to detail on row click"],
            ["ตาราง 9 คอลัมน์หลักครบ", "ปีเป็น required เมื่อใช้ /sbpgi/document", "ยอดขายไม่ครบ 60 วันแสดงแดง", "pagination คง filter เดิม"],
            ["ค้นหาด้วย docNo", "filter status", "เปิด detail", "empty result", "abnormal row"],
        ),
        Topic(
            "FE/LLDD-FE-Create-Document",
            "LLDD FE - Create Document",
            "FE",
            4.9,
            42,
            FE_OWNER_KITTISAK,
            "หน้าสร้างเอกสารประกันรายได้ — **มติ 2026-08-06: ไม่มีฟอร์มฝั่ง SBP** main card เป็น iframe ของหน้าสร้างเอกสารระบบ FS ตรง ๆ + หมายเหตุ 4 ขั้นตอนใต้ iframe · `POST /sbpgi/document` เรียกโดย pipeline/service token",
            ["k2-create-01.png"],
            ["🔴 **มติ 2026-08-06 — หน้านี้ไม่มีฟอร์มและไม่มีแท็บฝั่ง SBP**", "main card = iframe ของหน้าสร้างเอกสารระบบ FS ตรง ๆ (เหมือน `k2-create.html`)", "หมายเหตุ 4 ขั้นตอน (verbatim จากหน้าจอ K2 เดิม) อยู่ใต้ iframe นอกกรอบ", "`POST /sbpgi/document` เป็น pipeline/service-token ไม่ใช่ฟอร์ม FE — ต้นทางสร้างที่ FS แล้วรอ SBP Statement ส่งกลับ (~1 วัน)", "การคีย์/ปรับข้อมูลร้านตาม SDD GI ทำที่หน้าเอกสาร (`PUT /sbpgi/document/{docNo}`) ไม่ใช่หน้านี้", "⚠️ หัวข้อ 5.1-5.6 (SBP mirror form + FS bridge) เป็นดีไซน์ก่อนมติ — เก็บไว้เป็นทางเลือกสำรอง **ไม่อยู่ในขอบเขต 8 ชม.**"],
            [
                ("source", "MANUAL|FS", "required", "แสดง section ตาม source; payload ใช้ชื่อ field `source`"),
                ("activeTab", "MANUAL|FS_IFRAME", "required UI state", "เลือก tab สร้างเอกสารทั่วไปหรือเอกสารจาก FS"),
                ("fsIframeUrl", "URL", "required for FS tab", "อ่านจาก config; ใช้โหลด hidden iframe ของ FS"),
                ("fsFieldMap", "array", "required after iframe load", "metadata ของ input/select/textarea ที่อ่านจาก iframe เพื่อ render SBP mirror form"),
                ("fsMirrorValues", "object", "required for FS tab", "state ของ form ฝั่ง SBP ที่ sync เข้า hidden iframe เมื่อ change/submit"),
                ("impactedStoreCode", "string 5 digits", "required", "ค้นหาด้วย popup ร้านถูกกระทบ; คง leading zero"),
                ("impactedStoreName", "string", "readonly after select", "เติมอัตโนมัติหลังเลือกร้าน"),
                ("newStoreCode", "string 5 digits", "required", "เลือกร้านเปิดใหม่จาก popup; ส่งรหัสร้านและคง leading zero"),
                ("impactMonth", "YYYY-MM", "required", "month picker; ทั้งแสดงและส่งเป็น ค.ศ."),
                ("statementPeriod", "YYYY-MM", "required for FS", "Period Statement จาก SRS SCR-02"),
                ("roundNo", "integer >= 1", "required/default 1", "ครั้งที่ของเอกสาร/งวดชดเชย"),
                ("reason", "text", "required for MANUAL/out-of-condition", "เหตุผลการสร้างเอกสารนอกเงื่อนไข; trim ก่อนส่ง"),
            ],
            [
                ("Search store", "แว่นขยาย", "GET /store/search (ระบบ SBP เดิม)", "เลือก impacted/new store"),
                ("Open FS tab", "tab เอกสารจาก FS", "Load hidden iframe from fsIframeUrl", "discover FS fields and render SBP mirror form"),
                ("Change FS mirror value", "input/select ใน SBP mirror form", "iframe value sync service", "ส่งค่าเข้า field ใน hidden iframe และ dispatch input/change"),
                ("Save draft", "ปุ่มบันทึก", "POST /api/v1/sbpgi/document", "สร้าง draft"),
                ("Submit", "ปุ่มส่งดำเนินการ", "POST /api/v1/sbpgi/document", "สร้างเอกสารและเริ่ม workflow"),
                ("Submit FS iframe", "ปุ่มส่งใน tab เอกสารจาก FS", "sync all mirror values + submit iframe form", "submit form ของ FS ใน hidden iframe"),
            ],
            [
                ApiSpec("GET", "/store/search (ระบบ SBP เดิม)", "ค้นหาร้านสำหรับ popup", {"q": "012", "type": "impacted"}, {"items": [{"storeCode": "01234", "storeName": "สาขาตัวอย่าง", "regionCode": "BN"}]}),
                # URL ของ FS iframe อ่านจาก backend config (env `FS_CREATE_DOCUMENT_URL`) — ไม่มี endpoint /configs แล้ว (2026-08-06)
                ApiSpec("POST", "/api/v1/sbpgi/document", "สร้างเอกสาร", {"source": "MANUAL", "impactMonth": "2026-07", "statementPeriod": "2026-07", "impactedStoreCode": "01234", "newStoreCode": "22864", "roundNo": 1, "reason": "สร้างเอกสารนอกเงื่อนไข"}, {"docNo": "2026/00001", "statusCode": "06", "message": "created"}),
            ],
            ["User opens create page", "Choose tab: สร้างเอกสารทั่วไป or เอกสารจาก FS", "For FS tab load hidden iframe and discover fields", "Render SBP mirror form from iframe field metadata", "Search/select store or input period/source", "On change sync value into hidden iframe", "Validate", "Submit SBP API or submit hidden FS iframe", "Navigate to detail or show FS submit result"],
            ["required fields ทำงาน", "docNo ได้จาก API for MANUAL", "FS tab loads iframe and renders mirror form", "changing SBP mirror field updates hidden iframe field", "FS submit syncs all values before iframe submit", "validation message ชัดเจน"],
            ["ไม่เลือก store", "period format ผิด", "submit success", "API duplicate error", "FS iframe load timeout", "FS field mapping missing", "FS submit callback success/error"],
        ),
        Topic(
            "FE/LLDD-FE-Document-Detail",
            "LLDD FE - Document Detail and Action",
            "FE",
            11.5,
            98,
            FE_OWNER_KITTISAK,
            "สร้างหน้าเอกสารรายละเอียดและ Action Panel โดยแสดงผลตาม role profile ของผู้ใช้ที่ login",
            ["k2-document-01.png", "k2-document-02.png", "k2-document-03.png"],
            ["Document header", "Store impact/new-store/factor sections", "Role-based visible/editable sections", "Action panel by role profile", "History/timeline", "Attachment upload/download", "Map/ALLMAP link"],
            common_doc_fields() + [
                ("result", "verbatim from actionOptions", "required on submit action", "FE แสดง radio ตาม `actionOptions` จาก API เท่านั้น · ไม่เลือกแล้วกดส่ง → popup **verbatim SRS**: `ท่านยังไม่เลือกผลการพิจารณา กรุณาเลือกข้อมูลก่อนกดส่งดำเนินการ`"),
                ("comment", "text", "required บาง result", "trim before submit · SRS บังคับ required เมื่อเลือกไม่ชดเชย แต่ไม่ได้ระบุข้อความ popup"),
                ("compensatePercent", "number", "sum = 100", "validate before save · ไม่ครบ 100% → popup `โปรดตรวจสอบ %ชดเชย ของท่าน รวมกันแล้วไม่เท่ากับ 100%`"),
                ("competitorCode", "select จาก master competitors", "required เมื่อเพิ่ม/แก้แถวคู่แข่ง", "ไม่เลือก → popup **verbatim SRS §10**: `กรุณาเลือกร้านคู่แข่งที่ท่านต้องการ`"),
                ("factor.startDate / factor.endDate", "date", "endDate >= startDate", "**กติกา SRS §11**: วันที่สิ้นสุดต้องเท่ากับหรือมากกว่าวันที่เริ่มต้น — ถ้าน้อยกว่าต้องแสดง Pop-up แจ้งเตือน (SRS ไม่ได้ระบุข้อความ ให้ยืนยันกับ BA ก่อน UAT)"),
            ],
            [
                ("Save section", "ปุ่มบันทึก", "PUT /api/v1/sbpgi/document/{docNo}", "save partial"),
                ("Submit action", "ปุ่มส่งดำเนินการ", "POST /api/v1/sbpgi/document/{docNo}/actions", "submit selected result and reload status"),
                ("Upload file", "เลือกไฟล์", "POST /api/v1/sbpgi/document/{docNo}/attachments", "append attachment"),
                ("Open sales", "ข้อมูลยอดขายเพิ่มเติม", "GET /api/v1/sbpgi/document/{docNo}/sales", "show chart/detail"),
            ],
            [
                ApiSpec("GET", "/api/v1/sbpgi/document/{docNo}", "โหลดรายละเอียดเอกสารพร้อม role profile สำหรับหน้า detail", {"docNo": "2026/00123"}, {"docNo": "2026/00123", "statusCode": "06", "viewerRbacRoleCode": "R-XX", "roleProfileCode": "P-06", "visibleSections": ["doc-header", "sec-sales", "sec-map", "sec-newstore", "sec-competitor", "sec-factor", "sec-attach", "sec-comp-history", "sec-decision-history", "sec-action"], "editableSections": [], "canUploadAttachment": True, "canAction": True, "actionOptions": [{"label": "เห็นควรไม่ชดเชย", "requireComment": True}, {"label": "หยุดชดเชยประกันรายได้", "requireComment": False}, {"label": "ส่งหน่วยงานส่งเสริมธุรกิจ SBP", "requireComment": False}, {"label": "ส่งเจ้าหน้าที่ SBP DSA ดำเนินการ", "requireComment": False}], "impactedStore": {"storeCode": "01234"}, "newStores": []}),
                ApiSpec("PUT", "/api/v1/sbpgi/document/{docNo}", "บันทึกส่วนย่อย เช่น ร้านเปิดใหม่/คู่แข่ง/ปัจจัย", {"newStores": [{"newStoreCode": "22864", "compensatePercent": 100}]}, {"message": "saved"}),
                ApiSpec("POST", "/api/v1/sbpgi/document/{docNo}/actions", "ส่งผลพิจารณาที่เลือกจาก actionOptions; ตัวอย่าง currentSection=01 จึงเปลี่ยนไป 02", {"result": "เห็นควรชดเชย", "comment": "เห็นควรชดเชยตามหลักเกณฑ์"}, {"statusCode": "02", "nextSection": "02", "message": "submitted"}),
                ApiSpec("POST", "/api/v1/sbpgi/document/{docNo}/attachments", "แนบไฟล์", {"file": "multipart/form-data <= 5MB"}, {"attachmentId": "att-001", "fileName": "evidence.pdf"}),
            ],
            ["Load document detail", "Render role profile from API", "User edits allowed sections only", "Validate fields and popup text", "Confirm action", "Submit selected result", "Reload detail/timeline/status"],
            ["ส่วน read-only แก้ไม่ได้", "% ชดเชยรวม 100", "action required result", "upload limit 5MB", "timeline reload หลัง submit"],
            ["เปิดเอกสาร", "save section", "submit without result", "submit approve", "upload too large", "timeline display"],
        ),
        Topic(
            "FE/LLDD-FE-Testing-Delivery",
            "LLDD FE - Testing and Delivery",
            "FE",
            3.5,
            30,
            FE_OWNER,
            "กำหนด regression, responsive pass, API payload adjustment และ delivery note สำหรับ FE",
            [],
            ["Manual regression", "Responsive pass", "API contract verification", "UAT defect retest", "Release gate", "Delivery evidence"],
            [
                ("viewport", "desktop/tablet/mobile", "must verify key pages", "ใช้ browser responsive mode"),
                ("apiContractVersion", "string", "must match BE", "บันทึกใน delivery note"),
                ("defectId", "string", "required for UAT fix", "trace กลับ defect log"),
            ],
            [
                ("Run route regression", "เริ่ม regression suite", "test matrix state", "บันทึก pass/fail/evidence ต่อ route และ role profile"),
                ("Change viewport", "เลือก desktop/tablet/mobile", "browser responsive state", "ตรวจ overflow, modal, table, chart และ navigation"),
                ("Verify API contract", "replay fixture/request", "schema comparison", "แสดง field/type/error mismatch แบบ trace กลับ endpoint"),
                ("Retest defect", "เลือก defectId ที่แก้แล้ว", "retest status transition", "แนบ before/after evidence และผล regression ที่เกี่ยวข้อง"),
                ("Evaluate release gate", "กดสรุป readiness", "build/typecheck/defect/secret/contract checks", "ได้ PASS หรือ BLOCKED พร้อมเหตุผลราย gate"),
                ("Prepare delivery evidence", "ปิดรอบทดสอบ", "delivery bundle state", "สร้าง test summary/known limitations/verification commands โดยไม่มี secret"),
            ],
            [],
            ["Run page-by-page regression", "Verify responsive", "Compare API payload", "Fix UAT defects", "Run build check", "Prepare delivery note"],
            ["ทุกหน้าหลักไม่มี layout broken", "API payload ตรง contract", "UAT defects critical/high ปิดแล้ว", "delivery note พร้อม"],
            ["desktop regression ครบทุก route หลัก", "tablet/mobile regression", "request/response schema mismatch ต้องเป็นศูนย์", "Critical/High defects ต้องปิด", "report preview/export parity", "action transition 06→08→01→02→03→99", "delivery evidence ไม่มี token/secret"],
        ),
        Topic(
            "FE/LLDD-FE-Report",
            "LLDD FE - Status Summary Report",
            "FE",
            5.9,
            50,
            FE_OWNER,
            "สร้างรายงานตรวจสอบประกันรายได้ตาม SDD สไลด์ 60 (7 ตัวกรอง / 14 คอลัมน์) พร้อมค้นหาข้อมูลและ Export Excel",
            ["k2-report-01.png", "k2-report-02.png"],
            ["Report filters (SDD slide 60 · 2026-08-06: สถานะ*|รหัสร้านถูกกระทบ · รหัสร้านเปิดกระทบ|ประเภทร้าน (รหัสจาก common_code · รหัสที่ 4 รอยืนยัน) · Period Statement From-To (date, ค.ศ.) เต็มแถว · ภาคเต็มแถว · ผลการพิจารณาเต็มแถว)", "Summary table (sortable 14 columns)", "ปุ่มออกผล 3 ตัว (Preview Report · Export Excel · Export CSV to Batch)", "Sample data verification"],
            [
                ("impactedStoreCode", "string 5 digits", "optional; numeric only when input", "คง leading zero; ปุ่มแว่นขยายเรียก popup เลือกร้านที่ถูกกระทบ"),
                ("impactedStoreName", "string", "readonly", "แสดงอัตโนมัติหลังเลือกรหัสร้าน; ไม่ส่งเป็น filter หลักถ้ามี storeCode"),
                ("newStoreCode", "string 5 digits", "optional; numeric only when input", "รหัสร้านเปิดกระทบ/ร้านเปิดใหม่; คง leading zero"),
                ("impactMonthFrom", "YYYY-MM", "optional; month picker", "ส่งและแสดงเป็น ค.ศ. เช่น 2026-05"),
                ("impactMonthTo", "YYYY-MM", "optional; month picker; must be >= from", "ถ้า from > to ให้แสดง validation ก่อน call API"),
                ("storeTypes", "array ของ BranchTypeFGIName", "optional multi select", "**ยืนยันจาก master จริงแล้ว (`ข้อมูล Master K2.xlsx` · ชีต `BranchTypeProfile` จาก `CPA_FRN_FGI`)**: ค่าที่ใช้คือคอลัมน์ `BranchTypeFGIName` มี **7 ค่าไม่ซ้ำ** — `A` (A-Mo) · `B` (B(1)) · `C` (C และ C(Retire CPALL)) · `D` (Type D — เดิมเรียก BGC) · `E` (B(2)) · `PTT` · `บริษัท` (Corporate) · ⚠️ **D กับ E เป็นคนละประเภทและมีจริงทั้งคู่** — เอกสารรุ่นก่อนที่แสดงเพียง 4 ตัวเลือก (A/B/C/E หรือ A/B/C/D) **ผิด** ทั้ง SDD สไลด์ 60 (แสดงบางส่วน) และ SRS (เขียน “พนักงาน” ซึ่ง**ไม่มีใน master**) · ยังคง**ห้าม hardcode** — โหลดจาก `GET /common/common-code` ของระบบ SBP เดิม แล้วใช้ 7 ค่านี้เป็น expected set ตอนทดสอบ"),
                ("status", "statusCode string", "required single select", "บังคับเลือก 1 สถานะก่อน Preview/Export; options มาจาก sps_store.workflow_status ของ @srm/glb-workflow (ตาราง document_statuses ของ SBPGI ถูกตัดแล้ว)"),
                ("resultCategory", "APPROVE|REJECT|CANCELLED|PENDING", "optional radio (status เท่านั้นที่บังคับ)", "**4 ค่า** — APPROVE=ประกันรายได้ · REJECT=ไม่ประกันรายได้ · **CANCELLED=ยกเลิกโดยระบบ (เพิ่ม 2026-08-10)** · PENDING/ไม่มีค่า=ยังไม่มีผล · CANCELLED มาจาก master จริง `DecisionProfile` decision 14 `CancelBySystem` (`DecisionResultName` = ยกเลิกโดยระบบ) ซึ่ง SDD สไลด์ 60 ไม่ได้แสดงไว้"),
                ("regions", "array ของ ZoneName", "optional multi select", "**ยืนยันจาก master จริงแล้ว (`ข้อมูล Master K2.xlsx` · ชีต `ZoneProfile`)**: **13 ภาค** — BN(10) · BW(20) · BE(30) · BG(40) · BS(70) · REU(81) · NEU(82) · RSU(83) · RSL(84) · RN(85) · RC(86) · REL(90) · NEL(92) (ตัวเลขในวงเล็บคือ `ZoneCode`) — ตรงกับรายการที่ prototype ใช้ **ครบทั้ง 13 ค่า** · รายการ 8 ค่าใน SRS (BE/BN/BS/BW/RC/RE/RN/RS) เป็นของเก่า **ไม่ต้องใช้** · ยังคง**ห้าม hardcode** — โหลดจาก `GET /store/all-regions` ของระบบ SBP เดิม"),
                ("statementPeriodFrom", "YYYY-MM", "optional month picker", "Period Statement From; ส่ง ค.ศ. format YYYY-MM"),
                ("statementPeriodTo", "YYYY-MM", "optional month picker; must be >= from", "Period Statement To; validate range ก่อน call API"),
                ("page", "integer", "default 1; >=1", "pagination ของ preview table"),
                ("size", "integer", "default 20; max 100", "BE จำกัด page size เพื่อกัน query หนัก"),
                ("resultTable.storeCode", "string 5 digits", "display only", "คอลัมน์ 1 รหัสร้านถูกกระทบ"),
                ("resultTable.storeName", "string", "display only", "คอลัมน์ 2 ชื่อร้านถูกกระทบ"),
                ("resultTable.region", "string", "display only", "คอลัมน์ 3 ภาค"),
                ("resultTable.storeType", "string", "display only", "คอลัมน์ 4 ประเภทร้าน"),
                ("resultTable.impactMonth", "MM/YYYY ค.ศ.", "display only", "คอลัมน์ 5 เดือน/ปีที่ถูกกระทบ"),
                ("resultTable.statementPeriod", "MM/YYYY ค.ศ.", "nullable", "คอลัมน์ 6 Period Statement"),
                ("resultTable.newStoreCode", "string 5 digits or '-'", "display only", "คอลัมน์ 7 รหัสร้านเปิดกระทบ"),
                ("resultTable.newStoreName", "string or '-'", "display only", "คอลัมน์ 8 ชื่อร้านเปิดกระทบ"),
                ("resultTable.newStoreRegion", "string or '-'", "display only", "คอลัมน์ 9 ภาค (ร้านเปิดกระทบ)"),
                ("resultTable.newStoreType", "string or '-'", "display only", "คอลัมน์ 10 ประเภทร้าน (ร้านเปิดกระทบ)"),
                ("resultTable.compensationAmount", "number #,##0.00", ">=0", "คอลัมน์ 11 ยอดเงินชดเชย; align right"),
                ("derived.salesDataDays", "integer", "<60 = abnormal", "ข้อมูลประกอบสำหรับ class flag-red; ไม่ใช่ waitingDays"),
                ("resultTable.roundNo", "integer", ">=1", "คอลัมน์ 12 ครั้งที่"),
                ("resultTable.createdDate", "DD/MM/YYYY ค.ศ.", "required", "คอลัมน์ 13 วันที่สร้าง"),
                ("resultTable.docNo", "YYYY/xxxxx", "required", "คอลัมน์ 14 เลขที่เอกสาร; ใช้เปิด detail/preview"),
            ],
            [
                ("เปิด popup ร้าน", "ปุ่มแว่นขยายข้างรหัสร้านที่ถูกกระทบ", "GET /store/search (ระบบ SBP เดิม)", "เลือก store แล้วเติม storeCode/storeName"),
                ("ค้นหาข้อมูล", "ปุ่ม ค้นหาข้อมูล", "GET /api/v1/sbpgi/report/status-summary", "validate status (required) และคู่รหัสร้าน แล้ว render summary line + table 14 columns"),
                ("เคลียร์ค่าเริ่มใหม่", "ปุ่มเคลียร์ค่าเริ่มใหม่", "client state", "reset filter, summary, table และ error message"),
                ("Export Excel", "ปุ่ม Export Excel ท้าย filter", "GET /api/v1/sbpgi/report/status-summary/export", "ส่ง filter ชุดเดียวกับการค้นหา แล้วดาวน์โหลดไฟล์ .xlsx 14 คอลัมน์"),
                ("Open detail", "คลิกเลขที่เอกสารหรือ row", "navigate /sbpgi/document/{docNo} หรือ preview modal", "เปิดเอกสารที่เกี่ยวข้อง"),
            ],
            [
                ApiSpec("GET", "/store/search (ระบบ SBP เดิม)", "Popup เลือกร้านที่ถูกกระทบ", {"q": "00788", "type": "impacted"}, {"items": [{"storeCode": "00788", "storeName": "รัตนอุทิศ ซ.13", "region": "BN", "storeType": "SBP Type B"}]}),
                ApiSpec("GET", "/api/v1/sbpgi/report/status-summary", "ค้นหาข้อมูลรายงานตรวจสอบประกันรายได้ (14 คอลัมน์ · SDD สไลด์ 60)", {"status": "06", "impactedStoreCode": "00788", "newStoreCode": "00990", "periodStatementFrom": "2026-06-01", "periodStatementTo": "2026-06-30", "storeTypes": ["A", "B"], "regions": ["RSU", "BN"], "result": "APPROVE", "page": 1, "size": 20}, {"page": 1, "size": 20, "total": 10, "summary": {"totalItems": 10, "totalCompensationAmount": 439100.0, "overThresholdItems": 3, "abnormalSalesItems": 2}, "items": [{"impactedStoreCode": "00788", "impactedStoreName": "รัตนอุทิศ ซ.13", "impactedRegion": "RSU", "impactedStoreType": "B", "impactMonth": "2026-05", "periodStatement": "2026-06-07", "newStoreCode": "00990", "newStoreName": "เซเว่นฯ รัตนาธิเบศร์ 12", "newRegion": "RSU", "newStoreType": "A", "compensationAmount": 48200.0, "roundNo": 1, "createdDate": "2026-06-12", "docNo": "2026/00123"}]}),
                ApiSpec("GET", "/api/v1/sbpgi/report/status-summary/export", "Export Excel ด้วย filter เดียวกับการค้นหา", {"sameAsSearch": True, "format": "xlsx"}, {"contentType": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "fileName": "insurance-verification-2026.xlsx"}),
            ],
            ["เปิดหน้า Report", "โหลด reference status/region/store type ถ้ามี API (ภาคใหม่แสดง checkbox อัตโนมัติ)", "ผู้ใช้ระบุ filter 7 ตัวตาม SDD สไลด์ 60", "Validate status (required) · คู่รหัสร้านถูกกระทบ-เปิดกระทบ · Period Statement บังคับเมื่อสถานะ = เสร็จสิ้นดำเนินการ", "กด ค้นหาข้อมูล แล้ว call report API", "แสดงวันที่เป็น ค.ศ. ตามระบบ SBP เดิม (ไม่แปลงเป็น พ.ศ. — ตัดสินใจ 2026-08-06)", "render summary line และ table 14 คอลัมน์", "กด Export Excel แล้วส่ง filter เดียวกันไป export API"],
            ["status เป็น required ตัวเดียวก่อนค้นหา/export (resultCategory เป็นตัวเลือก · SDD สไลด์ 60)", "ระบุ impactedStoreCode แล้วต้องระบุ newStoreCode ด้วย", "Period Statement เป็นช่วงวันที่ ค.ศ. และ from <= to", "ตารางแสดง 14 คอลัมน์ครบและ export ออกครบ 14 คอลัมน์", "ยอดเงิน format #,##0.00 และ total summary ตรงกับผลรวม API", "แถวข้อมูลยอดขายไม่ครบ 60 วันใช้ class flag-red โดยอิง derived.salesDataDays < 60", "export ใช้ filter เดียวกับการค้นหาล่าสุด"],
            ["ไม่เลือก status แล้วค้นหาต้อง block", "ระบุร้านถูกกระทบแต่ไม่ระบุร้านเปิดกระทบ ต้อง block", "periodStatementFrom > periodStatementTo ต้อง error REPORT_DATE_RANGE_INVALID", "สถานะ = เสร็จสิ้นดำเนินการ แต่ไม่ระบุ Period Statement ต้อง block", "ค้นหาด้วยร้านถูกกระทบ", "เลือกหลาย region/storeType", "render table 14 columns", "export xlsx", "empty result แสดง summary เป็น 0"],
        ),
        Topic(
            "FE/LLDD-FE-Master-Data",
            "LLDD FE - Master Data",
            "FE",
            6.5,
            55,
            FE_OWNER_KITTISAK,
            "สร้างหน้าจอ master ที่ SBPGI ดูแลเอง 2 หน้า: ปัจจัยภายนอก (SCR-09 · k2-factors.html) "
            "และรายชื่อแบรนด์ร้านคู่แข่ง (k2-competitors.html · รหัส 01-11 ไทย+อังกฤษ) — "
            "หน้าผู้ปฏิบัติงาน/สิทธิ์เมนู/ตั้งค่าระบบ ไม่อยู่ในขอบเขตแล้ว (ใช้ของระบบ SBP เดิม)",
            ["k2-factors-01.png", "k2-competitors-01.png"],
            ["External factor master (SCR-09)", "Competitor brand master", "CRUD modal", "Active/inactive toggle"],
            [
                ("factorCode", "string", "required · unique · ห้ามซ้ำ", "คีย์ของปัจจัยภายนอก — แก้ไม่ได้หลังสร้าง"),
                ("factorName", "string", "required", "ชื่อปัจจัยภายนอกที่แสดงในหน้าเอกสาร"),
                ("description", "text", "optional", "คำอธิบายเพิ่มเติม"),
                ("competitorCode", "string(30)", "required · unique · รหัส 01-11", "คีย์ของแบรนด์คู่แข่ง — feed dropdown ร้านคู่แข่งในหน้าเอกสาร"),
                ("nameTh", "string(200)", "required", "ชื่อแบรนด์ภาษาไทย"),
                ("nameEn", "string(200)", "required", "ชื่อแบรนด์ภาษาอังกฤษ (ระบบเดิมเก็บทั้งสองภาษา)"),
                ("remark", "string(500)", "optional", "คอลัมน์ รายละเอียดเพิ่มเติม ของหน้า k2-competitors.html"),
                ("active", "boolean", "default true", "ปิดใช้งานแทนการลบเมื่อถูกอ้างในเอกสารแล้ว"),
            ],
            [
                ("Add/Edit", "modal action", "POST/PUT ของ master ที่กำลังเปิดอยู่", "ปิด modal + reload table"),
                ("Delete", "ปุ่มถังขยะ + confirm", "DELETE master API", "409 ถ้าถูกอ้างในเอกสารแล้ว → ให้ปิด active แทน"),
                ("Toggle active", "switch ในตาราง", "PUT พร้อม active", "แถวที่ปิดใช้งานไม่ขึ้นใน dropdown ของหน้าเอกสาร"),
            ],
            [
                ApiSpec("GET", "/api/v1/sbpgi/master/factors", "SCR-09 list/filter ปัจจัยภายนอก",
                        {"q": "ถนน", "active": True, "page": 1, "size": 20},
                        {"page": 1, "size": 20, "total": 1,
                         "items": [{"factorCode": "F001", "factorName": "ก่อสร้างถนน", "description": "ผลกระทบจากการก่อสร้าง", "active": True}]}),
                ApiSpec("POST", "/api/v1/sbpgi/master/factors", "SCR-09 เพิ่มปัจจัยภายนอก",
                        {"factorCode": "F001", "factorName": "ก่อสร้างถนน", "description": "ผลกระทบจากการก่อสร้าง", "active": True},
                        {"factorCode": "F001", "created": True}),
                ApiSpec("PUT", "/api/v1/sbpgi/master/factors/{code}", "SCR-09 แก้ไขปัจจัยภายนอก",
                        {"factorName": "ก่อสร้างถนนระยะยาว", "description": "กระทบการเข้าร้าน", "active": True},
                        {"factorCode": "F001", "updated": True}),
                ApiSpec("DELETE", "/api/v1/sbpgi/master/factors/{code}", "SCR-09 ลบปัจจัยภายนอกที่ยังไม่ถูกอ้างในเอกสาร",
                        {}, {"factorCode": "F001", "deleted": True}),
                ApiSpec("GET", "/api/v1/sbpgi/master/competitors", "list แบรนด์คู่แข่ง (master 11 รายการ)",
                        {"active": True},
                        {"total": 11, "items": [{"competitorCode": "01", "nameTh": "แฟมิลี่มาร์ท", "nameEn": "FamilyMart", "remark": "", "active": True}]}),
                ApiSpec("POST", "/api/v1/sbpgi/master/competitors", "เพิ่มแบรนด์คู่แข่ง",
                        {"competitorCode": "12", "nameTh": "ร้านตัวอย่าง", "nameEn": "Sample Shop", "remark": "", "active": True},
                        {"competitorCode": "12", "created": True}),
                ApiSpec("PUT", "/api/v1/sbpgi/master/competitors/{code}", "แก้ไขแบรนด์คู่แข่ง",
                        {"nameTh": "แฟมิลี่มาร์ท", "nameEn": "FamilyMart", "remark": "ปรับชื่อ", "active": True},
                        {"competitorCode": "01", "updated": True}),
                ApiSpec("DELETE", "/api/v1/sbpgi/master/competitors/{code}", "ลบแบรนด์คู่แข่งที่ยังไม่ถูกอ้างใน document_competitors",
                        {}, {"competitorCode": "12", "deleted": True}),
            ],
            ["Open master page", "Load table", "Open modal", "Validate required/unique", "Call API", "Reload table"],
            [
                "factorCode และ competitorCode ห้ามซ้ำ (409 CODE_DUPLICATE)",
                "nameTh และ nameEn ของคู่แข่งบังคับทั้งคู่",
                "ลบไม่ได้ถ้าถูกอ้างใน document_external_factors / document_competitors → 409 ให้ปิด active แทน",
                "ไม่มี reason และไม่มี audit log (ยกเลิกระบบ audit ของ master 2026-08-07)",
            ],
            [
                "duplicate factorCode",
                "duplicate competitorCode",
                "competitor ไม่ใส่ nameEn ต้อง block",
                "ลบ factor ที่ถูกอ้างในเอกสารแล้ว ต้องได้ 409",
                "toggle active แล้ว dropdown ในหน้าเอกสารต้องไม่แสดงแถวนั้น",
            ],
        ),
        Topic(
            "BE/LLDD-BE-API-Common-Contracts",
            "LLDD BE - API Common Contracts",
            "BE",
            2.4,
            20,
            BE_OWNER_BUTSABA,
            "กำหนดสัญญากลางของ REST API ทุกเส้นเพื่อไม่ให้ endpoint รายตัวตีความต่างกัน: transport/auth/error/format/pagination/action/RBAC/audit/idempotency",
            [],
            [
                "Base URL, content type, charset and request tracing",
                "Auth/JWT platform validation and service-token exception",
                "Standard success envelopes for list/detail/mutation",
                "Standard error envelope and HTTP status mapping",
                "Field format for date/month/docNo/storeCode/amount/percent",
                "Document action input/output contract",
                "RBAC/menu permission and editable section guard",
                "Audit/reason and idempotency rules",
            ],
            [
                ("Base URL", "/api/v1", "required", "ทุก endpoint ใช้ prefix นี้"),
                ("Content-Type", "application/json; charset=utf-8", "required for JSON", "multipart เฉพาะ attachments"),
                ("Authorization", "Bearer <JWT>", "required for user endpoints", "validate signature/expiry/role; platform provides token"),
                ("X-Service-Token", "opaque service token", "required for internal workflow/batch callbacks", "ใช้กับ /sbpgi/workflow/instances และ external callback ที่ไม่ใช่ user JWT"),
                ("X-Request-Id", "uuid/string", "optional but logged", "ถ้าไม่ส่ง BE generate แล้วคืนใน log/trace"),
                ("ErrorEnvelope", "{code,message}", "message Thai verbatim", "ห้ามเพิ่ม error shape อื่นใน endpoint รายตัว"),
                ("PageResponse<T>", "{page,size,total,items}", "page>=1 size<=100", "ใช้กับทุก GET list"),
                ("MutationResponse", "{message}", "message optional for simple save", "ถ้า workflow action ใช้ ActionResponse แทน"),
                ("docNo", "YYYY/xxxxx ค.ศ.", "path/query", "URL encode slash ตาม client/router; service ประกอบกลับเป็น docNo"),
                ("storeCode/newStoreCode", "string 5 digits", "preserve leading zero", "ห้ามใช้ numeric id แทนรหัสร้านใน payload"),
                ("date/month", "ISO-8601 ค.ศ.", "YYYY-MM-DD / YYYY-MM", "FE แสดง ค.ศ. เป็นค่าเริ่มต้น (buddhistEra=false) · แปลง พ.ศ. เฉพาะ component ที่เปิด flag"),
                ("amount/percent", "number", "2 decimal", "format display อยู่ FE; BE validate precision/range"),
                ("result", "verbatim from actionOptions", "required for /actions", "ต้องเป็นค่าที่ BE ส่งมาใน role profile ของเอกสารนั้น"),
                ("ActionResponse", "{statusCode,nextSection,message}", "required for /actions", "FE resolve label จาก /sbpgi/lookup/document-statuses; mutation response ไม่คืน label ไทยซ้ำ"),
                ("reason", "text", "ไม่บังคับแล้ว (ยกเลิกระบบ audit ของ master 2026-08-07)", "ไม่มีปลายทางเก็บ — ถ้าส่งมาให้ละเว้น"),
            ],
            [
                ("Authenticate user endpoint", "middleware", "auth.verifyJwt", "req.user = employeeId/roleCode/sectionCode"),
                ("Authorize menu/role", "middleware/service", "rbac.requireMenu/requireRole", "403 FORBIDDEN เมื่อไม่มีสิทธิ์"),
                ("Validate request", "controller", "zod schema", "400 VALIDATION envelope"),
                ("Return list", "repository/service", "PageResponse<T>", "pagination shape เดียวกัน"),
                ("Submit document action", "service", "documentAction.service.submit", "return ActionResponse"),
                ("Write audit", "transaction", "audit.service.write", "reason/updated_by/old_value/new_value"),
                ("Handle idempotency", "service", "requestId/business key", "duplicate returns existing result or 409 per endpoint rule"),
            ],
            [
                ApiSpec("ALL", "/api/v1/sbpgi/*", "Standard error envelope", None, {"code": "VALIDATION", "message": "ข้อความภาษาไทยตรงตาม SRS"}),
                ApiSpec("GET", "/api/v1/sbpgi/*", "Standard list envelope เมื่อ endpoint เป็นรายการ", {"page": 1, "size": 20}, {"page": 1, "size": 20, "total": 0, "items": []}),
                ApiSpec("POST", "/api/v1/sbpgi/document/{docNo}/actions", "**อ้างอิงเท่านั้น — เจ้าของ endpoint นี้คือ LLDD-BE-API-Document-Workflow-Actions (Tunyatorn)** · ยกมาเป็นตัวอย่างสัญญา action กลางที่ทุกเส้นต้องยึด (ตัวอย่าง currentSection=01 จึงเปลี่ยนไป 02)", {"result": "เห็นควรชดเชย", "comment": "เห็นควรชดเชยตามหลักเกณฑ์"}, {"statusCode": "02", "nextSection": "02", "message": "submitted"}),
            ],
            [
                "Request enters logging middleware and request id is attached",
                "BffUserGuard ตรวจ x-api-key แล้ว map BFF header เป็น user context — 🔴 SBPGI ไม่ออก/ไม่ตรวจ JWT เอง (login อยู่ที่ Cognito ฝั่ง BFF) · เส้น service-token ใช้ API key แยก",
                "RBAC guard checks role/menu/current workflow task owner",
                "Validate params/query/body with shared schema conventions",
                "Service executes business rule and document action if relevant",
                "Mutation writes domain row and audit/reason in the same transaction",
                "Controller maps result to standard envelope or throws AppError",
                "Error handler maps all failures to `{code,message}` only",
            ],
            [
                "ทุก endpoint ต้องใช้ common contract นี้",
                "ไม่มี endpoint คืน error shape อื่นนอกจาก `{code,message}`",
                "401/403/404/409/422/413/415 mapping คงที่และ test ได้",
                "GET list ทุกเส้นคืน `{page,size,total,items}`",
                "/actions รับ `{result,comment}` เท่านั้นและคืน `{statusCode,nextSection,message}`",
                "RBAC ใช้ role/menu/current task owner ฝั่ง BE เป็น source of truth",
                "workflow action ต้องเขียน consideration_logs · master mutation ไม่มี audit แล้ว (ยกเลิกระบบ audit ของ master 2026-08-07)",
            ],
            ["missing JWT 401", "role forbidden 403", "validation error 400", "not found 404", "duplicate conflict 409", "list envelope", "action transition envelope", "audit reason required", "service token endpoint"],
        ),
        Topic(
            "BE/LLDD-BE-API-Document-List-Search",
            "LLDD BE - API Document List and Search",
            "BE",
            3.3,
            28,
            BE_OWNER_BUTSABA,
            "ออกแบบ APIs สำหรับงานรอดำเนินการและค้นหาเอกสารที่เกี่ยวข้อง",
            [],
            ["Inbox tasks API", "Document search API", "Pagination", "Status/year filter", "Abnormal row support"],
            common_doc_fields() + [
                ("year", "ค.ศ. YYYY", "required for /sbpgi/document", "ไม่ระบุคืน 400 ตาม SRS · BE ผ่าน toAD() เผื่อ client ส่ง พ.ศ."),
                ("page/size", "integer", "page>=1 size<=100", "pagination"),
            ],
            [
                ("Inbox tasks", "GET", "task.service.searchOpenTasks", "return waiting list"),
                ("Document search", "GET", "document.service.search", "return related list"),
            ],
            [
                ApiSpec("GET", "/api/v1/sbpgi/document/tasks", "Inbox tasks API", {"sectionCode": "06", "page": 1, "size": 20}, {"items": [{"docNo": "2026/00123", "waitingDays": 3}]}),
                ApiSpec("GET", "/api/v1/sbpgi/document", "Document search API", {"year": 2026, "storeCode": "00788", "status": "06", "page": 1}, {"items": [{"docNo": "2026/00123", "statusCode": "06"}]}),
            ],
            ["Read JWT section/role", "Validate year for documents", "Build filter query", "Join impacted_stores", "Return page result"],
            ["year missing fails for /sbpgi/document", "leading zero storeCode preserved", "pagination returns total", "status filter works"],
            ["tasks by section", "documents missing year", "store search", "empty result"],
        ),
        Topic(
            "BE/LLDD-BE-API-Document-Create-Update",
            "LLDD BE - API Document Create and Update",
            "BE",
            4.2,
            36,
            BE_OWNER_BUTSABA,
            "ออกแบบ APIs สำหรับสร้างเอกสารใหม่และบันทึกส่วนย่อยของเอกสาร",
            [],
            ["Create document", "Duplicate guard", "Running doc number", "Partial update", "Business validation"],
            common_doc_fields() + [
                ("requestId", "string", "optional", "ใช้ trace request; duplicate guard หลักเป็น business key"),
                ("source", "MANUAL|FS", "required", "แยกแหล่งสร้างเอกสาร"),
            ],
            [
                ("Create document", "POST", "document.service.create", "create doc + first workflow task"),
                ("Update document section", "PUT", "document.service.updateSections", "save editable sections"),
            ],
            [
                ApiSpec("POST", "/api/v1/sbpgi/document", "Create document API", {"impactedStoreCode": "00788", "impactMonth": "2026-06", "source": "MANUAL", "newStoreCode": "00990", "roundNo": 1, "reason": "manual create", "requestId": "uuid"}, {"docNo": "2026/00124", "statusCode": "06"}),
                ApiSpec("PUT", "/api/v1/sbpgi/document/{docNo}", "Update document partial sections", {"newStores": [{"newStoreCode": "00990", "compensatePercent": 60, "sourceSystem": "ALLMAP"}, {"newStoreCode": "01180", "compensatePercent": 40, "sourceSystem": "USER"}]}, {"message": "saved"}),
            ],
            ["Validate required fields", "Check duplicate store/month", "Generate docNo", "Insert compensation_documents", "Open workflow task", "Save section updates in transaction"],
            ["duplicate business key returns 409", "docNo format YYYY/xxxxx", "compensatePercent sum=100", "requestId trace does not replace business duplicate guard"],
            ["create success", "create duplicate", "update allocation invalid", "permission denied section"],
        ),
        Topic(
            "BE/LLDD-BE-API-Document-Detail-Aggregate",
            "LLDD BE - API Document Detail Aggregate",
            "BE",
            4.5,
            38,
            BE_OWNER_BUTSABA,
            "ออกแบบ aggregate API สำหรับโหลดรายละเอียดเอกสารครบทุก section ให้หน้า FE detail",
            [],
            ["Document aggregate query", "Role profile output", "Store impact/new-store/factor mapping", "Compensation summary", "Related master lookup"],
            common_doc_fields() + [
                ("docNo", "YYYY/xxxxx", "required path param", "หาเอกสารและ section ทั้งหมด"),
                ("visibleSections/editableSections", "array", "computed by BE", "FE render ตาม key ที่ส่งมาเท่านั้น"),
                ("actionOptions", "array", "computed by BE", "radio options + requireComment สำหรับ action panel"),
            ],
            [
                ("Get detail", "GET", "documentAggregate.service.getByDocNo", "return 12 sections"),
                ("Get lookup", "GET", "lookup service", "return status/competitors/factors"),
            ],
            [
                ApiSpec("GET", "/api/v1/sbpgi/document/{docNo}", "Document aggregate API", {"docNo": "2026/00123"}, {"docNo": "2026/00123", "statusCode": "06", "viewerRbacRoleCode": "R-XX", "roleProfileCode": "P-06", "visibleSections": ["doc-header", "sec-sales", "sec-map", "sec-newstore", "sec-competitor", "sec-factor", "sec-attach", "sec-comp-history", "sec-decision-history", "sec-action"], "editableSections": [], "canUploadAttachment": True, "canAction": True, "actionOptions": [{"label": "เห็นควรไม่ชดเชย", "requireComment": True}], "impactedStore": {"storeCode": "00788"}, "newStores": []}),
                ApiSpec("GET", "/api/v1/sbpgi/master/competitors", "**อ้างอิงเท่านั้น — เจ้าของ endpoint นี้คือ LLDD-BE-API-Report-and-Master-Data (Peerakorn)** · เอกสารนี้เป็นผู้ใช้: อ่าน master คู่แข่งมาทำ dropdown ในหน้าเอกสาร", {"q": "lotus"}, {"items": [{"competitorCode": "C007", "competitorName": "Lotus Express"}]}),
            ],
            ["Validate docNo", "Load header", "Load child sections", "Compute role profile", "Map to FE response shape", "Return aggregate"],
            ["404 when doc not found", "role profile output matches FE Document Detail spec", "nullable section returns empty array", "amount/date formatting source consistent"],
            ["detail success", "detail not found", "role profile output", "empty child sections"],
        ),
        Topic(
            "BE/LLDD-BE-API-Document-Workflow-Actions",
            "LLDD BE - API Document Workflow Actions",
            "BE",
            4.0,
            34,
            BE_OWNER,
            "ออกแบบ APIs สำหรับรับผลพิจารณา ตรวจสิทธิ์ action และบันทึก audit/consideration log",
            [],
            ["Submit action", "Action owner guard", "Amount threshold reference", "Send back result", "Audit and email rule"],
            [
                ("docNo", "YYYY/xxxxx", "required", "path param"),
                ("result", "verbatim from actionOptions", "required", "ต้องเป็นค่าที่ API detail ส่งมาให้ผู้ใช้ในเอกสารนั้น"),
                ("comment", "text", "required for return/reject", "trim ก่อนบันทึก"),
            ],
            [
                ("Submit action", "POST", "documentAction.service.submit", "submit result and update status"),
                ("Write audit", "transaction", "considerationLog.repository.insert", "record action history"),
                ("Send email", "async", "notification.service.sendByStatusRule", "notify next owner"),
            ],
            [
                ApiSpec("POST", "/api/v1/sbpgi/document/{docNo}/actions", "Document action API ตัวอย่างเมื่อ currentSection=01 จึงเปลี่ยนไป 02", {"result": "เห็นควรชดเชย", "comment": "เห็นควรชดเชยตามหลักเกณฑ์"}, {"statusCode": "02", "nextSection": "02", "message": "submitted"}),
                ApiSpec("GET", "/api/v1/sbpgi/document/{docNo}/timeline", "**อ้างอิงเท่านั้น — เจ้าของ endpoint นี้คือ LLDD-BE-API-Attachment-Sales-Timeline (Peerakorn)** · เอกสารนี้อ้างเพราะ action ที่ส่งผลพิจารณาเป็นตัวเขียน consideration_logs ที่ timeline อ่าน", {"docNo": "2026/00123"}, {"items": [{"section": "06", "result": "ชดเชย"}]}),
            ],
            ["Lock current action task", "Validate owner and selected result against actionOptions", "Apply server-side business rule", "Update document/task", "Insert consideration_logs", "Trigger email"],
            ["non-owner returns 403", "missing result returns exact SRS message", "invalid result for this role profile returns 422", "duplicate submit blocked by current open task lock", "audit written in same transaction"],
            ["submit compensate", "submit not compensate", "send back", "invalid result", "duplicate action"],
        ),
        Topic(
            "BE/LLDD-BE-API-Workflow-Instances",
            "LLDD BE - Workflow Engine and API Workflow Instances",
            "BE",
            3.3,
            28,
            BE_OWNER,
            "ออกแบบ Workflow Engine ภายในและ POST /api/v1/sbpgi/workflow/instances สำหรับเปิด workflow จาก Job 8b แทน K2 REST StartInstance โดยเป็นเจ้าของ Gen Flow Gate W/Y/N",
            [],
            [
                "Internal Workflow Engine API only",
                "No FE screen and no Flow page work",
                "Gen Flow Gate W/Y/N owner",
                "Require compensation document created by Job 8",
                "Create workflow instance and first task section 06",
                "Idempotency and rerun behavior for Job 8b",
            ],
            [
                ("impactProcessId", "integer/string", "required", "อ้าง fgi_impact_processes และ compensation_documents ที่ Job 8 สร้างแล้ว"),
                ("sourceJobNo", "string", "required fixed 8b", "ใช้ trace รอบรันใน application log (structured) — ไม่มีตาราง job_run_histories แล้ว"),
                ("requestId", "uuid", "required", "idempotency key ต่อ impactProcessId + sourceJobNo"),
                ("workflow_generation_status", "W|Y|N", "computed", "W=ข้อมูลยังไม่พร้อมเพื่อ rerun, Y=เปิด workflow สำเร็จ, N=ไม่เข้าเกณฑ์ถาวร"),
                ("branchType/distanceKm", "enum/number|null", "required by gate", "branch นอกเซ็ตหรือระยะเกินตั้ง N; ระยะยังไม่มีค่าคง W"),
                ("growthRateDiff", "number|null", "<= -10 required by gate", "NULL คง W; ค่ามากกว่า -10 ตั้ง N แบบถาวร"),
                ("dvUserId/juristic", "string|null", "DV required; juristic must differ", "DV ว่างหรือ juristic เดียวกันตั้ง N; juristic ยังไม่พร้อมคง W"),
                ("salesStatus", "Y|N", "required by gate", "ค่าอื่นคง W และคืน 422"),
            ],
            [
                ("Open workflow", "POST", "workflowInstance.service.openFromImpact", "ผ่าน gate แล้วสร้าง/คืน instance"),
                ("Check status", "GET", "/api/v1/sbpgi/workflow/instances/{id}", "อ่าน instance status"),
                ("Summary", "GET", "/api/v1/sbpgi/workflow/summary", "ตัวเลข W/Y/N และงานค้างต่อ section"),
            ],
            [
                ApiSpec("POST", "/api/v1/sbpgi/workflow/instances", "เปิด workflow ภายในจาก impact process; เรียกโดย Job 8b ผ่าน service token ไม่ใช่ FE", {"impactProcessId": 901234, "sourceJobNo": "8b", "requestId": "job8b-901234-256907"}, {"docNo": "2026/00123", "instanceId": "WF-2026-00123", "workflowGenerationStatus": "Y", "firstSection": "06", "statusCode": "06", "status": "รอฝ่าย SBP DSA ดำเนินการ"}),
                ApiSpec("GET", "/api/v1/sbpgi/workflow/instances/{id}", "อ่านสถานะ workflow instance", {"id": "WF-2026-00123"}, {"instanceId": "WF-2026-00123", "docNo": "2026/00123", "status": "ACTIVE", "currentSection": "06"}),
                ApiSpec("GET", "/api/v1/sbpgi/workflow/summary", "สรุป W/Y/N และงานค้างต่อ section สำหรับ monitor", {"period": "2026-07"}, {"workflowGeneration": {"W": 12, "Y": 342, "N": 8}, "openTasksBySection": [{"sectionCode": "06", "count": 24}]}),
            ],
            [
                "Validate service token and idempotency key",
                "Load impact process and current workflow_generation_status",
                "Reject if status is already Y and return existing doc/instance idempotently",
                "Evaluate Gen Flow Gate in one service: status W, branch type allowlist, DV present, juristic different, growth_rate_diff <= -10, sales_status in Y/N",
                "If branch type is outside allowlist, distance exceeds threshold, DV is missing, juristic is the same, or growth_rate_diff > -10, update workflow_generation_status=N and return 200 with permanent-skip reason",
                "If distance/juristic/growth data is NULL or sales_status is not ready, keep workflow_generation_status=W and return 422 reason so Job 8b can rerun",
                "If gate passes, require compensation_documents from Job 8, open workflow via @srm/glb-workflow (initializeWorkflow + addPreApprover at state 06 — function names confirmed 2026-08-14 from the library's own LLDD, sheet Detail), then update fgi_impact_processes.workflow_generation_status=Y in one transaction",
                "Enqueue notification summary outside transaction after commit",
            ],
            [
                "ไม่มี FE screen หรือ Flow page deliverable เพิ่มจาก LLDD นี้",
                "Job 8b ต้องเรียก API/service นี้และไม่ duplicate Gen Flow Gate",
                "ไม่เรียก K2 REST StartInstance และไม่สร้างไฟล์ BPM06001O/2O/3O",
                "ผ่าน gate แล้ว transaction ต้องมี document + instance + first task + Y ครบ หรือ rollback ทั้งหมด",
                "fail ถาวร (branch type, distance over threshold, missing DV, same juristic, growth not met) ต้องตั้ง N; เฉพาะข้อมูล distance/juristic/growth/sales status ยังไม่พร้อมจึงคง W",
                "idempotent rerun ไม่สร้าง docNo/instance/task ซ้ำ",
            ],
            ["gate pass creates workflow", "branch type/distance over threshold sets N", "distance NULL keeps W", "missing DV sets N", "same juristic sets N", "growth NULL keeps W but growth > -10 sets N", "sales status NULL keeps W", "duplicate request returns existing instance", "transaction rollback on task insert failure", "service token missing returns 401"],
        ),
        Topic(
            "BE/LLDD-BE-API-Attachment-Sales-Timeline",
            "LLDD BE - API Attachment Sales and Timeline",
            "BE",
            3.5,
            30,
            BE_OWNER_PEERAKORN,
            "ออกแบบ APIs สำหรับไฟล์แนบ ข้อมูลยอดขายเพิ่มเติม และ timeline/history",
            [],
            ["Attachment metadata", "Upload/download adapter", "Sales 4 windows", "Timeline query", "File validation"],
            common_doc_fields() + [
                ("file", "multipart", "<=5MB", "validate extension and content type"),
                ("sectionCode", "string", "required on upload", "บันทึกว่าแนบในขั้นไหน"),
            ],
            [
                ("Upload attachment", "POST multipart", "attachment.service.upload", "store file and metadata"),
                ("Download attachment", "GET", "attachment.service.download", "stream file"),
                ("Get sales", "GET", "sales.service.getDocumentSales", "return sales windows"),
            ],
            [
                ApiSpec("POST", "/api/v1/sbpgi/document/{docNo}/attachments", "Upload attachment API", {"file": "multipart <= 5MB", "sectionCode": "06"}, {"attachId": 771, "fileName": "evidence.pdf"}),
                ApiSpec("GET", "/api/v1/sbpgi/document/{docNo}/attachments/{attachId}/download", "ดาวน์โหลดไฟล์แนบรายไฟล์ผ่าน BE — ตรวจสิทธิ์เอกสาร + attachment ต้องเป็นของ docNo + scan_status=CLEAN ก่อน stream", {}, {"contentType": "application/pdf", "note": "binary stream · ไฟล์จริงอยู่บน S3 ผ่าน service ของระบบ SBP เดิม"}),
                ApiSpec("GET", "/api/v1/sbpgi/document/{docNo}/attachments/download-all", "ดาวน์โหลดไฟล์แนบทั้งหมดเป็น .zip — ไม่มีไฟล์ที่ผ่าน scan เลยตอบ 404 (ไม่คืน zip เปล่า)", {}, {"contentType": "application/zip", "fileName": "2026-00123-attachments.zip"}),
                ApiSpec("GET", "/api/v1/sbpgi/document/{docNo}/sales", "Sales detail API", {"docNo": "2026/00123"}, {"growthRateDiff": -12.45, "totalWorkingDays": 60, "windows": [{"label": "ก่อนเปิด 15 วัน", "rows": []}]}),
                ApiSpec("GET", "/api/v1/sbpgi/document/{docNo}/timeline", "Timeline/history API", {"docNo": "2026/00123"}, {"items": []}),
            ],
            ["Validate docNo/permission", "Validate file size/type", "Store file metadata", "Load sales summary and transactions", "Return timeline ordered by action time"],
            ["file >5MB returns 413", "unsupported file type returns 415", "sales windows are ordered", "timeline newest/oldest order matches FE expectation"],
            ["upload success", "upload too large", "download missing file", "sales not found", "timeline empty"],
        ),
        Topic(
            "BE/LLDD-BE-API-Lookup",
            "LLDD BE - API Lookup",
            "BE",
            4.7,
            40,
            BE_OWNER,
            "ออกแบบ APIs กลุ่ม lookup ที่ใช้ร่วมทุกหน้าจอของ SBP Mall",
            [],
            ["Lookup APIs", "Auth endpoints are platform reference only"],
            [
                ("q", "string", "optional", "ใช้ค้นหา ร้าน (store/mas_store ระบบเดิม) · พนักงาน (business_user ระบบเดิม) · คู่แข่ง (competitors ของ SBPGI)"),
                ("type", "impacted|new", "required for /store/search (ระบบ SBP เดิม)", "เลือกแหล่งร้านถูกกระทบ/ร้านเปิดใหม่"),
                ("roleCode", "00-10", "required for permission", "group ของ auth-backend (ระบบเดิม) — SBPGI ไม่มีตาราง roles"),
                ("menuCode", "string", "required for permission", "menu ของ auth-backend (ระบบเดิม) — SBPGI ไม่มีตาราง menus"),
                ("templateCode", "EM-01..EM-08", "required", "email template key"),
                ("reason", "text", "ไม่บังคับแล้ว", "ไม่มีปลายทางเก็บ (ยกเลิกระบบ audit ของ master 2026-08-07)"),
            ],
            [
                ("Store lookup", "GET", "lookup.service.searchStores", "return impacted/new stores"),
                ("Employee lookup", "GET", "employee backend เดิมของระบบ SBP (ไม่ใช่ endpoint ของ SBPGI)", "return business_user for operator popup"),
                ("Permission save", "PUT", "rbac.service.saveMenuPermission", "update can_access and audit"),
                ("Email template save/reset", "PUT/POST", "notificationTemplate.service", "update/reset template and audit"),
            ],
            [
                ApiSpec("GET", "/store/search (ระบบ SBP เดิม)", "ค้นหาร้านสำหรับ popup", {"q": "00788", "type": "impacted"}, {"items": [{"storeCode": "00788", "storeName": "รัตนอุทิศ ซ.13"}]}),
                ApiSpec("GET", "/api/v1/sbpgi/lookup/document-statuses", "รายการสถานะเอกสาร verbatim", {}, {"items": [{"statusCode": "06", "statusName": "รอฝ่าย SBP DSA ดำเนินการ"}]}),
                ApiSpec("GET", "/api/v1/sbpgi/lookup/workflow-sections", "รายการ section 5 ขั้น", {}, {"items": [{"sectionCode": "06", "sectionName": "ฝ่าย SBP DSA"}]}),
            ],
            ["Validate query", "Read/write table by domain", "Return standard envelope for list endpoints"],
            ["status label ต้องเป็น verbatim", "permission mutation ต้อง audit", "SBPGI เรียก email-lib ส่งอีเมลเอง โดยใช้เลข template จาก workflow_route.email_id (ปิด DP-5 · 2026-08-14)", "Auth Group 1 เป็น platform/external reference ไม่ใช่งาน implement ใน LLDD นี้"],
            ["store lookup", "status lookup", "permission save without reason", "email template reset", "(ตัดออก) audit log search"],
        ),
        Topic(
            "BE/LLDD-BE-API-Report-and-Master-Data",
            "LLDD BE - API Report and Master Data",
            "BE",
            5.6,
            48,
            BE_OWNER_PEERAKORN,
            "ออกแบบ APIs สำหรับรายงานตรวจสอบประกันรายได้ และ Master Data ที่ SBPGI ดูแลเอง (ปัจจัยภายนอก + รายชื่อคู่แข่ง)",
            [],
            ["Report query service", "Excel export (14 columns, SDD slide 60)", "Operator/factor CRUD", "Report filters"],
            [
                ("year", "ค.ศ. YYYY", "required for report", "return 400 if missing · BE ผ่าน toAD() เผื่อ client ส่ง พ.ศ."),
                ("status", "statusCode string", "required", "6 สถานะเอกสาร; verbatim จาก sps_store.workflow_status ของ @srm/glb-workflow"),
                ("result", "APPROVE|REJECT|CANCELLED|PENDING", "optional for report (บังคับเฉพาะ status)", "maps to consideration_logs.result_category ล่าสุด · CANCELLED = ยกเลิกโดยระบบ (เพิ่ม 2026-08-10)"),
                ("region", "array/string", "optional", "13 region codes; multi-select"),
                ("storeType", "array ของ BranchTypeFGIName", "optional", "**7 ค่า** `A B C D E PTT บริษัท` (ยืนยันจาก master `BranchTypeProfile` ของ `CPA_FRN_FGI` 2026-08-10) · multi-select · **ห้าม hardcode** ให้โหลดจาก `GET /common/common-code` ของระบบ SBP เดิม"),
                ("impactedStoreCode", "string 5 digits", "optional", "คง leading zero"),
                ("newStoreCode", "string 5 digits", "optional", "คง leading zero"),
                ("reason", "text", "required mutation", "audit reason"),
                ("page/size", "integer", "page>=1 size<=100", "pagination"),
            ],
            [
                ("Report preview", "GET", "report.service.search", "paginated rows"),
                ("Report export", "GET", "report.service.exportCsv", "csv stream"),
                ("Master mutation", "POST/PUT/DELETE", "master.service.save", "อัปเดต row ของ master"),
            ],
            [
                ApiSpec("GET", "/api/v1/sbpgi/report/status-summary", "รายงานตรวจสอบประกันรายได้", {"year": 2026, "status": "06", "result": "APPROVE", "region": ["RSU"], "storeType": ["A"], "impactedStoreCode": "00788", "newStoreCode": "00990", "page": 1, "size": 20}, {"page": 1, "size": 20, "total": 0, "items": []}),
                ApiSpec("GET", "/api/v1/sbpgi/report/status-summary/export", "Export Excel", {"year": 2026, "status": "06", "result": "APPROVE", "region": ["RSU"], "storeType": ["A"], "impactedStoreCode": "00788", "newStoreCode": "00990"}, {"fileName": "insurance-verification-2026.xlsx"}),
                ApiSpec("GET", "/api/v1/sbpgi/master/factors", "อ่านปัจจัยภายนอก", {"q": "ก่อสร้าง", "active": True, "page": 1, "size": 20}, {"page": 1, "size": 20, "total": 1, "items": [{"factorCode": "ROAD", "factorName": "ก่อสร้างถนน", "description": "ปิดช่องทางจราจร", "active": True}]}),
                ApiSpec("POST", "/api/v1/sbpgi/master/factors", "สร้างปัจจัยภายนอก", {"factorCode": "ROAD", "factorName": "ก่อสร้างถนน", "description": "ปิดช่องทางจราจร", "active": True, "reason": "เพิ่มปัจจัยใหม่"}, {"factorCode": "ROAD", "factorName": "ก่อสร้างถนน", "active": True}),
                ApiSpec("PUT", "/api/v1/sbpgi/master/factors/{code}", "แก้ปัจจัยภายนอก", {"factorName": "ก่อสร้างและปิดถนน", "description": "ปิดช่องทางจราจรบางส่วน", "active": True, "reason": "ปรับคำอธิบาย"}, {"factorCode": "ROAD", "factorName": "ก่อสร้างและปิดถนน", "active": True}),
                ApiSpec("GET", "/api/v1/sbpgi/master/competitors", "master แบรนด์คู่แข่ง 11 รายการ (รหัส 01-11) — เป็นแหล่งของ dropdown ร้านคู่แข่งในหน้าเอกสารด้วย", {"q": "108"}, {"items": [{"code": "01", "nameTh": "108 Shop", "nameEn": "108 Shop", "remark": "", "isActive": True}]}),
                ApiSpec("POST", "/api/v1/sbpgi/master/competitors", "เพิ่มแบรนด์คู่แข่ง — code/nameTh/nameEn บังคับ · รหัสซ้ำตอบ 409", {"code": "12", "nameTh": "ร้านคู่แข่งรายใหม่", "nameEn": "New Competitor", "remark": ""}, {"code": "12", "message": "saved"}),
                ApiSpec("PUT", "/api/v1/sbpgi/master/competitors/{code}", "แก้ชื่อ/สถานะ — ห้ามแก้ code เพราะถูกอ้างจาก document_competitors", {"nameTh": "ลอว์สัน 108", "nameEn": "Lawson 108", "remark": "", "isActive": True}, {"message": "saved"}),
                ApiSpec("DELETE", "/api/v1/sbpgi/master/competitors/{code}", "ลบแบรนด์คู่แข่ง — ถูกอ้างในเอกสารแล้วตอบ 409", {}, {"message": "deleted"}),
                ApiSpec("DELETE", "/api/v1/sbpgi/master/factors/{code}", "ลบปัจจัยภายนอกที่ไม่ถูกใช้งาน", {"reason": "ยกเลิกค่าทดสอบ"}, {"factorCode": "ROAD", "deleted": True}),
            ],
            ["Validate filter", "Build query", "Apply pagination/export mode", "Return rows or CSV", "For mutations validate reason and write audit"],
            ["missing year/status/result fails", "export uses same filters as preview", "master edit requires reason", "config locked value cannot edit"],
            ["report missing year", "report export", "factor duplicate", "operator audit", "config locked"],
        ),
        Topic(
            "BE/LLDD-BE-Job-Batch-Email-SRM",
            "LLDD BE - Job Batch and Email Integration",
            "BE",
            6.4,
            54,
            BE_OWNER_PEERAKORN,
            "ออกแบบ Backend contracts สำหรับ batch runner (อ่าน config จาก backend), interface tracking/pending ACK และ Notification Service (ส่งผ่าน @gosoft-sbp/email-lib) — ไม่มี Job Admin API, Email Template API (2026-08-06) และไม่มี SRM inbound adapter แล้ว (2026-08-07)",
            [],
            ["Interface tracking และ pending ACK APIs (3 เส้น)", "Job runner guard และ application log", "Notification adapter ผ่าน @gosoft-sbp/email-lib", "STA ACK callback", "ไม่มี Batch Job Admin API และไม่มี inbound endpoint ของ SRM"],
            [
                ("jobNo", "string", "required", "maps to job registry"),
                ("sourceRefNo", "string", "required for SRM", "idempotency key"),
                ("templateCode", "EM-xx", "required", "email template key"),
                ("transactionId", "uuid", "generated", "integration log key"),
            ],
            [
                ("Run job", "POST", "jobRunner.run", "queued/run history"),
                ("Receive SRM", "POST", "srmIntegration.ingest", "transaction result"),
                ("Preview email", "POST", "emailTemplate.render", "merged subject/body"),
            ],
            [
                ApiSpec("GET", "/api/v1/sbpgi/interface/tracking", "ค้นสถานะ interface ตาม dataset/business key/status/ช่วงเวลา", {"dataName": "COMPENSATE_INIT_I", "status": "SENT", "pending": True, "sentFrom": "2026-07-01T00:00:00+07:00", "sentTo": "2026-07-22T23:59:59+07:00", "page": 1, "size": 20}, {"page": 1, "size": 20, "total": 1, "items": [{"trackingId": 9912, "dataName": "COMPENSATE_INIT_I", "direction": "OUT", "businessKey": "2026/00098", "docNo": "2026/00098", "fileName": "COMPENSATE_INIT_I_25690722.dat", "status": "SENT", "sentAt": "2026-07-20T17:02:00+07:00", "ackedAt": None, "returnCode": None, "ageHours": 41}]}),
                ApiSpec("GET", "/api/v1/sbpgi/interface/pending-ack", "รายการ ACK ค้างตาม watchdog rule อายุอย่างน้อย 1 วัน", {"thresholdHours": 24, "dataName": "COMPENSATE_INIT_I", "page": 1, "size": 20}, {"page": 1, "size": 20, "total": 1, "count": 1, "items": [{"trackingId": 9912, "dataName": "COMPENSATE_INIT_I", "businessKey": "2026/00098", "docNo": "2026/00098", "fileName": "COMPENSATE_INIT_I_25690722.dat", "sentAt": "2026-07-20T17:02:00+07:00", "ageHours": 41, "returnCode": None}]}),
                ApiSpec("POST", "/api/v1/sbpgi/interface/sta/ack", "STA ACK callback ให้ Job 10 เป็น safety net", {"transactionId": "TX-001", "returnCode": "A", "receivedAt": "2026-07-20T10:00:00+07:00"}, {"message": "acknowledged"}),
                # 2026-08-07: ตัด ApiSpec `POST /api/v1/integrations/srm/income-guarantee` ออก —
                # "SRM" ไม่ใช่ระบบต้นทาง เป็นเพียง prefix ของชื่อ resource (srm-sps-spsap-*) ·
                # SDD GI สไลด์ 75-77 ว่างเปล่า · ไม่มีเส้นนี้ใน 29 เส้นของ api.md ·
                # และหน้าที่ซ้ำกับ POST /sbpgi/document ที่ pipeline ใช้สร้างเอกสารอยู่แล้ว
            ],
            ["Receive request", "Validate schema", "Check idempotency", "Process records", "Log success/failure", "Return summary"],
            ["job run guard prevents duplicate running job", "email preview renders variables", "failed records include detail", "ไม่มี inbound endpoint ของ SRM แล้ว (ตัด 2026-08-07) — เอกสารต้องไม่อ้างถึงอีก"],
            ["run job", "run duplicate", "interface tracking filter", "pending ACK watchdog", "STA ACK callback", "email preview"],
        ),
        *new_be_design_topics(),
    ]
    base.extend(document_detail_role_topics())
    db_map = {
        "BE/LLDD-BE-API-Document-List-Search": [
            ("workflow_transaction / workflow_approver (@srm/glb-workflow)", "R", "อ่าน inbox ผ่าน getPendingFlowByUser() · เฉพาะ section 06 ต้อง union เอกสารที่จบด้วย หยุดชดเชยประกันรายได้ เข้ามาด้วย (stoppedReopenable)"),
            ("compensation_documents", "R", "ค้นเอกสารตาม year/status/store"),
            ("impacted_stores", "R", "ชื่อร้าน ภาค และข้อมูลร้าน"),
            ("fgi_impact_sales_summaries", "R", "flag ข้อมูลผิดปกติ/ยอดขายไม่ครบ 60 วัน"),
            ("consideration_logs", "R", "ผลการพิจารณาสุดท้าย — คัดเอกสารที่จบด้วย หยุดชดเชยประกันรายได้ เข้าคิวของ section 06 (SDD สไลด์ 46 ข้อ 1.9)"),
        ],
        "BE/LLDD-BE-API-Document-Create-Update": [
            ("compensation_documents", "R/W", "สร้างหัวเอกสารและแก้ไข section หลัก"),
            ("workflow_transaction / workflow_approver (@srm/glb-workflow)", "W (ผ่าน lib)", "เปิด workflow งานแรกตอนสร้างเอกสารด้วย initializeWorkflow() + addPreApprover() — **ห้าม INSERT ตรง**"),
            ("document_new_stores", "R/W", "ร้านเปิดใหม่และ % ชดเชย"),
            ("document_competitors", "R/W", "ร้านคู่แข่งในเอกสาร"),
            ("document_running_numbers", "R/W", "ตัวนับเลขเอกสารต่อปี ค.ศ. — ออกเลข YYYY/xxxxx แบบ atomic (INSERT … ON CONFLICT DO UPDATE … RETURNING)"),
            ("document_cost_details", "R/W", "ยอดชดเชยแยกรายเดือน/รายร้านเปิดใหม่ (cost_year/cost_month · cost_target · cost_amount · _n / _nc)"),
            ("document_external_factors", "R/W", "ปัจจัยภายนอกในเอกสาร"),
            ("compensation_documents unique guard", "R", "กัน duplicate ด้วย business key: impact_process_id หรือ source + impacted_store_code + impact_month + new_store_code + round_no"),
        ],
        "BE/LLDD-BE-API-Document-Detail-Aggregate": [
            ("compensation_documents", "R", "หัวเอกสาร สถานะ และ section ปัจจุบัน"),
            ("impacted_stores", "R", "ข้อมูลร้านถูกกระทบ"),
            ("document_new_stores", "R", "ร้านเปิดใหม่และ compensate_percent"),
            ("document_competitors", "R", "คู่แข่ง"),
            ("document_external_factors", "R", "ปัจจัยภายนอก"),
            ("document_attachments", "R", "metadata ไฟล์แนบ"),
            ("consideration_logs", "R", "timeline/history"),
        ],
        "BE/LLDD-BE-API-Document-Workflow-Actions": [
            ("workflow_transaction / workflow_history / workflow_approver (@srm/glb-workflow)", "R (เขียนผ่าน lib)", "eventWorkflow() เดิน state + บันทึก history"),
            ("compensation_documents", "W", "อัปเดต status/current_section/result"),
            ("consideration_logs", "W", "บันทึกผลพิจารณาและ comment"),
            ("workflow_transaction (@srm/glb-workflow)", "R (เขียนผ่าน lib)", "กัน action ซ้ำด้วย getTransaction/getPermissionEvents ก่อน eventWorkflow — ห้าม UPDATE ตรง"),
        ],
        "BE/LLDD-BE-API-Workflow-Instances": [
            ("fgi_impact_processes / fgi_impact_stores", "R/W", "อ่านข้อมูล impact และอัปเดต workflow_generation_status W/Y/N"),
            ("compensation_documents", "R/W", "create-if-missing จาก impact process และผูก docNo"),
            ("workflow_transaction (@srm/glb-workflow)", "W (โดย lib)", "initializeWorkflow() แทน K2 StartInstance — ห้าม INSERT ตรง"),
            ("workflow_approver (@srm/glb-workflow)", "W (ผ่าน lib)", "addPreApprover() ปักผู้รับงาน state 06 — **ห้าม INSERT ตรง**"),
            ("workflow_status / workflow_state (@srm/glb-workflow · sps_store)", "R", "lookup statusCode/status และ state แรก — ตาราง document_statuses/workflow_sections ของ SBPGI ถูกตัดแล้ว"),
            ("interface_transactions", "W", "บันทึกผลเรียกจาก Job 8b · ตาราง job_run_histories ถูกตัด 2026-08-06 — ผลการรันไปที่ application log"),
        ],
        "BE/LLDD-BE-API-Attachment-Sales-Timeline": [
            ("document_attachments", "R/W", "metadata ไฟล์แนบและ section ที่แนบ"),
            ("compensation_documents", "R", "ตรวจเอกสารและ impact_process_id"),
            ("fgi_impact_sales_summaries", "R", "หัวข้อมูลยอดขาย growth_rate_diff/total_working_days"),
            ("sales_transactions", "R", "ยอดขายรายวัน 4 windows"),
            ("consideration_logs", "R", "timeline/history"),
        ],
        "BE/LLDD-BE-API-Lookup": [
            ("impacted_stores (SBPGI) / store · mas_store · sevenshop (SBP เดิม)", "R", "store picker — SBPGI ไม่มีตาราง stores ของตัวเอง"),
            ("workflow_status / workflow_state (@srm/glb-workflow · sps_store)", "R", "lookup สถานะ verbatim และ 5 ขั้น 06/08/01/02/03 — ไม่สร้างตารางของ SBPGI เอง"),
            ("business_user (SBP เดิม)", "R", "popup ค้นหาพนักงาน — SBPGI ไม่มีตาราง employees"),
            ("auth-backend groups / menus / permissions (ระบบเดิม)", "R", "RBAC/menu matrix — จัดการที่หน้า /setting/manage-user-rights เดิม · SBPGI อ่านผ่าน header x-user-permissions เท่านั้น ไม่มีตารางของตัวเอง"),
            ("email_template (SBP เดิม)", "R", "template — SBPGI อ่านผ่าน lib เท่านั้น ไม่แก้ของระบบเดิม"),
            ("email_sent (SBP เดิม)", "W (โดย email-lib)", "log การส่ง — SBPGI เรียก sendEmail() ด้วยเลข template จาก workflow_route.email_id แล้ว lib เขียนแถวให้เอง (DP-5 · 2026-08-14) · ⚠️ คอลัมน์ผู้ส่งคือ send_by ไม่ใช่ sent_by"),
        ],
        "BE/LLDD-BE-API-Report-and-Master-Data": [
            ("compensation_documents", "R", "แหล่งข้อมูลรายงานและ filter status/year"),
            ("compensation_histories", "R", "ยอดเงินชดเชยและงวด statement"),
            ("consideration_logs", "R", "ผลพิจารณาล่าสุด APPROVE/REJECT"),
            ("auth-backend group + scope (business_user_group) / prepared approver ของ @srm/glb-workflow", "R", "ผู้ปฏิบัติงาน — ตาราง operator_assignments ถูกตัด 2026-08-05"),
            ("external_factors", "R/W", "master ปัจจัยภายนอก"),
            ("competitors", "R/W", "master แบรนด์คู่แข่ง 11 รายการ (code 01-11 · name_th · name_en · remark) — feed dropdown ร้านคู่แข่งของหน้าเอกสาร"),
            ("document_competitors", "R", "ตรวจว่าแบรนด์ถูกอ้างในเอกสารก่อนลบ (409)"),
            ("mas_param (SBP)", "R", "ค่ากำหนดกลางของระบบ SBP เดิม — **อ่านอย่างเดียว** (หน้า Global Config ของ SBPGI ถูกลบ 2026-08-06 · ระบบเดิมเป็นผู้แก้)"),
        ],
        "BE/LLDD-BE-Job-Batch-Email-SRM": [
            ("(backend config: config file/env)", "R", "enabled, cron, params ของ batch — ตาราง job_configs ถูกตัด 2026-08-06 ไม่มีหน้าจอควบคุม"),
            ("(application log แบบ structured)", "W", "ประวัติการรันและสถานะล่าสุด — ตาราง job_run_histories ถูกตัด 2026-08-06"),
            ("interface_transactions", "R/W", "tracking file/API interface และ ACK"),
            ("email_template (SBP)", "R", "subject_format/body_format ของระบบ SBP เดิม — อ่านอย่างเดียว"),
            ("email_sent (SBP)", "W (โดย email-lib)", "log การส่งของ batch — lib เขียนให้เอง"),
        ],
    }
    for topic in base:
        # FE-Testing-Delivery ใช้โครงเอกสารของตัวเอง (Test Suite Matrix / Release Gate)
        # ไม่ผ่าน renderer กลาง จึงไม่มีการวาด flow PNG — อย่าตั้ง flow_diagram ให้ชี้ไฟล์ที่ไม่มีจริง
        if (not topic.flow_diagram and not is_batch_monitor_doc(topic.file)
                and topic.file != "FE/LLDD-FE-Testing-Delivery"):
            topic.flow_diagram = f"LLDD/assets/flows/{sanitize_filename(topic.file)}.png"
        if not topic.db_tables and topic.file in db_map:
            topic.db_tables = db_map[topic.file]
    base.extend(be_job_topics())
    for topic in base:
        if topic.file in HIGH_LEVEL_ESTIMATES:
            topic.hours = HIGH_LEVEL_ESTIMATES[topic.file]
        # 2026-08-11: ถอด buffer ออก — ชั่วโมงที่แสดงคือค่าประเมินตรง ๆ ไม่บวกส่วนเผื่อ
        topic.base_hours = topic.hours
        topic.buffer = 0.0
        topic.days = round(topic.hours / HOURS_PER_DAY, 1)
    return base


MAIN_INDEX_ORDER: dict[str, int] = {
        "FE/LLDD-FE-Integration-Contracts": 5,
        "FE/LLDD-FE-Foundation": 10,
        "FE/LLDD-FE-Document-Lists": 30,
        "FE/LLDD-FE-Create-Document": 40,
        "FE/LLDD-FE-Document-Detail": 50,
        "FE/LLDD-FE-Document-Detail-Role-06-SBP-DSA": 51,
        "FE/LLDD-FE-Document-Detail-Role-08-SBP-DSA-Officer": 52,
        "FE/LLDD-FE-Document-Detail-Role-01-Business-Promotion": 53,
        "FE/LLDD-FE-Document-Detail-Role-02-GM-Business-Promotion": 54,
        "FE/LLDD-FE-Document-Detail-Role-03-AVP-SBP": 55,
        "FE/LLDD-FE-Report": 60,
        "FE/LLDD-FE-Master-Data": 70,
        "FE/LLDD-FE-Testing-Delivery": 90,
        "BE/LLDD-BE-Database-Structure": 100,
        "BE/LLDD-BE-Data-Migration-Cutover": 101,
        "BE/LLDD-BE-Integration-SBP-Platform": 102,
        # 2026-08-10: ย้ายมาจากลำดับ 156 (เดิมอยู่ท้ายกลุ่ม API) — เอกสารนี้นิยาม
        # state/status/route/part ของ @srm/glb-workflow ที่ BE-API-Document-Workflow-Actions
        # และ BE-API-Workflow-Instances ต้องใช้ ตารางเดิมจึงให้ผู้บริโภคเริ่มก่อนผู้นิยาม
        # (Workflow-Actions 21–27/08 · Workflow-Instances 24–27/08 vs Definition 27–31/08)
        # ลำดับนี้ทำให้เป็น blocker สัปดาห์แรกพร้อมอีก 3 ฉบับ ตรงกับ objective ของเอกสารเอง
        "BE/LLDD-BE-Workflow-Engine-Definition": 103,
        "BE/LLDD-BE-API-Common-Contracts": 105,
        "BE/LLDD-BE-API-Document-List-Search": 120,
        "BE/LLDD-BE-API-Document-Create-Update": 130,
        "BE/LLDD-BE-API-Document-Detail-Aggregate": 140,
        "BE/LLDD-BE-API-Document-Workflow-Actions": 150,
        "BE/LLDD-BE-API-Workflow-Instances": 155,
        "BE/LLDD-BE-API-Attachment-Sales-Timeline": 160,
        "BE/LLDD-BE-API-Lookup": 170,
        "BE/LLDD-BE-API-Report-and-Master-Data": 180,
        "BE/LLDD-BE-Job-Batch-Email-SRM": 190,
}


def main_index_ordered(all_topics: list[Topic]) -> list[Topic]:
    """ลำดับที่ใช้คำนวณตารางเวลา — ต้องเป็นลำดับเดียวกันทุกที่ ไม่งั้นวันที่จะเพี้ยน"""
    return sorted(all_topics, key=lambda t: MAIN_INDEX_ORDER.get(t.file, 999))


def main_doc_blocks(all_topics: list[Topic]) -> list[dict[str, Any]]:
    ordered = main_index_ordered(all_topics)
    counted_topics = [t for t in ordered if not is_document_detail_role_doc(t.file)]
    be_jobs = [t for t in counted_topics if "/Jobs/" in t.file]
    high_level = [t for t in counted_topics if "/Jobs/" not in t.file]
    role_docs = [t for t in ordered if is_document_detail_role_doc(t.file)]
    fe = [t for t in high_level if t.track == "FE"]
    be = [t for t in high_level if t.track == "BE"]
    schedule = build_topic_schedule(counted_topics)   # ตรวจว่า dependency ไม่ขัดกัน
    steps = dependency_steps(counted_topics)
    rows = [
        [
            t.track,
            t.title.replace("LLDD ", ""),
            (f"{total_hours(t)}" if not unit_test_hours(t)
             else f"**{total_hours(t)}** (impl {t.hours} + test {unit_test_hours(t)})"),
            f"{steps[t.file]}",
            t.owner,
            Path(t.file).name,
        ]
        for t in high_level
    ]
    owner_stats: dict[str, dict[str, Any]] = {}
    for topic in counted_topics:
        key = topic.owner
        owner_stats.setdefault(key, {"hours": 0, "base": 0, "tracks": set(), "topics": []})
        owner_stats[key]["hours"] += topic.hours
        owner_stats[key]["test"] = owner_stats[key].get("test", 0) + unit_test_hours(topic)
        owner_stats[key]["base"] += topic.base_hours
        owner_stats[key]["tracks"].add(topic.track)
        owner_stats[key]["topics"].append(topic.title.replace("LLDD FE - ", "").replace("LLDD BE - ", ""))
    owner_order = [
        FE_OWNER_KITTISAK,
        FE_OWNER,
        BE_OWNER_BUTSABA,
        BE_OWNER,
        BE_OWNER_PEERAKORN,
        BANK_BE_OWNER,
    ]
    continuity = {
        FE_OWNER_KITTISAK: "FE หน้าจอเอกสาร (สายลึกที่สุดของ FE): Document Detail/Action (+ role pack 5 ฉบับ) -> Master Data -> Create Document",
        FE_OWNER: "FE ที่ต่อกับระบบเดิม: Integration Contracts (auth/session/permission จาก BFF) -> Foundation (sidebar/header/menu gating ของ portal เดิม) -> Document Lists -> Report -> Testing/Delivery",
        BE_OWNER_BUTSABA: "BE เอกสาร/สัญญากลางของ SBPGI เอง: Common Contracts -> List/Search -> Create/Update -> Detail Aggregate -> Job 8 -> Job 8b (ตัวเรียก initializeWorkflow ของ flow แรก)",
        BE_OWNER: "BE ที่ต่อกับระบบเดิม + **เรียกใช้ engine**: Integration with SBP Platform -> Workflow Instances (initializeWorkflow) -> Workflow Actions (eventWorkflow = trigger event) -> Lookup -> Job 4, 6",
        BE_OWNER_PEERAKORN: "BE support/interface (ย้ายจากสาย FE 2026-08-07): Batch/Email -> Attachment/Sales/Timeline -> Report and Master Data -> Job 5, 7, 9, 10",
        BANK_BE_OWNER: "Migration DB + นิยาม workflow (มติ 2026-08-25): Database Structure -> Data Migration/Cutover (ORA FCS_FRN ฝั่ง Java + MSSQL CPA_FRN_FGI ฝั่ง K2) -> Workflow Engine Definition (สร้างข้อมูลใน DB ว่ามีกี่ step แต่ละ step ทำอะไร) -> Job 2, 3 (นำเข้า ALLMAP · สายข้อมูลเดียวกัน) · **ไม่รวม initializeWorkflow และ trigger event** ซึ่งเป็นของ BE คนอื่น",
    }
    owner_rows = []
    for key in owner_order:
        if key not in owner_stats:
            continue
        hours = owner_stats[key]["hours"]
        tests = owner_stats[key].get("test", 0)
        role = "FE & BE" if len(owner_stats[key]["tracks"]) > 1 else next(iter(owner_stats[key]["tracks"]))
        owner_rows.append([role, key, f"**{hours + tests}** (impl {hours} + test {tests})", continuity[key]])

    if set(owner_stats) != set(owner_order):
        raise ValueError("LLDD schedule must include all six developers")
    for owner, stats in owner_stats.items():
        work_weeks = (stats["hours"] + stats.get("test", 0)) / HOURS_PER_WEEK

    def summary_scope(topic: Topic, count: int) -> str:
        items = topic.scope[:count]
        return ", ".join(items)

    return [
        h(1, "1. Purpose"),
        p("เอกสารหลักนี้เป็น LLDD Index สำหรับ Phase #4 - 4.3 SBP Operating Management ประกันรายได้ โดยสรุปหัวข้อใหญ่ของงาน FE/BE เฉพาะระบบประกันรายได้ (SBP Mall) และเชื่อมไปยังเอกสาร LLDD รายละเอียดของแต่ละหัวข้อ"),
        h(1, "2. Scope"),
        bullets([
            "ครอบคลุมเฉพาะระบบประกันรายได้ (SBP Mall)",
            "งาน FE/BE ในเอกสารนี้นับเฉพาะหน้าจอ module SBP Mall และ API/Job/Service ที่รองรับระบบประกันรายได้เท่านั้น",
            "งานออกแบบ flow ระดับระบบและ schema ระดับองค์กรไม่ถูกนับซ้ำเป็นงานหน้าจอ FE",
            "รายละเอียดที่จำเป็นต่อการพัฒนา การตรวจรับ และการส่งมอบถูกรวมไว้ใน LLDD แต่ละฉบับ",
            "รูปหน้าจอในหัวข้อ FE ใช้อธิบายองค์ประกอบและพฤติกรรมที่ต้องพัฒนา",
            "ไม่รวมการพัฒนา Login/Auth ของ platform และกระบวนการภายนอกขอบเขต SBP Mall",
        ]),
        h(2, "2.1 Input / Progress / Output Contract"),
        table(["Stage", "Contract for implementation"], [
            ["Input", "Topic inventory, owner assignment, estimates, screenshots, API/job/database scope, and schedule assumptions for the SBP Mall income-guarantee work package."],
            ["Progress", "Use this index to sequence FE/BE work, confirm owner workload, locate detailed topic documents, and track dependency readiness before development starts."],
            ["Output", "A single implementation index with activity plan, owner workload, FE/BE summaries, job breakdown, dependencies, and deliverable checklist."],
        ]),
        h(1, "3. High Level Activity Plan"),
        table(["Track", "หัวข้อ", "ชั่วโมง (impl + unit test)", "ลำดับขั้น", "Owner", "เอกสารรายละเอียด"], rows),
        h(1, "4. Workload Balance and Continuity"),
        p("แผนนี้รวม owner ตามบุคคล (ปรับ 2026-08-07): ทีม 6 คนเหลือ FE 2 คนและ BE 4 คน โดย Peerakorn ย้ายจากสาย FE ไปสาย BE · Aphiwit เป็นเจ้าของ Database Structure + Data Migration/Cutover และ Job 2, 3, 4, 6, 8 · Peerakorn รับ Job 5, 7, 9, 10 · Tunyatorn รับ Job 8b เพราะเป็น job เดียวที่เรียก workflow engine และถือ Workflow Engine Definition อยู่แล้ว ชั่วโมงคิดที่ 5 วันต่อสัปดาห์และ 6 ชั่วโมงต่อวัน (30 ชั่วโมงต่อสัปดาห์) · ตัวเลขในตารางเป็นค่าประเมินตรง ๆ **ไม่มีส่วนเผื่อ (buffer)**"),
        table(["Role", "Owner", "ชั่วโมง (impl + unit test)", "Work Focus"], owner_rows),
        h(1, "5. FE Summary"),
        table(["FE Topic", "ชั่วโมง", "ลำดับขั้น", "Deliverable"], [[t.title.replace("LLDD FE - ", ""), t.hours, steps[t.file], summary_scope(t, 3)] for t in fe]),
        h(1, "6. Document Detail Role Pack"),
        p("เอกสารลูก 5 ฉบับนี้เป็นรายละเอียดแยกตาม role สำหรับอ่านประกอบ LLDD-FE-Document-Detail ไม่ถูกนับซ้ำใน activity plan/hour รวม"),
        table(["Role document", "Parent", "Hour allocation"], [[Path(t.file).name, "LLDD-FE-Document-Detail", "included in parent hours"] for t in role_docs]),
        h(1, "7. BE Summary"),
        table(["BE Topic", "ชั่วโมง", "ลำดับขั้น", "Deliverable"], [[t.title.replace("LLDD BE - ", ""), t.hours, steps[t.file], ", ".join(t.scope[:4])] for t in be]),
        h(1, "8. BE Batch Job Breakdown"),
        table(
            ["Job", "ชั่วโมง", "ลำดับขั้น", "Owner", "เอกสารรายละเอียด"],
            [
                [
                    t.title.replace("LLDD BE - ", ""),
                    t.hours,
                    steps[t.file],
                    t.owner,
                    Path(t.file).name,
                ]
                for t in be_jobs
            ],
        ),
        h(1, "9. Dependency"),
        table(["Dependency", "Owner", "ใช้โดย"], [
            ["Common API/FE contracts", "BE/FE", "LLDD-BE-API-Common-Contracts และ LLDD-FE-Integration-Contracts เป็นสัญญากลางของทุกหน้า FE และทุก service BE"],
            ["API contract", "BE/FE", "ทุกหน้า FE และทุก service BE"],
            ["Master Data contract", "FE/BE", "LLDD-FE-Master-Data ใช้ LLDD-BE-API-Report-and-Master-Data สำหรับปัจจัยภายนอกและรายชื่อคู่แข่ง (ไม่มี Operator/Menu Permission/System Config/Audit แล้ว — ใช้ระบบ SBP เดิม)"],
            ["Blocker ลำดับขั้นที่ 1", "BE", "LLDD-BE-Integration-SBP-Platform, LLDD-BE-Workflow-Engine-Definition, LLDD-BE-Database-Structure และสัญญากลางของ LLDD-API ต้องปิดก่อน เพราะเอกสาร BE ทุกฉบับอ้างอิง 4 ชิ้นนี้"],
            ["Auth/JWT platform และ menu service", "Platform/SSO/IAM", "FE Foundation เรียก /auth/profile + /users/current + /menus + /groups/current-user/permissions ของระบบเดิม; SBPGI รับตัวตนจาก BFF ผ่าน header x-api-key/x-user-id/x-user-group-id/x-user-permissions"],
            ["Mock/fixture data", "BE", "FE development และ SIT"],
            ["Screenshots/prototype", "FE", "UI implementation"],
            ["Business rules", "BA/BE", "validation/action/report"],
        ]),
        h(1, "10. Deliverable Checklist"),
        bullets(["Main LLDD Index", "Common contract LLDD สำหรับ API/FE integration", "LLDD-FE-Master-Data สำหรับปัจจัยภายนอกและรายชื่อคู่แข่ง", "Detailed FE LLDD per SBP Mall page group", "Detailed BE LLDD per SBP Mall API group and Jobs 2-10 + 8b", "Database Structure, Data Migration/Cutover, Integration with SBP Platform และ Workflow Engine Definition (เพิ่ม 2026-08-07)", "Screenshots embedded only for SBP Mall implementation pages", "Implementation flow diagrams embedded as reference, not Flow page deliverables"]),
    ]


def api_endpoint_groups() -> list[list[Any]]:
    return [
        # 29 เส้น · 6 กลุ่ม (ตรงกับ api.md และ plan-api.html) — Auth ถูกตัดทั้งกลุ่ม 2026-08-05 (ใช้ระบบ SBP เดิม)
        ["งาน & เอกสารประกันรายได้", "11", "GET /sbpgi/document/tasks, GET/POST/PUT /sbpgi/document*, POST /sbpgi/document/{docNo}/actions, attachments, sales, timeline", "core document workflow API"],
        ["Lookup / Reference", "2", "GET /sbpgi/lookup/document-statuses, /sbpgi/lookup/workflow-sections", "read-only reference ที่ไม่มีหน้าจอดูแล (ร้าน/ภาค/ประเภทสาขา ใช้ของระบบ SBP เดิม)"],
        ["Master Data", "8", "factors CRUD, competitors CRUD", "master ที่มีหน้าจอดูแลของตัวเอง (ไม่มี audit · ยกเลิกระบบ audit ของ master 2026-08-07)"],
        ["รายงาน", "2", "GET /sbpgi/report/status-summary, /export", "accounting search/export Excel (14 columns, SDD slide 60)"],
        ["Workflow ภายใน", "3", "POST /sbpgi/workflow/instances, GET /sbpgi/workflow/instances/{id}, /sbpgi/workflow/summary", "internal workflow engine for Job 8b"],
        ["Interface Tracking", "3", "GET /sbpgi/interface/tracking, GET /sbpgi/interface/pending-ack, POST /sbpgi/interface/sta/ack", "file tracking และ ACK (ตัด GET /dashboard/summary ออก 2026-08-06 · ตัด POST /integrations/srm/income-guarantee 2026-08-07)"],
    ]


def api_doc_text(text: Any) -> str:
    value = str(text if text is not None else "")
    replacements = {
        "workflow.md": "ตารางเส้นทาง workflow",
        "api.md": "API contract",
        "database.md": "Database contract",
        "plan-api.html": "API catalog",
        "plan-database.html": "Database catalog",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def plan_api_groups() -> list[dict[str, Any]]:
    return read_js_array_from_html("plan-api.html", "GROUPS")


def plan_api_sql_by_path() -> dict[str, str]:
    return read_js_object_from_html("plan-api.html", "SQL_BY_PATH")


def endpoint_method_path(endpoint: dict[str, Any]) -> str:
    return f"{endpoint.get('m', '')} {endpoint.get('p', '')}"


def is_batch_api_group(group: dict[str, Any]) -> bool:
    return api_doc_text(group.get("name", "")) == "Batch Job Admin"


def api_endpoint_detail_blocks(groups: list[dict[str, Any]], sql_by_path: dict[str, str]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = [h(1, "6. Detailed Endpoint Specification")]
    endpoint_no = 0
    for group_idx, group in enumerate(groups, start=1):
        endpoints = group.get("eps", [])
        batch_reference_only = is_batch_api_group(group)
        blocks.extend([
            h(2, f"6.{group_idx} {api_doc_text(group.get('name', ''))}"),
            table(
                ["Endpoint", "Method", "Path", "Summary"],
                [[idx, api_doc_text(ep.get("m", "")), api_doc_text(ep.get("p", "")), api_doc_text(ep.get("sum", ""))] for idx, ep in enumerate(endpoints, start=1)],
            ),
        ])
        if batch_reference_only:
            blocks.append(p("Batch Job Admin เป็น endpoint reference สำหรับ FE Batch Monitor เฉพาะ 2 tab: แบบฟอร์มพารามิเตอร์ และประวัติการรัน เท่านั้น; ไม่ออกแบบ flowchart การทำงาน, step-by-step batch flow หรือ Database ที่ใช้ใน LLDD API ฉบับรวม"))
        for endpoint_idx, endpoint in enumerate(endpoints, start=1):
            endpoint_no += 1
            key = endpoint_method_path(endpoint)
            blocks.extend([
                h(3, f"6.{group_idx}.{endpoint_idx} {api_doc_text(key)}"),
                p(api_doc_text(endpoint.get("sum", ""))),
                table(
                    ["Item", "Detail"],
                    [
                        ["Global No.", endpoint_no],
                        ["Method", api_doc_text(endpoint.get("m", ""))],
                        ["Path", api_doc_text(endpoint.get("p", ""))],
                        ["Group", api_doc_text(group.get("name", ""))],
                        ["Access / Role", api_doc_text(endpoint.get("roles", ""))],
                        ["Requirement Tag", api_doc_text(endpoint.get("refT", ""))],
                    ],
                ),
            ])
            if batch_reference_only:
                blocks.append(p("Scope note: รายละเอียด flow และ database ของ batch job ให้ดูเอกสาร BE/Runbook/Database reference แยก ไม่ใช่ tab หรือ deliverable ที่ต้องทำใน FE Batch Monitor"))
            else:
                blocks.extend([
                table(
                    ["Step", "Flow"],
                    [[idx, api_doc_text(step)] for idx, step in enumerate(endpoint.get("flow", []), start=1)],
                ),
                table(
                    ["DB Object", "R/W", "Usage"],
                    [[api_doc_text(row[0]), api_doc_text(row[1]), api_doc_text(row[2])] for row in endpoint.get("db", [])],
                ),
                ])
            blocks.extend([
                payload("Request / Query / Header", api_doc_text(endpoint.get("req", "(ไม่มี body)"))),
                payload("Response", api_doc_text(endpoint.get("res", ""))),
            ])
            if endpoint.get("resNote"):
                blocks.append(p(api_doc_text(endpoint["resNote"])))
            blocks.extend([
                table(
                    ["Error / Condition"],
                    [[api_doc_text(err)] for err in endpoint.get("err", [])],
                ),
            ])
            sql = sql_by_path.get(key)
            if sql and not batch_reference_only:
                blocks.extend([
                    p("SQL Reference"),
                    code(api_doc_text(sql), "sql"),
                ])
    return blocks


def lldd_api_blocks(all_topics: list[Topic]) -> list[dict[str, Any]]:
    be_api_topics = [t for t in all_topics if t.track == "BE" and t.file.startswith("BE/LLDD-BE-API")]
    groups = plan_api_groups()
    endpoint_total = sum(len(g.get("eps", [])) for g in groups)
    sql_by_path = plan_api_sql_by_path()
    return [
        h(1, "1. Purpose"),
        p("เอกสารนี้เป็น LLDD API ระดับรวมของระบบ SBPGI/SBP Mall ใช้เป็น master reference สำหรับ REST contract, auth, error, endpoint catalog, implementation pattern และ test scope ของ BE API LLDD รายกลุ่ม"),
        h(1, "2. Scope"),
        table(
            ["Item", "Detail"],
            [
                ["API base", "/api/v1"],
                ["Endpoint count", f"{endpoint_total} endpoints, {len(groups)} groups"],
                ["Detailed implementation docs", ", ".join(Path(t.file).name for t in be_api_topics)],
                ["Out of scope", "Login/Auth implementation ของ platform, SAP/SR process ภายนอก, abnormal-stores endpoints ที่ยัง comment รอตัดสินใจ"],
            ],
        ),
        h(2, "2.1 Input / Progress / Output Contract"),
        table(["Stage", "Contract for implementation"], [
            ["Input", "Endpoint catalog, auth mode, role/access rules, request/response payloads, error conditions, and SQL references from the API plan data."],
            ["Progress", "For each endpoint, apply middleware, bind DTO, validate, authorize, execute service transaction, map response, and pass errors through the centralized handler."],
            ["Output", "Normalized REST contract for implementation and testing: method/path, payload, response, errors, DB usage, and checklist coverage."],
        ]),
        h(1, "3. API Design Principles"),
        table(
            ["Rule", "Required behavior", "Developer note"],
            [
                ["Transport", "JSON UTF-8 ทุก endpoint; multipart เฉพาะ attachment upload", "FE shared API client เป็นจุดเดียวที่ตั้ง base URL/header"],
                ["Auth", "User endpoint ใช้ Bearer JWT; internal workflow/interface ใช้ service token/API key", "BE middleware ต้องแยก user token กับ service token ชัดเจน"],
                ["Status convention", "API ส่ง `statusCode`; FE resolve label จาก `/sbpgi/lookup/document-statuses`", "ห้ามส่ง label ไทยแทน code ใน field ที่กำหนดเป็น canonical code"],
                ["Role namespace", "`roleCode` = RBAC role, `sectionCode` = workflow section, `roleProfileCode` = P-06/P-08/P-01/P-02/P-03", "ป้องกันการชนความหมายของเลข 01/02/03/06/08"],
                ["Pagination", "GET list ใช้ `page,size` และคืน `{page,size,total,items}`", "size max 100 ตาม common contract"],
                ["Errors", "คืน `{code,message}`; message ภาษาไทยตาม SRS ถ้ามี", "FE แสดง message ตรง ๆ ไม่ paraphrase"],
                ["Mutation audit", "workflow action ลง consideration_logs เท่านั้น (ยกเลิกระบบ audit ของ master 2026-08-07 · jobs เขียน application log)", "mutation ที่ต้องมี reason ต้อง validate ก่อนเริ่ม transaction"],
            ],
        ),
        h(1, "4. Endpoint Catalog"),
        table(
            ["Group", "Count", "Endpoint pattern", "Implementation focus"],
            [[api_doc_text(g.get("name", "")), len(g.get("eps", [])), ", ".join(api_doc_text(e.get("p", "")) for e in g.get("eps", [])[:4]) + (" ..." if len(g.get("eps", [])) > 4 else ""), api_doc_text(g.get("refT", ""))] for g in groups],
        ),
        h(1, "5. Request Lifecycle"),
        table(
            ["Step", "API behavior", "Failure handling"],
            [
                ["1. Middleware", "ตรวจ correlationId/requestId, auth token, content type, payload size", "401/413/415 ก่อนเข้า service"],
                ["2. Controller", "รวม params/query/body เป็น DTO และเรียก service", "controller ไม่ใส่ business rule"],
                ["3. Validation", "required/format/enum/date/page/size/docNo/storeCode", "400/422 พร้อม code/message จาก catalog"],
                ["4. Authorization", "ตรวจ menu/RBAC/document participant/current task owner/service token", "403 หรือ 409 เมื่อ task เปลี่ยนแล้ว"],
                ["5. Transaction", "mutation เปิด transaction ใน service; read ใช้ read-only query", "rollback เมื่อ persist หรือ audit fail"],
                ["6. Mapper", "map domain object เป็น DTO ตาม API contract", "ไม่ expose objectKey/secret/internal raw row"],
                ["7. Response", "คืน JSON หรือ binary stream สำหรับ download", "error ผ่าน centralized error handler"],
            ],
        ),
        *api_endpoint_detail_blocks(groups, sql_by_path),
        h(1, "7. API Test Checklist"),
        table(
            ["Test group", "Required cases"],
            [
                ["Common contract", "401, 403, 404, 409, 422, pagination envelope, error `{code,message}`"],
                ["Document workflow", "create duplicate, submit no result, invalid result for role profile, current task conflict, threshold ≥ 100,000 -> AVP route (SDD GI)"],
                ["Attachment", "file >5MB, unsupported type, AV blocked, download not owner, download clean file"],
                ["Report", "year required, result required, CSV export with same filter as preview"],
                ["Job admin", "manual run when disabled, manual run while RUNNING, editable params only, run histories"],
                ["Security", "service token only endpoints, no objectKey/secret leak, audit reason required for mutations"],
            ],
        ),
        h(1, "8. Related LLDD"),
        table(["Document", "Use"], [[Path(t.file).name, t.objective] for t in be_api_topics]),
    ]


def database_table_catalog() -> list[list[Any]]:
    """20 target tables: 34 -> 24 (ตัด 10 ตารางที่ระบบ SBP เดิมมีอยู่แล้ว 2026-08-06)
    -> 22 (ตัด job_configs/job_run_histories พร้อม 2 tab ควบคุมของหน้า Batch Job)
    -> 21 (ตัด audit_logs 2026-08-07) -> 20 (ตัด decisions ไป common_code · มติ DP-9 2026-08-10).
    แถวที่ zone = "REF" เป็น schema reference สำหรับ dev เท่านั้น ไม่นับใน 20 ตาราง.
    """
    return [
        ["A", "fgi_impact_stores", "id", "impact_process_id, impacted_store_code", "impact pair; sales request and allocation data"],
        ["A", "fgi_impact_processes", "id", "impacted_store_code", "impact process hub and canonical workflow_generation_status"],
        ["A", "fgi_impact_sales_summaries", "id", "impact_process_id", "sales summary/growth rate"],
        ["A", "sales_transactions", "id", "sales_summary_id", "daily sales 4 windows x 15 days"],
        ["A", "fgi_impact_competitors", "id", "impact_process_id", "ALLMAP competitors"],
        ["A", "fcs_qssi_score", "id", "store_id + category + month + year", "QSSI scores — ⚠️ REUSE ตารางเดิมของ sps_store (เอกพจน์ · 23,958,780 แถว · มี import pipeline POST /performance/import-qssi ใช้งานอยู่) ห้ามสร้างใหม่ และห้ามใช้ชื่อพหูพจน์ fcs_qssi_scores"],
        ["A", "interface_transactions", "id", "impact_process_id/sales_summary_id/doc_no", "interface tracking replacement"],
        ["B", "compensation_documents", "doc_no", "impact_process_id, status_code, current_section_code", "document header/core"],
        ["B", "document_new_stores", "id", "doc_no, new_store_code", "new stores, compensate percent and amount"],
        ["B", "document_competitors", "id", "doc_no, competitor_code", "document competitors"],
        ["B", "document_external_factors", "id", "doc_no, factor_code", "document external factors"],
        ["B", "consideration_logs", "id", "doc_no", "approval/action history (decision code, result category, attachments)"],
        ["B", "document_attachments", "attach_id", "doc_no", "attachment metadata; file storage uses existing SBP S3 service"],
        ["B", "compensation_histories", "id", "store_code, ref_doc_no", "compensation history/accounting export"],
        ["B", "document_cost_details", "id", "doc_no, new_store_code", "monthly cost detail per new store (ImpactCostDetail)"],
        ["B", "document_running_numbers", "year", "-", "atomic YYYY/xxxxx running number"],
        ["C", "impacted_stores", "store_code", "store.store_id (SBP · varchar(10))", "SP impacted store subset"],
        
        ["C", "external_factors", "factor_code", "-", "external factor master"],
        ["C", "competitors", "competitor_code", "-", "competitor master"],
        ["-", "USE EXISTING SBP TABLES", "-", "-",
         "workflow engine 13 ตาราง ใน schema sps_store (workflow · workflow_version · workflow_state · workflow_status · workflow_event · workflow_route · "
         "workflow_group · workflow_group_map · workflow_transaction · workflow_history · workflow_approver · workflow_part · workflow_part_display) "
         "· store/mas_store/sevenshop · mas_zone · common_code · business_user · email_template + email_sent · mas_param · fcs_qssi_score "
         "— decided 2026-08-05/2026-08-06: do NOT recreate these in SBPGI"],
    ]


def database_ddl_sections() -> list[tuple[str, str]]:
    return [
        ("5.1 Zone C — Shared Master, RBAC, Config and Operations", """-- ❌ ไม่สร้างตาราง stores ใน SBPGI — ใช้ store / mas_store / sevenshop ของระบบ SBP เดิม (API: GET /store/search · /store/list · /store/detail)

CREATE TABLE impacted_stores (
    store_code VARCHAR(5) PRIMARY KEY,   -- ร้าน SP · master อยู่ที่ store/mas_store/sevenshop ของระบบเดิม
    dv_code VARCHAR(20), opt_dv_user_id VARCHAR(30), latitude NUMERIC(10,7), longitude NUMERIC(10,7),
    is_active BOOLEAN NOT NULL DEFAULT TRUE,
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
);

-- ❌ ไม่สร้างตาราง workflow_sections ใน SBPGI — ใช้ workflow_state / workflow_route ของ @srm/glb-workflow · วงเงินอนุมัติเก็บใน common_code (SBPGI_APPROVE_LIMIT)

-- ❌ ไม่สร้างตาราง document_statuses ใน SBPGI — ใช้ workflow_status ของ @srm/glb-workflow

-- ❌ ไม่สร้างตาราง roles ใน SBPGI — ใช้ auth-backend/ABS groups ของระบบ SBP เดิม (ตัดสินใจ 2026-08-05)

-- ❌ ไม่สร้างตาราง menus ใน SBPGI — ใช้ menus/permissions ของ auth-backend (ตัดสินใจ 2026-08-05)

-- ❌ ไม่สร้างตาราง menu_permissions ใน SBPGI — ใช้ permissions ต่อ URL ของ auth-backend (ตัดสินใจ 2026-08-05)

-- ❌ ไม่สร้างตาราง employees ใน SBPGI — ใช้ business_user / business_user_group ของระบบ SBP เดิม

ALTER TABLE impacted_stores
    -- opt_dv_user_id ไม่มี FK — ผู้ใช้อยู่ที่ business_user ของระบบ SBP เดิม (ตัด employees 2026-08-05)

-- ❌ ไม่สร้างตาราง operator_assignments ใน SBPGI — ใช้ group + scope ของ auth-backend + prepared approvers ของ @srm/glb-workflow (ตัดสินใจ 2026-08-05)

-- ❌ ไม่สร้างตาราง decisions ใน SBPGI — มติ DP-9 (2026-08-10) ย้ายไป common_code ของระบบ SBP เดิม
--    code_type='SBPGI_DECISION' · code_value=decision_code · code_name=decision_name
--    code_mapping=flow_name · other_value=result_name · remark=result_category+engine_event (จำกัด 50 ตัวอักษร)
--    FE อ่านผ่าน GET /common/common-code?codeType=SBPGI_DECISION (ตัดเส้น GET /decisions ออกแล้ว)
-- ⚠️ common_code ไม่มี PK และไม่มี unique constraint → กันรหัสซ้ำที่ระดับแอป
--    และต้องลงทะเบียน code_type ที่ common_code_type ก่อนใช้งาน

CREATE TABLE external_factors (
    factor_code VARCHAR(30) PRIMARY KEY,
    factor_name VARCHAR(200) NOT NULL, factor_remark VARCHAR(500),
    is_active BOOLEAN NOT NULL DEFAULT TRUE,
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
);

-- master แบรนด์คู่แข่ง 11 รายการ (รหัส 01-11) — ระบบเดิมเก็บชื่อทั้งไทยและอังกฤษ
-- คนละระดับกับ document_competitors ที่เก็บ "รายสาขา" พร้อมรหัสจาก ALLMAP (เช่น 4832, TD58_08)
CREATE TABLE competitors (
    competitor_code VARCHAR(30) PRIMARY KEY,
    name_th VARCHAR(200) NOT NULL,
    name_en VARCHAR(200) NOT NULL,
    remark VARCHAR(500),                 -- คอลัมน์ "รายละเอียดเพิ่มเติม" ของหน้า k2-competitors.html
    is_active BOOLEAN NOT NULL DEFAULT TRUE,
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
);

-- ❌ ไม่สร้างตาราง email_templates ใน SBPGI (ตัดสินใจ 2026-08-06) — ใช้ตาราง email_template ของระบบ SBP เดิม (email_template_id · subject_format · body_format) + email_sent

-- ❌ ไม่สร้างตาราง status_email_rules (ปิด DP-5 · แก้มติ 2026-08-14) — workflow ให้เลข template ผ่าน workflow_route.email_id แล้ว SBPGI เรียก email-lib ส่งเอง
--    อีเมลของ batch job (EM-07 error · EM-08 watchdog) ไม่ใช่ workflow event → ส่งผ่าน @gosoft-sbp/email-lib ของระบบเดิม
--    ผู้รับของ batch job อยู่ใน backend config (config file/env) ไม่ใช่ตารางของ SBPGI


-- ❌ ไม่สร้างตาราง user_accounts ใน SBPGI — ใช้ AWS Cognito + auth-backend — SBPGI รับตัวตนจาก header ของ BFF (ตัดสินใจ 2026-08-05)

-- ❌ ไม่สร้างตาราง system_configs ใน SBPGI (ตัดสินใจ 2026-08-06) — ใช้ตาราง mas_param ของระบบ SBP เดิม (param_name · param_value · ref_name · description · is_config · active_flag)"""),
        ("5.2 Zone A — Impact Pipeline, Sales and Interface", """CREATE TABLE fgi_impact_processes (
    id BIGSERIAL PRIMARY KEY,
    impacted_store_code VARCHAR(5) NOT NULL REFERENCES impacted_stores(store_code),
    impact_month CHAR(7) NOT NULL,   -- 'YYYY-MM' (ค.ศ.)
    impact_year INTEGER NOT NULL,    -- แตกจาก impact_month เพื่อ filter รายปีโดยไม่ต้อง substring
    process_status VARCHAR(30) NOT NULL, action_status VARCHAR(30),
    last_compensation_amount NUMERIC(14,2),
    workflow_generation_status CHAR(1) NOT NULL DEFAULT 'W' CHECK (workflow_generation_status IN ('W','Y','N')),
    -- ⬇ รับเข้าโครงตามมติ 2026-08-21 (gap F8) — ขนจาก ORA FGI_IMPACT_STORE_ON_PROCESS
    --   ใช้ตัดสิน "ประเภทเคส" ที่จุดเข้า flow และ auto-assign เจ้าของงานคนเดิม
    last_compensate_seq INTEGER NOT NULL DEFAULT 1,        -- รอบชดเชย (ขึ้นใหม่เมื่อเปิดเรื่องใหม่)
    last_compensate_seq_no INTEGER NOT NULL DEFAULT 1,     -- ครั้งที่ในรอบ · > 1 = เคสต่อเนื่อง
    start_compensate_month CHAR(7), start_compensate_year INTEGER,   -- กรอบงวดที่ชดเชยได้ (เริ่ม)
    end_compensate_month CHAR(7),   end_compensate_year INTEGER,     -- กรอบงวดที่ชดเชยได้ (จบ)
    -- ORA FGI_IMPACT_STORE_ON_PROCESS.FLAG_ACTION — โดเมนจริง Y/W/N (active = IN ('Y','W'))
    -- Job 6 ปิดรอบด้วย Y->N และพัก/รอจ่ายด้วย Y->W · CHECK เดิมที่รับแค่ ('Y','N') จะทำ migration ล้มทันทีที่เจอแถว 'W'
    flag_action CHAR(1) NOT NULL DEFAULT 'Y' CHECK (flag_action IN ('Y','W','N')),
    -- ช่องทางต้นทางของเคส (SDD GI สไลด์ 17 · 3 แหล่ง) — ORA FGI_IMPACT_STORE_ON_PROCESS.DATASOURCE
    --   ALM = ระบบดึงจาก ALLMAP (Job 2/3)   · STA = ระบบดึงจาก Franchise Statement (Job 5)   [ทั้งคู่มีในระบบเดิม]
    --   PRO = เชิงรุก  — OPT ประชุมพิจารณาแล้วเปิดเรื่อง (ต้นทางเอกสารอยู่ที่ All Memo)      [ใหม่ 2026-08-24]
    --   REA = เชิงรับ  — หน่วยงานอื่นแจ้งเข้ามาว่าร้านถูกกระทบ                                [ใหม่ 2026-08-24]
    -- ผลต่อ flow (SDD สไลด์ 47 · 49): ALM/STA = งานเข้ามาให้ จนท. SBP DSA เลือก · PRO/REA = เจ้าของงานต้องคีย์เอง
    -- ไม่ใส่ CHECK constraint — ระบบเดิมยังมีค่า HRS (HR feed) ปนอยู่ ถ้าบังคับโดเมนแคบจะ migrate ไม่ผ่าน
    datasource VARCHAR(5),
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_impact_process UNIQUE (impacted_store_code, impact_month)
);

-- รับเข้าโครงตามมติ 2026-08-21 (gap F1) — ขนจาก ORA FGI_IMPACT_STORE_COMPENSATE
-- ยอดชดเชย "รายงวด" ที่เกิดก่อนมีเอกสาร · จำเป็นเพื่อนับ "ยอด 0 ติดกันกี่เดือน" (กติกาเดือน 1-3 / เดือนที่ 4)
CREATE TABLE fgi_impact_compensations (
    id BIGSERIAL PRIMARY KEY,
    impact_process_id BIGINT NOT NULL REFERENCES fgi_impact_processes(id),
    impacted_store_code VARCHAR(5) NOT NULL REFERENCES impacted_stores(store_code),
    compensate_seq INTEGER NOT NULL,        -- รอบ (คู่กับ fgi_impact_processes.last_compensate_seq)
    compensate_seq_no INTEGER NOT NULL,     -- ครั้งที่ในรอบ
    compensate_month CHAR(7) NOT NULL,      -- งวดที่ชดเชย 'YYYY-MM' (ค.ศ.)
    compensate_year INTEGER NOT NULL,
    forecast_amount NUMERIC(14,2),          -- ระบบคำนวณ
    adjust_amount NUMERIC(14,2),            -- คนปรับ · ยอดที่ใช้จริง = COALESCE(adjust_amount, forecast_amount)
    compensate_status VARCHAR(5),
    compensate_comment VARCHAR(4000),
    stmt_month INTEGER, stmt_year INTEGER,  -- งวด statement
    approve_date DATE,
    created_by VARCHAR(100), created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    updated_by VARCHAR(100), updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_impact_compensation UNIQUE (impact_process_id, compensate_month)
);

CREATE TABLE fgi_impact_stores (
    id BIGSERIAL PRIMARY KEY,
    impact_process_id BIGINT NOT NULL REFERENCES fgi_impact_processes(id),
    impacted_store_code VARCHAR(5) NOT NULL REFERENCES impacted_stores(store_code),
    new_store_code VARCHAR(5) NOT NULL,   -- ร้านเปิดใหม่ · master ของระบบเดิม
    impact_month CHAR(7) NOT NULL, distance_km NUMERIC(8,3),
    sales_request_status CHAR(1) NOT NULL DEFAULT 'W' CHECK (sales_request_status IN ('W','P','Y','E')),
    forecast_compensate_percent NUMERIC(7,4), adjust_compensate_percent NUMERIC(7,4),
    forecast_compensation_amount NUMERIC(14,2), adjust_compensation_amount NUMERIC(14,2),
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_impact_store_pair UNIQUE (impacted_store_code, new_store_code, impact_month)
);

CREATE TABLE fgi_impact_sales_summaries (
    id BIGSERIAL PRIMARY KEY,
    impact_process_id BIGINT NOT NULL REFERENCES fgi_impact_processes(id),
    total_working_days INTEGER NOT NULL DEFAULT 0 CHECK (total_working_days >= 0),
    growth_rate_before NUMERIC(9,4), growth_rate_after NUMERIC(9,4), growth_rate_diff NUMERIC(9,4),
    sales_status CHAR(1) NOT NULL DEFAULT 'W' CHECK (sales_status IN ('W','Y','N','E')),
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_sales_summary_process UNIQUE (impact_process_id)
);

CREATE TABLE sales_transactions (
    id BIGSERIAL PRIMARY KEY,
    sales_summary_id BIGINT NOT NULL REFERENCES fgi_impact_sales_summaries(id) ON DELETE CASCADE,
    txn_date DATE NOT NULL, window_no SMALLINT NOT NULL CHECK (window_no BETWEEN 1 AND 4),
    sales_amount NUMERIC(14,2) NOT NULL, sales_diff NUMERIC(14,2),
    is_outlier BOOLEAN NOT NULL DEFAULT FALSE, source_checksum VARCHAR(64) NOT NULL,
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_sales_day_window UNIQUE (sales_summary_id, txn_date, window_no)
);

CREATE TABLE fgi_impact_competitors (
    id BIGSERIAL PRIMARY KEY,
    impact_process_id BIGINT NOT NULL REFERENCES fgi_impact_processes(id),
    competitor_code VARCHAR(30) NOT NULL REFERENCES competitors(competitor_code),
    name_th VARCHAR(200), branch_th VARCHAR(200), opened_date DATE, closed_date DATE,
    period_key CHAR(7) NOT NULL, updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_impact_competitor UNIQUE (impact_process_id, competitor_code, period_key)
);

-- ⚠️ fcs_qssi_score — ห้าม CREATE TABLE ใหม่ (ตรวจฐานจริง 2026-08-07)
--   ตารางนี้มีอยู่แล้วใน schema `sps_store` ชื่อ **เอกพจน์** `fcs_qssi_score`
--   มีข้อมูลจริง 23,958,780 แถว และมี import pipeline ทำงานอยู่
--   (`POST /performance/import-qssi` · staging `fcs_tmp_qssi_score` · `performance.service.ts`)
--   โครงคอลัมน์อ้างอิงด้านล่างเป็น target shape ที่ SBPGI ต้องการ — ต้องเทียบกับคอลัมน์จริงก่อน
--   ✅ DP-4 ปิดแล้ว 2026-08-24: อ่านอย่างเดียว ไม่แก้ constraint/index ของตารางเดิม
--   ห้ามใช้ชื่อพหูพจน์ `fcs_qssi_scores` ทุกกรณี
-- target shape (reference only — ห้ามรันเป็น DDL):
--   id BIGSERIAL PK · store_code VARCHAR(5) · category_code VARCHAR(30) · score_period CHAR(7)
--   · score_value NUMERIC(10,4) · source_file_name · source_checksum · updated_at
--   · UNIQUE (store_code, category_code, score_period)

CREATE TABLE interface_transactions (
    id BIGSERIAL PRIMARY KEY,
    -- run_id เป็น correlation id ของรอบรัน (มาจาก application log) — ไม่มี FK เพราะ job_run_histories ถูกตัด 2026-08-06
    run_id VARCHAR(50),
    -- direction: OUT = ส่งไฟล์ออกไประบบภายนอก (Job 4 → IAS · Job 6 → STA) · IN = รับไฟล์/ACK กลับ (Job 5 · callback ของ STA)
    --            INTERNAL = การส่งต่อ*ภายในระบบเดียวกัน* ที่มาแทนไฟล์ EAI เดิม (Jobs 7/8/9 เขียน DB ตรง — ไม่มี ACK ให้รอ จึงจบที่ status = COMPLETED)
    -- ชุดค่าปิด 9 ค่า เขียนโดย batch เท่านั้น (ไม่ใช่ input ของผู้ใช้) — ต้องล็อกเพราะ data_name เป็นส่วนหนึ่งของ
    -- UNIQUE ที่กันส่งซ้ำ และเป็นตัวกรองของ watchdog Job 10 · พิมพ์ผิดหนึ่งตัว = กันซ้ำไม่ทำงาน + watchdog เงียบ
    -- เพิ่ม interface ใหม่ = ALTER CONSTRAINT + อัปเดตตารางใน database.md พร้อมกัน
    data_name VARCHAR(80) NOT NULL CHECK (data_name IN (
        'IAS_SALES_REQUEST',                                    -- Job 4 -> IAS/MIS (OUT)
        'IMPACT_STORE_SALES',                                   -- Job 5 <- IAS/MIS (IN)
        'COMPENSATE_INIT_I','COMPENSATE_INIT_N',                -- Job 6 -> STA (OUT)
        'COMPENSATE_APPROVE_I','COMPENSATE_APPROVE_N',          -- Job 6 -> STA (OUT)
        'IMPACT_COMPETITOR','IMPACT_STORE','NEW_STORE'          -- Jobs 7/8/9 เขียน DB ตรง (INTERNAL)
    )),
    direction VARCHAR(10) NOT NULL CHECK (direction IN ('IN','OUT','INTERNAL')),
    status VARCHAR(20) NOT NULL CHECK (status IN ('READY','SENT','ACKED','COMPLETED','FAILED','FAILED_RETRY')),
    impact_process_id BIGINT REFERENCES fgi_impact_processes(id),
    sales_summary_id BIGINT REFERENCES fgi_impact_sales_summaries(id),
    doc_no VARCHAR(10), business_key VARCHAR(200) NOT NULL, period_key VARCHAR(20) NOT NULL,
    correlation_id VARCHAR(100), file_name VARCHAR(255), file_checksum VARCHAR(64),
    outbox_status VARCHAR(20), return_code VARCHAR(50), return_message VARCHAR(500),
    retry_count INTEGER NOT NULL DEFAULT 0, sent_at TIMESTAMP, acked_at TIMESTAMP,
    -- marker กัน watchdog (Job 10) ส่งอีเมลเตือนซ้ำในวันเดียวกัน — ย้ายมาจาก audit_logs ที่ยกเลิก 2026-08-07
    last_ack_notified_on DATE,
    purge_after TIMESTAMP, legal_hold BOOLEAN NOT NULL DEFAULT FALSE,
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP, completed_at TIMESTAMP,
    CONSTRAINT uq_interface_business UNIQUE (data_name, direction, business_key, period_key),
    CONSTRAINT ck_interface_typed_reference CHECK (num_nonnulls(impact_process_id, sales_summary_id, doc_no) >= 1)
);"""),
        ("5.3 Zone B — Document and Internal Workflow", """-- ✅ มติ DP-1 (2026-08-10 · ทางเลือก B): PK เป็น surrogate `id` · `doc_no` เป็น UNIQUE ไม่ใช่ PK
-- ⚠️ ผลที่ตามมาซึ่งยังต้องตัดสิน: ตารางลูก 8 ตัว (document_new_stores · document_competitors ·
--    document_external_factors · consideration_logs · document_attachments · document_cost_details ·
--    compensation_histories · interface_transactions) ยัง FK ด้วย doc_no แบบ NOT NULL
--    → แปลว่า "ต้องออก doc_no ให้เสร็จก่อนจึงบันทึกส่วนย่อยได้"
--    ถ้าธุรกิจต้องการสร้างเอกสารก่อนออกเลข ต้องเปลี่ยนตารางลูกไป FK ที่ id แทน (ยังไม่ตัดสิน)
--    referenceId ที่ส่งให้ @srm/glb-workflow = id (ตรงกับที่ระบบเดิมทำจริงใน cooperation-request/inform-evaluate)
--    doc_no อาจยังว่างตอนสร้างแถว แล้วออกเลขทีหลัง จึงเป็น NULL ได้
CREATE TABLE compensation_documents (
    id BIGSERIAL PRIMARY KEY,
    doc_no VARCHAR(10) UNIQUE,          -- YYYY/xxxxx (ปี ค.ศ.) · ออกจาก document_running_numbers
    year INTEGER, running_no INTEGER,   -- แตกจาก doc_no เพื่อ index/ค้นหา (NULL จนกว่าจะออกเลข)
    impact_process_id BIGINT NOT NULL UNIQUE REFERENCES fgi_impact_processes(id),
    impacted_store_code VARCHAR(5) NOT NULL REFERENCES impacted_stores(store_code),
    impact_month CHAR(7), new_store_code VARCHAR(5),   -- master ของระบบเดิม
    round_no INTEGER, loop_no INTEGER,  -- CompMainLoopNo / CompLoopNo — หน้าจอแสดง "รอบ 1 · ครั้งที่ 3"
    source VARCHAR(20) NOT NULL DEFAULT 'FS' CHECK (source IN ('FS','MANUAL')),
    status_code VARCHAR(2) NOT NULL,   -- ค่าจาก sps_store.workflow_status ของ engine
    current_section_code VARCHAR(2),   -- ค่าจาก sps_store.workflow_state ของ engine
    total_compensation_amount NUMERIC(14,2) NOT NULL DEFAULT 0,
    allmap_url VARCHAR(500),                   -- CompUrlMap — ปุ่ม Link To ALLMAP
    statement_id VARCHAR(50),                  -- CompStatementID — โยงกลับ SBP Statement ต้นทาง
    statement_date DATE,                       -- Period Statement (ค.ศ.) — ตัวกรอง/คอลัมน์ของรายงาน SDD สไลด์ 60
    account_year INTEGER, account_month INTEGER,   -- งวดบัญชี
    approver_snapshot JSONB,                   -- FC/Section/Manager/GM/AVP + ชื่อ/อีเมล ณ เวลาเปิดเอกสาร
    version_no INTEGER NOT NULL DEFAULT 1,
    created_by VARCHAR(30) NOT NULL, created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    updated_by VARCHAR(30), updated_at TIMESTAMP,
    CONSTRAINT uq_comp_year_running UNIQUE (year, running_no),
    CONSTRAINT uq_comp_business UNIQUE (source, impacted_store_code, impact_month, new_store_code, round_no)
);

ALTER TABLE interface_transactions
    ADD CONSTRAINT fk_interface_doc_no FOREIGN KEY (doc_no) REFERENCES compensation_documents(doc_no);

CREATE TABLE document_new_stores (
    id BIGSERIAL PRIMARY KEY,
    doc_no VARCHAR(10) NOT NULL REFERENCES compensation_documents(doc_no) ON DELETE CASCADE,
    new_store_code VARCHAR(5) NOT NULL,   -- ร้านเปิดใหม่ · master ของระบบเดิม
    distance_km NUMERIC(8,3), compensate_percent NUMERIC(7,4) NOT NULL CHECK (compensate_percent BETWEEN 0 AND 100),
    compensation_amount NUMERIC(14,2) NOT NULL DEFAULT 0,
    source_system VARCHAR(30) NOT NULL, updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_doc_new_store UNIQUE (doc_no, new_store_code)
);

CREATE TABLE document_competitors (
    id BIGSERIAL PRIMARY KEY,
    doc_no VARCHAR(10) NOT NULL REFERENCES compensation_documents(doc_no) ON DELETE CASCADE,
    competitor_code VARCHAR(30) NOT NULL REFERENCES competitors(competitor_code),
    name_th VARCHAR(200), branch_th VARCHAR(200), opened_date DATE, closed_date DATE, impact_date DATE,
    detail TEXT, remark TEXT, source_system VARCHAR(30) NOT NULL,
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_doc_competitor UNIQUE (doc_no, competitor_code)
);

CREATE TABLE document_external_factors (
    id BIGSERIAL PRIMARY KEY,
    doc_no VARCHAR(10) NOT NULL REFERENCES compensation_documents(doc_no) ON DELETE CASCADE,
    factor_code VARCHAR(30) NOT NULL REFERENCES external_factors(factor_code),
    date_from DATE, date_to DATE, detail TEXT, remark TEXT,
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_doc_factor UNIQUE (doc_no, factor_code, date_from),
    CONSTRAINT ck_doc_factor_dates CHECK (date_to IS NULL OR date_from IS NULL OR date_to >= date_from)
);

CREATE TABLE consideration_logs (
    id BIGSERIAL PRIMARY KEY,
    doc_no VARCHAR(10) NOT NULL REFERENCES compensation_documents(doc_no),
    section_code VARCHAR(2) NOT NULL,   -- ค่าจาก sps_store.workflow_state ของ engine
    result VARCHAR(100) NOT NULL,
    -- APPROVE=ประกันรายได้ · REJECT=ไม่ประกันรายได้ · CANCELLED=ยกเลิกโดยระบบ (decision 14 CancelBySystem) · PENDING=ยังไม่มีผล
    result_category VARCHAR(50) CHECK (result_category IN ('APPROVE','REJECT','CANCELLED','PENDING')),
    detail TEXT,
    consider_by VARCHAR(30) NOT NULL,   -- ผู้ใช้จาก business_user ของระบบเดิม
    action_datetime TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP, request_id VARCHAR(80)
);

CREATE TABLE document_attachments (
    attach_id BIGSERIAL PRIMARY KEY,
    doc_no VARCHAR(10) NOT NULL REFERENCES compensation_documents(doc_no),
    section_code VARCHAR(2) NOT NULL,   -- ค่าจาก sps_store.workflow_state ของ engine
    file_name VARCHAR(255) NOT NULL, mime_type VARCHAR(100) NOT NULL,
    file_size BIGINT NOT NULL CHECK (file_size <= 5242880),
    storage_provider VARCHAR(30) NOT NULL, bucket VARCHAR(120) NOT NULL,
    object_key VARCHAR(500) NOT NULL, sha256 VARCHAR(64) NOT NULL,
    scan_status VARCHAR(20) NOT NULL CHECK (scan_status IN ('PENDING','CLEAN','BLOCKED','FAILED')),
    scanned_at TIMESTAMP, scan_message VARCHAR(500), uploaded_by VARCHAR(30) NOT NULL,
    uploaded_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP, deleted_flag CHAR(1) NOT NULL DEFAULT 'N',
    CONSTRAINT uq_doc_attachment_hash UNIQUE (doc_no, sha256, deleted_flag)
);

CREATE TABLE compensation_histories (
    id BIGSERIAL PRIMARY KEY,
    store_code VARCHAR(5) NOT NULL,   -- master ของระบบเดิม
    ref_doc_no VARCHAR(10) REFERENCES compensation_documents(doc_no),
    submit_account_month CHAR(7) NOT NULL, compensate_amount NUMERIC(14,2) NOT NULL,
    accounting_status VARCHAR(30), external_ref VARCHAR(100),
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_compensation_history UNIQUE (store_code, ref_doc_no, submit_account_month)
);

CREATE TABLE document_cost_details (
    id BIGSERIAL PRIMARY KEY,
    doc_no VARCHAR(10) NOT NULL REFERENCES compensation_documents(doc_no) ON DELETE CASCADE,
    new_store_code VARCHAR(5) NOT NULL,
    cost_year SMALLINT NOT NULL, cost_month SMALLINT NOT NULL CHECK (cost_month BETWEEN 1 AND 12),
    cost_target_n NUMERIC(14,2), cost_amount_n NUMERIC(14,2),
    cost_target_nc NUMERIC(14,2), cost_amount_nc NUMERIC(14,2),
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_document_cost_detail UNIQUE (doc_no, new_store_code, cost_year, cost_month)
);

CREATE TABLE document_running_numbers (
    year SMALLINT PRIMARY KEY,   -- ปี ค.ศ. เท่านั้น (เช่น 2026) ห้ามเก็บ พ.ศ.
    last_running_no INTEGER NOT NULL DEFAULT 0 CHECK (last_running_no >= 0),
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
);
-- ⚠️ year เป็น "ค.ศ." (มติ 2026-08-06 · ทั้งระบบเป็น ค.ศ. — หน้าจอ K2 จริงก็ ค.ศ. เช่น 2026/01870)
--    ห้ามใช้ พ.ศ. · ถ้า client ส่ง พ.ศ. มา ให้ BE แปลงด้วย toAD(y) = y >= 2500 ? y - 543 : y ก่อนเสมอ
-- ออกเลขแบบ atomic (upsert กันกรณีปีใหม่ยังไม่มีแถว):
--   INSERT INTO document_running_numbers (year, last_running_no) VALUES (:ad_year, 1)
--   ON CONFLICT (year) DO UPDATE SET last_running_no = document_running_numbers.last_running_no + 1,
--                                    updated_at = CURRENT_TIMESTAMP
--   RETURNING last_running_no;   (row lock กันเลขชนเมื่อ batch และผู้ใช้สร้างพร้อมกัน)
--   doc_no = :ad_year || '/' || LPAD(last_running_no::text, 5, '0')

-- ❌ ไม่สร้างตาราง workflow_instances ใน SBPGI (ตัดสินใจ 2026-08-06) — ใช้ workflow_transaction ของ @srm/glb-workflow

-- ❌ ไม่สร้างตาราง workflow_tasks ใน SBPGI (ตัดสินใจ 2026-08-06) — ใช้ workflow_approver/workflow_transaction ของ @srm/glb-workflow"""),
        ("5.4 Required Indexes, Partial Uniqueness and Purge", """CREATE INDEX idx_document_status_section ON compensation_documents(status_code, current_section_code);
CREATE INDEX idx_document_impact_process ON compensation_documents(impact_process_id);
-- ❌ ไม่มี index ของ workflow_tasks/workflow_instances ใน SBPGI — ตารางทั้งสองถูกตัดไปแล้ว (2026-08-06)
--    งานค้าง/ผู้อนุมัติปัจจุบันอ่านจาก workflow_transaction + workflow_approver ของ @srm/glb-workflow (schema sps_store)
--    ⚠️ sps_store.workflow_transaction ไม่มี PK และไม่มี index เลย ทั้งที่มี 19,283 แถว (ตรวจ 2026-08-07)
--       -> ยังไม่ตัดสิน (DP-2) ว่าจะขอ sign-off ให้ทีมเจ้าของ library เพิ่ม PK/UNIQUE/index
--          หรือจะกันซ้ำ + ทำ index ที่ฝั่ง SBPGI เอง · ดู SBP/SBPGI-vs-existing-system.md หัวข้อ 4
CREATE INDEX idx_consideration_timeline ON consideration_logs(doc_no, action_datetime DESC);
CREATE INDEX idx_interface_pending ON interface_transactions(data_name, status, sent_at);
CREATE INDEX idx_interface_impact_process ON interface_transactions(impact_process_id);
CREATE INDEX idx_interface_sales_summary ON interface_transactions(sales_summary_id);
CREATE INDEX idx_interface_doc ON interface_transactions(doc_no);

-- index รองรับ FK ที่ PostgreSQL ไม่สร้างให้เอง (เพิ่ม 2026-08-24 หลังตรวจ FK coverage)
CREATE INDEX idx_impact_store_process ON fgi_impact_stores(impact_process_id);
CREATE INDEX idx_document_impacted_store ON compensation_documents(impacted_store_code);
CREATE INDEX idx_compensation_history_doc ON compensation_histories(ref_doc_no);
CREATE INDEX idx_impact_compensation_store ON fgi_impact_compensations(impacted_store_code);
CREATE INDEX idx_impact_competitor_code ON fgi_impact_competitors(competitor_code);
CREATE INDEX idx_document_competitor_code ON document_competitors(competitor_code);
CREATE INDEX idx_document_factor_code ON document_external_factors(factor_code);

-- index ที่หัวข้อ 6 (Index & Constraint) ระบุไว้ — เดิมมีแต่ในตารางสรุป ยังไม่ถูกสร้างจริง (เพิ่ม 2026-08-25)
CREATE INDEX idx_attachment_scan_status ON document_attachments(scan_status);
CREATE INDEX idx_consideration_result ON consideration_logs(result_category);

-- Retention worker: delete only terminal, expired, non-held rows in bounded batches.
WITH purge_candidates AS (
    SELECT id FROM interface_transactions
    WHERE status IN ('ACKED', 'COMPLETED')
      AND purge_after < CURRENT_TIMESTAMP
      AND legal_hold = FALSE
      AND data_name = ANY(:data_names)
    ORDER BY id
    LIMIT :batch_size
    FOR UPDATE SKIP LOCKED
)
DELETE FROM interface_transactions i
USING purge_candidates p
WHERE i.id = p.id
RETURNING i.id, i.data_name, i.business_key;"""),

    ]


def validate_schema_sql_contract() -> None:
    ddl = "\n".join(sql for _, sql in database_ddl_sections())
    # 2026-08-05/06: ตาราง RBAC, workflow, master ร้าน/ภาค/พนักงาน, email template และ config
    # ใช้ของระบบ SBP เดิม จึงไม่มี DDL ในเอกสารนี้ (เหลือเป็นบรรทัดหมายเหตุ "ไม่สร้างตาราง ... ใน SBPGI")
    required_ddl = {
        "sales_transactions": ["txn_date DATE", "window_no SMALLINT", "sales_amount NUMERIC", "sales_diff NUMERIC", "is_outlier BOOLEAN"],
        "consideration_logs": ["result VARCHAR(100)", "detail TEXT", "consider_by VARCHAR(30)", "action_datetime TIMESTAMP"],
    }
    reused_tables = [
        "stores", "zones", "branch_types", "employees", "workflow_sections", "document_statuses",
        "workflow_instances", "workflow_tasks", "email_templates", "system_configs",
        "roles", "menus", "menu_permissions", "user_accounts", "operator_assignments",
        # 2026-08-07: ล็อกเพิ่มตามข้อเท็จจริงที่ตรวจฐานจริง
        # engine 13 ตารางอยู่ใน schema sps_store (ไม่ใช่ 10 ตาราง และไม่ใช่ sps_auth)
        "workflow", "workflow_version", "workflow_state", "workflow_status", "workflow_event",
        "workflow_route", "workflow_group", "workflow_group_map", "workflow_transaction",
        "workflow_history", "workflow_approver", "workflow_part", "workflow_part_display",
        # fcs_qssi_score (เอกพจน์) มีอยู่จริงใน sps_store 23,958,780 แถว + import pipeline ใช้งานอยู่
        # (POST /performance/import-qssi + staging fcs_tmp_qssi_score) → ข้อเท็จจริง F4
        # ⚠️ การ์ดนี้เป็นการ "ล็อกชั่วคราวตามข้อเท็จจริง" ไม่ใช่การตัดสิน DP-4:
        #    ถ้าเจ้าของโครงการเลือกทางเลือก B ของ DP-4 (SBPGI สร้างตารางของตัวเอง)
        #    ต้องถอด "fcs_qssi_score" ออกจากลิสต์นี้ก่อน จึงจะเขียน DDL ลงเอกสารได้
        #    ชื่อพหูพจน์ fcs_qssi_scores ผิดเสมอ ไม่ผูกกับ DP-4
        "fcs_qssi_score", "fcs_qssi_scores",
        "mas_param", "common_code", "business_user", "email_template", "email_sent",
    ]
    # ตัดออกตามมติ 2026-08-07 — ไม่ได้ไป reuse ของระบบเดิม จึงแยกลิสต์และแยกข้อความ error
    # (จะเอา audit ของ master กลับมาหรือไม่ ยังเป็นข้อค้าง DP-12)
    removed_tables = ["audit_logs"]
    for table_name in reused_tables:
        if re.search(rf"CREATE TABLE {table_name} \(", ddl):
            raise ValueError(f"table must reuse the existing SBP system, DDL must not exist: {table_name}")
    for table_name in removed_tables:
        if re.search(rf"CREATE TABLE {table_name} \(", ddl):
            raise ValueError(
                f"table was removed from the target schema on 2026-08-07 (not reused), DDL must not exist: {table_name}"
                " — reinstating it is decision DP-12, see SBP/SBPGI-vs-existing-system.md §4"
            )
    for table_name, columns in required_ddl.items():
        table_match = re.search(rf"CREATE TABLE {table_name} \((.*?)\n\);", ddl, re.S)
        if not table_match:
            raise ValueError(f"missing DDL table: {table_name}")
        body = table_match.group(1)
        missing = [column for column in columns if column not in body]
        if missing:
            raise ValueError(f"DDL/SQL contract drift in {table_name}: missing {missing}")
    forbidden_ddl = ["can_view BOOLEAN", "password_hash", "secret_flag", "result_code VARCHAR", "considered_at TIMESTAMP"]
    for token in forbidden_ddl:
        if token in ddl:
            raise ValueError(f"obsolete DDL vocabulary remains: {token}")

    api_source = (ROOT / "plan-api.html").read_text(encoding="utf-8")
    required_api = [
        # 2026-08-07: ตัด token "mp.can_access = TRUE" ออก — endpoint /menu-permissions ถูกลบจาก
        # plan-api.html แล้ว (RBAC/เมนูใช้ auth-backend ของระบบ SBP เดิม · ตัดสินใจ 2026-08-05)
        # 2026-08-07: ตัด token "instance_status, started_at, started_by" ออก — SQL ที่ INSERT ลง
        # workflow_instances/workflow_tasks ถูกแทนด้วยการเรียก @srm/glb-workflow (schema sps_store) แล้ว
        "sps_store.workflow_transaction",
        "version_no = version_no + 1",
        "FROM fgi_impact_processes",
        "id AS tracking_id",
        "acked_at AS receive_date",
    ]
    for token in required_api:
        if token not in api_source:
            raise ValueError(f"API/DDL contract is missing canonical token: {token}")
    forbidden_api = [
        "u.password_hash",
        # ตารางที่ถูกตัดออกจากโครง 20 ตาราง — ห้ามกลับมาเป็น SQL/DDL ที่ execute ได้ใน plan-api.html
        "FROM workflow_tasks",
        "INTO workflow_tasks",
        "FROM workflow_instances",
        "INTO workflow_instances",
        "FROM document_statuses",
        "INSERT INTO workflow_instances (doc_no, status)",
        "INSERT INTO workflow_instances (instance_id, doc_no, status)",
        "UPDATE fgi_impact_stores SET workflow_generation_status",
        "WHERE tracking_id = :trackingId",
    ]
    for token in forbidden_api:
        if token in api_source:
            raise ValueError(f"obsolete API/DDL vocabulary remains: {token}")


def lldd_database_blocks(all_topics: list[Topic]) -> list[dict[str, Any]]:
    db_ref_topics = [t for t in all_topics if t.db_tables]
    return [
        h(1, "1. Purpose"),
        p("เอกสารนี้เป็น LLDD Database ระดับรวมของ target schema ระบบ SBPGI/SBP Mall ใช้เป็น reference สำหรับ BE API, Batch Job, migration, indexing, transaction และ data dictionary"),
        h(1, "2. Architecture Context"),
        bullets([
            "ระบบใหม่รวม EAI และ K2 เข้าเป็น SBPGI ใช้ฐานข้อมูลเดียวกัน",
            "ไม่มีไฟล์ BPM06001O/BPM06002O/BPM06003O ภายในเพื่อส่งเข้า K2; ใช้ FK จาก compensation_documents ไป impact_process แทน",
            "Workflow ใช้ engine กลาง `@srm/glb-workflow` **13 ตาราง ใน schema `sps_store`** (workflow · workflow_version · workflow_state · workflow_status · workflow_event · workflow_route · workflow_group · workflow_group_map · workflow_transaction · workflow_history · workflow_approver · workflow_part · workflow_part_display) แทน K2 engine ภายนอก — SBPGI ไม่สร้างตาราง workflow ของตัวเอง (ตัดสินใจ 2026-08-06 · แก้จำนวนตารางจาก 10 เป็น 13 และแก้ schema จาก sps_auth เป็น sps_store เมื่อ 2026-08-07)",
            "ตัดขั้นบัญชี 04/05 (SDD v7.5 — รวมเข้าการออกแบบแล้ว ไฟล์ต้นฉบับถูกลบจาก repo 2026-08-06); workflow ใช้ section 06/08/01/02/03; SDD ที่ยึดเป็นหลักคือ SDD GI 24/02/2026",
            "มาตรฐานชื่อ table/column เป็น English lower_snake_case",
            "ตาราง job_configs / job_run_histories ถูกตัดออกจาก target schema เมื่อ 2026-08-06 พร้อม 2 แท็บควบคุมของหน้า Batch Job — cron/พารามิเตอร์อยู่ใน backend config · ผลการรันเขียน application log + interface_transactions",
        ]),
        h(2, "2.1 Input / Progress / Output Contract"),
        table(["Stage", "Contract for implementation"], [
            ["Input", "Target table catalog, data zones, primary keys, foreign-key relationships, migration assumptions, index needs, and transaction boundaries."],
            ["Progress", "Use the data spine impact_process_id -> doc_no -> transaction_id -> approver_id (the last two live in sps_store.workflow_transaction / workflow_approver of @srm/glb-workflow, not in SBPGI tables) to implement APIs/jobs, then validate referential integrity and idempotency keys."],
            ["Output", "Data dictionary and implementation reference for schema creation, migration, indexing, transaction handling, and test data preparation."],
        ]),
        h(1, "3. Data Zones and Spine"),
        table(
            ["Zone", "Scope", "Core tables", "Owner usage"],
            [
                ["A", "FGI/FCS Impact Pipeline and external interfaces", "fgi_impact_processes, fgi_impact_stores, sales, interface_transactions", "Batch Jobs 1-7, IAS/ALLMAP/QSSI/STA tracking"],
                ["B", "K2 Document (workflow อยู่ที่ engine กลาง)", "compensation_documents, document_* tables, consideration_logs, compensation_histories", "Document APIs, workflow actions, FE detail/list/report"],
                ["C", "Master ที่ SBPGI เป็นเจ้าของ (RBAC/config/master ร้าน/ผลพิจารณา ใช้ของระบบ SBP เดิม)", "impacted_stores, external_factors, competitors", "Lookup, master maintenance, notification"],
            ],
        ),
        table(
            ["Order", "Key", "Meaning", "Used by"],
            [
                [1, "impact_process_id", "หนึ่งร้านถูกกระทบ + หนึ่งงวด", "FGI/FCS pipeline, Job 8/8b"],
                [2, "doc_no", "เอกสาร YYYY/xxxxx ปี ค.ศ.", "Document APIs, reports, attachments"],
                [3, "transaction_id (@srm/glb-workflow)", "workflow transaction ต่อเอกสาร — `reference_id` = `compensation_documents.id` (surrogate · DP-1 ปิดแล้ว 2026-08-17)", "Workflow engine ใน schema sps_store"],
                [4, "approver_id (@srm/glb-workflow)", "ผู้อนุมัติต่อ state — แทน task_id เดิม", "Inbox/current approver guard"],
                [5, "employee_id / user_id", "identity — มาจาก BFF header ไม่ใช่ตารางของ SBPGI", "lookup, assignment"],
            ],
        ),
        h(1, "4. Data Dictionary"),
        table(["Zone", "Table", "PK", "FK / relationship", "Role"], database_table_catalog()),
        h(2, "4.1 Canonical Column Contract"),
        table(
            ["Table", "Canonical columns used by DDL and SQL", "Rejected legacy vocabulary"],
            [
                ["workflow_transaction (@srm/glb-workflow)", "transaction_id, version_id, reference_id, current_state_id, current_status_id, current_approver", "instance_id/doc_no/instance_status ของตาราง workflow_instances ที่ถูกตัดไปแล้ว"],
                ["common_code (ระบบ SBP เดิม)", "code_type = SBPGI_APPROVE_LIMIT, code, code_value", "system_configs/approve_limit_amount ที่ถูกตัดไปแล้ว"],
                ["fcs_qssi_score (ระบบ SBP เดิม · เอกพจน์)", "store_id, category_code, period, score", "fcs_qssi_scores (พหูพจน์) — ห้ามใช้"],
                ["sales_transactions", "txn_date, window_no, sales_amount, sales_diff, is_outlier", "sale_date/window_code/net_sales"],
                ["consideration_logs", "result, result_category, detail, consider_by, action_datetime", "result_code/comment/considered_by/considered_at"],
                ["interface_transactions", "id, acked_at", "tracking_id/receive_date (API aliases only)"],
                ["fgi_impact_processes", "workflow_generation_status", "duplicate workflow flag on fgi_impact_stores"],
            ],
        ),
        h(1, "5. Executable DDL — 19 ตาราง (+ fcs_qssi_score ที่ reuse ของระบบ SBP เดิม = 20 ในโครง · + schema reference)"),
        p("หัวข้อ 5.1-5.4 เป็น PostgreSQL DDL ของ **20 ตารางในโครง SBPGI** เรียงตาม dependency พร้อม PK, typed FK, unique/check constraint และ index ที่จำเป็น ใช้เป็น migration baseline ได้โดยไม่ต้องเดา column เพิ่มเติม"),
        *[
            block
            for section_title, sql_text in database_ddl_sections()
            for block in (h(2, section_title), code(sql_text, "sql"))
        ],
        h(1, "6. Index and Constraint Plan"),
        table(
            ["Table", "Index / constraint", "Reason"],
            [
                ["compensation_documents", "UNIQUE (year, running_no), UNIQUE(source, impacted_store_code, impact_month, new_store_code, round_no), INDEX(status_code,current_section_code), INDEX(impact_process_id)", "docNo uniqueness, duplicate guard, list/inbox/report, pipeline trace"],
                ["workflow_transaction (@srm/glb-workflow · sps_store)", "ปัจจุบัน **ไม่มี PK และไม่มี index เลย** ทั้งที่มี 19,283 แถว (ตรวจ 2026-08-07) — ที่ต้องการคือ PK(transaction_id) + UNIQUE(version_id, reference_id) + INDEX(current_approver)", "current approver guard และ inbox · เป็นตารางของ library ไม่ใช่ของ SBPGI จึงต้องขอ sign-off (DP-2 · ยังไม่ตัดสิน)"],
                ["document_new_stores", "INDEX(doc_no) *(ได้จาก UNIQUE (doc_no, new_store_code))*, CHECK compensate_percent between 0 and 100", "detail load and allocation validation"],
                ["consideration_logs", "INDEX(doc_no, action_datetime DESC), INDEX(result_category)", "timeline/report result filter"],
                ["document_attachments", "INDEX(doc_no) *(ได้จาก UNIQUE ที่ขึ้นต้นด้วย doc_no)*, INDEX(scan_status), UNIQUE(doc_no, sha256, deleted_flag)", "attachment list/download/security"],
                ["interface_transactions", "INDEX(data_name,status), INDEX(impact_process_id), INDEX(doc_no)", "tracking and pending ACK"],
            ],
        ),
        h(1, "7. Transaction Rules"),
        table(
            ["Use case", "Transaction boundary", "Rollback rule"],
            [
                ["Create document", "docNo sequence lock (document_running_numbers) + compensation_documents + initializeWorkflow/addPreApprover ของ @srm/glb-workflow", "any fail rollback all; no partial document · engine อยู่คนละ DataSource จึงต้องมี compensating action เมื่อ commit ฝั่งใดฝั่งหนึ่งไม่ผ่าน"],
                ["Submit action", "ตรวจ current_approver จาก workflow_transaction + insert consideration_logs + eventWorkflow (เดิน state) + update compensation_documents", "duplicate/current approver conflict returns 409"],
                ["Auto-assign (SDD 46/48)", "06 เห็นควรไม่ชดเชย -> ปิดเอกสารและตั้งงานเดือนถัดไปให้เจ้าของงานคนเดิม ผ่าน addPreApprover · 06 หยุดชดเชยฯ -> เอกสารกลับเข้า GET /sbpgi/document/tasks ของ 06 ทันที (stoppedReopenable)", "เดือนที่กดเห็นควรไม่ชดเชย ต้องไม่พบเอกสารใน GET /sbpgi/document/tasks ของ 06 · เดือนถัดไปต้องพบพร้อม assignee คนเดิม"],
                ["Attachment upload", "metadata insert only after storage write and AV clean; objectKey never exposed", "storage/scan fail leaves no CLEAN metadata"],
                ["Job 4 IAS request", "durable file (fsync + atomic rename + checksum) ก่อน transaction W→P + outbox READY", "file fail คง W; DB fail rollback W→P/outbox; S3 upload fail retry transaction เดิม"],
                ["Interface ACK/purge", "ACK compare-and-set บน transaction เดิม; purge เฉพาะ terminal + purge_after + non-held", "pending/failed/unacked/legal-hold ห้ามลบ"],
                ["Master mutation", "update entity ใน transaction เดียว", "mutation fail ต้อง rollback ครบ"],
            ],
        ),
        h(1, "8. Seed Data"),
        table(
            ["Domain", "Required seed"],
            [
                ["workflow_state / workflow_status (@srm/glb-workflow)", "5 ขั้น 06, 08, 01, 02, 03 + state จบ flow · 6 สถานะเอกสาร (5 waiting + เสร็จสิ้น) — ลงทะเบียนที่ engine ไม่ใช่ตารางของ SBPGI"],
                ["(ไม่สร้าง) decisions", "ย้ายไป common_code ของระบบเดิม — มติ DP-9 2026-08-10"],
                ["competitors", "แบรนด์คู่แข่ง 11 รายการ รหัส 01-11 (ไทย + อังกฤษ)"],
                ["external_factors", "ปัจจัยภายนอกที่ใช้อยู่"],
                ["email_template (ระบบ SBP เดิม)", "EM-01..EM-08"],
                ["common_code / mas_param (ระบบ SBP เดิม)", "SBPGI_APPROVE_LIMIT: THRESHOLD=100000 (เกณฑ์เดียว · มติ 2026-08-18), impact radius 1/2 km, sales data threshold=60, growth rate threshold=-10"],
            ],
        ),
        h(1, "9. Migration and Verification Checklist"),
        table(
            ["Area", "Check"],
            [
                ["Naming", "all new tables/columns lower_snake_case"],
                ["Leading zero", "store_code/new_store_code stored as VARCHAR(5), never numeric"],
                ["docNo", "year/running_no/doc_no generated in DB transaction; concurrency test 20 parallel requests"],
                ["Workflow", "no active 04/05 accounting sections/statuses; ไม่มีตาราง workflow ของ SBPGI — ตรวจว่า state/route ถูกลงทะเบียนที่ engine ครบ"],
                ["Security", "no secrets in mas_param/backend config; storage objectKey not returned to FE"],
                ["External interface", "credential/certificate/private key อยู่ Secret Manager ผ่าน secretRef; TLS verify-full (HTTPS สำหรับ EAI S3 · AMQPS สำหรับ RabbitMQ ของ STA); ทดสอบ rotation และ invalid certificate/host key"],
                ["Tracking retention", "backfill typed FK/purge_after, validate FK, dry-run count แล้ว purge เฉพาะ ACKED/COMPLETED เป็น batch; reconcile count ก่อน/หลัง"],
                ["Data integrity", "FK/check constraints enabled before SIT; reject legacy invalid enum values"],
                ["Performance", "list/report/inbox queries explain plan uses indexes above"],
            ],
        ),
        h(1, "10. Related LLDD"),
        table(["Document", "DB usage"], [[Path(t.file).name, ", ".join(f"{row[0]}({row[1]})" for row in t.db_tables[:4])] for t in db_ref_topics]),
    ]


def reference_doc_links() -> list[dict[str, str]]:
    return [
        {
            "id": "LLDD-API",
            "title": "LLDD API - REST API and Integration Contract",
            "owner": "BE/FE",
            "scope": "REST conventions, endpoint catalog, request lifecycle, SQL/repository pattern",
            "base": "LLDD-API",
        },
        {
            "id": "LLDD-Database",
            "title": "LLDD Database - Target Schema and Data Dictionary",
            "owner": "BE/DB",
            "scope": "19-table target schema, data zones/spine, DDL reference, indexes, transaction rules, seed data",
            "base": "LLDD-Database",
        },
        {
            "id": "LLDD-To-Be",
            "title": "LLDD To-Be - SDD Traceability and Effort Allocation",
            "owner": "PM/BA",
            "scope": "สอบทานย้อนกลับ SDD GI หัวข้อ 1.9 To-Be -> เอกสาร FE/BE ที่ใช้ + ชั่วโมงต่อข้อ (implementation + unit test)",
            "base": "LLDD-To-Be",
        },
    ]


def render_reference_doc_rows() -> str:
    rows = []
    for doc in reference_doc_links():
        base = doc["base"]
        rows.append(
            "<tr>"
            f"<td><strong>{escape(doc['id'])}</strong><span>{escape(doc['title'])}</span></td>"
            f"<td>{escape(doc['owner'])}</td>"
            f"<td>{escape(doc['scope'])}</td>"
            "<td class=\"links\">"
            f"<a href=\"pdf/{escape(base)}.pdf\">PDF</a>"
            f"<a href=\"word/{escape(base)}.docx\">DOCX</a>"
            "</td>"
            "</tr>"
        )
    return "\n".join(rows)


def topic_links(topic: Topic, prefix: str = "") -> dict[str, str]:
    base = Path(topic.file)
    return {"pdf": f"{prefix}pdf/{base}.pdf", "docx": f"{prefix}word/{base}.docx"}


def doc_id(topic: Topic) -> str:
    return Path(topic.file).name.replace("LLDD-FE-", "FE-").replace("LLDD-BE-", "BE-")


def grouped_topics(all_topics: list[Topic]) -> dict[str, list[Topic]]:
    return {
        "fe_core": [t for t in all_topics if t.track == "FE" and t.file.startswith("FE/") and "Document-Detail-Role" not in t.file],
        "fe_roles": [t for t in all_topics if t.file.startswith("FE/LLDD-FE-Document-Detail-Role")],
        "be_api": [t for t in all_topics if t.track == "BE" and t.file.startswith("BE/LLDD-BE-API")],
        # BE ที่ไม่ใช่กลุ่ม API และไม่ใช่ Job — รวมเอกสารพื้นฐาน 4 ฉบับที่เพิ่ม 2026-08-07
        # (Database-Structure · Data-Migration-Cutover · Integration-SBP-Platform · Workflow-Engine-Definition)
        "be_ops": [
            t for t in all_topics
            if t.track == "BE" and t.file.startswith("BE/") and "/Jobs/" not in t.file
            and not t.file.startswith("BE/LLDD-BE-API")
        ],
        "be_jobs": [t for t in all_topics if "/Jobs/" in t.file],
    }


def render_doc_rows(topics_list: list[Topic]) -> str:
    rows = []
    for topic in topics_list:
        links = topic_links(topic)
        rows.append(
            "<tr>"
            f"<td><strong>{escape(doc_id(topic))}</strong><span>{escape(topic.title)}</span></td>"
            f"<td>{escape(topic.track)}</td>"
            f"<td>{escape(topic.owner)}</td>"
            f"<td>{estimate_html(topic)}</td>"
            "<td class=\"links\">"
            f"<a href=\"{escape(links['pdf'])}\">PDF</a>"
            f"<a href=\"{escape(links['docx'])}\">DOCX</a>"
            "</td>"
            "</tr>"
        )
    return "\n".join(rows)


def estimate_html(topic: Topic) -> str:
    if is_document_detail_role_doc(topic.file):
        return "included<br><small>in Document Detail</small>"
    ut = unit_test_hours(topic)
    if not ut:
        return f"{topic.hours}h"
    return f"{total_hours(topic)}h<br><small>impl {topic.hours} + test {ut}</small>"


def estimate_md(topic: Topic) -> str:
    if is_document_detail_role_doc(topic.file):
        return "included in Document Detail"
    ut = unit_test_hours(topic)
    if not ut:
        return f"{topic.hours}h"
    return f"{total_hours(topic)}h (impl {topic.hours} + test {ut})"


def build_main_index_csv(all_topics: list[Topic]) -> None:
    """Generate LLDD/Main-Index-FE-BE-Job.csv from the same data as
    main index §3 High Level Activity Plan + §8 BE Batch Job Breakdown.

    The CSV used to be maintained by hand and drifted badly (deleted documents,
    old names, pre-2026-08-07 owners, pre-reschedule dates). It is generated now
    so it can never drift from HIGH_LEVEL_ESTIMATES / JOB_ESTIMATES again.
    """
    counted = [t for t in main_index_ordered(all_topics) if not is_document_detail_role_doc(t.file)]
    schedule = build_topic_schedule(counted)
    high_level = [t for t in counted if "/Jobs/" not in t.file]
    be_jobs = [t for t in counted if "/Jobs/" in t.file]
    steps = dependency_steps(counted)
    lines = ["หัวข้อ,owner,ชั่วโมงรวม,implementation,unit test,ลำดับขั้น"]
    for topic in high_level + be_jobs:
        title = topic.title.replace("LLDD ", "")
        ut = unit_test_hours(topic)
        lines.append(
            f"{title},{topic.owner},{total_hours(topic)},{topic.hours},{ut},{steps[topic.file]}"
        )
    (OUT / "Main-Index-FE-BE-Job.csv").write_text("﻿" + "\n".join(lines) + "\n", encoding="utf-8")


def build_document_portal(all_topics: list[Topic]) -> None:
    groups = grouped_topics(all_topics)
    _billable = [t for t in all_topics if not is_document_detail_role_doc(t.file)]
    impl_hours = sum(t.hours for t in _billable)
    ut_hours = sum(unit_test_hours(t) for t in _billable)
    grand_hours = impl_hours + ut_hours
    main_pdf = "pdf/LLDD-Main-Index-Phase4-4-3-SBP-Operating-Management.pdf"
    html = f"""<!doctype html>
<html lang="th">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>LLDD Document Portal - SBP Income Guarantee</title>
  <style>
    :root {{ --ink:#1b2733; --muted:#66717f; --line:#d7e0ea; --head:#eef4fa; --accent:#1e6bb8; --soft:#f7fafc; }}
    body {{ margin:0; font-family: Arial, Tahoma, sans-serif; color:var(--ink); background:#fff; }}
    header {{ padding:28px 36px 22px; border-bottom:1px solid var(--line); background:linear-gradient(180deg,#f8fbff,#fff); }}
    h1 {{ margin:0 0 8px; font-size:28px; color:#0b2545; }}
    h2 {{ margin:28px 0 10px; font-size:20px; color:#174c7f; }}
    p {{ margin:6px 0; color:var(--muted); line-height:1.55; }}
    main {{ padding:22px 36px 42px; max-width:1280px; }}
    .quick {{ display:grid; grid-template-columns: repeat(4, minmax(0,1fr)); gap:12px; margin-top:18px; }}
    .quick a {{ display:block; padding:14px; border:1px solid var(--line); border-radius:8px; text-decoration:none; color:var(--ink); background:#fff; }}
    .quick strong {{ display:block; color:var(--accent); margin-bottom:4px; }}
    .summary {{ display:grid; grid-template-columns: repeat(4, minmax(0,1fr)); gap:10px; margin:18px 0 8px; }}
    .metric {{ border:1px solid var(--line); border-radius:8px; padding:12px; background:var(--soft); }}
    .metric b {{ display:block; font-size:22px; color:#0b2545; }}
    table {{ border-collapse:collapse; width:100%; table-layout:fixed; margin-bottom:18px; }}
    th, td {{ border:1px solid var(--line); padding:9px 10px; vertical-align:top; font-size:14px; }}
    th {{ text-align:left; background:var(--head); color:#24384b; }}
    td strong {{ display:block; color:#0b2545; }}
    td span {{ display:block; color:var(--muted); margin-top:3px; }}
    small {{ color:var(--muted); }}
    .links a {{ display:inline-block; margin:0 6px 6px 0; padding:5px 8px; border:1px solid #b9cbe0; border-radius:6px; text-decoration:none; color:#155c9f; background:#fff; }}
    .note {{ padding:12px 14px; border-left:4px solid var(--accent); background:#f3f8fd; color:#31465a; }}
    @media (max-width: 860px) {{ header, main {{ padding-left:18px; padding-right:18px; }} .quick, .summary {{ grid-template-columns:1fr; }} table {{ table-layout:auto; }} }}
  </style>
</head>
<body>
  <header>
    <h1>LLDD Document Portal</h1>
    <p>SBP Mall - ระบบประกันรายได้ | Phase 4.3 Operating Management</p>
    <div class="summary">
      <div class="metric"><b>{len(all_topics) + len(reference_doc_links())}</b><small>LLDD documents</small></div>
      <div class="metric"><b>{len(groups['fe_core']) + len(groups['fe_roles'])}</b><small>FE documents</small></div>
      <div class="metric"><b>{len(groups['be_api']) + len(groups['be_ops']) + len(groups['be_jobs'])}</b><small>BE/API/Job documents</small></div>
      <div class="metric"><b>{grand_hours}</b><small>estimated hours</small></div>
    </div>
    <div class="quick">
      <a href="{main_pdf}"><strong>Start here</strong>Main LLDD index PDF</a>
      <a href="pdf/LLDD-API.pdf"><strong>API LLDD</strong>REST contract and endpoint catalog</a>
      <a href="pdf/LLDD-Database.pdf"><strong>Database LLDD</strong>Target schema and data dictionary</a>
      <a href="#document-detail-roles"><strong>Document Detail Roles</strong>5 role-specific FE specs</a>
      <a href="#be-api"><strong>BE API</strong>Common and document APIs</a>
      <a href="#be-jobs"><strong>Batch Jobs</strong>Job 1-10 and 8b specs</a>
    </div>
  </header>
  <main>
    <p class="note">วิธีใช้: เอกสาร PDF อยู่ในโฟลเดอร์ pdf, Markdown อยู่ในโฟลเดอร์ md และเอกสาร Word อยู่ในโฟลเดอร์ word โดยแต่ละโฟลเดอร์คงโครงสร้าง FE/BE/Jobs เหมือนกัน</p>
    <p class="note">ขอบเขต 2026-08-07: ตัด <code>LLDD-FE-Overview</code> (หน้า Dashboard ยกเลิก) และ <code>LLDD-BE-API-Dashboard-Summary</code> (endpoint <code>/dashboard/summary</code> ตัดถาวร) · เพิ่ม 4 ฉบับ: <code>LLDD-BE-Database-Structure</code>, <code>LLDD-BE-Data-Migration-Cutover</code>, <code>LLDD-BE-Integration-SBP-Platform</code>, <code>LLDD-BE-Workflow-Engine-Definition</code> · เปลี่ยนชื่อ <code>FE-Master-Config</code> → <code>FE-Master-Data</code>, <code>BE-API-Lookup-RBAC-Email</code> → <code>BE-API-Lookup</code>, <code>BE-API-Report-Master-Config</code> → <code>BE-API-Report-and-Master-Data</code></p>
    <p class="note">ขอบเขต 2026-08-06: ตัดเอกสาร LLDD-FE-Batch-Monitor และ LLDD-FE-Email-Template ออกจากชุดส่งมอบ — หน้า Global Config และ Email Template ลบทั้งฟีเจอร์ (บริหารจัดการที่ระบบ SBP เดิม) ส่วนหน้า Batch Job Monitor พักไว้ก่อน ไม่ทำใน phase นี้ (batch job ยังรันปกติ แต่กำหนดพารามิเตอร์ใน backend config)</p>
    <p class="note">แผนทีม 6 คน (ปรับ 2026-08-07): {escape(FE_OWNER_KITTISAK)}, {escape(FE_OWNER)} (FE); {escape(BE_OWNER_BUTSABA)}, {escape(BE_OWNER)}, {escape(BE_OWNER_PEERAKORN)}, และ {escape(BANK_BE_OWNER)} (BE) โดย 1 week = {WORKDAYS_PER_WEEK} วัน, 1 วัน = {HOURS_PER_DAY} ชั่วโมง · เอกสารระบุเฉพาะชั่วโมงและลำดับขั้น ไม่ระบุวันที่</p>

    <h2 id="reference-docs">Reference Design Documents</h2>
    <table><thead><tr><th>Document</th><th>Owner</th><th>Scope</th><th>Open</th></tr></thead><tbody>
{render_reference_doc_rows()}
    </tbody></table>

    <h2 id="fe-core">FE Core Documents</h2>
    <table><thead><tr><th>Document</th><th>Track</th><th>Owner</th><th>Estimate</th><th>Open</th></tr></thead><tbody>
{render_doc_rows(groups['fe_core'])}
    </tbody></table>

    <h2 id="document-detail-roles">Document Detail Role Pack</h2>
    <p>เอกสารชุดนี้แยกจาก LLDD-FE-Document-Detail เพื่อให้อ่านง่ายตาม role ที่ login จริง</p>
    <table><thead><tr><th>Document</th><th>Track</th><th>Owner</th><th>Estimate</th><th>Open</th></tr></thead><tbody>
{render_doc_rows(groups['fe_roles'])}
    </tbody></table>

    <h2 id="be-api">BE API Documents</h2>
    <table><thead><tr><th>Document</th><th>Track</th><th>Owner</th><th>Estimate</th><th>Open</th></tr></thead><tbody>
{render_doc_rows(groups['be_api'] + groups['be_ops'])}
    </tbody></table>

    <h2 id="be-jobs">BE Batch Job Documents</h2>
    <table><thead><tr><th>Document</th><th>Track</th><th>Owner</th><th>Estimate</th><th>Open</th></tr></thead><tbody>
{render_doc_rows(groups['be_jobs'])}
    </tbody></table>
  </main>
</body>
</html>
"""
    (OUT / "index.html").write_text(html, encoding="utf-8")

    def md_rows(topics_list: list[Topic]) -> list[str]:
        rows = ["| Document | Owner | Estimate | PDF | DOCX |", "| --- | --- | --- | --- | --- |"]
        for topic in topics_list:
            links = topic_links(topic, "../")
            rows.append(
                f"| {doc_id(topic)} | {topic.owner} | {estimate_md(topic)} | "
                f"[PDF]({links['pdf']}) | [DOCX]({links['docx']}) |"
            )
        return rows

    lines = [
        "# LLDD Document Portal",
        "",
        "เปิดหน้า portal ใน browser หรือใช้รายการลิงก์ด้านล่าง.",
        "",
        f"- Main index: [PDF](../{main_pdf})",
        f"- Documents: {len(all_topics) + len(reference_doc_links())}",
        f"- Total estimate: {grand_hours} hours  (implementation {impl_hours} + unit test {ut_hours})",
        "- Unit test: BE/Job 30% · FE 25% ของชั่วโมง implementation · เอกสารสัญญา/ออกแบบไม่คิดแยก (ดู NO_UNIT_TEST_DOCS)",
        "- ขอบเขต 2026-08-07: ตัด `LLDD-FE-Overview` และ `LLDD-BE-API-Dashboard-Summary` · เพิ่ม `LLDD-BE-Database-Structure`, `LLDD-BE-Data-Migration-Cutover`, `LLDD-BE-Integration-SBP-Platform`, `LLDD-BE-Workflow-Engine-Definition` · เปลี่ยนชื่อ `FE-Master-Config` -> `FE-Master-Data`, `BE-API-Lookup-RBAC-Email` -> `BE-API-Lookup`, `BE-API-Report-Master-Config` -> `BE-API-Report-and-Master-Data`",
        "- ขอบเขต 2026-08-06: ตัด `LLDD-FE-Batch-Monitor` และ `LLDD-FE-Email-Template` ออกจากชุดส่งมอบ — หน้า Global Config/Email Template ลบทั้งฟีเจอร์ (ใช้ `mas_param`/`email_template` ของระบบ SBP เดิม) และหน้า Batch Job ย้ายไปกลุ่มเมนู Flow เหลือเฉพาะ Flowchart + Database ที่ใช้ (พารามิเตอร์อยู่ใน backend config)",
        f"- Plan: hours + dependency step only (no calendar dates) with 6-person team `{FE_OWNER_KITTISAK}`, `{FE_OWNER}` (FE) and `{BE_OWNER_BUTSABA}`, `{BE_OWNER}`, `{BE_OWNER_PEERAKORN}`, `{BANK_BE_OWNER}` (BE) — Peerakorn moved FE -> BE on 2026-08-07",
        f"- Working-time rule: 1 week = {WORKDAYS_PER_WEEK} days, 1 day = {HOURS_PER_DAY} hours ({HOURS_PER_WEEK:g} hours/week)",
        f"- Delivery target (2026-08-25): finish in **4 weeks** = {4 * HOURS_PER_WEEK:g} hours per person; team capacity 6 x {4 * HOURS_PER_WEEK:g} = {6 * 4 * HOURS_PER_WEEK:g} hours vs 824 hours of work ({824 / (6 * 4 * HOURS_PER_WEEK) * 100:.0f}% utilisation)",
        f"- Track ownership (2026-08-25, round 3): `Aphiwit <Bank> Khammoon` owns **database migration + ALL batch jobs + building the workflow definition** — Database Structure, Data Migration/Cutover (Oracle FCS_FRN from the Java side + SQL Server CPA_FRN_FGI from K2), Workflow Engine Definition (seeding 6 states / 12 routes into 10 of the engine tables, per `SBP/TSM-SRM-LLDD-SBP-workflow-1.2.md` section 4) and Jobs 2-10 + 8b. **Calling the engine stays with the other backend developers: `initializeWorkflow` (Workflow Instances) and `eventWorkflow` / trigger event (Workflow Actions).**",
        f"- \u26a0\ufe0f **Capacity warning:** that scope totals **296 hours for one person**, which is {296 / (4 * HOURS_PER_WEEK) :.1f}x the 4-week ceiling of {4 * HOURS_PER_WEEK:g} hours ({296 / HOURS_PER_DAY:.0f} working days at {HOURS_PER_DAY} h/day = about 7 weeks). Everyone else finishes well inside 4 weeks (92-123 hours). The plan does not fit until either the batch jobs are shared out or the deadline moves - see DECISIONS.",
        "",
        "## Reference Design Documents",
        "",
        "| Document | Owner | Scope | PDF | DOCX |",
        "| --- | --- | --- | --- | --- |",
        *[
            f"| {doc['id']} | {doc['owner']} | {doc['scope']} | [PDF](../pdf/{doc['base']}.pdf) | [DOCX](../word/{doc['base']}.docx) |"
            for doc in reference_doc_links()
        ],
        "",
        "## FE Core Documents",
        "",
        *md_rows(groups["fe_core"]),
        "",
        "## Document Detail Role Pack",
        "",
        *md_rows(groups["fe_roles"]),
        "",
        "## BE API Documents",
        "",
        *md_rows(groups["be_api"] + groups["be_ops"]),
        "",
        "## BE Batch Job Documents",
        "",
        *md_rows(groups["be_jobs"]),
        "",
    ]
    md_root = OUT / FORMAT_DIRS["md"]
    md_root.mkdir(parents=True, exist_ok=True)
    (md_root / "README.md").write_text("\n".join(lines), encoding="utf-8")


def parse_formats(value: str) -> set[str]:
    formats = {item.strip().lower() for item in value.split(",") if item.strip()}
    allowed = {"md", "docx", "pdf"}
    unknown = formats - allowed
    if not formats or unknown:
        raise argparse.ArgumentTypeError(f"formats must be a comma-separated subset of {sorted(allowed)}")
    return formats


def build_manifest() -> dict[str, list[str]]:
    return {
        "generated": [
            str(p.relative_to(ROOT))
            for p in sorted(OUT.rglob("*"))
            if p.suffix.lower() in {".html", ".md", ".docx", ".pdf", ".png"}
        ]
    }


def write_manifest() -> dict[str, list[str]]:
    manifest = build_manifest()
    (OUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest



# ---------------------------------------------------------------------------
# LLDD To-Be — สอบทานย้อนกลับจาก SDD GI หัวข้อ 1.9 To-Be Business Process
# ที่มา: SDD ปรับปรุงการชดเชยรายได้ในระบบ SBP GI (24/02/2026) สไลด์ 43-62
# ---------------------------------------------------------------------------
TOBE_ITEMS: list[tuple[str, str, str, list[str]]] = [
    ("TB-1", "แก้ไขระบบการสร้างเอกสารชดเชยรายได้", "สไลด์ 43 · 46-51", [
        "Step 1.0 — เพิ่ม**หน้าจองานค้าง** แสดงร้านที่เข้าเงื่อนไขการเปิดกระทบ · เลือกงานของตนเอง · filter · checkbox เลือกหลายเอกสาร · popup ยืนยัน",
        "Step 1.0 — เจ้าหน้าที่ SBP DSA **คีย์ข้อมูลร้านเปิดกระทบเอง** เพื่อให้ระบบคำนวณและ Adjust ยอดชดเชย",
        "Step 1.1 — ฝ่าย SBP DSA ตรวจสอบตัวเลข + ปุ่มพิจารณา 4 ปุ่ม (เห็นควรไม่ชดเชย · หยุดชดเชย · หน่วยงานส่งเสริมฯ SBP · เจ้าหน้าที่ SBP DSA)",
        "**เปิดเรื่องซ้ำได้เองโดยไม่ต้องเปิด SR** — สาขาที่เคยถูกปฏิเสธ/หยุดชดเชย เปิดใหม่ได้ในเดือนเดียวกันหรือเดือนถัดไป",
        "**auto-assign เจ้าของงานคนเดิม (สไลด์ 48)** — (ก) เคสต่อเนื่อง ระบบส่งงานให้เจ้าหน้าที่ SBP DSA **คนเดิม** อัตโนมัติ · (ข) ฝ่าย SBP DSA คลิก **เห็นควรไม่ชดเชย** เอกสารจบและ**ไม่แสดง**ในหน้ารอดำเนินการของ 06 ในเดือนนั้น แล้ว**เดือนถัดไป**ระบบดึงเข้ามาใหม่พร้อมเจ้าของงานคนเดิม",
        "**หยุดชดเชยประกันรายได้ (สไลด์ 46 ข้อ 1.9)** — เอกสารกลับมาแสดงใน**หน้ารอดำเนินการของ ฝ่าย SBP DSA (06) ทันทีในเดือนนั้น** ให้พิจารณาคำขอชดเชยอีกครั้งได้เอง · เปิดเอกสารแล้วเห็นข้อมูลเดิมครบ + แผงพิจารณาชุดเดียวกับสถานะ รอฝ่าย SBP DSA ดำเนินการ · บทบาทอื่นไม่เห็น",
        "**หลังพัฒนาไม่ต้องเปิด SR เพื่อลบข้อมูล (สไลด์ 46 · Note)** — กระบวนการลบ/แก้ข้อมูลเอกสารที่เดิมต้องเปิด SR (มี log ที่ TransectionDeleteStore พร้อม SRNumber) ต้องทำได้ในระบบ · **ยังไม่ระบุขอบเขต ต้องเคาะว่าอนุญาตให้ลบอะไรได้บ้างและใครลบได้**",
        "**แก้ไขครอบคลุมเอกสารทุก Type (สไลด์ 46 · Note)** — การเปลี่ยนแปลงข้อ TB-1 มีผลกับเอกสารทุกประเภทร้าน ไม่จำกัดเฉพาะ Type ใด Type หนึ่ง",
        "ยอดชดเชย 0: เดือน 1-3 คลิก หน่วยงานส่งเสริมธุรกิจ SBP · เดือน 4 คลิก หยุดชดเชยรายได้",
        "สิทธิ์การมองเห็น — เจ้าหน้าที่/ฝ่าย SBP DSA ดูรายละเอียดได้ทุกสาขา ไม่จำกัดเฉพาะงานที่รับผิดชอบ",
    ]),
    ("TB-2", "ปรับสิทธิ์การตรวจสอบยอดชดเชยรายได้และกระบวนการทำงานของทีมส่งเสริม", "สไลด์ 52-58", [
        "Step 2.1 — ลำดับตรวจสอบใหม่: หน่วยงานส่งเสริม (ผู้จัดการฝ่าย/ผู้เชี่ยวชาญ, เจ้าหน้าที่อาวุโส) → GM ส่งเสริม",
        "เปลี่ยนชื่อปุ่ม **ฝ่ายส่งเสริมธุรกิจ SBP → หน่วยงานส่งเสริมธุรกิจ SBP** (มีผลถึงชื่อสถานะเอกสารด้วย)",
        "เพิ่มสิทธิ์ **เจ้าหน้าที่อาวุโส** ให้ส่งต่อ Flow ให้ GM ส่งเสริมได้ (ระดับตำแหน่งอ้างจาก HR Connect)",
        "map ตำแหน่ง HR → กล่องอนุมัติ: AVP = ผู้ช่วยกรรมการผู้จัดการ · GM = ผู้จัดการทั่วไป/รอง/ผู้ช่วย · หน่วยงานส่งเสริม = ผู้จัดการฝ่าย/ผู้เชี่ยวชาญ, เจ้าหน้าที่อาวุโส",
        "**วงเงินอนุมัติ เกณฑ์เดียว 100,000 ต่อรายการ** — < 100,000 จบที่ GM · ≥ 100,000 ผ่าน GM แล้ว AVP อนุมัติ",
        "**เห็นควรไม่ชดเชยจบทันที** ทั้งระดับหน่วยงานส่งเสริมและ GM (เลิกตีกลับให้ SBP DSA รับทราบก่อน)",
    ]),
    ("TB-3", "ยกเลิก Process บัญชี SBP ในการ Approve ค่าใช้จ่าย + เมนูรายงานใน SBP Mall", "สไลด์ 59-62", [
        "ยกเลิกขั้นบัญชี Approve ยอดชดเชยรายได้ออกจาก workflow",
        "Step 3 — สร้าง **เมนูใหม่ใน SBP Mall: รายงานตรวจสอบประกันรายได้** ให้ทีมบัญชีดึงข้อมูลไปใช้ต่อเอง",
        "ตัวกรอง 7 ตัว — สถานะ (บังคับ) · รหัสร้านถูกกระทบ · รหัสร้านเปิดกระทบ · Period Statement (ค.ศ.) · ประเภทร้าน (checkbox) · ภาค (checkbox เพิ่มอัตโนมัติเมื่อมีภาคใหม่) · ผลการพิจารณา (radio)",
        "ปุ่ม ค้นหาข้อมูล · Export Excel · เคลียร์ค่าเริ่มต้น · Preview Report · Export CSV to Batch",
    ]),
    ("TB-0", "งานฐานรากที่ To-Be ทุกข้อใช้ร่วมกัน (ไม่ได้ระบุเป็นข้อใน SDD)", "—", [
        "โครงฐานข้อมูลเป้าหมาย 20 ตาราง + migration/cutover จากระบบเดิม",
        "pipeline FGI/FCS ที่ป้อนข้อมูลให้ทุก To-Be (Job 2-6, Job 10 และงาน interface/อีเมล)",
        "สัญญากลาง API/FE (envelope · error · auth · pagination) และ shell ของ portal",
        "master ที่ SBPGI ดูแลเอง (ปัจจัยภายนอก · แบรนด์คู่แข่ง) และงานทดสอบ/ส่งมอบ",
    ]),
]

# เอกสาร -> {รหัส To-Be: สัดส่วน %}  (แต่ละเอกสารต้องรวมได้ 100)
TOBE_ALLOCATION: dict[str, dict[str, int]] = {
    # ---- FE ----
    "FE/LLDD-FE-Foundation": {"TB-0": 100},
    "FE/LLDD-FE-Integration-Contracts": {"TB-0": 100},
    "FE/LLDD-FE-Testing-Delivery": {"TB-0": 100},
    "FE/LLDD-FE-Master-Data": {"TB-0": 100},
    "FE/LLDD-FE-Document-Lists": {"TB-1": 100},
    "FE/LLDD-FE-Create-Document": {"TB-1": 100},
    "FE/LLDD-FE-Document-Detail": {"TB-1": 60, "TB-2": 40},
    "FE/LLDD-FE-Document-Detail-Role-06-SBP-DSA": {"TB-1": 100},
    "FE/LLDD-FE-Document-Detail-Role-08-SBP-DSA-Officer": {"TB-1": 100},
    "FE/LLDD-FE-Document-Detail-Role-01-Business-Promotion": {"TB-2": 100},
    "FE/LLDD-FE-Document-Detail-Role-02-GM-Business-Promotion": {"TB-2": 100},
    "FE/LLDD-FE-Document-Detail-Role-03-AVP-SBP": {"TB-2": 100},
    "FE/LLDD-FE-Report": {"TB-3": 100},
    # ---- BE ----
    "BE/LLDD-BE-Database-Structure": {"TB-0": 100},
    "BE/LLDD-BE-Data-Migration-Cutover": {"TB-0": 100},
    "BE/LLDD-BE-API-Common-Contracts": {"TB-0": 100},
    "BE/LLDD-BE-Job-Batch-Email-SRM": {"TB-0": 100},
    "BE/Jobs/LLDD-BE-Job-2-ImportImpactStore": {"TB-0": 100},
    "BE/Jobs/LLDD-BE-Job-3-ImportImpactCompetitor": {"TB-0": 100},
    "BE/Jobs/LLDD-BE-Job-4-PrepareImpactStoreToIAS": {"TB-0": 100},
    "BE/Jobs/LLDD-BE-Job-5-ImportImpactSaleFromIAS": {"TB-0": 100},
    "BE/Jobs/LLDD-BE-Job-6-ExportImpactStoreToFS": {"TB-0": 100},
    "BE/Jobs/LLDD-BE-Job-10-NotifyNoReceiveData": {"TB-0": 100},
    "BE/Jobs/LLDD-BE-Job-7-SyncCompetitorToDocument": {"TB-1": 100},
    "BE/Jobs/LLDD-BE-Job-8-CreateCompensationDocument": {"TB-1": 100},
    "BE/Jobs/LLDD-BE-Job-8b-StartInternalWorkflow": {"TB-1": 100},
    "BE/Jobs/LLDD-BE-Job-9-SyncNewStoreToDocument": {"TB-1": 100},
    "BE/LLDD-BE-API-Document-List-Search": {"TB-1": 100},
    "BE/LLDD-BE-API-Document-Create-Update": {"TB-1": 100},
    "BE/LLDD-BE-API-Document-Detail-Aggregate": {"TB-1": 100},
    "BE/LLDD-BE-API-Attachment-Sales-Timeline": {"TB-1": 100},
    "BE/LLDD-BE-API-Document-Workflow-Actions": {"TB-1": 40, "TB-2": 60},
    "BE/LLDD-BE-API-Workflow-Instances": {"TB-1": 50, "TB-2": 50},
    "BE/LLDD-BE-Workflow-Engine-Definition": {"TB-2": 100},
    "BE/LLDD-BE-API-Lookup": {"TB-2": 100},
    "BE/LLDD-BE-Integration-SBP-Platform": {"TB-2": 100},
    "BE/LLDD-BE-API-Report-and-Master-Data": {"TB-3": 70, "TB-0": 30},
}


def tobe_split(topic: "Topic") -> dict[str, tuple[int, int]]:
    """คืน {รหัส To-Be: (ชม. implementation, ชม. unit test)} ของเอกสารนั้น

    ปัดให้ผลรวมของทุกข้อเท่ากับชั่วโมงจริงของเอกสารเสมอ (ข้อสุดท้ายรับเศษ)
    """
    alloc = TOBE_ALLOCATION.get(topic.file)
    if not alloc:
        raise SystemExit(f"TOBE_ALLOCATION ขาดเอกสาร {topic.file}")
    if sum(alloc.values()) != 100:
        raise SystemExit(f"TOBE_ALLOCATION ของ {topic.file} รวมได้ {sum(alloc.values())}% ไม่ใช่ 100%")
    ut = unit_test_hours(topic)
    keys = list(alloc)
    out: dict[str, tuple[int, int]] = {}
    imp_left, ut_left = topic.hours, ut
    for i, k in enumerate(keys):
        if i == len(keys) - 1:
            out[k] = (imp_left, ut_left)
        else:
            hi = round(topic.hours * alloc[k] / 100)
            hu = round(ut * alloc[k] / 100)
            out[k] = (hi, hu)
            imp_left -= hi
            ut_left -= hu
    return out


def tobe_blocks(all_topics: list["Topic"]) -> list[dict[str, Any]]:
    """LLDD To-Be — เขียนแบบ 'SDD สั่งอะไร · กี่ชั่วโมง · ใครทำ · ไปดูที่ LLDD ฉบับไหน'

    กติกาเวลา (มติ 2026-08-25): **นับเฉพาะงานที่ To-Be เพิ่มเข้ามาใหม่**
    งานฐานราก (TB-0) ที่ต้องทำอยู่แล้วไม่ว่าจะมี To-Be หรือไม่ ไม่นับเป็นเวลาของ To-Be
    """
    by_file = {t.file: t for t in all_topics}
    for t in all_topics:
        tobe_split(t)  # validate ทุกฉบับตั้งแต่ต้น
    billable = [t for t in all_topics if not is_document_detail_role_doc(t.file)]

    def short(owner: str) -> str:
        return owner.split("<")[1].split(">")[0] if "<" in owner else owner

    # code -> {"fe": ชม., "be": ชม., "owners": {ชื่อ: ชม.}, "rows": [...]}
    agg: dict[str, dict[str, Any]] = {
        code: {"fe": 0, "be": 0, "owners": {}, "rows": []} for code, *_ in TOBE_ITEMS
    }
    for t in billable:
        for code, (hi, hu) in tobe_split(t).items():
            hours = hi + hu
            if hours == 0:
                continue
            a = agg[code]
            a["fe" if t.track == "FE" else "be"] += hours
            a["owners"][short(t.owner)] = a["owners"].get(short(t.owner), 0) + hours
            a["rows"].append((t.track, t.file, short(t.owner), hours,
                              TOBE_ALLOCATION[t.file].get(code, 0)))

    tobe_codes = [c for c, *_ in TOBE_ITEMS if c != "TB-0"]
    add_fe = sum(agg[c]["fe"] for c in tobe_codes)
    add_be = sum(agg[c]["be"] for c in tobe_codes)
    base_fe, base_be = agg["TB-0"]["fe"], agg["TB-0"]["be"]

    def owner_line(code: str) -> str:
        fe = [(o, h) for o, h in agg[code]["owners"].items()
              if any(r[0] == "FE" and r[2] == o for r in agg[code]["rows"])]
        be = [(o, h) for o, h in agg[code]["owners"].items()
              if any(r[0] != "FE" and r[2] == o for r in agg[code]["rows"])]
        fmt = lambda xs: " · ".join(f"**{o}** {h} ชม." for o, h in sorted(xs, key=lambda x: -x[1])) or "—"
        return f"FE: {fmt(fe)}   |  BE: {fmt(be)}"

    blocks: list[dict[str, Any]] = [
        h(1, "1. เอกสารนี้ตอบอะไร"),
        p("ตอบคำถามเดียว: **SDD สั่งให้ทำอะไรเพิ่ม · ใช้เวลากี่ชั่วโมง · ใครทำฝั่ง FE ใครทำฝั่ง BE · "
          "รายละเอียดอยู่ใน LLDD ฉบับไหน** — ใช้คู่กับ SDD ปรับปรุงการชดเชยรายได้ในระบบ SBP GI (24/02/2026) หัวข้อ 1.9 To-Be Business Process"),
        p(f"⚠️ **กติกาการนับเวลา (มติ 2026-08-25):** นับ**เฉพาะงานที่ To-Be เพิ่มเข้ามาใหม่ = {add_fe + add_be} ชั่วโมง** เท่านั้น · "
          f"งานฐานรากที่ต้องทำอยู่แล้วไม่ว่าจะมี To-Be หรือไม่ ({base_fe + base_be} ชั่วโมง — โครงฐานข้อมูล, pipeline FGI/FCS, สัญญากลาง API/FE, shell ของ portal, งานทดสอบ/ส่งมอบ) "
          "แยกไว้ท้ายเอกสารและ**ไม่นับรวมเป็นเวลาของ To-Be**"),
        p("ชั่วโมงที่แสดงคือชั่วโมงเดียวกับที่ประกาศในแต่ละฉบับ (implementation + unit test) · "
          "เอกสารที่รับใช้ To-Be หลายข้อจะถูกแบ่งตามสัดส่วนในคอลัมน์ *สัดส่วนของฉบับ* จึงไม่มีการนับซ้ำ"),
        h(1, "2. สรุปงานที่ To-Be เพิ่มเข้ามา"),
        table(
            ["ข้อ", "SDD สไลด์", "ทำอะไร", "FE (ชม.)", "BE (ชม.)", "รวม (ชม.)"],
            [[code, src, title, str(agg[code]["fe"]), str(agg[code]["be"]),
              f"**{agg[code]['fe'] + agg[code]['be']}**"]
             for code, title, src, _ in TOBE_ITEMS if code != "TB-0"]
            + [["", "", "**รวมงานที่ To-Be เพิ่ม**", f"**{add_fe}**", f"**{add_be}**", f"**{add_fe + add_be}**"],
               ["TB-0", "—", "*(ฐานราก — ไม่นับเป็นเวลาของ To-Be)*",
                f"*{base_fe}*", f"*{base_be}*", f"*{base_fe + base_be}*"]],
        ),
    ]

    idx = 2
    for code, title, src, sdd_points in TOBE_ITEMS:
        if code == "TB-0":
            continue
        idx += 1
        a = agg[code]
        blocks += [
            h(1, f"{idx}. {code} · SDD {src} — {title}"),
            p(f"**ใช้เวลา {a['fe'] + a['be']} ชั่วโมง** (FE {a['fe']} + BE {a['be']})  ·  {owner_line(code)}"),
            h(2, f"{idx}.1 SDD สั่งให้ทำอะไร"),
        ]
        blocks += [bullets(list(sdd_points))]
        blocks += [
            h(2, f"{idx}.2 ทำที่เอกสาร LLDD ฉบับไหน"),
            table(
                ["เอกสาร LLDD", "สาย", "ผู้รับผิดชอบ", "ชม.", "สัดส่วนของฉบับ"],
                [[f"`{f.split('/')[-1]}`", tr, o, f"**{hh}**",
                  "เต็มฉบับ" if pct == 100 else f"{pct}% ของฉบับ"]
                 for tr, f, o, hh, pct in sorted(a["rows"], key=lambda r: (-r[3], r[1]))],
            ),
        ]

    idx += 1
    a0 = agg["TB-0"]
    blocks += [
        h(1, f"{idx}. TB-0 · งานฐานราก — **ไม่นับเป็นเวลาของ To-Be**"),
        p(f"งานชุดนี้ ({a0['fe'] + a0['be']} ชั่วโมง · FE {a0['fe']} + BE {a0['be']}) ต้องทำอยู่แล้วไม่ว่าจะมี To-Be หรือไม่ "
          "— SDD ไม่ได้ระบุเป็นข้อ และไม่ควรนับเป็นต้นทุนของการเปลี่ยนแปลงตาม To-Be · "
          "แสดงไว้เพื่อให้เห็นภาพรวมของชุดส่งมอบทั้งหมดเท่านั้น"),
    ]
    blocks += [bullets(list(next(pts for c, _, _, pts in TOBE_ITEMS if c == "TB-0")))]
    blocks += [
        table(
            ["เอกสาร LLDD", "สาย", "ผู้รับผิดชอบ", "ชม.", "สัดส่วนของฉบับ"],
            [[f"`{f.split('/')[-1]}`", tr, o, str(hh),
              "เต็มฉบับ" if pct == 100 else f"{pct}% ของฉบับ"]
             for tr, f, o, hh, pct in sorted(a0["rows"], key=lambda r: (-r[3], r[1]))],
        ),
        h(1, f"{idx + 1}. เอกสารที่รับใช้ To-Be มากกว่าหนึ่งข้อ"),
        p("เอกสารเหล่านี้ถูกแบ่งชั่วโมง จึงต้องอ่านคู่กันเมื่อวางแผนคน — ถ้าเลื่อนข้อใดข้อหนึ่ง เอกสารที่แชร์กันจะกระทบทั้งสองข้อ"),
        table(
            ["เอกสาร LLDD", "แบ่งให้ข้อไหนบ้าง", "เหตุผลที่ต้องแบ่ง"],
            [
                ["`LLDD-FE-Document-Detail`", "TB-1 60% · TB-2 40%",
                 "หน้าเดียวกันแต่คนละบทบาท — ฝั่ง DSA (คีย์งาน/ปรับยอด) เป็น TB-1 · ฝั่งส่งเสริม/GM/AVP (ปุ่มพิจารณา วงเงิน) เป็น TB-2"],
                ["`LLDD-BE-API-Document-Workflow-Actions`", "TB-1 40% · TB-2 60%",
                 "endpoint เดียวรับ 6-enum — เส้นทางของ DSA เป็น TB-1 · กติกาวงเงิน/ลำดับใหม่/เห็นควรไม่ชดเชยจบทันที เป็น TB-2"],
                ["`LLDD-BE-API-Workflow-Instances`", "TB-1 50% · TB-2 50%",
                 "เปิด instance ให้เอกสารใหม่ (TB-1) และลงทะเบียน version ที่ฝัง state/วงเงินใหม่ (TB-2)"],
                ["`LLDD-BE-API-Report-and-Master-Data`", "TB-3 70% · TB-0 30%",
                 "เส้นรายงาน (TB-3) กับเส้น master ปัจจัยภายนอก/คู่แข่ง (TB-0 ฐานราก) อยู่เอกสารเดียวกัน"],
            ],
        ),
        h(1, f"{idx + 2}. Related LLDD"),
        p("รายละเอียดการทำงานจริงอยู่ในเอกสารที่อ้างถึงข้างบน · ภาพรวมชุดส่งมอบดูที่ `LLDD/md/README.md` · "
          "สัญญา API ดูที่ `LLDD-API` · โครงฐานข้อมูลดูที่ `LLDD-Database`"),
    ]
    return blocks

def main() -> None:
    parser = argparse.ArgumentParser(description="Build LLDD deliverables")
    parser.add_argument("--formats", type=parse_formats, default={"md", "docx", "pdf"}, help="comma-separated: md,docx,pdf")
    args = parser.parse_args()
    formats: set[str] = args.formats
    OUT.mkdir(exist_ok=True)
    manifest_path = OUT / "manifest.json"
    if formats == {"md", "docx", "pdf"} and manifest_path.exists():
        try:
            previous = json.loads(manifest_path.read_text(encoding="utf-8")).get("generated", [])
            for rel in previous:
                target = ROOT / rel
                if target.is_file() and target.is_relative_to(OUT):
                    target.unlink()
        except Exception:
            pass
    validate_schema_sql_contract()
    all_topics = topics()
    render_all("LLDD Main Index - Phase 4.3 SBP Operating Management ประกันรายได้", main_doc_blocks(all_topics), OUT / "LLDD-Main-Index-Phase4-4-3-SBP-Operating-Management", formats)
    render_all("LLDD API - REST API and Integration Contract", lldd_api_blocks(all_topics), OUT / "LLDD-API", formats)
    render_all("LLDD Database - Target Schema and Data Dictionary", lldd_database_blocks(all_topics), OUT / "LLDD-Database", formats)
    render_all("LLDD To-Be - SDD Traceability and Effort Allocation", tobe_blocks(all_topics), OUT / "LLDD-To-Be", formats)
    for topic in all_topics:
        render_all(topic.title, topic_blocks(topic), OUT / topic.file, formats)
    if formats == {"md", "docx", "pdf"}:
        build_document_portal(all_topics)
        build_main_index_csv(all_topics)
    manifest = build_manifest()
    if formats == {"md", "docx", "pdf"}:
        manifest = write_manifest()
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
