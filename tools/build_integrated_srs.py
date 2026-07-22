#!/usr/bin/env python3
"""Build the integrated SBPGI SRS as DOCX and PDF from repository sources."""

from __future__ import annotations

import html as html_lib
import json
import os
import re
import textwrap
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

from lxml import html
from PIL import Image as PILImage

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    CondPageBreak,
    Flowable,
    Image,
    KeepTogether,
    LongTable,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "srs"
TMP = ROOT / "tmp"
DATA_FILE = TMP / "prototype_data.json"
HEADER_LOGO = TMP / "pdfs" / "extracted-images" / "img-002.png"
COVER_BADGE = TMP / "pdfs" / "extracted-images" / "img-000.png"

DOCX_OUT = OUT / "SRS-ระบบประกันรายได้-SBPGI-Integrated-v1.0.docx"
PDF_OUT = OUT / "pdf" / "SRS-ระบบประกันรายได้-SBPGI-Integrated-v1.0.pdf"
MD_OUT = OUT / "SRS-ระบบประกันรายได้-SBPGI-Integrated-v1.0.md"
SCREENSHOT_FULL_DIR = OUT / "screenshots" / "full"
SCREENSHOT_SLICE_DIR = OUT / "screenshots" / "slices"

TEMPLATE_CODE = ""
DOC_VERSION = "1.0"
RELEASE_DATE = "07/07/2026"
TARGET_TABLE_COUNT = 34

SCREEN_LABELS = {
    "index.html": "Overview / Dashboard",
    "flow-fgi.html": "Flow FGI/FCS",
    "k2-flow.html": "Flow K2",
    "plan-flow.html": "Integrated Target Flow",
    "fgi-database.html": "Database FGI/FCS",
    "k2-database.html": "Database K2",
    "plan-database.html": "Target Database Schema",
    "job-batch.html": "Batch Job Console",
    "k2-create.html": "Create Document",
    "k2-list-waiting.html": "Task Inbox",
    "k2-list-related.html": "Related Documents",
    "k2-list-abnormal.html": "Abnormal Assignment",
    "k2-document.html": "Document Detail",
    "k2-report.html": "Status Report",
    "k2-operators.html": "Operator Master",
    "k2-factors.html": "External Factor Master",
    "k2-permissions.html": "RBAC Matrix",
    "system-config.html": "Global Config",
    "plan-email.html": "Email Template",
    "plan-api.html": "API Specification",
}


def clean(value: Any) -> str:
    text = "" if value is None else str(value)
    text = html_lib.unescape(text)
    for filename, label in SCREEN_LABELS.items():
        text = re.sub(rf"\(\s*{re.escape(filename)}\s*\)", f"({label})", text)
        text = text.replace(filename, label)
    replacements = {
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2192": "->",
        "\u2190": "<-",
        "\u2265": ">=",
        "\u2264": "<=",
        "\u2260": "!=",
        "\u00d7": "x",
        "\u2026": "...",
        "\u2605": "",
        "\u25c4": "<-",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def element_text(node) -> str:
    return clean(" ".join(node.xpath(".//text()")))


def read_html(name: str):
    return html.fromstring((ROOT / name).read_text(encoding="utf-8"))


def parse_db_entities(name: str, classes: tuple[str, ...]) -> list[dict[str, Any]]:
    root = read_html(name)
    entities: list[dict[str, Any]] = []
    class_expr = " or ".join(
        f'contains(concat(" ", normalize-space(@class), " "), " {cls} ")' for cls in classes
    )
    for box in root.xpath(f"//div[{class_expr}]"):
        heading = box.xpath("./h3[1]")
        table = box.xpath(".//table[1]")
        if not heading or not table:
            continue
        names = [clean(x) for x in heading[0].xpath(".//code/text()") if clean(x)]
        title = " / ".join(names) or element_text(heading[0])
        source_tags = [
            clean(x)
            for x in heading[0].xpath('.//span[contains(@class,"src") or contains(@class,"db-tag")]/text()')
            if clean(x)
        ]
        headers = [element_text(x) for x in table[0].xpath(".//thead/tr[1]/th")]
        rows = []
        for tr in table[0].xpath(".//tbody/tr"):
            cells = [element_text(td) for td in tr.xpath("./th|./td")]
            if cells:
                rows.append(cells)
        notes = [
            element_text(p)
            for p in box.xpath("./p|.//p[contains(@class,'db-note')]")
            if element_text(p)
        ]
        entities.append(
            {"name": title, "source": " · ".join(source_tags), "headers": headers, "rows": rows, "notes": notes}
        )
    return entities


def parse_plan_database() -> list[list[str]]:
    root = read_html("plan-database.html")
    tables = root.xpath(
        '//h2[contains(.,"Data Dictionary")]/ancestor::div['
        'contains(concat(" ", normalize-space(@class), " "), " card ")][1]//table'
    )
    if not tables:
        raise RuntimeError("Plan database table not found")
    rows: list[list[str]] = []
    for tr in tables[0].xpath(".//tbody/tr"):
        cells = [element_text(td) for td in tr.xpath("./td")]
        if not cells:
            continue
        if len(cells) == 1:
            rows.append([cells[0], "", "", "", "", ""])
        else:
            rows.append(cells)
    return rows


def parse_plan_flow_migration() -> list[list[str]]:
    root = read_html("plan-flow.html")
    tables = root.xpath(
        '//h2[contains(.,"จุดเชื่อมต่อที่เปลี่ยน")]/ancestor::div['
        'contains(concat(" ", normalize-space(@class), " "), " card ")][1]//table'
    )
    if not tables:
        return []
    return [[element_text(td) for td in tr.xpath("./td")] for tr in tables[0].xpath(".//tbody/tr")]


def parse_markdown_table(path: Path, header_starts: str) -> list[list[str]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    start = next((i for i, line in enumerate(lines) if line.strip().startswith(header_starts)), -1)
    if start < 0:
        return []
    rows: list[list[str]] = []
    for line in lines[start + 2 :]:
        if not line.strip().startswith("|"):
            break
        rows.append([clean(cell) for cell in line.strip().strip("|").split("|")])
    return rows


def unique_texts(values: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for value in values:
        value = clean(value)
        if value and value not in seen:
            seen.add(value)
            out.append(value)
    return out


def target_job(job: dict[str, Any]) -> dict[str, Any]:
    """Align legacy batch-source entries with the SBPGI target architecture."""

    def normalize_target_text(value: Any) -> Any:
        if isinstance(value, str):
            for source, target in {
                "วันทำการ ≠ 60": "วันทำการ < 60",
                "วันทำการ != 60": "วันทำการ < 60",
                "total_working_days ≠ 60": "total_working_days < 60",
                "total_working_days != 60": "total_working_days < 60",
                "ไม่เท่ากับ 60": "น้อยกว่า 60",
            }.items():
                value = value.replace(source, target)
            return value
        if isinstance(value, list):
            return [normalize_target_text(item) for item in value]
        if isinstance(value, dict):
            return {key: normalize_target_text(item) for key, item in value.items()}
        return value

    no = str(job.get("no", ""))
    overrides: dict[str, dict[str, Any]] = {
        "1": {
            "params": [
                ["กำหนดการรัน (Cron)", "Monthly", "text", 1, "ตั้งเวลาใน scheduler ผ่าน deployment config"],
                ["งวดข้อมูล (เดือนที่รัน)", "07/2569", "text", 1, "ชื่อไฟล์ใช้เดือนปัจจุบัน แต่งวดใน DB คือเดือนก่อนหน้า"],
                ["SFTP endpoint alias", "qssi-monthly", "text", 0, "resolve host/port จาก environment; ไม่รับค่า host/port จาก request หรือ job_configs"],
                ["Secret reference", "secret/sbpgi/interfaces/qssi", "text", 0, "credential/private key อ่านจาก Secret Manager และบังคับ strict known_hosts"],
                ["Remote Directory", "/export/qssishare/onl/qssi/textfile/SBP/QSSI_Monthly/", "text", 1, "path เท่านั้น ไม่รวม credential"],
                ["Local Directory", "/appshare/SPS/FCS/interface_data/in/", "text", 1, "staging/quarantine path"],
                ["File Prefix (4 ไฟล์)", "mrs1trnf_, mrs2trnf_, mrs3trnf_, mrs5trnf_", "text", 0, ""],
                ["Encoding", "WINDOWS-874", "text", 0, ""],
                ["Batch Insert Size", "10000", "number", 1, "จำนวนแถวต่อรอบ insert"],
            ],
            "flow": [
                {"k": "start", "t": "เริ่ม"},
                {"k": "p", "t": "กำหนดงวดและ snapshot config", "d": "endpoint alias + secretRef; ไม่รับ host/port/credential จาก UI"},
                {"k": "io", "t": "เชื่อมต่อ SFTP ผ่าน qssi-monthly", "d": "Secret Manager + strict known_hosts"},
                {"k": "d", "t": "ดาวน์โหลดครบและ checksum ถูกต้อง?", "no": "quarantine / FAILED; ไม่แก้คะแนนเดิม", "noKind": "err", "d": ""},
                {"k": "p", "t": "parse และ validate 4 prefix / WINDOWS-874", "d": "reject ราย record พร้อมเหตุผล"},
                {"k": "p", "t": "transaction upsert fcs_qssi_scores", "d": "UNIQUE(store_code, category_code, score_period)"},
                {"k": "p", "t": "archive source และ reconcile count", "d": "checksum เดิมให้ SKIP"},
                {"k": "end", "t": "จบ"},
            ],
        },
        "3": {
            **job,
            "cls": "th.co.gosoft.fgi.main.ImportImpactCompetitor",
            "script": "/appstore/SPS/FGI/schedule/FGI_ImportCompetitor.sh",
            "params": [
                ["กำหนดการรัน (Cron)", "0 07 7 * *", "text", 1, "ใช้สคริปต์ /appstore/SPS/FGI/schedule/FGI_ImportCompetitor.sh; Operations ตรวจ deployment path และ owner permission ก่อนขึ้น production"],
                ["Argument (งวด)", "2569|06", "text", 1, "รูปแบบ YYYY|MM"],
                ["Chunk Size", "10000", "number", 1, "จำนวนแถวต่อรอบ insert"],
                ["Source View", "COMPETITOR_IMPACT_VIEW", "text", 0, "SELECT DISTINCT / map คอลัมน์ NAMT -> NAME_TH, BRANCHT -> BRANCH_TH"],
            ],
        },
        "4": {
            "desc": "สร้างไฟล์คำขอยอดขาย IAS/MIS แบบ durable ก่อนเปลี่ยนสถานะ W→P แล้วบันทึก transactional outbox เพื่อส่งซ้ำได้โดยไม่สร้างรายการซ้ำ",
            "params": [
                ["กำหนดการรัน (Cron)", "0 16 7-16 * *", "text", 1, "รันวันที่ 7-16 เวลา 16:00"],
                ["IAS SFTP endpoint alias", "ias-sales-request", "text", 0, "host/port resolve จาก environment; credential ใช้ secretRef และ strict known_hosts"],
                ["Secret reference", "secret/sbpgi/interfaces/ias", "text", 0, "ห้ามเก็บ password/private key ใน job_configs"],
                ["Output staging path", "/data/sbpgi/outbox/ias", "text", 1, "ต้องรองรับ temp file, fsync และ atomic rename"],
            ],
            "flow": [
                {"k": "start", "t": "เริ่ม"},
                {"k": "p", "t": "lock รายการ sales_request_status=W", "d": "FOR UPDATE SKIP LOCKED"},
                {"k": "p", "t": "สร้าง temporary file และ validate record count", "d": "ยังไม่เปลี่ยน W→P"},
                {"k": "p", "t": "fsync + atomic rename + SHA-256", "d": "ไฟล์ต้อง durable ก่อนเริ่ม DB transaction"},
                {"k": "p", "t": "transaction: update W→P + insert outbox READY", "d": "fail แล้ว rollback ทั้งสถานะและ outbox"},
                {"k": "io", "t": "dispatcher ส่ง SFTP ด้วย secretRef/strict known_hosts", "d": "retry จาก outbox transaction เดิม"},
                {"k": "end", "t": "จบ"},
            ],
            "tables": [
                ["fgi_impact_stores", "R/W", "lock candidate W และเปลี่ยนเป็น P หลัง durable file สำเร็จเท่านั้น"],
                ["fgi_impact_sales_summaries", "R/W", "สร้าง/ผูกหัวสรุปยอดขายใน transaction"],
                ["interface_transactions", "W", "transactional outbox READY/SENT/ACKED พร้อม checksum และ idempotency key"],
                ["job_run_histories", "W", "run status และ reconcile count"],
            ],
            "meta": {
                "trans": "durable file ก่อน; transaction เดียว update W→P + insert outbox READY; dispatcher ส่งภายหลัง",
                "rerun": "UNIQUE(data_name,direction,business_key,period_key) และ checksum เดิมไม่สร้าง request ซ้ำ",
                "mail": "Notification Service แจ้งเมื่อ durable write, DB transaction หรือ SFTP retry เกิน threshold",
                "risk": "Target remediation: ห้าม commit W→P ก่อน fsync/atomic rename/checksum สำเร็จ และห้ามส่ง SFTP โดยไม่มี outbox",
            },
        },
        "6": {
            "params": [
                ["กำหนดการรัน (Cron)", "0 17 * * *", "text", 1, "ทุกวัน 17:00"],
                ["dateStartInitToSTA", "7", "number", 1, "วันของเดือนที่เริ่มปล่อยสถานะ I, C"],
                ["numWaitPay", "3", "number", 1, "จำนวนงวดรอจ่าย"],
                ["หมวด QSSI ที่ตรวจ", "8, 9, 12, 1, 10, 16", "text", 0, "ต้องครบทั้ง 6 หมวดจากงวด max เดียว ในกรอบ 3 เดือน"],
                ["Output File", "FRBC0001_yyyyMMddHHmmss.txt (windows-874, 14 ฟิลด์, พ.ศ.)", "text", 0, "ฟิลด์ 3/5/6 เป็นวันที่แบบไทย/พุทธศักราช"],
                ["STA endpoint alias", "sta-compensation", "text", 0, "resolve host/port/TLS policy จาก environment; ห้าม editable endpoint"],
                ["Secret reference", "secret/sbpgi/interfaces/sta", "text", 0, "credential/certificate/private key จาก Secret Manager; TLS verify-full หรือ strict known_hosts"],
            ],
        },
        "7": {
            "name": "SyncCompetitorToDocument",
            "th": "บันทึกข้อมูลคู่แข่งเข้าเอกสาร",
            "tag": "DOCUMENT SERVICE",
            "tagCls": "doc",
            "phase": "B",
            "cls": "document.service.syncCompetitors",
            "script": "(internal scheduler / service)",
            "out": "document_competitors (DB)",
            "desc": "อ่านข้อมูลคู่แข่งล่าสุดจาก fgi_impact_competitors แล้วบันทึกเข้า document_competitors ผ่าน Document Service โดยตรง แทนการเขียนไฟล์ BPM06003O และ SFTP ไป BPM",
            "params": [
                ["กำหนดการรัน (Cron)", "30 17 7-31 * *", "text", 1, "ใช้รอบเดิม แต่ปลายทางเป็น DB ภายใน"],
                ["Target table", "document_competitors", "text", 0, "upsert ด้วย doc_no / competitor_code / source_system=ALM"],
                ["เงื่อนไขเลือกข้อมูล", "งวดคู่แข่งล่าสุดต่อร้าน + forecast เริ่มต้น + ยังไม่ sync", "text", 0, "คง business rule เดิม"],
            ],
            "flow": [
                {"k": "start", "t": "เริ่ม"},
                {"k": "p", "t": "อ่านคู่แข่งงวดล่าสุดต่อร้านจาก fgi_impact_competitors", "d": "dense rank ตามงวดต้นทางของคู่แข่ง"},
                {"k": "d", "t": "มี compensation_documents ของ impact_process_id แล้ว?", "no": "คงสถานะรอ sync / log pending", "noKind": "err", "d": ""},
                {"k": "p", "t": "upsert document_competitors", "d": "source_system=ALM, ผูก doc_no และ competitor_code"},
                {"k": "p", "t": "บันทึก interface_transactions เป็น INTERNAL_DB_WRITE", "d": "ไม่สร้างไฟล์ BPM06003O"},
                {"k": "end", "t": "จบ"},
            ],
            "tables": [
                ["fgi_impact_competitors", "R", "ข้อมูลคู่แข่งล่าสุดจาก Job 3"],
                ["compensation_documents", "R", "หา doc_no จาก impact_process_id"],
                ["document_competitors", "W", "บันทึกคู่แข่งเข้าเอกสารโดยตรง"],
                ["interface_transactions", "W", "tracking ภายใน type=INTERNAL_DB_WRITE"],
            ],
            "rels": [
                "แทน legacy BPM06003O ด้วย document_competitors.doc_no -> compensation_documents",
                "ไม่มี SFTP/EAI/BPM ภายในใน target architecture",
            ],
            "meta": {
                "trans": "DB transaction ครอบการ upsert document_competitors + tracking",
                "rerun": "idempotent ด้วย doc_no + competitor_code + source_system",
                "mail": "ส่ง error ผ่าน Notification Service กลางเมื่อ sync ล้มเหลว",
                "risk": "ห้าม re-implement การเขียนไฟล์ BPM06003O หรือ SFTP ไป BPM; legacy file เป็น reference เท่านั้น",
            },
        },
        "8": {
            "name": "CreateCompensationDocument",
            "th": "สร้างเอกสารประกันรายได้อัตโนมัติ",
            "tag": "DOCUMENT SERVICE",
            "tagCls": "doc",
            "phase": "B",
            "cls": "document.service.createFromImpact",
            "script": "(internal scheduler / service)",
            "out": "compensation_documents (DB)",
            "desc": "สร้าง compensation_documents จาก impact profile และข้อมูลชดเชยในฐานข้อมูลเดียวกัน แทนการเขียนไฟล์ BPM06001O และ SFTP ไป compensateflow; ไม่เรียก workflow โดยตรง",
            "params": [
                ["กำหนดการรัน (Cron)", "30 17 7-31 * *", "text", 1, "ใช้รอบเดิม แต่ปลายทางเป็น DB ภายใน"],
                ["Target table", "compensation_documents", "text", 0, "สร้าง doc_no YYYY/xxxxx และผูก impact_process_id"],
                ["เงื่อนไขเลือกข้อมูล", "สถานะ I + forecast + ยังไม่สร้างเอกสาร", "text", 0, "Gen Flow Gate อยู่ที่ Job 8b / Workflow Engine"],
                ["ข้อห้ามเชิงสถาปัตยกรรม", "ห้ามสร้างไฟล์ BPM06001O, ห้าม SFTP, ห้ามเรียก K2 REST", "text", 0, "ใช้ Document Service + DB transaction เท่านั้น"],
            ],
            "flow": [
                {"k": "start", "t": "เริ่ม"},
                {"k": "p", "t": "query impact profile สถานะ I + forecast + ยังไม่สร้างเอกสาร", "d": "ใช้ impact_process_id เป็น idempotency key"},
                {"k": "d", "t": "ข้อมูลผู้อนุมัติ/ร้าน/ยอดชดเชยครบ?", "no": "บันทึก reject reason / ไม่สร้างเอกสาร", "noKind": "err", "d": ""},
                {"k": "p", "t": "generate doc_no YYYY/xxxxx", "d": "running ต่อปี พ.ศ."},
                {"k": "p", "t": "insert compensation_documents", "d": "ผูก impact_process_id และสถานะเริ่มต้น"},
                {"k": "p", "t": "บันทึก interface_transactions เป็น INTERNAL_DB_WRITE", "d": "ไม่สร้างไฟล์ BPM06001O"},
                {"k": "end", "t": "จบ - workflow เปิดโดย Job 8b / POST /workflows/instances"},
            ],
            "tables": [
                ["fgi_impact_stores", "R/W", "อ่าน candidate และอัปเดตสถานะสร้างเอกสาร"],
                ["fgi_impact_processes", "R", "hub รอบชดเชย"],
                ["compensation_documents", "W", "สร้างหัวเอกสารแทนไฟล์ BPM06001O"],
                ["interface_transactions", "W", "tracking ภายใน type=INTERNAL_DB_WRITE"],
            ],
            "rels": [
                "compensation_documents.impact_process_id -> fgi_impact_processes แทน BPM06001O",
                "Job 8 สร้างเอกสารเท่านั้น; Job 8b เป็นตัวเปิด Workflow ภายใน",
            ],
            "meta": {
                "trans": "DB transaction เดียวครอบ generate doc_no + insert document + tracking",
                "rerun": "idempotent ด้วย impact_process_id; เจอ doc เดิมให้ skip และคืนสถานะ already_created",
                "mail": "Notification Service แจ้ง error/pending ตาม config",
                "risk": "ห้ามนำ logic SFTP compensateflow หรือ K2 StartInstance กลับมาใช้ใน target design",
            },
        },
        "8b": {
            "name": "StartInternalWorkflow",
            "th": "เปิด Workflow ภายใน",
            "tag": "WORKFLOW ENGINE",
            "tagCls": "k2",
            "phase": "B",
            "cls": "workflow.service.startFromImpact",
            "script": "(internal scheduler / service token)",
            "cron": "after-job-8",
            "cronTh": "trigger หลัง Job 8 สร้างเอกสารสำเร็จ; manual rerun ได้ตาม period",
            "out": "workflow_instances / workflow_tasks (DB)",
            "desc": "คัดรายการที่ผ่าน Gen Flow Gate แล้วเรียก Workflow Engine ภายในผ่าน POST /api/v1/workflows/instances แทน K2 REST StartInstance; เกณฑ์ W/Y/N เดิมยังคงใช้สำหรับ reconcile",
            "params": [
                ["Scheduler", "หลัง Job 8 สร้างเอกสารสำเร็จ; manual rerun ตาม period", "text", 1, "แยกเพื่อ rerun ได้อิสระ; Operations ตรวจ deployment schedule/queue เท่านั้น"],
                ["Workflow API", "POST /api/v1/workflows/instances", "text", 0, "internal service token; ไม่ใช่ K2 REST"],
                ["เกณฑ์ Growth Rate", "growth_rate_diff <= -10", "number", 0, "คง business rule เดิม"],
                ["Branch Type ผ่าน Gate", "FAM, FB1, FC1, FB2, FVB, FVC", "text", 0, "นอกเซ็ตหรือระยะทางเกินเกณฑ์ให้ตั้ง N"],
                ["เงื่อนไข Gate อื่น", "workflow_generation_status=W · DV ไม่ว่าง · juristic ต่างกัน · sales_status in {Y,N}", "text", 0, "DV หาย, นิติบุคคลเดียวกัน หรือ growth ไม่ถึงเกณฑ์เป็น N; distance/juristic/growth/sales status ที่ยังไม่มีค่าเท่านั้นจึงคง W"],
            ],
            "flow": [
                {"k": "start", "t": "เริ่ม"},
                {"k": "p", "t": "อ่าน candidate ที่มี compensation_documents แล้วและ workflow_generation_status=W", "d": ""},
                {"k": "d", "t": "พบเงื่อนไขไม่ผ่านถาวร?", "no": "ไม่พบ - ตรวจความพร้อมของข้อมูลต่อ", "d": "branch type, distance, missing DV, same juristic หรือ growth > -10 -> N"},
                {"k": "d", "t": "ข้อมูล Gate พร้อมครบ?", "no": "distance/juristic/growth เป็น NULL หรือ sales status ยังไม่พร้อม -> คง W", "d": "คง W เฉพาะข้อมูลต้นทางที่ยังรอเติมเพื่อให้ rerun ได้"},
                {"k": "io", "t": "POST /api/v1/workflows/instances", "d": "service token ภายใน ไม่ใช้ HTTP Basic Auth/K2 REST"},
                {"k": "p", "t": "insert workflow_instances + workflow_tasks แรก Section 06", "d": ""},
                {"k": "p", "t": "workflow_generation_status = Y", "d": "เปิด workflow สำเร็จ"},
                {"k": "io", "t": "ส่งอีเมลสรุปราย DV ผ่าน Notification Service", "d": ""},
                {"k": "end", "t": "จบ"},
            ],
            "tables": [
                ["fgi_impact_stores", "R/W", "อ่าน candidate + เขียน W/Y/N"],
                ["compensation_documents", "R/W", "ยืนยันเอกสารจาก Job 8 หรือสร้างถ้ายังไม่มีตาม idempotency"],
                ["workflow_instances", "W", "เปิด instance ภายใน"],
                ["workflow_tasks", "W", "สร้าง task แรก Section 06"],
                ["status_email_rules", "R", "ผู้รับอีเมลตามสถานะ"],
            ],
            "rels": [
                "แทน K2 REST StartInstance ด้วย Workflow Engine ภายใน",
                "ไม่มี Basic Auth/HTTP endpoint legacy ใน runtime target",
            ],
            "meta": {
                "trans": "DB transaction ครอบ create instance/task + update W/Y/N",
                "rerun": "idempotent ด้วย doc_no/impact_process_id; ตรวจ workflow_instances เดิมก่อนสร้างใหม่",
                "mail": "อีเมลราย DV ผ่าน Notification Service",
                "risk": "ห้ามเรียก K2 REST endpoint legacy; เก็บไว้เป็น reference migration เท่านั้น",
            },
        },
        "9": {
            "name": "SyncNewStoreToDocument",
            "th": "บันทึกร้านเปิดใหม่เข้าเอกสาร",
            "tag": "DOCUMENT SERVICE",
            "tagCls": "doc",
            "phase": "B",
            "cls": "document.service.syncNewStores",
            "script": "(internal scheduler / service)",
            "out": "document_new_stores (DB)",
            "desc": "อ่านโปรไฟล์ร้านเปิดใหม่และค่า forecast/adjust แล้วบันทึกเข้า document_new_stores ผ่าน Document Service โดยตรง แทนการเขียนไฟล์ BPM06002O และ SFTP ไป impactprofile",
            "params": [
                ["กำหนดการรัน (Cron)", "30 17 7-31 * *", "text", 1, "ใช้รอบเดิม แต่ปลายทางเป็น DB ภายใน"],
                ["Target table", "document_new_stores", "text", 0, "upsert ด้วย doc_no / new_store_code"],
                ["กฎ Forecast / Percent", "NVL(adjust_n, forecast_n)", "text", 0, "ค่า adjust มาก่อน forecast เสมอ; NULL หรือค่านอกช่วง 0..100 ต้อง reject ก่อน upsert"],
                ["เงื่อนไขเลือกข้อมูล", "ร้านเปิดใหม่ สถานะ I + forecast + ยังไม่ sync", "text", 0, ""],
            ],
            "flow": [
                {"k": "start", "t": "เริ่ม"},
                {"k": "p", "t": "query ร้านเปิดใหม่ สถานะ I + forecast + ยังไม่ sync", "d": ""},
                {"k": "d", "t": "มี compensation_documents ของ impact_process_id แล้ว?", "no": "คงสถานะรอ sync / log pending", "noKind": "err", "d": ""},
                {"k": "d", "t": "compensate_percent ครบและอยู่ในช่วง 0..100 ทุกแถว?", "no": "COMPENSATE_PERCENT_INVALID + rollback ก่อน upsert/prune", "noKind": "err", "d": "COALESCE(adjust_compensate_percent, forecast_compensate_percent) ต้องไม่เป็น NULL"},
                {"k": "p", "t": "upsert document_new_stores", "d": "forecast/percent = NVL(adjust_n, forecast_n)"},
                {"k": "p", "t": "validate allocation percent รวมต่อ doc_no", "d": "ต้องรวมได้ 100 ก่อน submit workflow"},
                {"k": "p", "t": "บันทึก interface_transactions เป็น INTERNAL_DB_WRITE", "d": "ไม่สร้างไฟล์ BPM06002O"},
                {"k": "end", "t": "จบ"},
            ],
            "tables": [
                ["fgi_impact_stores", "R", "โปรไฟล์ร้านเปิดใหม่และค่า forecast/adjust รายงวด"],
                ["compensation_documents", "R", "หา doc_no จาก impact_process_id"],
                ["document_new_stores", "W", "บันทึกร้านเปิดใหม่เข้าเอกสารโดยตรง"],
                ["interface_transactions", "W", "tracking ภายใน type=INTERNAL_DB_WRITE"],
            ],
            "rels": [
                "แทน legacy BPM06002O ด้วย document_new_stores.doc_no -> compensation_documents",
                "ไม่มี SFTP/EAI/BPM ภายในใน target architecture",
            ],
            "meta": {
                "trans": "validate percent ไม่เป็น NULL และอยู่ 0..100 ก่อน; DB transaction ครอบ upsert document_new_stores + tracking; พบค่าผิดให้ rollback ก่อน prune",
                "rerun": "idempotent ด้วย doc_no + new_store_code",
                "mail": "ส่ง error ผ่าน Notification Service กลางเมื่อ sync ล้มเหลว",
                "risk": "ห้าม re-implement การเขียนไฟล์ BPM06002O หรือ SFTP ไป BPM; legacy file เป็น reference เท่านั้น",
            },
        },
        "10": {
            "out": "อีเมลเตือน UTF-8 + pending ACK dashboard",
            "desc": "งาน safety net ตรวจ interface_transactions หา ACK จาก STA ที่ยังค้างเกิน 1 วัน หลังเพิ่ม POST /api/v1/interfaces/sta/ack ให้ STA callback ตรง; ส่งอีเมล UTF-8 ผ่าน Notification Service กลาง",
            "params": [
                ["กำหนดการรัน (Cron)", "0 07 * * *", "text", 1, "ทุกวัน 07:00; เป็น safety net หลัง STA callback"],
                ["Pending threshold", ">= 1 วัน", "number", 1, "เตือนเมื่อยังไม่มี ACK หลังครบ threshold"],
                ["data_name ที่เฝ้าดู", "COMPENSATE_INIT_I, COMPENSATE_APPROVE_I", "text", 0, "เฉพาะฝั่ง STA - ไม่เฝ้า dataset ของ BPM"],
                ["Encoding", "UTF-8", "text", 0, "แทน TIS-620 เดิมตาม Notification Service กลาง"],
            ],
            "flow": [
                {"k": "start", "t": "เริ่ม"},
                {"k": "p", "t": "อ่าน interface_transactions ฝั่ง STA ที่ยังไม่มี ACK และอายุ >= threshold", "d": ""},
                {"k": "d", "t": "พบรายการค้าง?", "no": "จบการทำงาน", "noKind": "end", "d": ""},
                {"k": "io", "t": "ส่งอีเมล UTF-8 ผ่าน Notification Service", "d": "ผู้รับตาม config/status_email_rules"},
                {"k": "p", "t": "แสดงรายการใน /interfaces/pending-ack", "d": "POST /interfaces/sta/ack เป็นเส้นทางหลักเมื่อ STA ตอบกลับ"},
                {"k": "end", "t": "จบ"},
            ],
            "tables": [
                ["interface_transactions", "R", "pending ACK จาก STA และสถานะล่าสุด"],
                ["email_templates", "R", "template EM-08 watchdog ACK"],
                ["status_email_rules", "R", "ผู้รับอีเมล"],
            ],
            "meta": {
                "trans": "read-only; callback /interfaces/sta/ack เป็นผู้เขียน ACK หลัก",
                "rerun": "รันซ้ำได้; ต้องไม่ส่งอีเมลซ้ำถ้ามี sent marker ในรอบเดียวกัน",
                "mail": "Notification Service UTF-8",
                "risk": "ห้ามกลับไปใช้ TIS-620/hardcoded recipient; Job 10 เป็น safety net ไม่ใช่ primary ACK path",
            },
        },
    }
    updated = dict(job)
    if no in overrides:
        updated.update(overrides[no])
        if no in {"7", "8", "8b", "9", "10"}:
            updated["run"] = job.get("run", ["-", "pending", "-"])
            updated["hist"] = [
                [updated["run"][0], "-", "-", updated["out"], "ok", "target design - DB/internal workflow"],
            ]
    return normalize_target_text(updated)


SRS_JOB_USER_CATALOG: dict[str, dict[str, str]] = {
    "1": {
        "purpose": "นำเข้าคะแนน QSSI รายเดือนเพื่อใช้ประกอบการคำนวณและตรวจเงื่อนไขการชดเชย",
        "input": "ไฟล์คะแนน QSSI รายเดือน 4 ชุดจาก SFTP, งวดเดือนที่ต้องประมวลผล, และหมวดคะแนนที่ระบบกำหนด",
        "summary": "ระบบอ่านไฟล์ ตรวจรูปแบบและงวดข้อมูล คัดรายการล่าสุดต่อร้าน/หมวดคะแนน แล้วปรับปรุงคะแนน QSSI ของงวดนั้นให้เป็นชุดล่าสุด",
        "output": "คะแนน QSSI ของร้านถูกบันทึกพร้อมใช้งานสำหรับงานส่ง Statement และรายงานผลการประมวลผลแสดงจำนวนไฟล์/จำนวนรายการ/สถานะสำเร็จหรือผิดพลาด",
        "visible": "Admin ติดตามได้จาก Batch Job Console และประวัติการรัน; ผู้ใช้ธุรกิจเห็นผลผ่านข้อมูลประกอบเอกสาร/รายงาน",
    },
    "2": {
        "purpose": "นำเข้าคู่ร้านที่ได้รับผลกระทบจากร้านเปิดใหม่ เพื่อสร้างฐานข้อมูลการพิจารณาชดเชย",
        "input": "ข้อมูลงวดเดือนและข้อมูลร้านจาก ALLMAP ที่ระบุร้านเปิดใหม่ ร้านถูกกระทบ ระยะทาง รัศมี โซน และประเภทสาขา",
        "summary": "ระบบคัดเลือกร้านที่เข้าเกณฑ์ ตรวจซ้ำตามงวดและคู่ร้าน แล้วบันทึกเป็นรายการผลกระทบตั้งต้นสำหรับ pipeline ประกันรายได้",
        "output": "รายการร้านถูกกระทบและร้านเปิดใหม่ถูกสร้าง/ปรับสถานะให้พร้อมสำหรับการขอยอดขายและการคำนวณต่อไป",
        "visible": "Admin เห็นจำนวนรายการที่นำเข้าและสถานะรอบล่าสุด; ทีมงานเห็นข้อมูลเป็น candidate ของเอกสารในขั้นต่อไป",
    },
    "3": {
        "purpose": "นำเข้าข้อมูลคู่แข่งรอบล่าสุดของร้านที่ได้รับผลกระทบ",
        "input": "งวดปี/เดือนและข้อมูลคู่แข่งจาก ALLMAP เช่น รหัสคู่แข่ง ชื่อ สาขา โซน วันที่เปิด/ปิด",
        "summary": "ระบบตรวจว่างวดนั้นเคยนำเข้าหรือยัง คัดข้อมูลคู่แข่งที่เกี่ยวข้อง แล้วบันทึกเข้าฐานข้อมูลคู่แข่งของร้านถูกกระทบ",
        "output": "ข้อมูลคู่แข่งพร้อมถูกนำไปแสดงในเอกสารประกันรายได้หลังระบบสร้างเอกสาร",
        "visible": "Admin ตรวจได้จาก run history; ผู้พิจารณาเห็นคู่แข่งในหน้าเอกสารเมื่อ sync สำเร็จ",
    },
    "4": {
        "purpose": "ส่งคำขอข้อมูลยอดขายรายวันไปยัง IAS/MIS สำหรับร้านที่ต้องใช้ยอดขายประกอบการคำนวณ",
        "input": "รายการร้านที่รอข้อมูลยอดขาย, วันที่เปิดร้านใหม่, งวดที่ต้องตรวจ, และพารามิเตอร์รอบส่งไฟล์",
        "summary": "ระบบคัดรายการที่ครบเงื่อนไข สร้างไฟล์คำขอยอดขาย ส่งออกไปยัง IAS/MIS และบันทึกสถานะว่ารอผลตอบกลับ",
        "output": "ไฟล์คำขอยอดขายถูกส่งออก และรายการที่เกี่ยวข้องถูกตั้งสถานะรอข้อมูลขายกลับมา",
        "visible": "Admin เห็นชื่อไฟล์ จำนวนรายการ และสถานะส่งออก; งานที่ยังรอยอดขายไม่ควรถูกสร้างเอกสารก่อนครบข้อมูล",
    },
    "5": {
        "purpose": "รับยอดขายจาก IAS/MIS แล้วคำนวณผลกระทบยอดขายก่อน/หลังร้านเปิดใหม่",
        "input": "ไฟล์ยอดขายรายวันจาก IAS/MIS, รายการร้านที่เคยส่งคำขอ, และกฎจำนวนวัน/ช่วงเวลาที่ต้องเปรียบเทียบ",
        "summary": "ระบบอ่านยอดขาย แยกช่วงก่อนและหลังเปิดร้านใหม่ทั้งปีก่อนหน้าและปีปัจจุบัน คำนวณอัตราเติบโตและผลต่าง แล้วตรวจความครบของวันทำการ",
        "output": "สรุปยอดขายและค่า growth rate ถูกบันทึก; รายการที่ข้อมูลไม่ครบหรือผิดเงื่อนไขถูกแยกให้ตรวจสอบก่อนเดิน workflow",
        "visible": "ผู้ใช้เห็นผลผ่านสถานะข้อมูลผิดปกติ/ข้อมูลพร้อมสร้างเอกสาร และ Admin เห็นจำนวน success/reject ใน run history",
    },
    "6": {
        "purpose": "ส่งข้อมูลชดเชยที่ผ่านเงื่อนไขไปยังระบบ Statement/บัญชี",
        "input": "เอกสารหรือรายการชดเชยที่อนุมัติแล้ว, ข้อมูล QSSI ที่เกี่ยวข้อง, และสถานะรายการที่ต้องส่ง Statement",
        "summary": "ระบบคัดรายการที่พร้อมส่ง ตรวจเงื่อนไขสำคัญ สร้างข้อมูลส่งออกไป STA และบันทึก tracking เพื่อรอการตอบกลับ",
        "output": "รายการชดเชยถูกส่งไป STA/Statement และระบบมีรายการติดตาม ACK สำหรับ reconcile",
        "visible": "ทีมบัญชีและ Admin เห็นสถานะส่งออก/รอ ACK ผ่านรายงานและ Batch Job Console",
    },
    "7": {
        "purpose": "บันทึกข้อมูลคู่แข่งที่เกี่ยวข้องเข้าเอกสารประกันรายได้",
        "input": "ข้อมูลคู่แข่งล่าสุดของร้านถูกกระทบและเอกสารประกันรายได้ที่สร้างแล้ว",
        "summary": "ระบบจับคู่ข้อมูลคู่แข่งกับเอกสารที่เกี่ยวข้อง และบันทึกเข้ารายการคู่แข่งของเอกสารโดยไม่ให้ซ้ำ",
        "output": "หน้าเอกสารมีข้อมูลคู่แข่งครบสำหรับผู้พิจารณาใช้ประกอบการตัดสินใจ",
        "visible": "ผู้พิจารณาเห็นข้อมูลคู่แข่งในหน้าเอกสาร; Admin เห็นจำนวนรายการที่ sync สำเร็จหรือรอเอกสาร",
    },
    "8": {
        "purpose": "สร้างเอกสารประกันรายได้อัตโนมัติจากข้อมูลร้านที่ผ่านเงื่อนไข",
        "input": "ข้อมูล impact process, ร้านถูกกระทบ, ร้านเปิดใหม่, ยอดชดเชยตั้งต้น, และสถานะพร้อมสร้างเอกสาร",
        "summary": "ระบบตรวจว่าข้อมูลหลักครบหรือยัง สร้างเลขเอกสาร ผูกเอกสารกับ impact process และกันการสร้างเอกสารซ้ำ",
        "output": "เกิดเอกสารประกันรายได้พร้อมสถานะเริ่มต้น เพื่อรอเปิด workflow และเติมข้อมูลประกอบจาก job อื่น",
        "visible": "ผู้ใช้เห็นเอกสารใหม่ในรายการเมื่อสิทธิ์และ workflow พร้อม; Admin เห็นจำนวนเอกสารที่สร้าง/ข้ามเพราะมีอยู่แล้ว",
    },
    "8b": {
        "purpose": "เปิด workflow ภายในสำหรับเอกสารที่ผ่านเงื่อนไข Gen Flow Gate",
        "input": "เอกสารที่สร้างแล้ว, สถานะรอเปิด workflow, เงื่อนไข branch type, DV, นิติบุคคล, growth rate และ sales status",
        "summary": "ระบบตรวจเงื่อนไข Gen Flow Gate ถ้าผ่านจะสร้าง workflow instance และ task แรก ถ้าไม่ผ่านจะคง/ปรับสถานะตามสาเหตุเพื่อให้ตรวจสอบหรือรันซ้ำได้",
        "output": "เอกสารถูกส่งเข้าสู่ workflow และมี task ให้ผู้รับผิดชอบดำเนินการ หรือถูกคงสถานะรอแก้ไขเมื่อยังไม่ครบเงื่อนไข",
        "visible": "ผู้รับผิดชอบเห็นงานใน Inbox; Admin เห็นรายการผ่าน/ไม่ผ่าน gate และเหตุผลใน run history",
    },
    "9": {
        "purpose": "บันทึกร้านเปิดใหม่และสัดส่วนชดเชยเข้าเอกสารประกันรายได้",
        "input": "ข้อมูลร้านเปิดใหม่ ค่า forecast/adjust และเอกสารที่เกี่ยวข้องกับ impact process",
        "summary": "ระบบจับคู่ร้านเปิดใหม่กับเอกสาร บันทึกยอด/เปอร์เซ็นต์ชดเชย และตรวจว่าข้อมูลรวมพร้อมให้ผู้ใช้พิจารณาต่อ",
        "output": "หน้าเอกสารมีรายการร้านเปิดใหม่พร้อมยอดและเปอร์เซ็นต์ชดเชยสำหรับตรวจสอบ",
        "visible": "ผู้พิจารณาเห็นร้านเปิดใหม่ในหน้าเอกสาร; Admin เห็นจำนวนรายการ sync สำเร็จหรือรอเอกสาร",
    },
    "10": {
        "purpose": "เฝ้าระวังรายการส่ง Statement ที่ยังไม่ได้รับผลตอบกลับจาก STA",
        "input": "รายการ interface ที่ส่งไป STA แล้วแต่ยังไม่มี ACK/ผลตอบกลับเกินระยะเวลาที่กำหนด",
        "summary": "ระบบค้นหารายการค้าง จัดกลุ่มตามประเภทข้อมูลและไฟล์/ช่องทางส่ง แล้วส่งแจ้งเตือนให้ผู้เกี่ยวข้องติดตาม",
        "output": "เกิดอีเมลหรือรายการแจ้งเตือน pending ACK เพื่อให้ทีมงานตรวจสอบกับระบบปลายทาง",
        "visible": "Admin และทีมบัญชีเห็นรายการค้างผ่าน dashboard/report และได้รับการแจ้งเตือนตาม rule",
    },
}


def page_inventory(name: str) -> dict[str, Any]:
    root = read_html(name)
    title = element_text(root.xpath("//h1[1]")[0]) if root.xpath("//h1[1]") else name
    labels = unique_texts(
        [element_text(x) for x in root.xpath("//label")]
        + [element_text(x) for x in root.xpath('//span[contains(@class,"flabel")]')]
    )
    actions = unique_texts([element_text(x) for x in root.xpath("//button")])
    tables: list[dict[str, Any]] = []
    for table in root.xpath("//table"):
        headers = unique_texts([element_text(x) for x in table.xpath(".//thead/tr[1]/th")])
        if headers:
            tables.append({"id": table.get("id", ""), "headers": headers})
    return {"file": name, "title": title, "labels": labels, "actions": actions, "tables": tables}


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    widths: list[float] | None = None
    path: Path | None = None
    caption: str = ""


class Model:
    def __init__(self):
        self.blocks: list[Block] = []
        self.figure_counter = 0

    def heading(self, text: str, level: int = 1):
        self.blocks.append(Block("heading", clean(text), level=level))

    def para(self, text: str):
        self.blocks.append(Block("paragraph", clean(text)))

    def bullet(self, text: str, level: int = 0):
        self.blocks.append(Block("bullet", clean(text), level=level))

    def number(self, text: str, level: int = 0):
        self.blocks.append(Block("number", clean(text), level=level))

    def table(self, headers: list[str], rows: list[list[Any]], widths: list[float] | None = None):
        clean_rows = [[clean(cell) for cell in row] for row in rows]
        self.blocks.append(Block("table", headers=[clean(h) for h in headers], rows=clean_rows, widths=widths))

    def code(self, text: str):
        normalized = text
        for source, target in {
            "\u2010": "-",
            "\u2011": "-",
            "\u2012": "-",
            "\u2013": "-",
            "\u2014": "-",
            "\u2192": "->",
            "\u2190": "<-",
            "\u2265": ">=",
            "\u2264": "<=",
            "\u2260": "!=",
            "\u00d7": "x",
            "\u2026": "...",
        }.items():
            normalized = normalized.replace(source, target)
        self.blocks.append(Block("code", normalized.strip()))

    def note(self, text: str):
        self.blocks.append(Block("note", clean(text)))

    def pagebreak(self):
        self.blocks.append(Block("pagebreak"))

    def image(self, path: Path, caption: str = ""):
        self.blocks.append(Block("image", path=path, caption=caption))

    def figure(self, path: Path, title: str):
        self.figure_counter += 1
        self.image(path, f"รูปที่ {self.figure_counter}: {clean(title)}")


SRS_ARTIFACT_LABELS = {
    "workflow.md": "Workflow design",
    "api.md": "API design",
    "database.md": "Database design",
    "plan-api.html": "API specification screen",
    "plan-database.html": "Database design screen",
    "plan-flow.html": "Integrated flow screen",
    "plan-email.html": "Email template screen",
    "job-batch.html": "Batch job console",
    "system-config.html": "Global config screen",
    "index.html": "Portal screen",
    "flow-fgi.html": "FGI/FCS flow screen",
    "k2-flow.html": "K2 flow screen",
    "fgi-database.html": "FGI/FCS database screen",
    "k2-database.html": "K2 database screen",
    "k2-create.html": "Create Document screen",
    "k2-document.html": "Document Detail screen",
    "k2-report.html": "Status Report screen",
    "k2-list-waiting.html": "Task Inbox screen",
    "k2-list-related.html": "Related Documents screen",
    "k2-list-abnormal.html": "Abnormal Data screen",
    "k2-operators.html": "Operator Master screen",
    "k2-factors.html": "External Factor Master screen",
    "k2-permissions.html": "RBAC Matrix screen",
}


def scrub_srs_text(value: Any) -> str:
    text = str(value)
    for source, target in SRS_ARTIFACT_LABELS.items():
        text = text.replace(source, target)
    text = text.replace("LLDD/BE/Jobs", "detailed batch design package")
    text = text.replace("LLDD API/FE contracts", "API and FE detailed contracts")
    text = text.replace("LLDD-FE-Document-Detail", "Document Detail role design")
    text = text.replace("LLDD", "detailed design")
    text = text.replace("Reference Email Template", "Email Template")
    text = text.replace("platform reference", "platform service")
    text = text.replace("Lookup / Reference", "Lookup")
    text = text.replace("k2-list-related", "เอกสารที่เกี่ยวข้อง")
    text = text.replace("k2-report", "รายงานสรุปสถานะ")
    text = text.replace("หน้าจอ 3.1.1", "หน้าสิทธิ์การเข้าถึงเมนู")
    text = text.replace("หน้า 3.1.8 / 3.1.9", "หน้ากำหนดผู้ปฏิบัติงานและหน้ากำหนดปัจจัยภายนอก")
    text = text.replace("หน้า 3.1.8", "หน้ากำหนดผู้ปฏิบัติงาน")
    text = text.replace("ในเอกสาร SRS (จัดการในระบบ BPM เดิม)", "ในระบบ BPM เดิม")
    text = re.sub(r"([\w./\-\u0E00-\u0E7F]+)\.md\b", lambda m: Path(m.group(1)).name.replace("-", " "), text)
    text = re.sub(r"([\w./\-\u0E00-\u0E7F]+)\.html\b", lambda m: Path(m.group(1)).name.replace("-", " ") + " screen", text)
    return text


def scrub_srs_model(model: Model) -> Model:
    for block in model.blocks:
        block.text = scrub_srs_text(block.text)
        block.headers = [scrub_srs_text(header) for header in block.headers]
        block.rows = [[scrub_srs_text(cell) for cell in row] for row in block.rows]
        block.caption = scrub_srs_text(block.caption)
    return model


def screenshot_slices(html_file: str, max_height: int = 1500) -> list[Path]:
    source = SCREENSHOT_FULL_DIR / html_file.replace(".html", ".png")
    if not source.exists():
        raise RuntimeError(f"Missing captured image: {source}")
    SCREENSHOT_SLICE_DIR.mkdir(parents=True, exist_ok=True)
    stem = source.stem
    with PILImage.open(source) as image:
        width, height = image.size
        parts = max(1, (height + max_height - 1) // max_height)
        part_height = (height + parts - 1) // parts
        paths: list[Path] = []
        for index in range(parts):
            top = index * part_height
            bottom = min(height, (index + 1) * part_height)
            output = SCREENSHOT_SLICE_DIR / f"{stem}-{index + 1:02d}.png"
            image.crop((0, top, width, bottom)).save(output, "PNG", optimize=True)
            paths.append(output)
    return paths


def add_screen_capture(model: Model, html_file: str, title: str):
    slices = screenshot_slices(html_file)
    for index, image_path in enumerate(slices, 1):
        suffix = f" - ส่วนที่ {index}/{len(slices)}" if len(slices) > 1 else ""
        model.figure(
            image_path,
            f"{title}{suffix}",
        )


def build_model() -> Model:
    data = json.loads(DATA_FILE.read_text(encoding="utf-8"))
    jobs = [target_job(job) for job in data["jobs"]]
    api_groups = data["apiGroups"]
    endpoint_total = sum(len(group["eps"]) for group in api_groups)
    api_group_total = len(api_groups)
    migration_rows = parse_plan_flow_migration()

    model = Model()
    model.heading("1. SRS Overview and Scope", 1)
    model.heading("1.1 Purpose", 2)
    model.para(
        "เอกสารนี้กำหนดความต้องการของระบบประกันรายได้ SBPGI แบบรวม ครอบคลุมกระบวนการนำเข้าข้อมูลผลกระทบและยอดขาย "
        "การคำนวณ การสร้างเอกสาร การพิจารณาอนุมัติ การรายงาน การส่ง Statement และการติดตามผลการทำงานของระบบ"
    )
    model.para(
        "ขอบเขตงานพัฒนาของ FE/BE อยู่ที่ระบบประกันรายได้ (SBP Mall) และบริการภายในที่ระบุในเอกสารนี้เท่านั้น "
        "รูป Flow ตารางข้อมูล และภาพหน้าจอที่อยู่ใน SRS ถือเป็นส่วนหนึ่งของคำอธิบายระบบ แต่ไม่เพิ่มขอบเขตนอกเหนือจาก requirement ที่ระบุ"
    )
    model.note(
        "หมายเหตุเวอร์ชัน: เอกสาร v1.0 เป็น baseline เริ่มต้นสำหรับการพัฒนาและตรวจรับระบบประกันรายได้ SBPGI"
    )
    model.heading("1.2 Requirement classification", 2)
    model.table(
        ["Tag", "ความหมาย", "การใช้งาน"],
        [
            ["REQ", "ข้อกำหนดของระบบที่ได้รับอนุมัติใน SRS ฉบับนี้", "ต้องพัฒนาและทดสอบตามข้อความที่กำหนด"],
            ["SYS", "ข้อกำหนดร่วมด้านสถาปัตยกรรม ข้อมูล ความปลอดภัย และการปฏิบัติการ", "ใช้กับองค์ประกอบที่เกี่ยวข้องทั้งหมด"],
            ["PROTO", "พฤติกรรมหรือข้อมูลตัวอย่างใน prototype", "ใช้ยืนยัน UX ไม่ใช่ข้อมูล Production"],
            ["OPEN", "ประเด็นขัดแย้งหรือยังไม่ตัดสินใจ", "ห้ามถือเป็นข้อยุติจนกว่าจะมีผู้อนุมัติ"],
        ],
        [0.12, 0.38, 0.5],
    )
    model.heading("1.3 Baseline and change control", 2)
    for text in [
        "SRS v1.0 ฉบับนี้เป็น baseline เดียวสำหรับกำหนดขอบเขต พัฒนา ทดสอบ และตรวจรับระบบ",
        "ข้อความ ตาราง รูป และ acceptance criteria ภายใน SRS มีผลร่วมกัน หากมีความขัดแย้งให้เปิดประเด็นตัดสินใจก่อนพัฒนา",
        "รายละเอียดเชิงออกแบบต้องไม่เพิ่ม ลด หรือเปลี่ยน requirement โดยไม่มีการอนุมัติ change request",
        "รายการที่ระบุ OPEN ยังไม่ถือเป็นขอบเขตที่อนุมัติจนกว่าจะมีข้อยุติและปรับ baseline",
        "ข้อมูลตัวอย่างและพฤติกรรม prototype ใช้ยืนยัน UX เท่านั้น ต้องไม่ถูกนำไปใช้เป็นข้อมูล Production",
        f"ขอบเขต API ใน SRS ประกอบด้วย {endpoint_total} endpoints / {api_group_total} กลุ่ม โดยบริการยืนยันตัวตนเป็นบริการ platform กลาง",
    ]:
        model.bullet(text)

    model.heading("1.4 How to read this document", 2)
    model.para(
        "เอกสารจัดลำดับจากภาพรวมธุรกิจไปสู่ข้อกำหนดที่ใช้พัฒนาและตรวจรับ เพื่อให้ Business, FE, BE, QA และ Operations "
        "ใช้ baseline เดียวกันได้โดยไม่ต้องตีความรายละเอียดเชิงออกแบบเป็น requirement เพิ่มเติม"
    )
    model.table(
        ["ผู้อ่าน", "หัวข้อที่ควรเริ่ม", "สิ่งที่ต้องใช้จากเอกสาร"],
        [
            ["Business / Product Owner", "1, 2, 3.1, 5 และ 6", "ยืนยันขอบเขต กฎธุรกิจ เกณฑ์ตรวจรับ และประเด็นที่ต้องตัดสินใจ"],
            ["Frontend / UX", "3.4, 3.5 และ 4", "หน้าจอ ข้อมูลที่แสดง การกระทำ ข้อความตอบกลับ สิทธิ์ และพฤติกรรม responsive"],
            ["Backend / Integration", "3.1, 3.2, 3.3, 3.5 และ 4", "workflow, data controls, batch, interface capability, audit และ reliability"],
            ["QA / UAT", "3, 4 และ 5", "เงื่อนไขก่อนทดสอบ ผลลัพธ์ที่คาดหวัง กฎยอมรับ และ traceability"],
            ["Operations", "2.4, 3.3, 4 และ 6", "schedule, monitoring, rerun/reconcile, availability และ open decision"],
        ],
        [0.2, 0.3, 0.5],
    )
    model.note(
        "หลักการตีความ: ข้อความที่ระบุว่า ‘ระบบต้อง’ หรืออยู่ภายใต้ REQ/SYS/acceptance ถือเป็นข้อกำหนดที่ต้องทดสอบได้ "
        "ส่วน OPEN ต้องได้รับอนุมัติก่อนนำไปพัฒนา"
    )
    model.heading("1.5 Assumptions, Constraints and Sign-off", 2)
    model.table(
        ["ID", "Type", "Statement", "Validation / approval gate"],
        [
            ["ASM-001", "Assumption", "Platform SSO/AD/LDAP ยืนยัน credential และส่ง employee identity ที่เชื่อถือได้ให้ SBPGI; SBPGI ไม่เก็บ password hash", "ผ่าน integration/security test กับ platform identity"],
            ["ASM-002", "Assumption", "QSSI, ALLMAP, IAS/MIS, STA, SAP และ SMTP ให้บริการตาม interface window และ data contract ที่อนุมัติ", "ผ่าน connectivity และ golden-file test ก่อน UAT"],
            ["ASM-003", "Assumption", "ข้อมูลสาขามี region/branch type/nิติบุคคล/DV ที่เพียงพอสำหรับ candidate selection และ Gen Flow Gate", "รายงาน reject/missing master ต้องเป็นศูนย์หรือได้รับ waiver"],
            ["CON-001", "Constraint", "Store code เป็นข้อความ 5 หลักและต้องรักษาเลขศูนย์นำหน้าใน DB, API, file และ UI", "contract/golden-file test"],
            ["CON-002", "Constraint", "ระบบใช้ workflow 5 ขั้น 06/08/01/02/03 และสถานะเอกสาร 6 ค่า 06/08/01/02/03/99", "lookup/transition test"],
            ["CON-003", "Constraint", "Secret, password, private key, token และ connection credential ต้องอยู่นอก source/config ธรรมดาและส่งผ่าน TLS", "secret scan และ deployment evidence"],
            ["CON-004", "Constraint", "ข้อความ error ภาษาไทยและผลพิจารณาที่กำหนดเป็น verbatim ต้องไม่ถูกเปลี่ยนโดย FE", "contract/UI test"],
        ],
        [0.12, 0.14, 0.48, 0.26],
    )
    model.table(
        ["Sign-off role", "Approval scope", "Required before"],
        [
            ["Business Owner / Product Owner", "ขอบเขต กฎรัศมี กฎยอดขาย/เงินชดเชย และ OPEN decisions", "Development baseline / UAT"],
            ["Solution / Data Architect", "API, data ownership, migration, transaction และ integration", "Schema/API freeze"],
            ["Security", "Identity, RBAC, secret management, TLS, attachment และ audit", "Production readiness"],
            ["QA / UAT", "Requirement coverage, acceptance evidence และ regression", "Release approval"],
            ["Operations", "Batch schedule, monitoring, rerun/reconcile, backup/restore และ runbook", "Go-live"],
        ],
        [0.24, 0.5, 0.26],
    )

    model.pagebreak()
    model.heading("2. System Overview", 1)
    model.heading("2.1 Product perspective", 2)
    model.para(
        "ระบบประกันรายได้ใช้บริหารการชดเชยรายได้ของร้าน Store Partner ที่ได้รับผลกระทบจากร้านเปิดใหม่ "
        "โดยรับข้อมูลผลกระทบ ยอดขาย และคะแนน QSSI ประมวลผลเงื่อนไข สร้างเอกสาร "
        "เดิน workflow อนุมัติ และส่งผลชดเชยไปยังระบบบัญชี/Statement"
    )
    model.heading("2.2 Target architecture", 2)
    model.table(
        ["Layer", "องค์ประกอบ", "หน้าที่"],
        [
            ["Frontend", "Web SPA จากต้นแบบหน้าจอ", "Dashboard, K2 forms, report, batch monitor และ administration"],
            ["Backend", "RBAC, Document, Workflow, Batch Scheduler, Interface, Report/Notification", "ให้บริการ REST API /api/v1 และ orchestration ภายใน; Auth token/menu มาจาก platform กลาง"],
            ["Database", "Schema รวม Zone A/B/C", "เก็บ pipeline, เอกสาร/workflow, master/config และ audit"],
            ["External", "QSSI, ALLMAP, IAS/MIS, STA, SAP, SMTP", "คง file/SFTP/API ตามขอบเขตระบบภายนอก"],
        ],
        [0.15, 0.35, 0.5],
    )
    model.note(
        "SYS: ระบบต้องรวมการสร้างเอกสารและ workflow ไว้ภายใน SBPGI โดยใช้ DB transaction และ Workflow Engine ภายใน "
        "ห้ามสร้างไฟล์ BPM06001O/BPM06002O/BPM06003O หรือเรียก K2 StartInstance ใน runtime ใหม่"
    )
    model.heading("2.3 User roles", 2)
    model.table(
        ["Role code", "Role", "ขอบเขต"],
        [
            ["00", "Default", "ผู้ดำเนินการในแบบฟอร์ม"],
            ["01", "Admin", "เห็นทุกเมนูและจัดการข้อมูลทั้งหมด"],
            ["02", "HQ", "HQ Support และงานบริหารข้อมูล"],
            ["03", "User Admin", "ผู้ดูแลระบบระดับผู้ใช้งาน"],
            ["04", "Report Admin", "รายงานและรายงานสรุป"],
            ["05", "Assign Job", "แจกงานข้อมูลผิดปกติ"],
            ["06", "Report Admin Special", "เรียกดูเอกสารทั้งหมด"],
            ["10", "UserViewer", "อ่านเอกสารตามรายการที่ได้รับสิทธิ์"],
        ],
        [0.12, 0.28, 0.6],
    )
    model.note(
        "Role code ในหัวข้อนี้เป็นรหัสกลุ่มสิทธิ์การใช้งานของ RBAC เท่านั้น; ห้ามนำไปตีความเป็นรหัสหน่วยงาน/ขั้นตอนการพิจารณา "
        "หน้า Document Detail ต้องประเมินสิทธิ์การมองเห็น แก้ไข และดำเนินการจาก role, section และ task owner ปัจจุบัน"
    )
    model.heading("2.4 External interfaces", 2)
    model.table(
        ["System", "Direction", "Mechanism", "Requirement"],
        [
            ["QSSI", "Inbound", "SFTP, mrs* 4 files", "WINDOWS-874; คะแนน 6 หมวด 8,9,12,1,10,16"],
            ["ALLMAP", "Inbound", "SQL Server views / link", "คู่ร้านถูกกระทบ ร้านคู่แข่ง และ POI map"],
            ["IAS/MIS", "Outbound/Inbound", "AMS06001O / AMS06001I", "ยอดขาย 4 windows x 15 days"],
            ["STA", "Outbound/Inbound", "FRBC0001 + ACK/API callback", "ส่งผลชดเชยและเฝ้าระวัง ACK"],
            ["SAP", "Downstream via STA", "Accounting posting", "รับรายการเมื่อ STA approve"],
            ["SMTP", "Outbound", "E-mail", "แจ้งผู้ดำเนินการ เตือนงานค้าง และ batch errors"],
        ],
        [0.14, 0.16, 0.28, 0.42],
    )
    model.heading("2.5 Business outcomes and scope boundary", 2)
    model.para(
        "ผลลัพธ์ปลายทางของระบบคือการเปลี่ยนข้อมูลผลกระทบและยอดขายให้เป็นเอกสารชดเชยที่อนุมัติ ตรวจสอบย้อนหลัง "
        "และส่งต่อบัญชีได้ครบถ้วน โดยไม่เพิ่มหน้าจอหรือระบบย่อยนอกขอบเขตที่ระบุใน SRS"
    )
    model.table(
        ["ประเภท", "อยู่ในขอบเขต", "อยู่นอกขอบเขต"],
        [
            ["Business process", "นำเข้าข้อมูล คำนวณ สร้างเอกสาร พิจารณา อนุมัติ รายงาน และส่ง Statement", "การเปลี่ยนกฎของ QSSI, ALLMAP, IAS/MIS, STA หรือ SAP ภายนอกระบบ"],
            ["Application", "หน้าจอ SBP Mall, API, Document/Workflow Service, Batch Scheduler, Notification และ audit", "การพัฒนาระบบ workflow/integration เดิม, Login/SSO platform กลาง และเครื่องมือออกแบบระบบ"],
            ["Data", "ข้อมูลประมวลผล เอกสาร workflow master/config interface tracking และไฟล์แนบ", "การเปลี่ยน ownership หรือโครงสร้างข้อมูลต้นทางของระบบภายนอก"],
            ["Delivery evidence", "ผลทดสอบตาม acceptance, interface golden file, audit trail และ run/reconcile evidence", "prototype data และภาพหน้าจอเป็นข้อมูล production"],
        ],
        [0.16, 0.44, 0.4],
    )

    model.pagebreak()
    model.heading("3. Specific Requirements", 1)
    model.para(
        "หัวข้อนี้เป็นข้อกำหนดที่ใช้ส่งต่อให้ทีมพัฒนาและ QA โดยเรียงตามลำดับการทำงานจริง: flow, data, batch, screen และ API "
        "ทุกส่วนต้องอ่านร่วมกับ Non-Functional Requirements และ Acceptance/Traceability ไม่ควรตรวจรับจากภาพหน้าจอเพียงอย่างเดียว"
    )
    model.heading("3.0 Atomic Requirement Register", 2)
    requirement_rows = [
        ["REQ-BUS-001", "ระบบต้องคัดร้านเปิดใหม่ในกรุงเทพฯ/ปริมณฑลที่อยู่ห่างร้านถูกกระทบไม่เกิน 1 กิโลเมตร", "candidate selection boundary test ที่ 0.999/1.000/1.001 กม."],
        ["REQ-BUS-002", "ระบบต้องคัดร้านเปิดใหม่ในต่างจังหวัดที่อยู่ห่างร้านถูกกระทบไม่เกิน 2 กิโลเมตร", "candidate selection boundary test ที่ 1.999/2.000/2.001 กม."],
        ["REQ-BUS-003", "ระบบต้องเปิด workflow เฉพาะรายการที่ Gen Flow Gate ทุกเงื่อนไขผ่าน", "Job 8b/API gate test ครบ Y/W/N"],
        ["REQ-BUS-004", "ระบบต้อง flag รายการที่ยอดขายมีวันทำการน้อยกว่า 60 วันและแสดงเป็นแถวผิดปกติ", "list/report test ที่ 59/60 วัน"],
        ["REQ-BUS-005", "ระบบต้องปฏิเสธการบันทึกเมื่อผลรวมเปอร์เซ็นต์ชดเชยของร้านเปิดใหม่ไม่เท่ากับ 100%", "validation test ต่ำกว่า/เท่ากับ/มากกว่า 100"],
        ["REQ-BUS-006", "ยอดชดเชยไม่เกิน 100,000 บาทต้องสิ้นสุดที่ Section 02; ยอดเกิน 100,000 บาทต้องผ่าน Section 03 ก่อนสิ้นสุด", "routing boundary test 99,999.99/100,000/100,000.01"],
        ["REQ-DOC-001", "ระบบต้องสร้างเลขเอกสารรูป YYYY/xxxxx โดยใช้ปี พ.ศ. และ running แยกต่อปี", "uniqueness/format/concurrency test"],
        ["REQ-DOC-002", "ระบบต้องป้องกันเอกสารซ้ำต่อ business key และ impact process", "duplicate/idempotency test"],
        ["REQ-DOC-003", "ระบบต้องเก็บความสัมพันธ์ impact_process_id -> doc_no -> instance_id -> task_id ให้ trace ได้", "referential-integrity trace"],
        ["REQ-WFL-001", "ระบบต้องอนุญาต action เฉพาะ current task owner ที่ผ่าน RBAC และ record access", "authorization test 401/403/409"],
        ["REQ-WFL-002", "ระบบต้องบันทึกผลพิจารณา เหตุผล ผู้กระทำ เวลา สถานะก่อน/หลัง และ correlation id ของทุก transition", "audit trace sample"],
        ["REQ-WFL-003", "ระบบต้องใช้ optimistic concurrency และคืน STALE_VERSION เมื่อ version เอกสารถูกเปลี่ยนแล้ว", "parallel update test"],
        ["REQ-INT-001", "Job 4 ต้องสร้าง durable file สำเร็จก่อน commit W เป็น P และ outbox READY", "failure injection ก่อน/หลัง fsync"],
        ["REQ-INT-002", "Interface callback ต้องอัปเดต tracking เดิมแบบ compare-and-set และงาน purge ต้องลบเฉพาะ terminal/expired/non-held", "ACK race และ retention test"],
        ["REQ-INT-003", "ระบบต้องใช้ typed FK สำหรับ interface transaction และรักษา business key/idempotency key", "schema constraint/rerun test"],
        ["REQ-SEC-001", "ระบบต้องไม่เก็บ password hash หรือ credential ของ platform identity ภายใน user account ของ SBPGI", "schema/secret scan"],
        ["REQ-SEC-002", "การเชื่อมต่อภายนอกต้องอ่าน secret จาก Secret Manager และบังคับ TLS/host verification", "deployment/security evidence"],
        ["REQ-FIL-001", "ไฟล์แนบต้องไม่เกิน 5 MB ผ่าน type/AV scan และดาวน์โหลดได้เฉพาะผู้มีสิทธิ์เมื่อสถานะ CLEAN", "upload/download security test"],
        ["REQ-RPT-001", "รายงานหน้าจอและ CSV ต้องใช้ filter/dataset เดียวกันและมีข้อมูลครบ 19 คอลัมน์", "preview/export reconciliation"],
        ["REQ-OPS-001", "Jobs 1-10 และ 8b ต้องรองรับ rerun โดยไม่สร้างข้อมูลซ้ำและต้องรายงาน input/success/reject/skipped", "rerun/reconcile evidence"],
        ["REQ-SCR-001", "ระบบต้องมีหน้าจอ committed SCR-01 ถึง SCR-04 และ SCR-06 ถึง SCR-11 ตาม requirement รายหน้าจอ", "screen/UAT traceability"],
        ["SYS-API-001", f"ระบบต้องมี API capability {endpoint_total} endpoints ใน {api_group_total} กลุ่มตาม catalog", "OpenAPI/contract coverage"],
        ["SYS-DAT-001", "ระบบต้องมี logical data model 34 ตารางพร้อม PK/FK/constraint ที่บังคับกฎสำคัญ", "migration/schema test"],
        ["SYS-NFR-001", "ระบบต้องมี correlation log, metrics, alert และ audit ที่เชื่อม request/job/interface กับผลธุรกิจได้", "observability trace"],
    ]
    model.table(["Requirement ID", "Atomic shall statement", "Verification"], requirement_rows, [0.16, 0.58, 0.26])
    model.heading("3.1 Business Flow and System Diagrams", 2)
    model.note(
        "รูป Flow ในหัวข้อนี้เป็นส่วนหนึ่งของ SRS ใช้อธิบายลำดับการทำงานและเงื่อนไขทางธุรกิจ "
        "แต่ไม่ใช่หน้าจอผู้ใช้งานที่ต้องพัฒนา"
    )
    add_screen_capture(model, "flow-fgi.html", "Flow FGI/FCS - Batch Pipeline")
    add_screen_capture(model, "k2-flow.html", "Flow การพิจารณาและอนุมัติ")
    add_screen_capture(model, "plan-flow.html", "Flow ระบบเป้าหมายแบบรวม")
    model.pagebreak()
    model.heading("3.1.1 End-to-end flow", 3)
    stages = [
        ("A1", "นำเข้าคะแนน QSSI รายเดือน", "Job 1 รับ 4 ไฟล์ผ่าน SFTP, dedup และบันทึก fcs_qssi_scores"),
        ("A2", "นำเข้าคู่ร้านและคู่แข่ง", "Jobs 2-3 อ่าน ALLMAP ทุกวันที่ 7 และตั้ง verify_status ตามกฎ DENY/ON_PROCESS"),
        ("A3", "ขอยอดขายรายวัน", "Job 4 สร้าง AMS06001O วันที่ 7-16 เวลา 16:00"),
        ("A4", "รับยอดขายและคำนวณ", "Job 5 รับ AMS06001I, คำนวณ 4x15 วัน, outlier |sales_diff| >= 50"),
        ("B1", "สร้างเอกสารอัตโนมัติ", "Document Service สร้าง compensation_documents และรายการลูกโดยตรงใน DB"),
        ("B2", "เปิด workflow", "Workflow Engine เปิด instance เมื่อผ่าน Gen Flow Gate และเริ่ม Section 06"),
        ("C1", "SBP DSA ตรวจสอบ", "Section 06 และ 08 ตรวจข้อมูลและคำนวณเงินชดเชย"),
        ("C2", "ฝ่ายส่งเสริมธุรกิจปรับข้อมูล", "Section 01 แก้ร้านเปิดใหม่ คู่แข่ง ปัจจัย และตรวจ % ชดเชยรวม 100%"),
        ("C3", "GM/AVP อนุมัติ", "Section 02; ยอด > 100,000 ผ่าน Section 03 แล้วจบ, ยอด <= 100,000 จบที่ GM"),
        ("C4", "บัญชีตรวจสอบนอก workflow", "เมื่อเอกสารเสร็จสิ้น ทีมบัญชีใช้รายงาน SBP Mall และ Export CSV to Batch เพื่อกระทบ SAP"),
        ("D1", "ส่ง Statement", "Job 6 ส่ง FRBC0001 ไป STA เวลา 17:00 ทุกวัน"),
        ("D2", "ติดตาม ACK", "STA callback อัปเดต ACK และ Job 10 เป็น safety net เมื่อค้าง >= 1 วัน"),
    ]
    model.table(["Step", "Process", "Requirement"], stages, [0.1, 0.28, 0.62])
    model.heading("3.1.2 Gen Flow Gate", 3)
    for rule in [
        "คู่ร้านต้องผ่านกฎรัศมี: กรุงเทพฯ/ปริมณฑลไม่เกิน 1 กิโลเมตร และต่างจังหวัดไม่เกิน 2 กิโลเมตร",
        "workflow_generation_status ต้องเป็น W",
        "branch_type อยู่ใน FAM, FB1, FC1, FB2, FVB, FVC",
        "opt_dv_user_id ต้องไม่ว่าง",
        "นิติบุคคลของร้านเปิดใหม่ต้องต่างจากร้านถูกกระทบ",
        "growth_rate_diff ต้องน้อยกว่าหรือเท่ากับ -10",
        "sales_status ต้องเป็น Y หรือ N",
        "กรณี branch type ไม่เข้าเกณฑ์ให้สถานะ N; กรณีอื่นที่ยังไม่พร้อมให้คง W เพื่อแก้ไขและรันซ้ำ",
    ]:
        model.bullet(rule)
    model.heading("3.1.3 Document action requirements", 3)
    model.table(
        ["Requirement", "รายละเอียด"],
        [
            ["Action ownership", "ผู้ใช้ส่งผลพิจารณาได้เฉพาะเอกสาร/งานที่ตนมีสิทธิ์ดำเนินการตาม RBAC และ task ownership"],
            ["Result options", "ระบบต้องแสดงชุดผลพิจารณาที่อนุญาตสำหรับผู้ใช้จาก API/role profile ไม่ให้ FE คำนวณสิทธิ์เอง"],
            ["Status convention", "API mutation/action ต้องคืน statusCode เป็นค่ากลาง และ FE resolve label ไทยจาก document_statuses"],
            ["Amount approval rule", "ยอดเงินชดเชยรวม 100,000 บาทเป็น threshold ทางธุรกิจสำหรับชั้นอนุมัติตามลำดับที่กำหนดใน 3.1.3"],
            ["Audit", "ทุก action ต้องบันทึกผลพิจารณา ความคิดเห็น สถานะก่อน/หลัง ผู้กระทำ เวลา และ correlation id"],
            ["Notification", "เมื่อ action สำเร็จ ระบบต้องแจ้งผู้เกี่ยวข้องตาม e-mail rule/template ที่กำหนด"],
        ],
        [0.24, 0.76],
    )
    model.note(
        "ลำดับ workflow ที่ต้องรองรับคือ Section 06 -> 08 -> 01 -> 02; ยอดรวมไม่เกิน 100,000 บาทสิ้นสุดที่ Section 02 "
        "ส่วนยอดเกิน 100,000 บาทต้องส่งต่อ Section 03 ก่อนสิ้นสุด ระบบต้องคืน action ที่อนุญาตตาม role, section และ task owner ปัจจุบัน"
    )
    model.pagebreak()
    model.heading("3.1.4 Migration map", 3)
    if migration_rows:
        model.table(["Connection", "Legacy", "Target"], [row[:3] for row in migration_rows], [0.24, 0.32, 0.44])
    else:
        model.para("ระบบต้องแสดงลำดับการย้ายจากกระบวนการเดิมสู่บริการเป้าหมายให้ตรวจสอบได้")
    model.heading("3.1.5 Flow controls", 3)
    for rule in [
        "กฎ candidate selection ต้องใช้รัศมีไม่เกิน 1 กิโลเมตรสำหรับกรุงเทพฯ/ปริมณฑล และไม่เกิน 2 กิโลเมตรสำหรับต่างจังหวัด โดยรวมค่าขอบเขตเท่ากับเกณฑ์",
        "รายการที่ข้อมูลยอดขายไม่ครบ 60 วันต้องแสดงเป็นข้อมูลผิดปกติและแถวสีแดง",
        "ระบบต้องกันเปิดงาน/เอกสารซ้ำต่อ impact process/document",
        "บัญชีตรวจสอบยอดผ่านรายงาน SBP Mall และ Export CSV to Batch นอก workflow",
        "งานเตือนรายสัปดาห์ทำงานวันจันทร์ 10:00 และ escalation งานค้าง 30/45/60 วันต้องอ่านค่าจาก config",
        "การเปลี่ยนกฎธุรกิจ เช่น -10, 50, 60 วัน และ 100,000 บาท ต้องผ่าน Business sign-off",
        "ทุก action ต้องบันทึก consideration_logs, ผู้กระทำ, เวลา, สถานะก่อน/หลัง และ correlation id",
    ]:
        model.bullet(rule)
    if (ROOT / "Flow ประกันรายได้.png").exists():
        model.figure(ROOT / "Flow ประกันรายได้.png", "Approve Flow เดิม ใช้ประกอบการเทียบพฤติกรรม")

    model.pagebreak()
    model.heading("3.2 Data Requirements and Logical Data Model", 2)
    model.note(
        "หัวข้อนี้กำหนดข้อมูลที่ระบบต้องเก็บ ตรวจสอบ และเชื่อมโยงเพื่อรองรับธุรกรรมและการตรวจสอบย้อนหลัง "
        "ชื่อทางกายภาพของตาราง/คอลัมน์สามารถกำหนดในขั้นตอนออกแบบได้ แต่ต้องรักษาความสัมพันธ์และข้อควบคุมใน SRS"
    )
    model.heading("3.2.1 Data subjects", 3)
    model.table(
        ["Data subject", "Requirement"],
        [
            ["Impact processing", "ระบบต้องเก็บคู่ร้านถูกกระทบ/ร้านเปิดใหม่ งวดผลกระทบ ผล QSSI ยอดขาย และสถานะการสร้าง workflow ให้ตรวจสอบย้อนกลับได้"],
            ["Compensation document", "ระบบต้องเก็บหัวเอกสาร เลขเอกสาร ร้านเปิดใหม่ คู่แข่ง ปัจจัยภายนอก เงินชดเชย ไฟล์แนบ และประวัติการพิจารณา"],
            ["Workflow", "ระบบต้องเก็บ instance/task/current section/status/assignee เพื่อควบคุมงานค้างและ audit ทุก transition"],
            ["Master/config", "ระบบต้องเก็บ role/menu/permission, ผู้ปฏิบัติงาน, external factors, email rules/templates และ system config ที่ใช้ร่วมกัน"],
            ["Interface tracking", "ระบบต้องเก็บสถานะไฟล์/callback/batch run เพื่อ reconcile งานภายนอกและ rerun ได้โดยไม่สร้างข้อมูลซ้ำ"],
        ],
        [0.26, 0.74],
    )
    model.heading("3.2.2 Data controls", 3)
    for rule in [
        "Store code ต้องเก็บเป็น varchar(5) เพื่อรักษา leading zero",
        "doc_no ต้อง unique และรูปแบบ YYYY/xxxxx; running แยกต่อปี",
        "ข้อมูลหนึ่งเอกสารต้อง trace ได้ครบจาก impact process ไปยัง document, workflow instance และ task ปัจจุบัน",
        "% ชดเชยของร้านเปิดใหม่ต่อเอกสารต้องรวมเท่ากับ 100%",
        "สถานะเอกสาร, section, role และ workflow task ต้องอ้าง lookup กลางเพื่อไม่ให้ label/code ปนกัน",
        "ค่าธุรกิจที่ถูก lock ต้องแก้ผ่าน UI/API ไม่ได้หากไม่มี Business sign-off",
        "ระบบต้องรองรับ concurrency control เมื่อมีการแก้เอกสาร/workflow พร้อมกัน",
        "ทุก master mutation ต้องบันทึก audit_logs ค่าเดิม ค่าใหม่ เหตุผล ผู้แก้ และเวลา",
        "Timestamp ภายใน DB ใช้ UTC; UI แสดง Asia/Bangkok และปี พ.ศ. ตามข้อยุติด้าน format",
    ]:
        model.bullet(rule)
    model.heading("3.2.3 Logical data relationships", 3)
    model.table(
        ["Data area", "Key relationship", "Requirement"],
        [
            ["Impact processing", "impact_process_id เชื่อมร้านถูกกระทบ ร้านเปิดใหม่ คะแนน และยอดขาย", "หนึ่งรอบประมวลผลต้องตรวจสอบข้อมูลนำเข้า สถานะ และผลการคำนวณย้อนหลังได้"],
            ["Compensation document", "doc_no เชื่อมหัวเอกสาร ร้านเปิดใหม่ คู่แข่ง ปัจจัย ไฟล์แนบ และยอดชดเชย", "doc_no ต้อง unique และข้อมูลลูกทุกประเภทต้องไม่หลุดจากหัวเอกสาร"],
            ["Workflow", "instance_id และ task_id เชื่อมเอกสาร ขั้นตอน ผู้รับผิดชอบ และประวัติ action", "ต้องทราบ current task และทุก transition ของเอกสารได้ตลอดเวลา"],
            ["Master/config", "role, menu, section, operator, factor, template และ config key", "ค่ากลางต้องมี version/status และ audit เมื่อเปลี่ยนแปลง"],
            ["Interface tracking", "run_id, transaction id และ correlation id", "ต้องเชื่อมไฟล์ callback batch run และผลลัพธ์ธุรกิจเพื่อ reconcile/rerun ได้"],
        ],
        [0.2, 0.34, 0.46],
    )
    model.heading("3.2.4 Required remediation", 3)
    remediation = [
        ["P0", "Job 4 transaction", "ใช้ transaction/outbox ไม่ให้ W->P commit ก่อนสร้างไฟล์สำเร็จ"],
        ["P0", "Secrets/TLS", "ย้าย credential ไป Secret Manager และบังคับ TLS"],
        ["P0", "Tracking purge", "แก้ SQL purge data_name และทำ migration/test"],
        ["P1", "Polymorphic FK", "ใช้ typed FK ใน interface_transactions"],
        ["P1", "NULL growth rate", "ส่งรอตรวจสอบแทน auto-accept; ต้องมี Business sign-off"],
        ["P1", "Master joins", "รายงาน reject/reconcile แทนการทำแถวหายเงียบ"],
        ["P1", "Golden files", "ทดสอบ encoding วันที่ พ.ศ. delimiter และ field count ทุก interface"],
    ]
    model.table(["Priority", "Issue", "Target requirement"], remediation, [0.12, 0.28, 0.6])

    model.heading("3.3 Batch Job Requirements", 2)
    model.note(
        "SRS ส่วนนี้อธิบายงาน Batch Job ในระดับที่ผู้ใช้ธุรกิจและผู้ดูแลระบบต้องเข้าใจ: "
        "แต่ละ job ทำเพื่ออะไร รับข้อมูลหรือเงื่อนไขอะไร ระบบทำอะไรโดยสรุป และผลลัพธ์ที่ต้องเห็นคืออะไร "
        "ไม่ลงรายละเอียด coding, SQL, class/script หรือ transaction ภายใน"
    )
    model.heading("3.3.1 Batch console", 3)
    add_screen_capture(model, "job-batch.html", "Batch Job Console - 11 jobs")
    model.para(
        "หน้า Batch Job Console สำหรับ Admin แสดง pipeline A-E, รายการ 11 entry points, "
        "สถานะรอบล่าสุด/ถัดไป, เปิดปิดงาน, พารามิเตอร์ที่อนุญาตให้แก้, manual run, ลำดับงาน และ run history"
    )
    model.table(
        ["Job", "Name", "Thai name", "Phase", "Schedule", "Output"],
        [[j["no"], j["name"], j["th"], j["phase"], f"{j['cron']} ({j['cronTh']})", j.get("out", "")] for j in jobs],
        [0.07, 0.18, 0.24, 0.07, 0.22, 0.22],
    )
    model.heading("3.3.2 Common controls", 3)
    for rule in [
        "สิทธิ์จัดการ Batch Job เป็น Admin 01 เท่านั้น",
        "Manual run ต้องระบุงวดข้อมูลและสร้าง run_id; API ตอบ 202 Accepted",
        "ห้ามรัน job เดียวกันซ้อน และต้องป้องกัน shared temp resource ของ Job 1",
        "แก้ไขได้เฉพาะ parameter ที่ระบุ editable; business constants ต้องถูก lock",
        "ทุกการเปิด/ปิด แก้ parameter และ manual run ต้องบันทึก audit",
        "run history ต้องเก็บ start/end, status, row count, file, error, correlation id และผู้สั่งรัน",
        "การ re-run ต้องปฏิบัติตาม runbook ของแต่ละ job โดยตรวจ DB, tracking, backup และปลายทางก่อน",
    ]:
        model.bullet(rule)
    model.heading("3.3.3 Job business requirement catalog", 3)
    for idx, job in enumerate(jobs, start=1):
        catalog = SRS_JOB_USER_CATALOG.get(str(job["no"]), {})
        model.heading(f"3.3.3.{idx} Job {job['no']} - {job['th']}", 4)
        model.table(
            ["หัวข้อ", "รายละเอียด"],
            [
                ["เป้าหมาย", catalog.get("purpose", job.get("desc", ""))],
                ["รับข้อมูล/เงื่อนไข", catalog.get("input", "งวดข้อมูลและพารามิเตอร์ของงาน")],
                ["ระบบทำอะไรโดยสรุป", catalog.get("summary", job.get("desc", ""))],
                ["ผลลัพธ์ที่ต้องได้", catalog.get("output", job.get("out", ""))],
                ["ผู้ใช้ติดตามได้จาก", catalog.get("visible", "ติดตามได้จาก Batch Job Console และ run history")],
            ],
            [0.22, 0.78],
        )
    model.heading("3.3.4 Required job outcomes", 3)
    for rule in [
        "ทุก job ต้องแสดงสถานะล่าสุดและประวัติการรันให้ Admin ตรวจสอบได้",
        "ผลลัพธ์ของ job ต้องตรวจนับได้ เช่น จำนวนไฟล์ จำนวนรายการที่อ่าน สำเร็จ ข้าม รอข้อมูล หรือผิดพลาด",
        "เมื่อ job ล้มเหลว ต้องมีข้อความสาเหตุที่ผู้ดูแลระบบใช้ติดตามกับทีมที่เกี่ยวข้องได้",
        "เมื่อไม่มีข้อมูลให้ประมวลผล ระบบต้องบันทึกเป็น no data หรือ skipped อย่างชัดเจน ไม่ถือว่าเป็น error โดยอัตโนมัติ",
        "job ที่ส่งหรือรับข้อมูลจากระบบภายนอกต้องมีสถานะติดตามปลายทาง เช่น รอ ACK, ได้รับ ACK, หรือค้างเกินกำหนด",
        "การรันซ้ำต้องไม่ทำให้เอกสาร รายการร้าน คู่แข่ง ยอดขาย หรือข้อมูล Statement ซ้ำ",
    ]:
        model.bullet(rule)

    model.pagebreak()
    model.heading("3.4 K2 Screen Requirements", 2)
    model.note(
        "Committed implementation scope ของหน้าจอ SBP Mall คือ 10 หน้า + 1 placeholder/deferred: "
        "SCR-05 ข้อมูลผิดปกติ / แจกงานยังเป็น OPEN item ใช้อธิบายกฎยอดขายไม่ครบ 60 วันเท่านั้น "
        "และไม่ถูกนับเป็นงานสร้างหน้า FE/BE จนกว่าจะมีคำตัดสิน keep/drop"
    )
    screens = [
        {
            "id": "SCR-01",
            "file": "index.html",
            "name": "Overview / Dashboard",
            "purpose": "แสดงงานค้าง ร้านที่เข้าเกณฑ์ ยอดชดเชย ข้อมูลผิดปกติ กราฟ และทางลัดตามสิทธิ์",
            "actors": "ทุก role ที่ login",
            "rules": [
                "ตัวเลขต้อง aggregate จากข้อมูลจริงและรองรับ cache ไม่เกิน 5 นาที",
                "ทางลัดและ sidebar ต้องสร้างตาม menu_permissions",
                "ค่าชื่อผู้ใช้/role ต้องมาจาก JWT ไม่ใช้ข้อมูล mock",
            ],
        },
        {
            "id": "SCR-02",
            "file": "k2-create.html",
            "name": "สร้างเอกสาร",
            "purpose": "สร้างเอกสารนอกเงื่อนไขอัตโนมัติ หรือส่งสร้างผ่าน FS",
            "actors": "HQ 02, User Admin 03 และผู้ที่ได้รับสิทธิ์",
            "rules": [
                "Manual tab ต้องระบุรหัสร้าน เดือน/ปี ร้านเปิดใหม่ และเหตุผล",
                "FS tab ต้องระบุรหัสร้าน เดือน/ปี และ Period Statement",
                "ตรวจ duplicate ร้าน+งวดก่อนสร้าง",
                "ออกเลขเอกสารอัตโนมัติและเปิด workflow Section 06",
            ],
        },
        {
            "id": "SCR-03",
            "file": "k2-list-waiting.html",
            "name": "เอกสารรอดำเนินการ",
            "purpose": "Task inbox แสดงเฉพาะ OPEN task ที่ผู้ใช้/section ปัจจุบันต้องดำเนินการ",
            "actors": "ผู้ดำเนินการ workflow",
            "rules": [
                "filter ด้วยข้อความ สถานะ ภาค ประเภทร้าน วันที่ ยอดขายลด เงินชดเชย และวันค้าง",
                "คลิกแถวเปิดเอกสาร; งานข้อมูลยอดขายไม่ครบ 60 วันเป็นแถวแดง",
                "Role switcher เป็น prototype aid เท่านั้น Production ใช้ JWT/assignment จริง",
            ],
        },
        {
            "id": "SCR-04",
            "file": "k2-list-related.html",
            "name": "เอกสารที่เกี่ยวข้อง",
            "purpose": "แสดงเอกสารทั้งหมดที่ผู้ใช้เคยมีส่วนร่วม โดยแก้ไขได้เฉพาะงานที่อยู่ในสิทธิ์ปัจจุบัน",
            "actors": "ผู้ใช้งานทั่วไปตามสิทธิ์",
            "rules": [
                "filter และ columns เหมือนหน้ารอดำเนินการ",
                "เอกสารนอก task ปัจจุบันต้องเป็น read-only",
                "ผลการค้นหาต้องจำกัดตาม role และ record-level access",
            ],
        },
        {
            "id": "SCR-05",
            "file": "k2-list-abnormal.html",
            "name": "ข้อมูลผิดปกติ / แจกงาน (placeholder/deferred)",
            "purpose": "Placeholder สำหรับค้นหาและมอบหมายรายการผิดปกติ โดยใช้กฎยอดขายไม่ครบ 60 วัน",
            "actors": "Assign Job 05 และ Admin",
            "rules": [
                "OPEN: ไม่เป็นหน้าจอ committed ใน scope FE/BE รอบนี้",
                "ถ้าเปิด scope ในอนาคต ต้องรองรับ multi-select และแจกงานเฉพาะรายการที่เลือก",
                "ถ้าเปิด scope ในอนาคต ต้องแสดงสาเหตุ ผู้รับผิดชอบ และสถานะ assignment",
                "OPEN: เมนูนี้และ API 2 เส้นถูก comment ไว้ รอคำตัดสิน keep/drop",
            ],
        },
        {
            "id": "SCR-06",
            "file": "k2-document.html",
            "name": "เอกสารข้อมูลร้านถูกกระทบ",
            "purpose": "หน้าหลักสำหรับดู แก้ คำนวณ พิจารณา แนบไฟล์ และเดิน workflow",
            "actors": "ผู้ดำเนินการตาม Section และผู้มีสิทธิ์อ่าน",
            "rules": [
                "แสดงหัวเอกสาร ร้านถูกกระทบ ร้านเปิดใหม่ แผนที่ คู่แข่ง ปัจจัย เอกสารแนบ ชดเชย ประวัติ และผลพิจารณา",
                "สิทธิ์แก้ไขต้องประเมินต่อ section/role; ส่วนอื่นเป็น read-only",
                "% ชดเชยร้านเปิดใหม่รวมต้องเท่ากับ 100%",
                "วันที่สิ้นสุดปัจจัยต้องไม่ก่อนวันที่เริ่มต้น",
                "ไฟล์แนบไม่เกิน 5 MB และต้องบันทึก section/uploader/time",
                "ส่งดำเนินการต้องเลือกผล; ข้อความ popup ต้องตรงตาม SRS",
            ],
        },
        {
            "id": "SCR-07",
            "file": "k2-report.html",
            "name": "รายงานสรุปสถานะ",
            "purpose": "ค้นหา แสดงกราฟ/ผล 19 คอลัมน์ และ Export CSV to Batch",
            "actors": "Admin 01, HQ 02, Report Admin 04, Report Admin Special 06",
            "rules": [
                "บังคับระบุปีและคืนเฉพาะรายการที่มีเลขเอกสาร",
                "ประเภทร้านและภาคเลือกหลายค่า; สถานะและผลพิจารณาเลือกหนึ่งค่า",
                "ผลและ CSV Export to Batch ต้องใช้ dataset/เงื่อนไขเดียวกัน",
                "บัญชีใช้รายงานนี้เพื่อตรวจยอดและกระทบ SAP นอก workflow หลังเอกสารเสร็จสิ้น",
                "แถวข้อมูลยอดขายไม่ครบ 60 วันต้องเป็นสีแดง",
            ],
        },
        {
            "id": "SCR-08",
            "file": "k2-operators.html",
            "name": "กำหนดผู้ปฏิบัติงาน",
            "purpose": "จัดการผู้ปฏิบัติงานต่อ section และ zone",
            "actors": "Admin 01, HQ 02, User Admin 03",
            "rules": [
                "ชื่อพนักงานและตำแหน่งเป็น required; เลือกพนักงานจาก popup",
                "แสดงภาคเมื่อเป็นตำแหน่งส่งเสริมธุรกิจพันธมิตรฯ",
                "เพิ่ม/แก้/ลบต้องบันทึก audit และเหตุผลเมื่อแก้ไข",
            ],
        },
        {
            "id": "SCR-09",
            "file": "k2-factors.html",
            "name": "กำหนดปัจจัยภายนอก",
            "purpose": "จัดการ external factor master และประวัติแก้ไข",
            "actors": "Admin 01, HQ 02, User Admin 03",
            "rules": [
                "factor_code และ factor_name เป็น required; factor_code ห้ามซ้ำ",
                "แก้ได้เฉพาะชื่อและรายละเอียด; ต้องระบุเหตุผล",
                "ทุก mutation ต้องบันทึก audit_logs",
            ],
        },
        {
            "id": "SCR-10",
            "file": "k2-permissions.html",
            "name": "สิทธิ์การเข้าถึงเมนู",
            "purpose": "แสดงและบริหาร RBAC 8 role ต่อ main menu และ master forms",
            "actors": "Admin และผู้ดูแลสิทธิ์ที่ได้รับมอบหมาย",
            "rules": [
                "sidebar ต้องอิงสิทธิ์จาก backend ไม่ใช่ซ่อนเฉพาะฝั่ง FE",
                "API ต้องตรวจ role/record access ซ้ำทุก request",
                "การเปลี่ยนสิทธิ์ต้อง audit และมีผลกับ token/session ตามนโยบาย",
            ],
        },
        {
            "id": "SCR-11",
            "file": "system-config.html",
            "name": "ตั้งค่าระบบ (Global Config)",
            "purpose": "จัดการค่ากำหนดกลางที่ใช้ร่วมทั้งระบบ เช่น รัศมีผลกระทบ เกณฑ์ข้อมูล วงเงินอนุมัติ timeout และ notification switch",
            "actors": "Admin และผู้ดูแลระบบที่ได้รับมอบหมาย",
            "rules": [
                "config_key ต้องเป็น dot notation และห้ามซ้ำ",
                "value_type ต้อง validate ค่า NUMBER, STRING, BOOLEAN, JSON หรือ CRON ก่อนบันทึก",
                "ค่าที่ is_editable=false เป็นค่าคงที่ทางธุรกิจ แก้หรือลบผ่าน UI/API ไม่ได้",
                "ห้ามเก็บ secret เช่น password, API key หรือ connection string ใน system_configs",
                "ทุกการเพิ่ม แก้ ลบ และ invalidate cache ต้องบันทึก audit_logs พร้อมเหตุผล",
            ],
        },
    ]
    screen_outcomes = {
        "SCR-01": "ผู้ใช้เห็นสถานะและงานสำคัญตามสิทธิ์ พร้อมเปิดรายการเป้าหมายจากทางลัดได้",
        "SCR-02": "ระบบสร้างเอกสารเพียงหนึ่งรายการต่อร้าน/งวด ออกเลขเอกสาร และเปิดงานเริ่มต้นสำเร็จ",
        "SCR-03": "ผู้ใช้เปิดดำเนินการเฉพาะ task ที่ตนรับผิดชอบและเห็นรายการผิดปกติอย่างชัดเจน",
        "SCR-04": "ผู้ใช้ค้นและเปิดเอกสารที่เกี่ยวข้องได้ โดยรายการนอก task ปัจจุบันเป็น read-only",
        "SCR-05": "ยังไม่มีผลลัพธ์ที่ commit; หากอนุมัติ scope ระบบต้องบันทึกผู้รับผิดชอบและสถานะ assignment",
        "SCR-06": "ข้อมูลที่แก้ไขถูกตรวจสอบ บันทึก audit และเปลี่ยนสถานะ workflow ตาม action ที่ได้รับอนุญาต",
        "SCR-07": "ผลบนหน้าจอและไฟล์ CSV ตรงกันภายใต้ filter เดียวกันและนำไปตรวจสอบบัญชีได้",
        "SCR-08": "assignment ต่อ section/zone มีผลกับการแจก task และตรวจสอบประวัติการเปลี่ยนแปลงได้",
        "SCR-09": "external factor master พร้อมใช้งานในเอกสารและตรวจสอบผู้แก้/เหตุผลย้อนหลังได้",
        "SCR-10": "เมนูและ API บังคับสิทธิ์จาก policy เดียวกันและการเปลี่ยนสิทธิ์มี audit",
        "SCR-11": "ค่ากำหนดที่ผ่าน validation ถูกเผยแพร่ให้บริการที่เกี่ยวข้องโดยไม่เปิดเผย secret",
    }
    for screen in screens:
        inv = page_inventory(screen["file"])
        model.heading(f"{screen['id']} {screen['name']}", 3)
        add_screen_capture(model, screen["file"], screen["name"])
        model.table(
            ["Item", "Requirement"],
            [
                ["Purpose", screen["purpose"]],
                ["Actor", screen["actors"]],
                ["Pre-condition", "ผ่านการยืนยันตัวตนจาก platform กลาง และมีสิทธิ์เมนู/ข้อมูล"],
                ["Post-condition / expected outcome", screen_outcomes[screen["id"]]],
                ["Scope status", "Deferred / OPEN" if screen["id"] == "SCR-05" else "Committed"],
            ],
            [0.2, 0.8],
        )
        if inv["labels"]:
            model.heading("Input / filter fields", 4)
            model.para(" · ".join(inv["labels"]))
        if inv["tables"]:
            model.heading("Displayed tables", 4)
            for table_info in inv["tables"]:
                label = table_info["id"] or "Table"
                model.bullet(f"{label}: " + " | ".join(table_info["headers"]))
        if inv["actions"]:
            model.heading("Actions", 4)
            model.para(" · ".join(inv["actions"]))
        model.heading("Business rules / acceptance", 4)
        for rule in screen["rules"]:
            model.bullet(rule)
    model.heading("3.4.13 Notification template requirements", 3)
    model.note(
        "ระบบต้องรองรับการจัดการเนื้อหาและกฎการส่ง notification ตามรายการในหัวข้อนี้ "
        "โดยหน้าจอจัดการ template เป็นส่วนหนึ่งของขอบเขตผู้ดูแลระบบ"
    )
    add_screen_capture(model, "plan-email.html", "Email Template Administration")
    model.table(
        ["Item", "Requirement"],
        [
            ["Purpose", "จัดการเนื้อหาอีเมล 8 template ของ Notification Service และจุดส่ง workflow/batch"],
            ["Scope status", "Committed - ครอบคลุมหน้าจอผู้ดูแล template และพฤติกรรม Notification Service"],
        ],
        [0.2, 0.8],
    )
    for rule in [
        "รองรับ template EM-01 ถึง EM-08 ครอบคลุม workflow transition, reminder, escalation, batch error และ STA ACK watchdog",
        "แก้ไขได้เฉพาะ subject/body และตัวแปร merge ที่รองรับของ template นั้น",
        "From/To/Cc ต้องล็อกตาม status_email_rules หรือ config ต่อ job ไม่ให้แก้ใน template",
        "ต้องรีเซ็ตกลับ Default ได้ทั้งราย template และทั้งหมด",
        "ทุกการแก้ไขหรือรีเซ็ตต้องบันทึก audit_logs พร้อมเหตุผล",
    ]:
        model.bullet(rule)
    model.heading("3.4.14 Shared UI contract", 3)
    for rule in [
        "ทุกหน้าจอต้องมี metadata สำหรับ page, nav, module, breadcrumb, sidebar mount และ main content",
        "Header/sidebar ถูกสร้างโดย shared shell; ห้ามทำซ้ำในแต่ละหน้า",
        "Schema modal อ้างชื่อ table header แบบ exact match; การเปลี่ยน label ต้องแก้ mapping และทดสอบ add/view/edit/delete",
        "รองรับ desktop และ responsive layout; ตารางกว้างต้องเลื่อนแนวนอนโดยไม่ตัดข้อมูล",
        "ข้อความ popup/validation ภาษาไทยและ source tag (FGI/FCS), (K2), (ใหม่) ต้องคงตามข้อกำหนด",
    ]:
        model.bullet(rule)

    model.heading("3.5 API Requirements", 2)
    model.note(
        "หัวข้อนี้กำหนด capability ของ API วิธีเรียกใช้ สิทธิ์ และพฤติกรรมร่วมที่ต้องตรวจรับ "
        "บริการ Auth Group 1 จัดหาโดย platform กลางและไม่อยู่ในขอบเขตการพัฒนา Login/SSO ของ SBP Mall"
    )
    model.heading("3.5.1 Interface requirements", 3)
    model.table(
        ["Topic", "Requirement"],
        [
            ["User identity and access", "ระบบต้องตรวจสิทธิ์ผู้ใช้ทุกหน้าจอและทุกการเปลี่ยนข้อมูลตาม role/menu/current task owner"],
            ["Consistent user feedback", "ข้อความ error, popup และ validation ที่มีใน SRS ต้องแสดงตรงตัวและไม่ตีความใหม่ในแต่ละหน้าจอ"],
            ["Document action", "ระบบต้องรับผลพิจารณาจากผู้ถือสิทธิ์ปัจจุบัน ตรวจ result ที่อนุญาต และคืน statusCode ตาม convention กลาง"],
            ["Search and report lists", "รายการค้นหาและรายงานต้องรองรับข้อมูลจำนวนมากโดยแบ่งหน้า/จำกัดผลลัพธ์ตามสิทธิ์"],
            ["Lookup data", "สถานะเอกสาร, workflow section, role/menu และ master data ต้องมีแหล่งข้อมูลกลางเพื่อให้ FE/BE ใช้ค่าเดียวกัน"],
            ["Audit", "ทุกการเปลี่ยนข้อมูลต้องบันทึกผู้กระทำ เวลา เหตุผล/ผลพิจารณา และค่าก่อน/หลังตามโดเมนที่เกี่ยวข้อง"],
            ["Duplicate prevention", "การสร้างเอกสาร เปิด workflow รับ callback และสั่ง batch ต้องป้องกันข้อมูลซ้ำจากการรันซ้ำหรือกดซ้ำ"],
            ["Contract consistency", "ทุก endpoint ต้องใช้รูปแบบ payload, field naming, status code, error envelope, pagination และ security mechanism ตามข้อกำหนดร่วมใน 3.5.3"],
        ],
        [0.2, 0.8],
    )
    model.heading("3.5.2 Endpoint catalog", 3)
    catalog_rows = []
    for group in api_groups:
        group_name = (
            group["name"]
            .replace(" (platform reference)", " (platform service)")
            .replace("Lookup / Reference", "Lookup")
            .replace("ข้อมูลอ้างอิง (Lookup)", "ข้อมูล Lookup")
        )
        for ep in group["eps"]:
            catalog_rows.append([group_name, ep["m"], ep["p"], ep["roles"], ep["sum"]])
    model.table(
        ["Group", "Method", "Path", "Roles", "Purpose"],
        catalog_rows,
        [0.16, 0.08, 0.22, 0.15, 0.39],
    )
    model.heading("3.5.3 API contract requirements", 3)
    for rule in [
        "Request/response JSON ใช้ camelCase และ Content-Type application/json; file download ต้องระบุ content type และ filename ที่ถูกต้อง",
        "ผลสำเร็จต้องคืน HTTP status ที่สอดคล้องกับการทำงาน เช่น 200, 201, 202 หรือ 204 และ payload ต้องมีข้อมูลที่ FE ใช้อัปเดตหน้าจอได้",
        "ข้อผิดพลาดต้องคืนโครงสร้างกลางอย่างน้อย code, message และ correlationId; validation error ต้องระบุ field ที่ไม่ผ่านเมื่อทำได้",
        "List/search/report ต้องรองรับ page, size, sort และ filter ที่ระบุ พร้อม totalElements/totalPages หรือ cursor ที่มีความหมายเทียบเท่า",
        "วันที่เวลาใน API ใช้ ISO 8601 และ UTC; UI แปลงเป็น Asia/Bangkok ส่วนรอบเดือน/ปีต้องระบุรูปแบบใน field อย่างชัดเจน",
        "Endpoint ที่สร้างเอกสาร เปิด workflow ส่ง action รับ callback หรือสั่ง batch ต้องรองรับ duplicate guard/idempotency",
        "ทุก request ต้องตรวจ token, role, menu permission, record access และ current task owner ที่ฝั่ง server ก่อนอ่านหรือเปลี่ยนข้อมูล",
        "Mutation ต้องบันทึก actor, เวลา, correlationId, เหตุผลหรือผลพิจารณา และค่าก่อน/หลังตามโดเมนที่เกี่ยวข้อง",
    ]:
        model.bullet(rule)

    model.pagebreak()
    model.heading("4. Non-Functional Requirements", 1)
    model.para(
        "ข้อกำหนดในหัวข้อนี้ใช้กับทุกหน้าจอ API batch และ interface เว้นแต่ระบุเป็นอย่างอื่น ค่าใดที่ยังไม่มีตัวเลขอนุมัติ "
        "ต้องถูกติดตามเป็น OPEN item และห้ามสมมติเป็น production SLA"
    )
    model.heading("4.1 Operational quality", 2)
    model.table(
        ["Category", "Requirement", "Verification / evidence"],
        [
            ["Performance", "รองรับผู้ใช้พร้อมกันเฉลี่ย 80 คน สูงสุด 100 คน; interaction ปกติตอบภายใน 30 วินาที; API list/report ต้องกำหนด SLA แยกก่อน production", "ผล load test ตาม workload ที่อนุมัติ พร้อม percentile, error rate และ resource usage"],
            ["Availability", "บริการ 7x24 ยกเว้น maintenance window; Batch Scheduler ต้อง resume/reconcile หลัง restart", "restart/failover test และหลักฐาน reconcile งานที่ค้าง"],
            ["Reliability", "Transaction ที่สำเร็จต้อง durable; error ต้องไม่เขียนข้อมูลบางส่วน; file interface ต้อง reconcile row/file/tracking", "failure injection, transaction rollback และ rerun/idempotency test"],
            ["Backup/Recovery", "กำหนด RPO/RTO, backup DB/config/object files และทดสอบ restore อย่างน้อยตามรอบองค์กร", "restore drill พร้อมเวลาจริงและรายการข้อมูลที่ตรวจคืน"],
            ["Observability", "Metrics/log/trace สำหรับ API, batch, workflow, interface ACK, queue lag และ e-mail failure พร้อม alert threshold", "monitoring dashboard, alert test และ correlation trace"],
        ],
        [0.16, 0.52, 0.32],
    )
    model.heading("4.2 Security and product quality", 2)
    model.table(
        ["Category", "Requirement", "Verification / evidence"],
        [
            ["Security", "SSO/AD หรือ LDAP, JWT อายุจำกัด, refresh token revoke, least privilege, secrets vault, TLS, API key rotation และ server-side RBAC", "security test, dependency/secret scan และหลักฐาน server-side authorization"],
            ["Auditability", "บันทึก login, document mutation, workflow action, master change, job action และ external callback พร้อม actor/time/correlation id", "trace sample จาก request/run ไปยัง audit log และผลลัพธ์ปลายทาง"],
            ["Usability", "รองรับ Chrome รุ่นองค์กร, ภาษาไทย, keyboard focus, responsive table/modal และข้อความ validation ตรงตาม SRS", "browser/responsive/keyboard test และ UAT ตามข้อความที่กำหนด"],
            ["Maintainability", "แยก FE/BE, OpenAPI 3.0 contract, configuration versioning, migration scripts และ automated tests สำหรับ business rules", "contract validation, migration rehearsal และ automated test report"],
            ["Portability", "Deployment ต้องไม่ผูก credential/path กับเครื่อง; ใช้ environment/config/secret manager", "deploy ด้วย environment ใหม่โดยไม่แก้ source code"],
        ],
        [0.16, 0.52, 0.32],
    )

    model.pagebreak()
    model.heading("5. Acceptance and Traceability", 1)
    model.para(
        "การตรวจรับต้องยืนยันทั้งผลลัพธ์ทางธุรกิจ สิทธิ์ ความถูกต้องของข้อมูล และหลักฐานตรวจสอบย้อนหลัง "
        "รายการต่อไปนี้เป็นเกณฑ์สำคัญขั้นต่ำและต้องเชื่อมกับ test case/UAT evidence ในรอบส่งมอบ"
    )
    model.heading("5.1 High-priority acceptance criteria", 2)
    for rule in [
        "เอกสารหนึ่งรายการ trace ได้ครบ impact_process_id -> doc_no -> instance_id -> task_id",
        "กฎ threshold 100,000 บาทใช้กับชั้นอนุมัติถูกต้องทั้งค่าต่ำกว่า เท่ากับ และสูงกว่า",
        "หน้า Document Detail แสดง visible/editable/action options ตาม role profile ของผู้ใช้จริงและไม่มี role switcher ใน production",
        "ผลรวม % ชดเชย 100% ถูกตรวจทั้ง FE และ BE",
        "ร้านยอดขายไม่ครบ 60 วันถูก flag ใน inbox/report และมีเหตุผลตรวจสอบย้อนกลับ",
        "Jobs 1-10/8b รันซ้ำตาม runbook โดยไม่สร้างข้อมูลซ้ำหรือสูญหาย",
        f"API capability {endpoint_total} endpoints ใน scope ต้องผ่าน authorization, validation, audit, duplicate guard/idempotency, pagination และ error-contract test; Auth Group 1 เป็น platform service",
        "ข้อมูล export/import ทุก interface ผ่าน golden-file test เรื่อง encoding/date/delimiter/field count",
        "หน้าจอรายงานและ CSV Export to Batch ให้ผลตรงกันภายใต้ filter เดียวกัน",
    ]:
        model.bullet(rule)
    model.heading("5.2 Traceability matrix", 2)
    trace_rows = [
        ["REQ-BUS-001/002", "Impact radius", "กฎรัศมี 1 กม. กรุงเทพฯ/ปริมณฑล และ 2 กม. ต่างจังหวัด", "3.0, 3.1.2, Job 2"],
        ["REQ-BUS-003", "Gen Flow Gate", "gate SQL และผล Y/W/N", "3.0, 3.1.2, Job 8b"],
        ["REQ-BUS-004", "Abnormal sales", "เกณฑ์ 60 วันและแถวผิดปกติ", "3.0, SCR-03/04/07"],
        ["REQ-BUS-005", "Allocation", "ผลรวมเปอร์เซ็นต์ชดเชยเท่ากับ 100%", "3.0, SCR-06, Job 9"],
        ["REQ-BUS-006", "Approval threshold", "routing ที่ 100,000 บาท", "3.0, 3.1.3, SCR-06"],
        ["REQ-DOC-001/002/003", "Document integrity", "เลขเอกสาร duplicate guard และ data spine", "3.0, 3.2, SCR-02/06"],
        ["REQ-WFL-001/002/003", "Workflow integrity", "ownership, audit และ optimistic concurrency", "3.0, 3.1.3, 3.2"],
        ["REQ-INT-001/002/003", "Interface reliability", "durable file/outbox, ACK/purge และ typed FK", "3.0, 3.2.4, 3.3"],
        ["REQ-SEC-001/002", "Identity and secrets", "platform identity, Secret Manager และ TLS", "1.5, 3.0, 4.2"],
        ["REQ-FIL-001", "Attachment", "5 MB, type/AV scan และ authorization", "3.0, SCR-06, 3.5"],
        ["REQ-RPT-001", "Report export", "19 columns และ preview/export reconciliation", "3.0, SCR-07"],
        ["REQ-OPS-001", "Batch rerun", "idempotency และ run reconciliation", "3.0, 3.3"],
        ["REQ-SCR-001", "Committed screens", "SCR-01..04 และ SCR-06..11", "3.4"],
        ["SYS-API-001", "API capability", f"{endpoint_total} endpoints / {api_group_total} groups", "3.5"],
        ["SYS-DAT-001", "Data model", "34 tables and integrity controls", "3.2"],
        ["SYS-NFR-001", "Observability", "correlation/metrics/alert/audit evidence", "4"],
        ["FLOW-01", "Batch pipeline", "ขั้นตอนนำเข้า คำนวณ สร้างเอกสาร ส่ง Statement และติดตาม ACK", "3.1, 3.3"],
        ["FLOW-02", "Approval workflow", "Section 06 -> 08 -> 01 -> 02 และ Section 03 ตามวงเงิน", "3.1.1, 3.1.3"],
        ["DATA-01", "Logical data model", "Data subjects, relationships, controls และ remediation", "3.2"],
        ["JOB-01", "Batch Job Console", "11 entry points, common controls และผลลัพธ์ที่ตรวจรับได้", "3.3"],
        ["K2-01", "Overview / Dashboard", "Dashboard", "SCR-01"],
        ["K2-02", "Create Document", "Create document", "SCR-02"],
        ["K2-03", "Task Inbox", "Task inbox", "SCR-03"],
        ["K2-04", "Related Documents", "Related documents", "SCR-04"],
        ["K2-05", "Abnormal Assignment", "Abnormal assignment", "SCR-05 / OPEN"],
        ["K2-06", "Document Detail", "Document detail/action", "SCR-06"],
        ["K2-07", "Status Report", "Status report", "SCR-07"],
        ["K2-08", "Operator Master", "Operator master", "SCR-08"],
        ["K2-09", "External Factor Master", "External factor master", "SCR-09"],
        ["K2-10", "RBAC Matrix", "RBAC matrix", "SCR-10"],
        ["K2-11", "Global Config", "Global system configuration", "SCR-11"],
        ["EMAIL-01", "Email Template", "หน้าจอผู้ดูแล template และกฎ Notification Service", "3.4.13"],
        ["API-01", "REST API", f"Capability catalog {endpoint_total} endpoints และข้อกำหนด contract กลาง", "3.5"],
    ]
    model.table(["ID", "Requirement area", "Scope coverage", "SRS section"], trace_rows, [0.14, 0.28, 0.4, 0.18])

    model.pagebreak()
    model.heading("6. Decisions and Open Items", 1)
    model.para(
        "หัวข้อนี้แยกมติที่ปิดแล้วออกจากประเด็นที่ยังเปิด เพื่อให้ทีมพัฒนาไม่ต้องอนุมานจากรายละเอียดเชิงออกแบบ "
        "รายการ CLOSED ถือเป็น baseline ของ SRS ฉบับนี้ ส่วน OPEN ยังห้ามนำไปพัฒนาเป็นข้อยุติโดยอัตโนมัติ"
    )
    model.heading("6.1 Closed decisions", 2)
    closed_items = [
        ["OPEN-01", "CLOSED", "22/07/2026", "เลขเอกสารใช้ปี พ.ศ. รูป YYYY/xxxxx และเก็บ be_year/running_no เพื่อ uniqueness; วันที่/เดือนใน API และฐานข้อมูลเชิงเวลาใช้ ISO-8601 ปี ค.ศ.; FE แปลงเป็น พ.ศ. เฉพาะการแสดงผล"],
        ["OPEN-03", "CLOSED", "22/07/2026", "Job 8b ใช้ event/dependency trigger หลัง Job 8 สร้างเอกสารสำเร็จ ไม่ใช้เวลา wall-clock คงที่; Operations สั่ง manual rerun ตาม period ได้ โดยใช้ run lock และ idempotency key เดิม"],
    ]
    model.table(["ID", "Status", "Effective date", "Baseline decision"], closed_items, [0.1, 0.12, 0.16, 0.62])
    model.heading("6.2 Open decisions required", 2)
    model.para("รายการต่อไปนี้ยังไม่ถือเป็น requirement ที่อนุมัติ เมื่อได้ข้อยุติต้องบันทึกผล วันที่มีผล และปรับ baseline ก่อนพัฒนาส่วนที่เกี่ยวข้อง")
    open_items = [
        ["OPEN-02", "Abnormal screen", "ตัดสินใจ keep/drop หน้าจอและ API 2 เส้น พร้อมปรับ role 05", "ขอบเขต FE/BE, API และ UAT"],
        ["OPEN-04", "NULL growth_rate", "อนุมัติรอตรวจสอบแทน auto-accept หรือกำหนดกฎใหม่", "การคัดรายการและ workflow generation"],
        ["OPEN-05", "Legacy date routing", "ยืนยันเงื่อนไข routing สำหรับร้านก่อน/หลัง 1/10/2557", "routing และผลพิจารณา"],
        ["OPEN-06", "NFR SLA/RPO/RTO", "กำหนด SLA API/report/batch และ RPO/RTO production", "capacity, HA, backup และ acceptance"],
        ["OPEN-07", "File retention", "กำหนด retention, encryption และ purge ของ attachment/interface/archive", "storage, compliance และ recovery"],
        ["OPEN-08", "Permission matrix", "ยืนยัน menu/master/record permission ต่อ role", "sidebar, API authorization และ UAT"],
    ]
    model.table(["ID", "Topic", "Decision required", "Impact if unresolved"], open_items, [0.1, 0.18, 0.44, 0.28])

    model.pagebreak()
    model.heading("7. Appendices", 1)
    model.para("ภาคผนวกรวบรวมคำย่อและหลักการระบุ requirement เพื่อให้ business, development และ test evidence ใช้ความหมายเดียวกัน")
    model.heading("7.1 Definitions and abbreviations", 2)
    model.table(
        ["Term", "Definition"],
        [
            ["SBPGI", "Target integrated system for FGI/FCS processing and K2-style documents/workflow"],
            ["SP / Store Partner", "ร้าน Franchise ที่อยู่ในขอบเขตประกันรายได้"],
            ["FGI/FCS", "Legacy batch domains for impact and QSSI data"],
            ["K2", "Legacy BPM/workflow platform and original K2 SRS scope"],
            ["STA", "Statement/accounting interface system"],
            ["IAS/MIS", "Sales data interface"],
            ["QSSI", "Monthly score source"],
            ["ALLMAP", "Store/competitor/map source"],
            ["Gen Flow Gate", "ชุดเงื่อนไขก่อนสร้าง/เปิด workflow"],
        ],
        [0.24, 0.76],
    )
    model.heading("7.2 Requirement conventions", 2)
    for convention in [
        "REQ ใช้กับข้อกำหนดเชิงหน้าที่และกฎธุรกิจที่ต้องทดสอบได้",
        "SYS ใช้กับข้อกำหนดร่วมด้านสถาปัตยกรรม ข้อมูล ความปลอดภัย และการปฏิบัติการ",
        "PROTO ระบุข้อมูลหรือพฤติกรรมตัวอย่างที่ใช้ยืนยัน UX แต่ไม่ใช่ข้อมูล Production",
        "OPEN ระบุประเด็นที่ยังไม่อนุมัติและต้องไม่ถูกนำไปพัฒนาเป็นข้อยุติโดยอัตโนมัติ",
        "Acceptance evidence ต้องเชื่อมกลับมายัง section หรือ requirement area ในตาราง traceability",
    ]:
        model.bullet(convention)
    return model


def make_md(model: Model):
    lines = [
        "# SOFTWARE REQUIREMENT SPECIFICATION",
        "",
        "## ระบบประกันรายได้ SBPGI",
        "",
        f"Version {DOC_VERSION}",
        "",
        "> เอกสารฉบับนี้เป็น baseline แบบ self-contained สำหรับการพัฒนา ทดสอบ และตรวจรับระบบ",
        "",
    ]
    for block in model.blocks:
        if block.kind == "heading":
            lines.extend(["", f"{'#' * min(6, block.level)} {block.text}", ""])
        elif block.kind == "paragraph":
            lines.extend([block.text, ""])
        elif block.kind == "bullet":
            lines.append(f"- {block.text}")
        elif block.kind == "number":
            lines.append(f"1. {block.text}")
        elif block.kind == "note":
            lines.extend([f"> {block.text}", ""])
        elif block.kind == "code":
            lines.extend(["```text", block.text, "```", ""])
        elif block.kind == "table":
            lines.append("| " + " | ".join(block.headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(block.headers)) + " |")
            for row in block.rows:
                cells = [cell.replace("|", "\\|").replace("\n", "<br>") for cell in row]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
        elif block.kind == "image" and block.path:
            lines.extend([f"![{block.caption}]({block.path.name})", ""])
        elif block.kind == "pagebreak":
            lines.extend(["", "---", ""])
    MD_OUT.parent.mkdir(parents=True, exist_ok=True)
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def docx_set_font(run, name="Cordia New", size: float | None = None, bold: bool | None = None, color: RGBColor | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, total_twips: int, proportions: list[float] | None):
    cols = len(table.columns)
    if proportions is None or len(proportions) != cols:
        proportions = [1 / cols] * cols
    norm = sum(proportions)
    widths = [round(total_twips * x / norm) for x in proportions]
    widths[-1] += total_twips - sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, side: str, color: str, size: int = 8, space: int = 1):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def new_docx_numbering_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "300")
    p_pr.append(tabs)
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_docx_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def configure_docx_styles(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.8)
    sec.bottom_margin = Cm(2.8)
    sec.left_margin = Cm(1.75)
    sec.right_margin = Cm(1.75)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.75)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Cordia New"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Cordia New")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Cordia New")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Cordia New")
    normal.font.size = Pt(14)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.widow_control = True

    for name, size, before, after in [
        ("Heading 1", 18, 12, 6),
        ("Heading 2", 16, 10, 5),
        ("Heading 3", 15, 8, 4),
        ("Heading 4", 14, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Cordia New"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Cordia New")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Cordia New")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Cordia New")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    for name in ["List Bullet", "List Bullet 2", "List Number", "List Number 2"]:
        style = styles[name]
        style.font.name = "Cordia New"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Cordia New")
        style.font.size = Pt(14)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.keep_together = False
        style.paragraph_format.widow_control = True

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def build_docx(model: Model):
    doc = Document()
    configure_docx_styles(doc)
    section = doc.sections[0]

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Cm(17.5))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    set_table_widths(header_table, 9920, [0.5, 0.5])
    if HEADER_LOGO.exists():
        p = header_table.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(HEADER_LOGO), width=Cm(3.95))
    p = header_table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(TEMPLATE_CODE)
    docx_set_font(r, size=10)
    for cell in header_table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)
    hp = header.add_paragraph()
    hp.paragraph_format.space_after = Pt(0)
    set_paragraph_border(hp, "bottom", "FF0000", 8, 0)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(2)
    set_paragraph_border(fp, "top", "FF0000", 8, 0)
    ft = footer.add_table(rows=3, cols=2, width=Cm(17.5))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    ft.autofit = False
    set_table_widths(ft, 9920, [0.68, 0.32])
    left = [
        "Confidential - Restricted Circulation Only",
        "© 2023 Gosoft (Thailand) Co., Ltd prepared for internal use only",
        "Printed copies of this Document are not controlled and will not be updated.",
    ]
    for idx, text in enumerate(left):
        p = ft.cell(idx, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        docx_set_font(r, size=8.5, color=RGBColor(255, 0, 0) if idx == 0 else RGBColor(50, 50, 50))
        p = ft.cell(idx, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        if idx == 0:
            r = p.add_run(TEMPLATE_CODE)
            docx_set_font(r, size=8.5, color=RGBColor(30, 70, 255))
        elif idx == 2:
            r = p.add_run("Page ")
            docx_set_font(r, size=8.5)
            add_field(p, "PAGE")
            r = p.add_run(" of ")
            docx_set_font(r, size=8.5)
            add_field(p, "NUMPAGES")
    for row in ft.rows:
        for cell in row.cells:
            set_cell_margins(cell, 0, 0, 0, 0)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SOFTWARE REQUIREMENT SPECIFICATION")
    docx_set_font(r, size=22)
    p.paragraph_format.space_after = Pt(26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ระบบประกันรายได้ SBPGI")
    docx_set_font(r, size=20, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"VERSION {DOC_VERSION.upper()}")
    docx_set_font(r, size=17)
    p.paragraph_format.space_after = Pt(42)
    if COVER_BADGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(COVER_BADGE), width=Cm(5.2))
    for _ in range(4):
        doc.add_paragraph()
    for line in [
        "Gosoft (Thailand) CO.,LTD.",
        "313 CP Tower, 24th Fl., Silom Road",
        "Bangrak, Bangkok 10500.",
        "Website: www.gosoft.co.th",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        docx_set_font(r, size=12)
    doc.add_page_break()

    # Version history
    doc.add_heading("Document Version History", level=1)
    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    headers = ["Version Number", "Release Date", "Created By", "Detail", "Reviewed by", "Authorized by"]
    values = [
        DOC_VERSION,
        RELEASE_DATE,
        "SBPGI Project Team",
        "Initial integrated SBPGI baseline for SBP Mall scope, internal workflow, batch, API and operational requirements",
        "",
        "",
    ]
    for i, value in enumerate(headers):
        table.cell(0, i).text = value
    for i, value in enumerate(values):
        table.cell(1, i).text = value
    set_table_widths(table, 9920, [0.11, 0.12, 0.15, 0.34, 0.14, 0.14])
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx == 0:
                set_cell_shading(cell, "E7E7E7")
            for p in cell.paragraphs:
                for r in p.runs:
                    docx_set_font(r, size=11, bold=(ridx == 0))
    doc.add_page_break()

    doc.add_heading("Sign Off", level=1)
    sign = doc.add_table(rows=6, cols=4)
    sign.style = "Table Grid"
    sign_rows = [
        ["Organization", "Role", "Name / Signature", "Date"],
        ["Gosoft", "Reviewed By - Service Manager", "", ""],
        ["Gosoft", "Approved By - Project Manager", "", ""],
        ["Client", "Reviewed By", "", ""],
        ["Client", "Approved By", "", ""],
        ["Architecture/Operations", "Reviewed By", "", ""],
    ]
    for i, row in enumerate(sign_rows):
        for j, value in enumerate(row):
            sign.cell(i, j).text = value
    set_table_widths(sign, 9920, [0.22, 0.32, 0.28, 0.18])
    set_repeat_table_header(sign.rows[0])
    for ridx, row in enumerate(sign.rows):
        for cell in row.cells:
            set_cell_margins(cell, 120, 100, 240 if ridx else 100, 100)
            if ridx == 0:
                set_cell_shading(cell, "E7E7E7")
            for p in cell.paragraphs:
                for r in p.runs:
                    docx_set_font(r, size=11, bold=(ridx == 0))
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("รายการหัวข้อหลักของ SRS; หัวข้อระดับย่อยทั้งหมดแสดงใน Navigation Pane ของโปรแกรมเอกสาร")
    docx_set_font(r, size=12)
    for block in model.blocks:
        if block.kind != "heading" or block.level > 2:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55 if block.level == 2 else 0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_together = True
        r = p.add_run(block.text)
        docx_set_font(r, size=11.5, bold=(block.level == 1))
    doc.add_page_break()

    current_docx_num_id: int | None = None
    for block in model.blocks:
        if block.kind != "number":
            current_docx_num_id = None
        if block.kind == "heading":
            p = doc.add_heading(block.text, level=min(4, block.level))
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            p.paragraph_format.widow_control = True
        elif block.kind == "paragraph":
            p = doc.add_paragraph()
            if block.text in {"Request", "Response", "Errors", "Flow:", "Reference table access:"}:
                p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = False
            p.paragraph_format.widow_control = True
            r = p.add_run(block.text)
            docx_set_font(r, size=14)
        elif block.kind == "bullet":
            style = "List Bullet 2" if block.level else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.keep_together = False
            p.paragraph_format.widow_control = True
            r = p.add_run(block.text)
            docx_set_font(r, size=14)
        elif block.kind == "number":
            if current_docx_num_id is None:
                current_docx_num_id = new_docx_numbering_id(doc)
            p = doc.add_paragraph()
            apply_docx_numbering(p, current_docx_num_id)
            p.paragraph_format.keep_together = False
            p.paragraph_format.widow_control = True
            r = p.add_run(block.text)
            docx_set_font(r, size=14)
        elif block.kind == "note":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.right_indent = Cm(0.15)
            p.paragraph_format.keep_together = True
            p.paragraph_format.widow_control = True
            set_paragraph_border(p, "left", "FF0000", 18, 4)
            r = p.add_run(block.text)
            docx_set_font(r, size=13.5)
            set_cell_shading  # keep linter quiet about helper locality
        elif block.kind == "code":
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            set_table_widths(table, 9920, [1.0])
            cell = table.cell(0, 0)
            set_cell_shading(cell, "F2F2F2")
            set_cell_margins(cell, 100, 120, 100, 120)
            set_row_cant_split(table.rows[0])
            p = cell.paragraphs[0]
            p.paragraph_format.keep_together = True
            p.paragraph_format.widow_control = True
            for idx, line in enumerate(block.text.splitlines() or [""]):
                if idx:
                    p.add_run().add_break()
                r = p.add_run(line)
                docx_set_font(r, name="Arial", size=8.5)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        elif block.kind == "table":
            if not block.headers:
                continue
            cols = len(block.headers)
            table = doc.add_table(rows=1, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for idx, header in enumerate(block.headers):
                table.cell(0, idx).text = header
            for row_data in block.rows:
                row = table.add_row()
                for idx in range(cols):
                    row.cells[idx].text = row_data[idx] if idx < len(row_data) else ""
            set_table_widths(table, 9920, block.widths)
            set_repeat_table_header(table.rows[0])
            for ridx, row in enumerate(table.rows):
                set_row_cant_split(row)
                for cidx, cell in enumerate(row.cells):
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell)
                    if ridx == 0:
                        set_cell_shading(cell, "E7E7E7")
                    for p in cell.paragraphs:
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.keep_together = True
                        p.paragraph_format.widow_control = True
                        for r in p.runs:
                            docx_set_font(r, size=10.5, bold=(ridx == 0))
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(1)
        elif block.kind == "image" and block.path and block.path.exists():
            with PILImage.open(block.path) as img:
                ratio = img.height / img.width
            width = Cm(17.0)
            height = width * ratio
            if height > Cm(19.5):
                height = Cm(19.5)
                width = height / ratio
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.add_run().add_picture(str(block.path), width=width, height=height)
            if block.caption:
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.keep_together = True
                r = cp.add_run(block.caption)
                docx_set_font(r, size=11)
        elif block.kind == "pagebreak":
            doc.add_page_break()

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Software Requirement Specification - ระบบประกันรายได้ SBPGI"
    doc.core_properties.subject = "Self-contained SBP Mall software requirements"
    doc.core_properties.author = "SBPGI Project Team"
    doc.core_properties.comments = "Self-contained SRS baseline; open items require approval before implementation."
    doc.save(DOCX_OUT)


PDF_FONT = "Tahoma"
PDF_FONT_BOLD = "Tahoma-Bold"
PDF_BODY = 9.0
PAGE_W, PAGE_H = A4
LEFT = 18 * mm
RIGHT = 18 * mm
TOP = 26 * mm
BOTTOM = 28 * mm
CONTENT_W = PAGE_W - LEFT - RIGHT


def register_pdf_fonts():
    regular = Path("/System/Library/Fonts/Supplemental/Tahoma.ttf")
    bold = Path("/System/Library/Fonts/Supplemental/Tahoma Bold.ttf")
    fallback = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")
    if regular.exists() and bold.exists():
        pdfmetrics.registerFont(TTFont(PDF_FONT, str(regular)))
        pdfmetrics.registerFont(TTFont(PDF_FONT_BOLD, str(bold)))
    elif fallback.exists():
        pdfmetrics.registerFont(TTFont(PDF_FONT, str(fallback)))
        pdfmetrics.registerFont(TTFont(PDF_FONT_BOLD, str(fallback)))
    else:
        raise RuntimeError("Thai-capable TrueType font not found")


def ptext(text: str) -> str:
    return html_lib.escape(clean(text)).replace("\n", "<br/>")


class SrsDocTemplate(BaseDocTemplate):
    def __init__(self, filename, **kwargs):
        super().__init__(filename, **kwargs)
        frame = self._make_frame()
        self.addPageTemplates([PageTemplate(id="SRS", frames=[frame], onPage=draw_pdf_furniture)])

    def _make_frame(self):
        from reportlab.platypus import Frame

        return Frame(LEFT, BOTTOM, CONTENT_W, PAGE_H - TOP - BOTTOM, id="normal")

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph):
            style_name = flowable.style.name
            if style_name.startswith("Heading"):
                level = int(style_name.replace("Heading", "")) - 1
                text = flowable.getPlainText()
                key = f"h{self.seq.nextf('heading')}"
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(text, key, level=level, closed=False)
                self.notify("TOCEntry", (level, text, self.page, key))


def draw_pdf_furniture(canv: canvas.Canvas, doc):
    if doc.page == 1:
        return
    canv.saveState()
    if HEADER_LOGO.exists():
        canv.drawImage(str(HEADER_LOGO), LEFT, PAGE_H - 18 * mm, width=39 * mm, height=7.3 * mm, preserveAspectRatio=True, mask="auto")
    canv.setFont(PDF_FONT, 7.3)
    canv.setFillColor(colors.HexColor("#333333"))
    canv.drawRightString(PAGE_W - RIGHT, PAGE_H - 13.5 * mm, TEMPLATE_CODE)
    canv.setStrokeColor(colors.red)
    canv.setLineWidth(0.6)
    canv.line(LEFT, PAGE_H - 20.2 * mm, PAGE_W - RIGHT, PAGE_H - 20.2 * mm)

    y = 18.2 * mm
    canv.line(LEFT, y + 3.2 * mm, PAGE_W - RIGHT, y + 3.2 * mm)
    canv.setFont(PDF_FONT, 6.6)
    canv.setFillColor(colors.red)
    canv.drawString(LEFT, y, "Confidential - Restricted Circulation Only")
    canv.setFillColor(colors.HexColor("#2346FF"))
    canv.drawRightString(PAGE_W - RIGHT, y, TEMPLATE_CODE)
    canv.setFillColor(colors.HexColor("#444444"))
    canv.setFont(PDF_FONT, 6.3)
    canv.drawString(LEFT, y - 5.0 * mm, "© 2023 Gosoft (Thailand) Co., Ltd prepared for internal use only")
    canv.drawString(LEFT, y - 9.2 * mm, "Printed copies of this Document are not controlled and will not be updated.")
    canv.restoreState()


class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self._saved_page_states)
        for page_num, state in enumerate(self._saved_page_states, 1):
            self.__dict__.update(state)
            if page_num > 1:
                self.setFont(PDF_FONT, 6.5)
                self.setFillColor(colors.HexColor("#444444"))
                self.drawRightString(PAGE_W - RIGHT, 9.0 * mm, f"Page {page_num} of {page_count}")
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)


def pdf_styles():
    styles = getSampleStyleSheet()
    common = dict(fontName=PDF_FONT, textColor=colors.black, wordWrap="CJK")
    out = {
        "Body": ParagraphStyle("Body", parent=styles["Normal"], fontSize=PDF_BODY, leading=12, spaceAfter=4, **common),
        "Bullet": ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=PDF_BODY, leading=12, leftIndent=12, firstLineIndent=-7, bulletIndent=0, spaceAfter=2.5, **common),
        "Number": ParagraphStyle("Number", parent=styles["Normal"], fontSize=PDF_BODY, leading=12, leftIndent=12, firstLineIndent=-7, spaceAfter=2.5, **common),
        "Note": ParagraphStyle("Note", parent=styles["Normal"], fontSize=8.6, leading=11.5, leftIndent=8, rightIndent=4, spaceBefore=4, spaceAfter=6, borderColor=colors.red, borderWidth=0.7, borderPadding=5, backColor=colors.HexColor("#FAFAFA"), **common),
        "Code": ParagraphStyle("Code", parent=styles["Code"], fontName=PDF_FONT, fontSize=6.7, leading=8.5, wordWrap="CJK", textColor=colors.HexColor("#222222")),
        "Label": ParagraphStyle("Label", parent=styles["Normal"], fontName=PDF_FONT_BOLD, fontSize=PDF_BODY, leading=12, spaceBefore=3, spaceAfter=3, keepWithNext=True, textColor=colors.black, wordWrap="CJK"),
        "Heading1": ParagraphStyle("Heading1", parent=styles["Heading1"], fontName=PDF_FONT_BOLD, fontSize=13, leading=16, spaceBefore=9, spaceAfter=5, keepWithNext=True, wordWrap="CJK"),
        "Heading2": ParagraphStyle("Heading2", parent=styles["Heading2"], fontName=PDF_FONT_BOLD, fontSize=11.5, leading=14, spaceBefore=8, spaceAfter=4, keepWithNext=True, wordWrap="CJK"),
        "Heading3": ParagraphStyle("Heading3", parent=styles["Heading3"], fontName=PDF_FONT_BOLD, fontSize=10.2, leading=13, spaceBefore=7, spaceAfter=3, keepWithNext=True, wordWrap="CJK"),
        "Heading4": ParagraphStyle("Heading4", parent=styles["Heading4"], fontName=PDF_FONT_BOLD, fontSize=9.4, leading=12, spaceBefore=5, spaceAfter=2.5, keepWithNext=True, wordWrap="CJK"),
        "Table": ParagraphStyle("Table", parent=styles["Normal"], fontName=PDF_FONT, fontSize=6.8, leading=8.6, wordWrap="CJK", textColor=colors.HexColor("#222222")),
        "TableHead": ParagraphStyle("TableHead", parent=styles["Normal"], fontName=PDF_FONT_BOLD, fontSize=7, leading=8.8, alignment=TA_CENTER, wordWrap="CJK"),
        "Small": ParagraphStyle("Small", parent=styles["Normal"], fontName=PDF_FONT, fontSize=7.4, leading=9.5, wordWrap="CJK"),
        "CoverTitle": ParagraphStyle("CoverTitle", parent=styles["Title"], fontName=PDF_FONT, fontSize=17, leading=21, alignment=TA_CENTER, spaceAfter=18),
        "CoverThai": ParagraphStyle("CoverThai", parent=styles["Title"], fontName=PDF_FONT_BOLD, fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=6),
        "CoverVersion": ParagraphStyle("CoverVersion", parent=styles["Title"], fontName=PDF_FONT, fontSize=13, leading=17, alignment=TA_CENTER, spaceAfter=30),
        "Center": ParagraphStyle("Center", parent=styles["Normal"], fontName=PDF_FONT, fontSize=9.5, leading=13, alignment=TA_CENTER, wordWrap="CJK"),
    }
    return out


def pdf_table(block: Block, st) -> LongTable:
    cols = len(block.headers)
    proportions = block.widths if block.widths and len(block.widths) == cols else [1 / cols] * cols
    total = sum(proportions)
    widths = [CONTENT_W * x / total for x in proportions]
    data = [[Paragraph(ptext(h), st["TableHead"]) for h in block.headers]]
    for row in block.rows:
        data.append([Paragraph(ptext(row[i] if i < len(row) else ""), st["Table"]) for i in range(cols)])
    table = LongTable(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E7E7E7")),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3.5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3.5),
                ("TOPPADDING", (0, 0), (-1, -1), 3.2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3.2),
            ]
        )
    )
    return table


def build_pdf(model: Model):
    register_pdf_fonts()
    st = pdf_styles()
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = SrsDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=LEFT,
        rightMargin=RIGHT,
        topMargin=TOP,
        bottomMargin=BOTTOM,
        title="Software Requirement Specification - ระบบประกันรายได้ SBPGI",
        author="SBPGI Project Team",
        subject="Self-contained SBP Mall software requirements",
    )
    story: list[Flowable] = []

    # Cover
    story.extend([Spacer(1, 39 * mm), Paragraph("SOFTWARE REQUIREMENT SPECIFICATION", st["CoverTitle"])])
    story.append(Paragraph("ระบบประกันรายได้ SBPGI", st["CoverThai"]))
    story.append(Paragraph(f"VERSION {DOC_VERSION.upper()}", st["CoverVersion"]))
    if COVER_BADGE.exists():
        story.append(Image(str(COVER_BADGE), width=48 * mm, height=49 * mm, hAlign="CENTER"))
    story.append(Spacer(1, 37 * mm))
    for line in [
        "Gosoft (Thailand) CO.,LTD.",
        "313 CP Tower, 24th Fl., Silom Road",
        "Bangrak, Bangkok 10500.",
        "Website: www.gosoft.co.th",
    ]:
        story.append(Paragraph(line, st["Center"]))
    story.append(PageBreak())

    story.append(Paragraph("Document Version History", st["Heading1"]))
    version_block = Block(
        "table",
        headers=["Version Number", "Release Date", "Created By", "Detail", "Reviewed by", "Authorized by"],
        rows=[
            [
                DOC_VERSION,
                RELEASE_DATE,
                "SBPGI Project Team",
                "Initial integrated SBPGI baseline for SBP Mall scope, internal workflow, batch, API and operational requirements",
                "",
                "",
            ]
        ],
        widths=[0.11, 0.12, 0.15, 0.34, 0.14, 0.14],
    )
    story.append(pdf_table(version_block, st))
    story.append(PageBreak())

    story.append(Paragraph("Sign Off", st["Heading1"]))
    sign_block = Block(
        "table",
        headers=["Organization", "Role", "Name / Signature", "Date"],
        rows=[
            ["Gosoft", "Reviewed By - Service Manager", "", ""],
            ["Gosoft", "Approved By - Project Manager", "", ""],
            ["Client", "Reviewed By", "", ""],
            ["Client", "Approved By", "", ""],
            ["Architecture/Operations", "Reviewed By", "", ""],
        ],
        widths=[0.22, 0.32, 0.28, 0.18],
    )
    story.append(pdf_table(sign_block, st))
    story.append(PageBreak())

    story.append(Paragraph("Table of Contents", st["Heading1"]))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle(name="TOC1", fontName=PDF_FONT, fontSize=9, leading=13, leftIndent=0, firstLineIndent=0),
        ParagraphStyle(name="TOC2", fontName=PDF_FONT, fontSize=8.5, leading=12, leftIndent=12, firstLineIndent=0),
        ParagraphStyle(name="TOC3", fontName=PDF_FONT, fontSize=8, leading=11, leftIndent=24, firstLineIndent=0),
        ParagraphStyle(name="TOC4", fontName=PDF_FONT, fontSize=7.5, leading=10, leftIndent=36, firstLineIndent=0),
    ]
    story.append(toc)
    story.append(PageBreak())

    pending_code_label: str | None = None
    pending_table_label: str | None = None
    section_labels = {"Errors", "Flow:", "Reference table access:"}
    pdf_number_counter = 0
    for block in model.blocks:
        if block.kind != "number":
            pdf_number_counter = 0
        if pending_code_label is not None and block.kind != "code":
            story.append(Paragraph(ptext(pending_code_label), st["Body"]))
            pending_code_label = None
        if pending_table_label is not None and block.kind != "table":
            story.append(Paragraph(ptext(pending_table_label), st["Label"]))
            pending_table_label = None
        if block.kind == "heading":
            min_space = {1: 32 * mm, 2: 25 * mm, 3: 18 * mm, 4: 14 * mm}.get(min(4, block.level), 18 * mm)
            story.append(CondPageBreak(min_space))
            story.append(Paragraph(ptext(block.text), st[f"Heading{min(4, block.level)}"]))
        elif block.kind == "paragraph":
            if block.text in {"Request", "Response"}:
                pending_code_label = block.text
            elif block.text == "Reference table access:":
                pending_table_label = block.text
            elif block.text in section_labels:
                story.append(CondPageBreak(24 * mm))
                story.append(Paragraph(ptext(block.text), st["Label"]))
            else:
                story.append(Paragraph(ptext(block.text), st["Body"]))
        elif block.kind == "bullet":
            story.append(Paragraph(ptext(block.text), st["Bullet"], bulletText="•"))
        elif block.kind == "number":
            pdf_number_counter += 1
            story.append(Paragraph(ptext(f"{pdf_number_counter}. {block.text}"), st["Number"], bulletText=""))
        elif block.kind == "note":
            story.append(Paragraph(ptext(block.text), st["Note"]))
        elif block.kind == "code":
            lines = []
            for line in block.text.splitlines() or [""]:
                wrapped = textwrap.wrap(line, width=100, replace_whitespace=False, drop_whitespace=False) or [""]
                lines.extend(wrapped)
            code_para = Paragraph("<br/>".join(html_lib.escape(line) for line in lines), st["Code"])
            code_table = Table([[code_para]], colWidths=[CONTENT_W])
            code_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F2F2F2")),
                        ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#888888")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            code_items: list[Flowable] = []
            if pending_code_label is not None:
                code_items.append(Paragraph(ptext(pending_code_label), st["Body"]))
                pending_code_label = None
            code_items.extend([code_table, Spacer(1, 4)])
            story.append(KeepTogether(code_items))
        elif block.kind == "table":
            if pending_table_label is not None:
                story.append(CondPageBreak(20 * mm))
                story.append(KeepTogether([Paragraph(ptext(pending_table_label), st["Label"]), pdf_table(block, st), Spacer(1, 5)]))
                pending_table_label = None
            else:
                story.append(CondPageBreak(16 * mm))
                story.extend([pdf_table(block, st), Spacer(1, 5)])
        elif block.kind == "image" and block.path and block.path.exists():
            with PILImage.open(block.path) as img:
                ratio = img.height / img.width
            width = CONTENT_W
            height = width * ratio
            if height > 175 * mm:
                height = 175 * mm
                width = height / ratio
            items: list[Flowable] = [Image(str(block.path), width=width, height=height, hAlign="CENTER")]
            if block.caption:
                items.append(Paragraph(ptext(block.caption), st["Center"]))
            story.append(KeepTogether(items))
        elif block.kind == "pagebreak":
            story.append(PageBreak())
    if pending_code_label is not None:
        story.append(Paragraph(ptext(pending_code_label), st["Body"]))
    doc.multiBuild(story, canvasmaker=NumberedCanvas)


def main():
    if not DATA_FILE.exists():
        raise SystemExit(f"Missing {DATA_FILE}; run: node tmp/extract_js_data.mjs")
    if not HEADER_LOGO.exists() or not COVER_BADGE.exists():
        raise SystemExit("Missing extracted PDF template images; run pdfimages first")
    OUT.mkdir(parents=True, exist_ok=True)
    model = scrub_srs_model(build_model())
    make_md(model)
    build_docx(model)
    build_pdf(model)
    print(DOCX_OUT)
    print(PDF_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
